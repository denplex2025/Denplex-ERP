from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, Query, Request, Response, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import bcrypt
import jwt
import uuid
import base64
import io
import csv
import asyncio
import secrets
import pyotp
import phonenumbers
import httpx
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr, ConfigDict
from typing import List, Optional, Literal, Any, Dict
from datetime import datetime, timezone, timedelta
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from twilio.rest import Client as TwilioClient
from twilio.base.exceptions import TwilioRestException
from email.message import EmailMessage
from email.utils import parseaddr, parsedate_to_datetime
from email.header import decode_header
import smtplib
import imaplib
import email as emaillib
import re
import ssl
import hashlib
from bson import ObjectId
from io import BytesIO
from cryptography.fernet import Fernet, InvalidToken
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# ---------------- PDF font registration ----------------
# Helvetica (reportlab default) lacks proper rupee glyph + many UTF-8 symbols.
# Try DejaVuSans (present on most Ubuntu/Debian + Railway), then a bundled
# fallback in backend/fonts/, then give up and use Helvetica.
_PDF_FONT_REGULAR = "Helvetica"
_PDF_FONT_BOLD = "Helvetica-Bold"

def _ensure_font_files_on_disk():
    """Best-effort: if backend/fonts/DejaVu*.ttf is missing, fetch from CDN.
    Runs once at startup; idempotent. Failures are non-fatal (we fall back to Helvetica)."""
    try:
        import urllib.request
        target_dir = Path(__file__).parent / "fonts"
        target_dir.mkdir(parents=True, exist_ok=True)
        urls = {
            "DejaVuSans.ttf": "https://cdn.jsdelivr.net/npm/dejavu-fonts-ttf@2.37.3/ttf/DejaVuSans.ttf",
            "DejaVuSans-Bold.ttf": "https://cdn.jsdelivr.net/npm/dejavu-fonts-ttf@2.37.3/ttf/DejaVuSans-Bold.ttf",
        }
        for name, url in urls.items():
            tgt = target_dir / name
            if tgt.exists() and tgt.stat().st_size > 100_000:
                continue
            try:
                urllib.request.urlretrieve(url, str(tgt))
            except Exception:
                pass
    except Exception:
        pass

def _try_register_pdf_fonts():
    global _PDF_FONT_REGULAR, _PDF_FONT_BOLD
    candidates = [
        # Liberation Sans is metric-compatible with Arial/Helvetica (matches the Vyapar look)
        # and includes the ₹ glyph. Preferred when present; otherwise fall back to DejaVu.
        ("LiberationSans", "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
         "LiberationSans-Bold", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
        ("LiberationSans", "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
         "LiberationSans-Bold", "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
        ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "DejaVuSans-Bold", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("DejaVuSans", str(Path(__file__).parent / "fonts" / "DejaVuSans.ttf"),
         "DejaVuSans-Bold", str(Path(__file__).parent / "fonts" / "DejaVuSans-Bold.ttf")),
        ("NotoSans", str(Path(__file__).parent / "fonts" / "NotoSans-Regular.ttf"),
         "NotoSans-Bold", str(Path(__file__).parent / "fonts" / "NotoSans-Bold.ttf")),
    ]
    import os as _os
    for name, path, name_b, path_b in candidates:
        try:
            if _os.path.exists(path) and _os.path.exists(path_b):
                pdfmetrics.registerFont(TTFont(name, path))
                pdfmetrics.registerFont(TTFont(name_b, path_b))
                # Override Helvetica so table body cells (which inherit reportlab's default)
                # also pick up ₹ and other Unicode glyphs.
                try:
                    pdfmetrics.registerFont(TTFont("Helvetica", path))
                    pdfmetrics.registerFont(TTFont("Helvetica-Bold", path_b))
                except Exception:
                    pass
                _PDF_FONT_REGULAR = name
                _PDF_FONT_BOLD = name_b
                try:
                    import logging as _l
                    _l.getLogger(__name__).info("PDF fonts registered: %s / %s (also aliased as Helvetica)", name, name_b)
                except Exception:
                    pass
                return name
        except Exception:
            continue
    return _PDF_FONT_REGULAR

# Order matters: try to download then register
_ensure_font_files_on_disk()
_try_register_pdf_fonts()

MONGO_URL = os.environ['MONGO_URL']
DB_NAME = os.environ['DB_NAME']
JWT_SECRET = os.environ.get('JWT_SECRET', 'dev-secret')
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')

client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

app = FastAPI(title="Denplex ERP")
api = APIRouter(prefix="/api")
security = HTTPBearer(auto_error=False)

# Trial role write-guard middleware: trial users cannot mutate (DELETE/PUT/PATCH) on /api/*
@app.middleware("http")
async def trial_write_guard(request: Request, call_next):
    if request.url.path.startswith("/api") and request.method in ("DELETE", "PUT", "PATCH"):
        auth = request.headers.get("authorization", "")
        if auth.lower().startswith("bearer "):
            try:
                token = auth.split(" ", 1)[1]
                payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
                if payload.get("role") == "trial":
                    from fastapi.responses import JSONResponse
                    return JSONResponse(
                        status_code=403,
                        content={"detail": "Trial accounts have view + create access only. Editing and deleting are disabled during the 1-month trial. Please contact admin@denplex.co for a full license."},
                    )
            except Exception:
                pass
    return await call_next(request)

# Audit middleware: log all successful writes (POST/PUT/PATCH/DELETE) on /api/* by authenticated users
@app.middleware("http")
async def audit_writes_mw(request: Request, call_next):
    response = await call_next(request)
    try:
        path = request.url.path
        method = request.method
        if (method in ("POST", "PUT", "PATCH", "DELETE")
                and path.startswith("/api/")
                and not path.startswith("/api/auth/")
                and not path.startswith("/api/audit")
                and response.status_code < 400):
            auth = request.headers.get("authorization", "")
            user_name = "anonymous"
            if auth.lower().startswith("bearer "):
                try:
                    token = auth.split(" ", 1)[1]
                    p = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
                    u = await db.users.find_one({"id": p["sub"]}, {"_id": 0, "name": 1, "email": 1})
                    if u:
                        user_name = u.get("name") or u.get("email") or "unknown"
                except Exception:
                    pass
            fwd = request.headers.get("x-forwarded-for", "")
            ip = fwd.split(",")[0].strip() if fwd else (request.client.host if request.client else "")
            await db.audit_logs.insert_one({
                "id": str(uuid.uuid4()),
                "user": user_name,
                "action": method.lower(),
                "entity": "api",
                "entity_id": path,
                "details": {"status": response.status_code},
                "ip": ip,
                "user_agent": (request.headers.get("user-agent", "") or "")[:300],
                "created_at": datetime.now(timezone.utc).isoformat(),
            })
    except Exception as e:
        logger.warning("audit writes mw failed: %s", e)
    return response

ROLES = ["admin", "manager", "production", "qc", "accountant", "ca", "sales", "design", "employee", "trial"]

# ---------------- helpers ----------------
def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def new_id() -> str:
    return str(uuid.uuid4())

def hash_password(pw: str) -> str:
    return bcrypt.hashpw(pw.encode(), bcrypt.gensalt()).decode()

def verify_password(pw: str, h: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode(), h.encode())
    except Exception:
        return False

def create_token(uid: str, role: str) -> str:
    payload = {"sub": uid, "role": role, "exp": datetime.now(timezone.utc) + timedelta(days=7)}
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

async def get_current_user(creds: Optional[HTTPAuthorizationCredentials] = Depends(security)) -> Dict[str, Any]:
    if not creds:
        raise HTTPException(401, "Not authenticated")
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=["HS256"])
    except Exception:
        raise HTTPException(401, "Invalid token")
    user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0, "password": 0, "totp_secret": 0})
    if not user:
        raise HTTPException(401, "User not found")
    # enforce trial expiry
    exp = user.get("trial_expires_at")
    if exp:
        try:
            exp_dt = datetime.fromisoformat(exp.replace("Z", "+00:00"))
            if datetime.now(timezone.utc) > exp_dt:
                raise HTTPException(403, "Trial expired. Please contact admin@denplex.co to extend your access.")
        except HTTPException:
            raise
        except Exception:
            pass
    return user

async def require_can_edit(request: Request, user=Depends(get_current_user)) -> Dict[str, Any]:
    # Trial users cannot DELETE/PUT/PATCH on most endpoints
    if user.get("role") == "trial" and request.method in ("DELETE", "PUT", "PATCH"):
        raise HTTPException(403, "Trial accounts have view + create access only. Editing and deleting are disabled during the 1-month trial.")
    return user

def require_roles(*allowed: str):
    async def checker(user: Dict[str, Any] = Depends(get_current_user)):
        if user["role"] not in allowed and user["role"] != "admin":
            raise HTTPException(403, "Insufficient permissions")
        return user
    return checker

async def next_seq(name: str) -> int:
    doc = await db.counters.find_one_and_update(
        {"_id": name},
        {"$inc": {"seq": 1}},
        upsert=True,
        return_document=True,
    )
    return doc["seq"] if doc else 1

async def gen_code(prefix: str, name: str) -> str:
    n = await next_seq(name)
    year = datetime.now(timezone.utc).strftime("%y")
    # Use the saved Document-Masters prefix for this doc type when configured (e.g. "2627/" -> "2627/0001")
    try:
        doc = await db.settings.find_one({"_id": "masters_prefixes"}, {"_id": 0})
        pref = ((doc or {}).get("value") or {}).get(name)
        if pref:
            return f"{pref}{n:04d}"
    except Exception:
        pass
    return f"{prefix}-{year}-{n:04d}"

# ---------------- Models ----------------
class RegisterIn(BaseModel):
    name: str
    email: EmailStr
    password: str
    role: Literal["admin", "manager", "production", "qc", "accountant", "ca", "sales", "design", "employee", "trial"] = "employee"
    unit: Optional[str] = "Unit 1"

class TrialRequestIn(BaseModel):
    name: str
    company: str
    phone: str
    email: EmailStr
    gstin: Optional[str] = ""
    business_type: Optional[str] = ""
    purpose: Optional[str] = ""

class TrialApproveIn(BaseModel):
    note: Optional[str] = ""

class LoginIn(BaseModel):
    email: EmailStr
    password: str
    totp_code: Optional[str] = ""

class Customer(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    name: str
    contact_person: Optional[str] = ""
    phone: Optional[str] = ""
    email: Optional[str] = ""
    gstin: Optional[str] = ""
    address: Optional[str] = ""
    customer_type: Literal["repeat", "one_time"] = "one_time"
    orders_count: int = 0
    created_at: str = Field(default_factory=now_iso)

class Lead(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    name: str
    company: Optional[str] = ""
    phone: Optional[str] = ""
    email: Optional[str] = ""
    source: Optional[str] = "manual"  # manual, b2b, website
    requirement: Optional[str] = ""
    status: Literal["new", "contacted", "qualified", "converted", "lost"] = "new"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class Supplier(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    name: str
    contact_person: Optional[str] = ""
    phone: Optional[str] = ""
    email: Optional[str] = ""
    gstin: Optional[str] = ""
    address: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class PartRevision(BaseModel):
    """A single revision entry in a Part's history."""
    revision: str                              # Internal revision (e.g. "Rev A", "01", "B2")
    customer_revision: Optional[str] = ""      # Customer's revision label for the same change (e.g. "ECN 1234 / Rev D")
    customer_change_ref: Optional[str] = ""    # Customer's ECN/PCO/change document reference
    effective_date: str = Field(default_factory=now_iso)
    change_reason: Optional[str] = ""          # Why this rev was made (customer change, internal optimization, etc.)
    drawing_pdf_b64: Optional[str] = ""        # Snapshot of drawing at this rev (base64)
    step_file_b64: Optional[str] = ""          # Snapshot of STEP at this rev (base64)
    drawing_filename: Optional[str] = ""       # Original filename for reference
    step_filename: Optional[str] = ""
    created_by: Optional[str] = ""
    notes: Optional[str] = ""

class PartMaster(BaseModel):
    """Central part identity. Every WO, BOM, inventory entry references a Part.
    The single source of truth for what each component IS."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    part_number: str                                       # Internal Denplex part number — primary identifier
    customer_part_number: Optional[str] = ""               # Customer's part number for cross-reference
    name: str                                              # Descriptive name
    description: Optional[str] = ""
    customer_id: Optional[str] = ""
    customer_name: Optional[str] = ""
    # Engineering specs
    material: Optional[str] = ""                           # e.g. "EN31", "SS316", "AISI 4140"
    material_grade: Optional[str] = ""                     # e.g. "Hardened", "Forged", "Annealed"
    process: List[str] = []                                # ["Turning", "Milling", "Grinding", "Heat Treatment"]
    cycle_time_minutes: float = 0                          # Standard cycle time per piece
    weight_kg: float = 0                                   # Finished piece weight
    raw_material_size: Optional[str] = ""                  # e.g. "Ø50 x 200mm bar", "Plate 100x100x10"
    raw_material_qty_per_part: float = 0                   # For material planning (e.g. 0.5 kg / part)
    # Inspection & tooling
    inspection_plan: Optional[str] = ""                    # Critical dimensions, gauges, tolerances
    critical_dimensions: List[Dict[str, str]] = []         # [{dim: "Ø25 ±0.01", gauge: "Snap gauge G1"}]
    tools_required: List[str] = []                         # Tools/fixtures/inserts needed
    # Revisions
    current_revision: Optional[str] = ""                   # e.g. "Rev B"
    revisions: List[PartRevision] = []                     # Full history
    # Current revision files (also stored in latest revision but mirrored here for fast access)
    drawing_pdf_b64: Optional[str] = ""
    step_file_b64: Optional[str] = ""
    drawing_filename: Optional[str] = ""
    step_filename: Optional[str] = ""
    # Sourcing — how the part is produced
    sourcing: Literal["manufactured", "bought_out", "ready_made"] = "manufactured"
    # Status
    is_active: bool = True
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class InventoryItem(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    sku: str
    name: str
    category: Optional[str] = "raw"  # raw, wip, finished, tool, consumable
    uom: str = "pcs"                 # base unit
    secondary_unit: Optional[str] = ""        # optional alternate unit (e.g. Box)
    conversion_factor: float = 0              # 1 secondary unit = conversion_factor base units
    qty_on_hand: float = 0
    qty_in_process: float = 0
    qty_by_location: dict = Field(default_factory=dict)   # {"Vatva": 12, "Santej": 3} — per-location on-hand
    reorder_level: float = 0
    unit_cost: float = 0
    sale_price: float = 0             # default selling rate — line items can pull from this instead of manual entry
    purchase_price: float = 0         # default buying rate
    hsn: Optional[str] = ""
    gst_rate: float = 18.0
    location: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class StockMovement(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    item_id: str
    item_sku: str
    item_name: str
    type: Literal["in", "out", "adjust", "in_process", "transfer"]
    qty: float
    location: Optional[str] = ""          # which location this movement affects
    to_location: Optional[str] = ""       # for transfers: destination
    ref: Optional[str] = ""
    notes: Optional[str] = ""
    by_user: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class StockTransfer(BaseModel):
    model_config = ConfigDict(extra="ignore")
    item_id: str
    from_location: str
    to_location: str
    qty: float
    notes: Optional[str] = ""

DEFAULT_LOCATIONS = ["Vatva", "Santej"]

# ---------------------------------------------------------------------------
# Standard Parts Library — reusable catalog of common bought-out components
# (fasteners, pneumatics, bearings, standard hardware) so they don't have to
# be re-typed for every BOM / inventory entry. Seeded once on first read.
# ---------------------------------------------------------------------------
PART_LIB_CATEGORIES = ["Fastener", "Pneumatic", "Bearing", "Standard", "Electrical", "Other"]

class PartLibraryItem(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    category: str = "Standard"          # one of PART_LIB_CATEGORIES
    name: str                            # e.g. "Socket Head Cap Screw M6x20"
    standard: Optional[str] = ""         # e.g. "ISO 4762", "DIN 625"
    size: Optional[str] = ""             # e.g. "M6x20", "6205", "1/4\" BSP"
    material: Optional[str] = ""         # e.g. "SS304", "Grade 12.9", "Brass"
    uom: str = "Nos"
    hsn: Optional[str] = ""
    gst_rate: float = 18.0
    unit_cost: float = 0
    notes: Optional[str] = ""
    seeded: bool = False                 # True for built-in catalog rows
    created_at: str = Field(default_factory=now_iso)

def _seed_part(category, name, standard="", size="", material="", uom="Nos", hsn="", gst_rate=18.0, unit_cost=0):
    return {"id": new_id(), "category": category, "name": name, "standard": standard, "size": size,
            "material": material, "uom": uom, "hsn": hsn, "gst_rate": gst_rate, "unit_cost": unit_cost,
            "notes": "", "seeded": True, "created_at": now_iso()}

STANDARD_PARTS_SEED = [
    # Fasteners (HSN 7318 typical)
    _seed_part("Fastener", "Socket Head Cap Screw M5x16", "ISO 4762", "M5x16", "Grade 12.9", hsn="7318", unit_cost=3),
    _seed_part("Fastener", "Socket Head Cap Screw M6x20", "ISO 4762", "M6x20", "Grade 12.9", hsn="7318", unit_cost=4),
    _seed_part("Fastener", "Socket Head Cap Screw M8x25", "ISO 4762", "M8x25", "Grade 12.9", hsn="7318", unit_cost=6),
    _seed_part("Fastener", "Socket Head Cap Screw M10x30", "ISO 4762", "M10x30", "Grade 12.9", hsn="7318", unit_cost=9),
    _seed_part("Fastener", "Hex Bolt M8x40", "ISO 4014", "M8x40", "Grade 8.8", hsn="7318", unit_cost=5),
    _seed_part("Fastener", "Hex Bolt M10x50", "ISO 4014", "M10x50", "Grade 8.8", hsn="7318", unit_cost=8),
    _seed_part("Fastener", "Hex Nut M8", "ISO 4032", "M8", "Grade 8", hsn="7318", unit_cost=2),
    _seed_part("Fastener", "Hex Nut M10", "ISO 4032", "M10", "Grade 8", hsn="7318", unit_cost=3),
    _seed_part("Fastener", "Plain Washer M8", "ISO 7089", "M8", "SS304", hsn="7318", unit_cost=1),
    _seed_part("Fastener", "Spring Washer M8", "IS 3063", "M8", "Spring Steel", hsn="7318", unit_cost=1),
    _seed_part("Fastener", "Dowel Pin Ø6x20", "ISO 2338", "Ø6x20", "Hardened Steel", hsn="7318", unit_cost=7),
    _seed_part("Fastener", "Dowel Pin Ø8x25", "ISO 2338", "Ø8x25", "Hardened Steel", hsn="7318", unit_cost=9),
    _seed_part("Fastener", "Grub Screw M6x10", "ISO 4029", "M6x10", "Grade 45H", hsn="7318", unit_cost=3),
    # Pneumatics (HSN 8412 / 8481 typical)
    _seed_part("Pneumatic", "Pneumatic Cylinder Ø32x100 (DA)", "ISO 15552", "Ø32x100", "Aluminium", uom="Nos", hsn="8412", unit_cost=1200),
    _seed_part("Pneumatic", "Pneumatic Cylinder Ø40x150 (DA)", "ISO 15552", "Ø40x150", "Aluminium", uom="Nos", hsn="8412", unit_cost=1650),
    _seed_part("Pneumatic", "5/2 Solenoid Valve 1/4\" BSP", "", "1/4\" BSP", "Aluminium", uom="Nos", hsn="8481", unit_cost=900),
    _seed_part("Pneumatic", "One-Touch Fitting 6mm-1/8\"", "", "6mm-1/8\"", "Brass", uom="Nos", hsn="8481", unit_cost=35),
    _seed_part("Pneumatic", "One-Touch Fitting 8mm-1/4\"", "", "8mm-1/4\"", "Brass", uom="Nos", hsn="8481", unit_cost=45),
    _seed_part("Pneumatic", "Polyurethane Tube 6mm OD", "", "6mm OD", "PU", uom="Mtr", hsn="3917", unit_cost=18),
    _seed_part("Pneumatic", "Polyurethane Tube 8mm OD", "", "8mm OD", "PU", uom="Mtr", hsn="3917", unit_cost=24),
    _seed_part("Pneumatic", "Flow Control Valve 1/4\" BSP", "", "1/4\" BSP", "Brass", uom="Nos", hsn="8481", unit_cost=140),
    _seed_part("Pneumatic", "FRL Unit 1/4\"", "", "1/4\"", "Aluminium", uom="Nos", hsn="8421", unit_cost=1400),
    # Bearings (HSN 8482)
    _seed_part("Bearing", "Deep Groove Ball Bearing 6200", "DIN 625", "6200", "Chrome Steel", uom="Nos", hsn="8482", unit_cost=90),
    _seed_part("Bearing", "Deep Groove Ball Bearing 6202", "DIN 625", "6202", "Chrome Steel", uom="Nos", hsn="8482", unit_cost=110),
    _seed_part("Bearing", "Deep Groove Ball Bearing 6205", "DIN 625", "6205", "Chrome Steel", uom="Nos", hsn="8482", unit_cost=160),
    _seed_part("Bearing", "Linear Bearing LM12UU", "", "LM12UU", "Chrome Steel", uom="Nos", hsn="8482", unit_cost=180),
    _seed_part("Bearing", "Thrust Bearing 51105", "DIN 711", "51105", "Chrome Steel", uom="Nos", hsn="8482", unit_cost=210),
    # Standard hardware / locating
    _seed_part("Standard", "Toggle Clamp Horizontal 200kg", "", "GH-201", "Steel", uom="Nos", hsn="8205", unit_cost=320),
    _seed_part("Standard", "Toggle Clamp Vertical 340kg", "", "GH-101", "Steel", uom="Nos", hsn="8205", unit_cost=360),
    _seed_part("Standard", "Locating Pin (Diamond) Ø10", "", "Ø10", "Hardened Steel", uom="Nos", hsn="8466", unit_cost=240),
    _seed_part("Standard", "Locating Pin (Round) Ø10", "", "Ø10", "Hardened Steel", uom="Nos", hsn="8466", unit_cost=220),
    _seed_part("Standard", "Rest Pad Flat Ø16", "", "Ø16", "Hardened Steel", uom="Nos", hsn="8466", unit_cost=150),
    _seed_part("Standard", "Compression Spring Ø10x30", "", "Ø10x30", "Spring Steel", uom="Nos", hsn="7320", unit_cost=12),
    _seed_part("Standard", "Knob / Hand Lever M10", "", "M10", "Thermoplastic", uom="Nos", hsn="3926", unit_cost=85),
    _seed_part("Electrical", "Proximity Sensor M12 PNP NO", "", "M12", "", uom="Nos", hsn="8536", unit_cost=420),
    _seed_part("Electrical", "Limit Switch (Roller Lever)", "", "", "", uom="Nos", hsn="8536", unit_cost=180),
]

class BOMLine(BaseModel):
    # Either reference a legacy inventory item (item_id) OR a Part Master entry (component_part_id).
    # Going forward, component_part_id is preferred. item_id is kept for back-compat.
    item_id: Optional[str] = ""
    item_name: Optional[str] = ""
    component_part_id: Optional[str] = ""      # Reference to PartMaster
    component_part_number: Optional[str] = ""
    component_part_name: Optional[str] = ""
    qty: float
    uom: str = "pcs"
    scrap_factor_pct: float = 0                # Extra material allowance (5% = 5)
    sourcing: Optional[str] = ""               # Mirror of Part's sourcing; helps UI
    notes: Optional[str] = ""

class BOMRevision(BaseModel):
    """A single revision entry in a BOM's history. Snapshots the lines at the time."""
    revision: str                              # e.g. "Rev A", "01", "B"
    effective_date: str = Field(default_factory=now_iso)
    change_reason: Optional[str] = ""
    lines_snapshot: List[Dict[str, Any]] = []  # Frozen copy of BOM lines at this revision
    drawing_pdf_b64: Optional[str] = ""        # Optional assembly drawing for this rev
    drawing_filename: Optional[str] = ""
    customer_revision: Optional[str] = ""      # Customer's matching rev label, if any
    customer_change_ref: Optional[str] = ""
    created_by: Optional[str] = ""
    notes: Optional[str] = ""

class BOM(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    product_name: str
    description: Optional[str] = ""
    design_code: Optional[str] = ""
    solidworks_url: Optional[str] = ""
    parent_part_id: Optional[str] = ""
    parent_part_number: Optional[str] = ""
    revision: Optional[str] = "Rev A"          # Current/active BOM revision label
    revision_history: List[BOMRevision] = []   # Full audit trail
    is_default: bool = True
    bom_type: Literal["assembly", "subassembly", "standard_lib"] = "assembly"
    is_active: bool = True
    drawing_pdf_b64: Optional[str] = ""        # Current assembly drawing
    drawing_filename: Optional[str] = ""
    lines: List[BOMLine] = []
    created_at: str = Field(default_factory=now_iso)

class WorkOrder(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    customer_id: Optional[str] = ""
    customer_name: Optional[str] = ""
    bom_id: Optional[str] = ""
    product: str
    part_number: Optional[str] = ""
    qty: float
    po_ref: Optional[str] = ""
    status: Literal["planned", "in_progress", "qc", "completed", "on_hold", "cancelled"] = "planned"
    priority: Literal["low", "medium", "high"] = "medium"
    start_date: Optional[str] = ""
    due_date: Optional[str] = ""
    notes: Optional[str] = ""
    progress: int = 0
    created_at: str = Field(default_factory=now_iso)

class JobCard(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    work_order_id: str
    work_order_code: Optional[str] = ""
    operation: str
    machine: Optional[str] = ""
    operator: Optional[str] = ""
    qty_planned: float = 0
    qty_done: float = 0
    status: Literal["pending", "in_progress", "done"] = "pending"
    started_at: Optional[str] = ""
    finished_at: Optional[str] = ""
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class QuoteLine(BaseModel):
    description: str
    qty: float
    rate: float
    gst_rate: float = 18.0

class Quotation(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    customer_id: str
    customer_name: str
    date: str = Field(default_factory=now_iso)
    valid_until: Optional[str] = ""
    lines: List[QuoteLine] = []
    subtotal: float = 0
    gst_total: float = 0
    total: float = 0
    status: Literal["draft", "sent", "accepted", "rejected"] = "draft"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class ExtraCharge(BaseModel):
    """Document-level additional charge (Freight / Packaging / Adjustment). Lump sum, non-taxable."""
    name: str
    amount: float = 0

class POLine(BaseModel):
    description: str
    item_code: Optional[str] = ""
    hsn: Optional[str] = ""
    qty: float
    unit: Optional[str] = "Nos"
    rate: float
    discount_pct: float = 0.0
    discount_amount: float = 0.0
    gst_rate: float = 18.0

class PurchaseOrder(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    supplier_id: str
    supplier_name: str
    supplier_gstin: Optional[str] = ""
    date: str = Field(default_factory=now_iso)
    delivery_date: Optional[str] = ""
    reference: Optional[str] = ""
    place_of_supply: Optional[str] = ""
    is_interstate: bool = False
    terms_text: Optional[str] = ""
    round_off: float = 0
    lines: List[POLine] = []
    subtotal: float = 0
    gst_total: float = 0
    total: float = 0
    status: Literal["draft", "sent", "received", "cancelled"] = "draft"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class SaleOrder(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    customer_id: str
    customer_name: str
    customer_gstin: Optional[str] = ""
    date: str = Field(default_factory=now_iso)
    due_date: Optional[str] = ""
    delivery_date: Optional[str] = ""
    po_number: Optional[str] = ""
    reference: Optional[str] = ""
    place_of_supply: Optional[str] = ""
    is_interstate: bool = False
    terms_text: Optional[str] = ""
    round_off: float = 0
    lines: List[POLine] = []
    subtotal: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    gst_total: float = 0
    total: float = 0
    status: Literal["draft", "sent", "confirmed", "cancelled"] = "draft"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class VendorBill(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None              # supplier's bill number (editable; else auto)
    supplier_id: str
    supplier_name: str
    supplier_gstin: Optional[str] = ""
    date: str = Field(default_factory=now_iso)
    due_date: Optional[str] = ""
    reference: Optional[str] = ""
    place_of_supply: Optional[str] = ""
    is_interstate: bool = False
    terms_text: Optional[str] = ""
    round_off: float = 0
    tds: float = 0
    tds_rate: float = 0
    tds_section: Optional[str] = ""
    tcs: float = 0
    tcs_rate: float = 0
    extra_charges: List[ExtraCharge] = []
    charges_total: float = 0
    lines: List[POLine] = []
    subtotal: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    gst_total: float = 0
    total: float = 0
    status: Literal["unpaid", "paid", "partial", "overdue"] = "unpaid"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class InvoiceLine(BaseModel):
    description: str
    item_code: Optional[str] = ""        # SKU / part number, shown as "Item Code" col
    hsn: Optional[str] = ""
    qty: float
    unit: Optional[str] = "Nos"          # Mtr, Kg, Nos, etc.
    rate: float
    discount_pct: float = 0.0
    discount_amount: float = 0.0
    gst_rate: float = 18.0

class Invoice(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    customer_id: str
    customer_name: str
    customer_gstin: Optional[str] = ""
    place_of_supply: Optional[str] = ""
    is_interstate: bool = False
    date: str = Field(default_factory=now_iso)
    due_date: Optional[str] = ""
    # Optional Ship To override (when shipping address ≠ billing address)
    ship_to_name: Optional[str] = ""
    ship_to_address: Optional[str] = ""
    ship_to_gstin: Optional[str] = ""
    # Optional Bill From / Ship From overrides (defaults derived from company settings)
    bill_from_name: Optional[str] = ""
    bill_from_address: Optional[str] = ""
    ship_from_name: Optional[str] = ""        # e.g. "Unit - 1"
    ship_from_address: Optional[str] = ""
    # Optional PO meta fields
    po_number: Optional[str] = ""
    po_date: Optional[str] = ""
    purchaser_name: Optional[str] = ""
    payment_mode: Optional[str] = ""
    eway_bill_no: Optional[str] = ""
    eway_generated_at: Optional[str] = ""      # date the e-way bill was generated
    eway_distance_km: float = 0
    eway_over_dimensional: bool = False
    eway_valid_until: Optional[str] = ""        # auto-computed: generated + days (1 day / 200km, ODC 1/20)
    eway_details: dict = Field(default_factory=dict)   # full NIC form payload (dispatch/ship/transport/Part-B/value)
    # ---- e-Invoice (IRP / IRN) ----
    irn: Optional[str] = ""                     # Invoice Reference Number from the IRP
    ack_no: Optional[str] = ""                  # acknowledgement number
    ack_date: Optional[str] = ""                # acknowledgement date
    signed_qr: Optional[str] = ""               # signed QR payload (rendered as QR on the PDF)
    einvoice_status: Optional[str] = ""         # ""/"generated"/"cancelled"
    einvoice_details: dict = Field(default_factory=dict)
    invoice_type: Literal["gst", "non_gst", "export"] = "gst"
    payment_terms: Optional[str] = ""       # e.g. "Net 30"
    godown: Optional[str] = ""              # dispatch location (Vatva / Santej)
    custom_fields: dict = Field(default_factory=dict)   # configurable doc fields (Transport Name, Vehicle No, etc.)
    terms_text: Optional[str] = ""         # Terms & Conditions body
    round_off: float = 0
    tds: float = 0                         # TDS amount deducted (reduces total)
    tds_rate: float = 0
    tds_section: Optional[str] = ""
    tcs: float = 0                         # TCS amount collected (adds to total)
    tcs_rate: float = 0
    extra_charges: List[ExtraCharge] = []
    charges_total: float = 0
    lines: List[InvoiceLine] = []
    subtotal: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    total: float = 0
    status: Literal["draft", "sent", "paid", "overdue"] = "draft"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

# ---------------- Payment In / Out + Expenses (Money flow Phase A) ----------------
class ProformaInvoice(BaseModel):
    """Formal pre-invoice with terms — distinct from informal Estimate/Quotation.
    Convertible into a Sale Invoice once accepted."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    customer_id: str
    customer_name: str
    customer_gstin: Optional[str] = ""
    place_of_supply: Optional[str] = ""
    is_interstate: bool = False
    date: str = Field(default_factory=now_iso)
    valid_until: Optional[str] = ""
    ship_to_name: Optional[str] = ""
    ship_to_address: Optional[str] = ""
    ship_to_gstin: Optional[str] = ""
    po_number: Optional[str] = ""
    po_date: Optional[str] = ""
    purchaser_name: Optional[str] = ""
    payment_mode: Optional[str] = ""
    lines: List[InvoiceLine] = []
    subtotal: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    total: float = 0
    status: Literal["draft", "sent", "accepted", "rejected", "converted"] = "draft"
    converted_invoice_id: Optional[str] = ""
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class ReturnLine(BaseModel):
    description: str
    item_code: Optional[str] = ""
    hsn: Optional[str] = ""
    qty: float
    unit: Optional[str] = "Nos"
    rate: float
    gst_rate: float = 18.0
    reason: Optional[str] = ""

class SaleReturn(BaseModel):
    """Inventory + accounting reversal for sold goods returned by customer.
    Creates a Credit Note + optionally restores inventory."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    customer_id: str
    customer_name: str
    customer_gstin: Optional[str] = ""
    original_invoice_id: Optional[str] = ""
    original_invoice_code: Optional[str] = ""
    date: str = Field(default_factory=now_iso)
    lines: List[ReturnLine] = []
    subtotal: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    total: float = 0
    restore_inventory: bool = True
    credit_note_id: Optional[str] = ""
    reason: Optional[str] = ""
    status: Literal["draft", "issued", "settled"] = "draft"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class PurchaseReturn(BaseModel):
    """Reversal for goods returned to supplier. Creates a Debit Note + reduces inventory."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    supplier_id: str
    supplier_name: str
    supplier_gstin: Optional[str] = ""
    original_bill_id: Optional[str] = ""
    original_bill_code: Optional[str] = ""
    date: str = Field(default_factory=now_iso)
    lines: List[ReturnLine] = []
    subtotal: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    total: float = 0
    reduce_inventory: bool = True
    debit_note_id: Optional[str] = ""
    reason: Optional[str] = ""
    status: Literal["draft", "issued", "settled"] = "draft"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class PaymentAllocation(BaseModel):
    """A single allocation of a payment to an invoice/bill."""
    document_id: str
    document_code: Optional[str] = ""
    document_type: Literal["invoice", "vendor_bill", "expense", "credit_note", "debit_note"] = "invoice"
    amount: float                                 # cash applied to this document
    tds_amount: float = 0                         # TDS deducted by the party — settles the doc but is not cash received

class PaymentIn(BaseModel):
    """Money received from a customer. Can be unallocated, partially allocated,
    or fully allocated to one or more sale invoices / credit notes."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None  # e.g. PMT-IN-26-0001
    party_id: str
    party_name: str
    date: str = Field(default_factory=now_iso)
    amount: float                                # Total received in this payment
    allocated_amount: float = 0                  # Sum of allocations applied
    payment_type: Literal["Cash", "Bank Transfer", "UPI", "Cheque", "Card", "Other"] = "Cash"
    ref_no: Optional[str] = ""                   # cheque no / UPI ref / txn id
    bank_name: Optional[str] = ""
    notes: Optional[str] = ""
    allocations: List[PaymentAllocation] = []
    status: Literal["Unused", "Partially Used", "Used"] = "Unused"
    created_at: str = Field(default_factory=now_iso)

class PaymentOut(BaseModel):
    """Money paid to a supplier or for an expense. Same shape as PaymentIn."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    party_id: str
    party_name: str
    date: str = Field(default_factory=now_iso)
    amount: float
    allocated_amount: float = 0
    payment_type: Literal["Cash", "Bank Transfer", "UPI", "Cheque", "Card", "Other"] = "Cash"
    ref_no: Optional[str] = ""
    bank_name: Optional[str] = ""
    notes: Optional[str] = ""
    allocations: List[PaymentAllocation] = []
    status: Literal["Unused", "Partially Used", "Used"] = "Unused"
    created_at: str = Field(default_factory=now_iso)

class ExpenseCategory(BaseModel):
    """An expense bucket (Courier, Salary, Rent, etc.). Direct vs Indirect for accounting."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    name: str
    classification: Literal["direct", "indirect"] = "indirect"
    created_at: str = Field(default_factory=now_iso)

class Expense(BaseModel):
    """A business expense entry."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    category_id: str
    category_name: Optional[str] = ""
    party_id: Optional[str] = ""                 # Optional vendor
    party_name: Optional[str] = ""
    date: str = Field(default_factory=now_iso)
    amount: float
    paid_amount: float = 0                       # 0 = Unpaid, == amount = Paid
    payment_type: Literal["Cash", "Bank Transfer", "UPI", "Cheque", "Card", "Other"] = "Cash"
    ref_no: Optional[str] = ""
    notes: Optional[str] = ""
    status: Literal["Paid", "Unpaid", "Partial"] = "Unpaid"
    created_at: str = Field(default_factory=now_iso)


class QCReport(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    work_order_id: Optional[str] = ""
    work_order_code: Optional[str] = ""
    customer_id: Optional[str] = ""
    customer_name: Optional[str] = ""
    inspector: Optional[str] = ""
    inspection_date: str = Field(default_factory=now_iso)
    parameter: str
    spec: Optional[str] = ""
    measured: Optional[str] = ""
    result: Literal["pass", "fail", "rework"] = "pass"
    photos: List[str] = []  # base64 strings
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

# ============================================================================
# PHASE Q.QC.1 — DIMENSIONAL QC INSPECTION REPORTS (Backend Models + Endpoints)
# ============================================================================
# Add this section to backend/server.py after the existing QCReport model (around line 500)
# Includes: models, CRUD endpoints, PDF/Excel export helpers

# --- Models ---
class QCDimensionSpec(BaseModel):
    """One column/parameter in the dimensional QC report."""
    label: str                          # e.g., "250", "120 (+0.17/-0.12)", "Ø130 H7"
    nominal: Optional[float] = None     # parsed nominal value (250, 120, 130)
    tol_upper: Optional[float] = None   # upper tolerance (+0.17, +0.04, etc.)
    tol_lower: Optional[float] = None   # lower tolerance (-0.12, 0, etc.)
    unit: str = "mm"                    # measurement unit
    raw_spec: str                       # original spec string from drawing

class QCSampleRow(BaseModel):
    """One row of measurements (one sample piece, all dimensions)."""
    sample_no: int                      # 1-10
    measurements: List[Optional[float]] = []  # parallel array to dimensions
    result: Optional[str] = ""          # "pass" | "fail" | "" (unset)
    sign: Optional[str] = ""            # inspector signature/initials
    note: Optional[str] = ""            # per-row note

class QCInspection(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None          # auto-generated QCI-0001, QCI-0002, etc.
    # --- Header ---
    report_no: Optional[str] = ""
    inspection_date: str = Field(default_factory=now_iso)
    supplier_name: Optional[str] = ""   # supplier/vendor name
    invoice_no: Optional[str] = ""
    invoice_date: Optional[str] = ""
    part_number: Optional[str] = ""
    part_name: str                      # BEARING HOUSING, SHAFT, etc.
    drawing_name: Optional[str] = ""    # 600NL-0441-1, BRMG0100H01CHAS010
    drawing_pdf_b64: Optional[str] = None  # uploaded PDF (base64, optional)
    # --- Inspection Data ---
    dimensions: List[QCDimensionSpec] = []  # column headers (specs)
    samples: List[QCSampleRow] = []     # 10 sample rows with measurements
    # --- Overall ---
    overall_result: Optional[str] = "pending"  # pass | fail | pending
    inspector_name: Optional[str] = ""
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)
    created_by: Optional[str] = ""
    updated_at: Optional[str] = None

# --- Helpers: Tolerance Parsing ---
def parse_tolerance_spec(spec: str) -> tuple[Optional[float], Optional[float]]:
    """Parse tolerance string: '+0.17/-0.12', '±0.2', '+0.04/0', '-0.5', etc.
    Returns: (tol_upper, tol_lower)
    """
    if not spec or not spec.strip():
        return None, None
    spec = spec.strip().replace(" ", "")
    try:
        if "±" in spec:
            val = float(spec.split("±")[1])
            return val, -val
        elif "/" in spec:
            parts = spec.split("/")
            upper = float(parts[0].replace("+", "")) if parts[0] else 0
            lower = float(parts[1]) if len(parts) > 1 else 0
            return upper, lower
        else:
            val = float(spec.replace("+", ""))
            return val, 0
    except:
        return None, None

def extract_nominal_from_label(label: str) -> Optional[float]:
    """Try to extract nominal value from label like '120 (+0.17/-0.12)' → 120.0"""
    try:
        base = label.split("(")[0].strip() if "(" in label else label.strip()
        base = base.replace("Ø", "").replace("M", "").replace("H", "").replace("E", "").strip()
        if base and base[0].isdigit():
            return float(base.split()[0])
    except:
        pass
    return None

def check_tolerance(measured: Optional[float], nominal: Optional[float], tol_upper: Optional[float], tol_lower: Optional[float]) -> str:
    """Return 'pass' or 'fail' based on tolerance check. Returns '' if any value is None."""
    if measured is None or nominal is None or tol_upper is None or tol_lower is None:
        return ""
    try:
        upper_bound = nominal + tol_upper
        lower_bound = nominal + tol_lower
        if lower_bound <= measured <= upper_bound:
            return "pass"
        else:
            return "fail"
    except:
        return ""

# --- Q.QC.1 helper: Mongo doc normalization ---
def fixup(doc):
    """Normalize a MongoDB document for JSON return: stringify _id and mirror to id."""
    if not doc:
        return doc
    if "_id" in doc:
        doc["_id"] = str(doc["_id"])
        doc["id"] = doc["_id"]
    return doc

# --- Q.QC.3 helpers: dimension enrichment + auto pass/fail ---
def enrich_dimensions(dims):
    """Fill nominal + tolerances from the label / raw_spec when left blank by the user."""
    out = []
    for d in (dims or []):
        d = dict(d)
        if d.get("nominal") in (None, ""):
            nom = extract_nominal_from_label(d.get("label", "") or "")
            if nom is not None:
                d["nominal"] = nom
        if d.get("tol_upper") in (None, "") and d.get("tol_lower") in (None, ""):
            tu, tl = parse_tolerance_spec((d.get("raw_spec") or d.get("label") or ""))
            if tu is not None:
                d["tol_upper"] = tu
            if tl is not None:
                d["tol_lower"] = tl
        out.append(d)
    return out

def validate_inspection(doc):
    """Compute per-sample pass/fail and the overall result from measurements vs tolerances.
    Mutates and returns doc. Samples with no measurements are left untouched."""
    dims = doc.get("dimensions", []) or []
    samples = doc.get("samples", []) or []
    any_fail = False
    any_measured = False
    for s in samples:
        meas = s.get("measurements", []) or []
        has_val = False
        sample_fail = False
        for i, d in enumerate(dims):
            m = meas[i] if i < len(meas) else None
            if m is None or m == "":
                continue
            try:
                mv = float(m)
            except (TypeError, ValueError):
                continue
            has_val = True
            if check_tolerance(mv, d.get("nominal"), d.get("tol_upper"), d.get("tol_lower")) == "fail":
                sample_fail = True
        if has_val:
            any_measured = True
            s["result"] = "fail" if sample_fail else "pass"
            if sample_fail:
                any_fail = True
    doc["overall_result"] = ("fail" if any_fail else "pass") if any_measured else "pending"
    return doc

# --- Q.QC.2 config: AI drawing extraction (Claude vision) ---
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
ANTHROPIC_BASE_URL = os.environ.get("ANTHROPIC_BASE_URL", "https://api.anthropic.com").rstrip("/")
CAD_SERVICE_URL = os.environ.get("CAD_SERVICE_URL", "").rstrip("/")   # optional STEP geometry microservice
QC_VISION_MODEL = os.environ.get("QC_VISION_MODEL", "claude-sonnet-4-5")

QC_EXTRACT_PROMPT = (
    "You are a precision-machining QC engineer. Read this engineering drawing and extract every "
    "controlled dimension/parameter that should be inspected. Return ONLY a JSON array, no prose. "
    "Each element: {\"label\": str, \"nominal\": number|null, \"tol_upper\": number|null, "
    "\"tol_lower\": number|null, \"unit\": str, \"raw_spec\": str}. "
    "nominal is the basic size; tol_upper/tol_lower are signed deviations (e.g. +0.04 and -0.00, "
    "or +/-0.1 => +0.1/-0.1). For a fit code like H7/g6, convert to numeric deviations if the drawing "
    "shows them, otherwise leave tolerances null and keep the fit code in raw_spec. unit defaults to "
    "\"mm\". raw_spec is the exact text as drawn. Skip surface-finish/notes that are not measurable dimensions."
)

def _parse_dims_json(text: str):
    """Parse the model's JSON array out of its reply and normalize to QCDimensionSpec dicts."""
    import json
    t = (text or "").strip()
    if "```" in t:
        for p in t.split("```"):
            p = p.strip()
            if p.startswith("json"):
                p = p[4:].strip()
            if p.startswith("[") or p.startswith("{"):
                t = p
                break
    if not t.startswith("[") and "[" in t and "]" in t:
        t = t[t.index("["): t.rindex("]") + 1]
    try:
        raw = json.loads(t)
    except Exception:
        return []
    if isinstance(raw, dict):
        raw = raw.get("dimensions") or raw.get("data") or []
    dims = []
    for d in (raw if isinstance(raw, list) else []):
        if not isinstance(d, dict):
            continue
        dims.append({
            "label": str(d.get("label") or d.get("name") or "").strip(),
            "nominal": d.get("nominal"),
            "tol_upper": d.get("tol_upper"),
            "tol_lower": d.get("tol_lower"),
            "unit": d.get("unit") or "mm",
            "raw_spec": str(d.get("raw_spec") or d.get("label") or "").strip(),
        })
    return enrich_dimensions([d for d in dims if d["label"]])

# --- Auto-code generation ---
async def get_next_qc_inspection_code() -> str:
    """Generate next QCI code: QCI-0001, QCI-0002, etc."""
    last = await db.qc_inspections.find_one({}, sort=[("code", -1)])
    if not last or not last.get("code"):
        return "QCI-0001"
    try:
        num = int(last["code"].split("-")[1]) + 1
        return f"QCI-{num:04d}"
    except:
        return "QCI-0001"

# --- Endpoints ---
@api.post("/qc-inspections")
async def create_qc_inspection(req: QCInspection, claims: dict = Depends(get_current_user)):
    """Create a new dimensional QC inspection."""
    req.created_by = claims.get("sub", "unknown")
    req.code = await get_next_qc_inspection_code()
    req.created_at = now_iso()
    data = req.model_dump(by_alias=False)
    data["dimensions"] = enrich_dimensions(data.get("dimensions"))
    validate_inspection(data)  # Q.QC.3 auto pass/fail
    result = await db.qc_inspections.insert_one(data)
    return {"id": str(result.inserted_id), "code": req.code, "overall_result": data.get("overall_result")}

@api.get("/qc-inspections")
async def list_qc_inspections(claims: dict = Depends(get_current_user)):
    """List all dimensional QC inspections."""
    items = await db.qc_inspections.find().to_list(1000)
    return [fixup(i) for i in items]

@api.get("/qc-inspections/{qid}")
async def get_qc_inspection(qid: str, claims: dict = Depends(get_current_user)):
    """Get a single dimensional QC inspection (includes drawing PDF if present)."""
    doc = await db.qc_inspections.find_one({"_id": ObjectId(qid)})
    if not doc:
        raise HTTPException(status_code=404, detail="QC inspection not found")
    return fixup(doc)

@api.put("/qc-inspections/{qid}")
async def update_qc_inspection(qid: str, req: QCInspection, claims: dict = Depends(get_current_user)):
    """Update a dimensional QC inspection."""
    req.updated_at = now_iso()
    data = req.model_dump(by_alias=False)
    data["dimensions"] = enrich_dimensions(data.get("dimensions"))
    validate_inspection(data)  # Q.QC.3 auto pass/fail
    await db.qc_inspections.update_one({"_id": ObjectId(qid)}, {"$set": data})
    return {"ok": True, "overall_result": data.get("overall_result")}

@api.delete("/qc-inspections/{qid}")
async def delete_qc_inspection(qid: str, claims: dict = Depends(get_current_user)):
    """Delete a dimensional QC inspection."""
    await db.qc_inspections.delete_one({"_id": ObjectId(qid)})
    return {"ok": True}

@api.post("/qc-inspections/{qid}/drawing")
async def upload_qc_drawing(qid: str, file: UploadFile = File(...), claims: dict = Depends(get_current_user)):
    """Upload a drawing PDF to a QC inspection. Stores as base64."""
    content = await file.read()
    b64 = base64.b64encode(content).decode("utf-8")
    await db.qc_inspections.update_one(
        {"_id": ObjectId(qid)},
        {"$set": {"drawing_pdf_b64": b64}}
    )
    return {"ok": True, "size": len(content)}

@api.post("/qc-inspections/{qid}/validate")
async def validate_qc_inspection_endpoint(qid: str, claims: dict = Depends(get_current_user)):
    """Q.QC.3 - recompute per-sample pass/fail + overall result for a saved inspection."""
    doc = await db.qc_inspections.find_one({"_id": ObjectId(qid)})
    if not doc:
        raise HTTPException(status_code=404, detail="QC inspection not found")
    doc["dimensions"] = enrich_dimensions(doc.get("dimensions"))
    validate_inspection(doc)
    await db.qc_inspections.update_one(
        {"_id": ObjectId(qid)},
        {"$set": {"dimensions": doc["dimensions"], "samples": doc.get("samples", []),
                  "overall_result": doc.get("overall_result"), "updated_at": now_iso()}},
    )
    return fixup(doc)

@api.post("/qc-inspections/extract-dimensions")
async def extract_qc_dimensions(file: UploadFile = File(...), claims: dict = Depends(get_current_user)):
    """Q.QC.2 - read an engineering drawing (PDF or image) with Claude vision and return a list
    of inspectable dimensions to pre-fill a new inspection."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "AI drawing extraction is not configured. Set ANTHROPIC_API_KEY "
                                 "(optionally QC_VISION_MODEL / ANTHROPIC_BASE_URL) in the backend environment.")
    content = await file.read()
    if not content:
        raise HTTPException(400, "Empty file.")
    b64 = base64.b64encode(content).decode("utf-8")
    ctype = (file.content_type or "").lower()
    fname = (file.filename or "").lower()
    if "pdf" in ctype or fname.endswith(".pdf"):
        media_block = {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}}
    else:
        media = ctype if ctype.startswith("image/") else "image/png"
        media_block = {"type": "image", "source": {"type": "base64", "media_type": media, "data": b64}}
    payload = {
        "model": QC_VISION_MODEL,
        "max_tokens": 2500,
        "messages": [{"role": "user", "content": [media_block, {"type": "text", "text": QC_EXTRACT_PROMPT}]}],
    }
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=120) as cx:
            r = await cx.post(f"{ANTHROPIC_BASE_URL}/v1/messages", json=payload, headers=headers)
    except Exception as e:
        raise HTTPException(502, f"Could not reach the vision API: {e}")
    if r.status_code >= 400:
        raise HTTPException(502, f"Vision API error {r.status_code}: {r.text[:300]}")
    try:
        body = r.json()
        text = "".join(b.get("text", "") for b in body.get("content", []) if b.get("type") == "text")
    except Exception as e:
        raise HTTPException(502, f"Unexpected vision API response: {e}")
    dims = _parse_dims_json(text)
    return {"dimensions": dims, "count": len(dims)}

# --- Export Helpers ---
def _qc_header_footer(canvas, doc):
    """Denplex letterhead header + footer drawn on every QC report page."""
    from reportlab.lib.units import mm
    from reportlab.lib.utils import ImageReader
    from reportlab.lib.colors import HexColor
    W, H = doc.pagesize
    RED = HexColor("#CC0000"); BLACK = HexColor("#1A1A1A"); GREY = HexColor("#666666")
    canvas.saveState()
    # ---- HEADER ----
    try:
        logo = ImageReader(str(ROOT_DIR / "logo.png"))
        lw = 24*mm; lh = lw * 658.0/767.0
        canvas.drawImage(logo, 12*mm, H - 8*mm - lh, width=lw, height=lh,
                         preserveAspectRatio=True, mask='auto')
    except Exception:
        pass
    cx = 40*mm
    canvas.setFillColor(BLACK); canvas.setFont(_PDF_FONT_BOLD, 16)
    canvas.drawString(cx, H - 13*mm, "DENPLEX ENGINEERING COMPANY")
    canvas.setFillColor(GREY); canvas.setFont(_PDF_FONT_REGULAR, 8)
    canvas.drawString(cx, H - 17*mm, "Complete Engineering Solutions  -  Since 2015")
    canvas.setFillColor(RED); canvas.setFont(_PDF_FONT_BOLD, 7.5)
    canvas.drawString(cx, H - 21*mm, "JIGS & FIXTURES   |   3D PRINTING   |   CAD/CAM   |   PRECISION MACHINING")
    # right-aligned contact block
    canvas.setFillColor(BLACK); canvas.setFont(_PDF_FONT_REGULAR, 7)
    rx = W - 12*mm
    canvas.drawRightString(rx, H - 11*mm, "Shed No.20, Pushkar Mahadev Estate-1, Vatva, Ahmedabad-382445")
    canvas.drawRightString(rx, H - 14.5*mm, "+91 90333 38999   |   contact@denplex.co   |   www.denplex.co")
    canvas.setFont(_PDF_FONT_BOLD, 7)
    canvas.drawRightString(rx, H - 18*mm, "GST: 24AALFD1671P1Z2")
    # red rule under header
    canvas.setStrokeColor(RED); canvas.setLineWidth(1.4)
    canvas.line(12*mm, H - 24*mm, W - 12*mm, H - 24*mm)
    # ---- FOOTER ----
    canvas.setFillColor(BLACK)
    canvas.rect(0, 8*mm, W, 11*mm, stroke=0, fill=1)
    canvas.setFillColor(HexColor("#DDDDDD")); canvas.setFont(_PDF_FONT_REGULAR, 7)
    canvas.drawString(12*mm, 13.6*mm,
        "Reg. Office: Shed No.4, Shriram Estate, Santej, Gandhinagar-382721")
    canvas.drawCentredString(W/2.0, 13.6*mm,
        "+91 90333 38999  -  contact@denplex.co  -  www.denplex.co")
    canvas.setFillColor(RED); canvas.setFont(_PDF_FONT_BOLD, 7)
    canvas.drawRightString(W - 12*mm, 13.6*mm, "GST: 24AALFD1671P1Z2")
    canvas.setFillColor(HexColor("#888888")); canvas.setFont(_PDF_FONT_REGULAR, 7)
    canvas.drawCentredString(W/2.0, 10*mm, "Complete Engineering Solutions  -  Since 2015  -  600+ Clients Served   |   Page %d" % canvas.getPageNumber())
    canvas.restoreState()


def build_qc_pdf(inspection: dict) -> bytes:
    """Denplex-branded dimensional Quality Inspection Report.
    Letterhead header/footer + sample-format table (specified dims row + up to 10
    actual-measurement rows + QC Result + Sign). Dimensions paginate in chunks."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer, SimpleDocTemplate
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    W, H = landscape(A4)
    RED = colors.HexColor("#CC0000"); BLACK = colors.HexColor("#1A1A1A")
    LGREY = colors.HexColor("#EFEFEF")
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=landscape(A4),
        topMargin=27*mm, bottomMargin=22*mm, leftMargin=10*mm, rightMargin=10*mm,
        title="Quality Inspection Report",
    )
    styles = getSampleStyleSheet()
    P = lambda t, s=8, b=False, c=BLACK, a=0: Paragraph(
        str(t), ParagraphStyle("p", parent=styles["Normal"], fontSize=s,
            leading=s+2, fontName=(_PDF_FONT_BOLD if b else _PDF_FONT_REGULAR),
            textColor=c, alignment=a))
    story = []

    # ---- Title ----
    story.append(Paragraph("Quality Inspection Report",
        ParagraphStyle("t", parent=styles["Title"], fontSize=15,
            fontName=_PDF_FONT_BOLD, textColor=BLACK, alignment=1, spaceAfter=4)))

    # ---- Report header grid ----
    hdr = [
        [P("Report No.", 8, True), P(inspection.get("report_no") or inspection.get("code") or "-"),
         P("Date", 8, True), P(inspection.get("inspection_date", "-"))],
        [P("Supplier", 8, True), P(inspection.get("supplier_name", "-")),
         P("Invoice No. & Date", 8, True),
         P(f"{inspection.get('invoice_no','-')}  /  {inspection.get('invoice_date','-')}")],
        [P("Part Name", 8, True), P(inspection.get("part_name", "-")),
         P("Part Number", 8, True), P(inspection.get("part_number", "-"))],
        [P("Drawing Name", 8, True), P(inspection.get("drawing_name", "-")),
         P("Inspector", 8, True), P(inspection.get("inspector_name", "-"))],
    ]
    htab = Table(hdr, colWidths=[34*mm, 95*mm, 38*mm, (W-20*mm-167*mm)])
    htab.setStyle(TableStyle([
        ("FONT", (0,0), (-1,-1), _PDF_FONT_REGULAR, 8),
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (0,-1), LGREY),
        ("BACKGROUND", (2,0), (2,-1), LGREY),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING", (0,0), (-1,-1), 3), ("BOTTOMPADDING", (0,0), (-1,-1), 3),
    ]))
    story.append(htab); story.append(Spacer(1, 4*mm))

    dims = inspection.get("dimensions", []) or []
    samples = (inspection.get("samples", []) or [])[:10]
    if not dims:
        dims = []
    CHUNK = 9
    chunks = [dims[i:i+CHUNK] for i in range(0, len(dims), CHUNK)] or [[]]

    for ci, chunk in enumerate(chunks):
        ndim = len(chunk)
        # header row: label + dim labels + QC Result + Sign
        row0 = [P("Dimensions & Tolerances<br/>Specified in the Drawing", 7.5, True, a=1)]
        row0 += [P(d.get("label", "-"), 8, True, a=1) for d in chunk]
        row0 += [P("QC Result<br/>PASS / FAIL", 7.5, True, a=1), P("Sign", 7.5, True, a=1)]
        # spec/tolerance row
        rowspec = [P("ACTUAL MEASURED DIMENSIONS", 7.5, True, RED, a=1)]
        rowspec += [P(d.get("raw_spec") or "-", 7, a=1) for d in chunk]
        rowspec += [P("", 7), P("", 7)]
        data = [row0, rowspec]
        # 10 sample rows
        for si in range(10):
            s = samples[si] if si < len(samples) else {}
            meas = s.get("measurements", []) or []
            cells = [P(str(si+1), 8, True, a=1)]
            for di in range(ndim):
                gi = ci*CHUNK + di
                v = meas[gi] if gi < len(meas) else None
                cells.append(P("" if v is None else str(v), 8, a=1))
            res = (s.get("result") or "").upper()
            rc = colors.green if res == "PASS" else (RED if res == "FAIL" else BLACK)
            cells.append(P(res or "", 8, True, rc, a=1))
            cells.append(P(s.get("sign") or "", 8, a=1))
            data.append(cells)

        label_w = 42*mm; tail_w = 24*mm + 16*mm
        dim_w = (W - 20*mm - label_w - tail_w) / max(ndim, 1)
        col_w = [label_w] + [dim_w]*ndim + [24*mm, 16*mm]
        t = Table(data, colWidths=col_w, rowHeights=[12*mm, 7*mm] + [7.6*mm]*10)
        t.setStyle(TableStyle([
            ("FONT", (0,0), (-1,-1), _PDF_FONT_REGULAR, 8),
            ("GRID", (0,0), (-1,-1), 0.5, colors.black),
            ("BACKGROUND", (0,0), (-1,0), LGREY),
            ("BACKGROUND", (0,1), (0,1), colors.HexColor("#FBE9E9")),
            ("SPAN", (0,1), (0,1)),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ]))
        story.append(t)
        # overall result under last chunk
        if ci == len(chunks) - 1:
            overall = (inspection.get("overall_result") or "pending").upper()
            oc = colors.green if overall == "PASS" else (RED if overall == "FAIL" else colors.HexColor("#B8860B"))
            story.append(Spacer(1, 3*mm))
            story.append(Paragraph(f"Overall Result: {overall}",
                ParagraphStyle("ov", parent=styles["Normal"], fontSize=12,
                    fontName=_PDF_FONT_BOLD, textColor=oc)))
        if ci < len(chunks) - 1:
            from reportlab.platypus import PageBreak
            story.append(PageBreak())

    doc.build(story, onFirstPage=_qc_header_footer, onLaterPages=_qc_header_footer)
    return buffer.getvalue()


def build_qc_excel(inspection: dict) -> bytes:
    """Generate Excel file matching the Denplex QC report template."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "QC Report"

    row = 1
    # --- Header ---
    ws[f"A{row}"] = "DENPLEX — Quality Inspection Report"
    ws[f"A{row}"].font = Font(bold=True, size=12)
    row += 2

    # Header info
    header_fields = [
        ("Report No.", inspection.get("report_no", "—")),
        ("Date", inspection.get("inspection_date", "—")),
        ("Supplier", inspection.get("supplier_name", "—")),
        ("Invoice No & Date", f"{inspection.get('invoice_no', '—')} / {inspection.get('invoice_date', '—')}"),
        ("Part Name", inspection.get("part_name", "—")),
        ("Part Number", inspection.get("part_number", "—")),
        ("Drawing Name", inspection.get("drawing_name", "—")),
        ("Inspector", inspection.get("inspector_name", "—")),
    ]
    for label, value in header_fields:
        ws[f"A{row}"] = label
        ws[f"B{row}"] = value
        ws[f"A{row}"].font = Font(bold=True, size=10)
        row += 1

    row += 1
    # --- Measurements Table ---
    dims = inspection.get("dimensions", [])
    samples = inspection.get("samples", [])

    # Headers
    headers = ["Sample #"] + [d.get("label", "?") for d in dims] + ["QC Result", "Sign"]
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=col_idx, value=header)
        cell.font = Font(bold=True, size=10)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    row += 1

    # Specs row
    specs = ["SPEC"] + [d.get("raw_spec", "—") for d in dims] + ["", ""]
    for col_idx, spec in enumerate(specs, start=1):
        cell = ws.cell(row=row, column=col_idx, value=spec)
        cell.font = Font(bold=True, size=9)
        cell.fill = PatternFill(start_color="E8E8E8", end_color="E8E8E8", fill_type="solid")
    row += 1

    # Sample rows
    for sample in samples[:10]:
        measurements = sample.get("measurements", [])
        result = sample.get("result", "")
        sign = sample.get("sign", "")
        sample_no = sample.get("sample_no", "")

        ws.cell(row=row, column=1, value=sample_no)
        for col_idx, m in enumerate(measurements, start=2):
            ws.cell(row=row, column=col_idx, value=m)
        ws.cell(row=row, column=len(dims) + 2, value=result.upper() if result else "")
        ws.cell(row=row, column=len(dims) + 3, value=sign)
        row += 1

    # Pad to 10 rows
    while row < 13 + len(header_fields):
        row += 1

    # Overall result
    row += 1
    ws[f"A{row}"] = "Overall Result:"
    ws[f"B{row}"] = inspection.get("overall_result", "pending").upper()
    ws[f"A{row}"].font = Font(bold=True, size=11)
    ws[f"B{row}"].font = Font(bold=True, size=11, color="FF0000" if inspection.get("overall_result") == "fail" else "00AA00")

    # Adjust column widths
    ws.column_dimensions["A"].width = 15
    for col_idx in range(2, len(headers) + 1):
        ws.column_dimensions[chr(64 + col_idx)].width = 12

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()

@api.get("/qc-inspections/{qid}/pdf")
async def export_qc_pdf(qid: str, claims: dict = Depends(get_current_user)):
    """Export QC inspection as PDF (Denplex template format)."""
    doc = await db.qc_inspections.find_one({"_id": ObjectId(qid)})
    if not doc:
        raise HTTPException(status_code=404, detail="QC inspection not found")

    pdf_bytes = build_qc_pdf(doc)
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=qc-{doc.get('code', 'report')}.pdf"}
    )

@api.get("/qc-inspections/{qid}/xlsx")
async def export_qc_xlsx(qid: str, claims: dict = Depends(get_current_user)):
    """Export QC inspection as Excel (Denplex template format)."""
    doc = await db.qc_inspections.find_one({"_id": ObjectId(qid)})
    if not doc:
        raise HTTPException(status_code=404, detail="QC inspection not found")

    xlsx_bytes = build_qc_excel(doc)
    return StreamingResponse(
        BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=qc-{doc.get('code', 'report')}.xlsx"}
    )


class DocumentMeta(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    name: str
    category: Optional[str] = "general"  # general, iso, drawing, qc, packaging
    linked_to: Optional[str] = ""  # work_order_id / invoice_id etc
    linked_type: Optional[str] = ""
    file_base64: str  # data URL
    mime: Optional[str] = ""
    size: int = 0
    uploaded_by: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class BillScanIn(BaseModel):
    image_base64: str  # raw base64 (no data URL prefix) or data URL
    mime: str = "image/jpeg"

# ---------------- P1: Accounting / HR / Marketing models ----------------
class Expense(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    date: str = Field(default_factory=now_iso)
    category: str = "general"
    description: str
    vendor: Optional[str] = ""
    amount: float = 0
    gst_rate: float = 0
    gst_amount: float = 0
    total: float = 0
    payment_mode: Optional[str] = "bank"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class Employee(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    name: str
    designation: Optional[str] = ""
    department: Optional[str] = ""
    phone: Optional[str] = ""
    email: Optional[str] = ""
    join_date: Optional[str] = ""
    monthly_salary: float = 0
    status: Literal["active", "inactive"] = "active"
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class Attendance(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    employee_id: str
    employee_name: Optional[str] = ""
    date: str
    status: Literal["present", "absent", "half_day", "leave"] = "present"
    hours: float = 8
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class Campaign(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    title: str
    channel: Literal["whatsapp", "instagram", "linkedin", "facebook", "email", "other"] = "whatsapp"
    content: Optional[str] = ""
    scheduled_for: Optional[str] = ""
    status: Literal["draft", "scheduled", "published"] = "draft"
    metrics: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class DocRevIn(BaseModel):
    file_base64: str
    notes: Optional[str] = ""

# ---------------- Auth ----------------
@api.post("/auth/register")
async def register(payload: RegisterIn, request: Request, current=Depends(get_current_user)):
    if current["role"] != "admin":
        raise HTTPException(403, "Only admin can register users")
    if await db.users.find_one({"email": payload.email.lower()}):
        raise HTTPException(400, "Email already exists")
    user = {
        "id": new_id(),
        "name": payload.name,
        "email": payload.email.lower(),
        "role": payload.role,
        "unit": (payload.unit or "Unit 1"),
        "password": hash_password(payload.password),
        "created_at": now_iso(),
    }
    await db.users.insert_one(user)
    await write_audit(
        current.get("name", "admin"),
        "user_created",
        "user",
        user["id"],
        {"email": user["email"], "role": user["role"], "unit": user["unit"]},
        request=request,
    )
    user.pop("_id", None); user.pop("password", None)
    return user

@api.post("/auth/login")
async def login(payload: LoginIn, request: Request):
    user = await db.users.find_one({"email": payload.email.lower()})
    if not user or not verify_password(payload.password, user["password"]):
        await write_audit(payload.email.lower(), "login_failed", "auth", "", {"reason": "invalid_credentials"}, request=request)
        raise HTTPException(401, "Invalid credentials")
    # Trial expiry check at login
    exp = user.get("trial_expires_at")
    if exp:
        try:
            exp_dt = datetime.fromisoformat(exp.replace("Z", "+00:00"))
            if datetime.now(timezone.utc) > exp_dt:
                await write_audit(user.get("name") or payload.email.lower(), "login_failed", "auth", user["id"], {"reason": "trial_expired"}, request=request)
                raise HTTPException(403, "Trial expired. Please contact admin@denplex.co to extend your access.")
        except HTTPException:
            raise
        except Exception:
            pass
    if user.get("totp_enabled"):
        if not payload.totp_code:
            raise HTTPException(401, "TOTP code required", headers={"X-2FA-Required": "1"})
        if not pyotp.TOTP(user.get("totp_secret", "")).verify(payload.totp_code, valid_window=1):
            await write_audit(user.get("name") or payload.email.lower(), "login_failed", "auth", user["id"], {"reason": "invalid_totp"}, request=request)
            raise HTTPException(401, "Invalid TOTP code")
    token = create_token(user["id"], user["role"])
    await write_audit(user.get("name") or payload.email.lower(), "login_success", "auth", user["id"], {"role": user["role"]}, request=request)
    return {
        "token": token,
        "user": {
            "id": user["id"], "name": user["name"], "email": user["email"], "role": user["role"],
            "trial_expires_at": user.get("trial_expires_at", ""),
        }
    }

@api.get("/auth/me")
async def me(user=Depends(get_current_user)):
    return user

class ChangePwIn(BaseModel):
    current_password: str
    new_password: str

@api.post("/auth/change-password")
async def change_password(payload: ChangePwIn, request: Request, user=Depends(get_current_user)):
    if len(payload.new_password) < 8:
        raise HTTPException(400, "New password must be at least 8 characters")
    u = await db.users.find_one({"id": user["id"]})
    if not u or not verify_password(payload.current_password, u["password"]):
        raise HTTPException(401, "Current password is incorrect")
    await db.users.update_one({"id": user["id"]}, {"$set": {"password": hash_password(payload.new_password)}})
    await write_audit(user["name"], "password_changed", "user", user["id"], request=request)
    return {"ok": True}

@api.get("/users")
async def list_users(user=Depends(require_roles("admin"))):
    return await db.users.find({}, {"_id": 0, "password": 0, "totp_secret": 0}).to_list(500)

# ---------------- Generic CRUD helpers ----------------
def serialize(d: Dict[str, Any]) -> Dict[str, Any]:
    d.pop("_id", None)
    return d

async def list_collection(coll, query: Dict = None, sort_key: str = "created_at", limit: int = 5000):
    cursor = coll.find(query or {}, {"_id": 0}).sort(sort_key, -1)
    return await cursor.to_list(limit)

# ---------------- Customers ----------------
@api.post("/customers")
async def create_customer(c: Customer, user=Depends(get_current_user)):
    doc = c.model_dump()
    doc["code"] = await gen_code("CUST", "customer")
    await db.customers.insert_one(doc)
    return serialize(doc)

@api.get("/customers")
async def list_customers(user=Depends(get_current_user)):
    return await list_collection(db.customers)

@api.put("/customers/{cid}")
async def update_customer(cid: str, c: Customer, user=Depends(get_current_user)):
    data = c.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.customers.update_one({"id": cid}, {"$set": data})
    return {"ok": True}

# ---------------- Recycle bin (soft delete + restore) ----------------
RECYCLE_TITLE_FIELDS = ["code", "name", "part_number", "sku", "title", "id"]
async def _recycle(coll: str, label: str, doc, user):
    """Snapshot a document into recycle_bin before it is hard-deleted, so it can be restored."""
    if not doc:
        return
    snap = {k: v for k, v in doc.items() if k != "_id"}
    title = next((str(snap[f]) for f in RECYCLE_TITLE_FIELDS if snap.get(f)), snap.get("id", ""))
    await db.recycle_bin.insert_one({
        "id": new_id(), "coll": coll, "label": label, "title": title,
        "doc": snap, "deleted_at": now_iso(),
        "deleted_by": (user.get("name") or user.get("email", "")) if isinstance(user, dict) else "",
    })

@api.get("/recycle-bin")
async def list_recycle_bin(user=Depends(get_current_user)):
    return await db.recycle_bin.find({}, {"_id": 0, "doc": 0}).sort("deleted_at", -1).to_list(2000)

@api.post("/recycle-bin/{rid}/restore")
async def restore_recycled(rid: str, user=Depends(require_roles("admin", "manager"))):
    rec = await db.recycle_bin.find_one({"id": rid})
    if not rec:
        raise HTTPException(404, "Item not found in recycle bin")
    coll = rec.get("coll")
    doc = {k: v for k, v in (rec.get("doc") or {}).items() if k != "_id"}
    if doc.get("id") and await db[coll].find_one({"id": doc["id"]}):
        raise HTTPException(400, "A live record with this id already exists")
    entries_backup = doc.pop("_entries_backup", None)
    await db[coll].insert_one(doc)
    if entries_backup:   # registers carry their rows along with the template
        clean = [{k: v for k, v in e.items() if k != "_id"} for e in entries_backup]
        if clean:
            await db.register_entries.insert_many(clean)
    await db.recycle_bin.delete_one({"id": rid})
    return {"ok": True, "restored_to": coll}

@api.delete("/recycle-bin/{rid}")
async def purge_recycled(rid: str, user=Depends(require_roles("admin", "manager"))):
    await db.recycle_bin.delete_one({"id": rid})
    return {"ok": True}

@api.delete("/recycle-bin")
async def empty_recycle_bin(user=Depends(require_roles("admin"))):
    r = await db.recycle_bin.delete_many({})
    return {"deleted": r.deleted_count}

@api.delete("/customers/{cid}")
async def del_customer(cid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.customers.find_one({"id": cid}, {"_id": 0})
    await _recycle("customers", "Customer", doc, user)
    await db.customers.delete_one({"id": cid})
    return {"ok": True}

# ---------------- Leads ----------------
@api.post("/leads")
async def create_lead(l: Lead, user=Depends(get_current_user)):
    doc = l.model_dump()
    await db.leads.insert_one(doc)
    return serialize(doc)

@api.get("/leads")
async def list_leads(user=Depends(get_current_user)):
    return await list_collection(db.leads)

@api.put("/leads/{lid}")
async def update_lead(lid: str, l: Lead, user=Depends(get_current_user)):
    data = l.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.leads.update_one({"id": lid}, {"$set": data})
    return {"ok": True}

@api.delete("/leads/{lid}")
async def del_lead(lid: str, user=Depends(get_current_user)):
    await db.leads.delete_one({"id": lid})
    return {"ok": True}

# ---------------- Suppliers ----------------
@api.post("/suppliers")
async def create_supplier(s: Supplier, user=Depends(get_current_user)):
    doc = s.model_dump()
    await db.suppliers.insert_one(doc)
    return serialize(doc)

@api.get("/suppliers")
async def list_suppliers(user=Depends(get_current_user)):
    return await list_collection(db.suppliers)

@api.put("/suppliers/{sid}")
async def update_supplier(sid: str, s: Supplier, user=Depends(get_current_user)):
    data = s.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.suppliers.update_one({"id": sid}, {"$set": data})
    return {"ok": True}

@api.delete("/suppliers/{sid}")
async def del_supplier(sid: str, user=Depends(get_current_user)):
    doc = await db.suppliers.find_one({"id": sid}, {"_id": 0})
    await _recycle("suppliers", "Supplier", doc, user)
    await db.suppliers.delete_one({"id": sid})
    return {"ok": True}

# ---------------- Inventory ----------------
@api.post("/inventory/items")
async def create_item(it: InventoryItem, user=Depends(get_current_user)):
    if await db.items.find_one({"sku": it.sku}):
        raise HTTPException(400, "SKU already exists")
    doc = it.model_dump()
    await db.items.insert_one(doc)
    return serialize(doc)

@api.get("/inventory/items")
async def list_items(user=Depends(get_current_user)):
    return await list_collection(db.items, sort_key="name")

@api.put("/inventory/items/{iid}")
async def update_item(iid: str, it: InventoryItem, user=Depends(get_current_user)):
    data = it.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.items.update_one({"id": iid}, {"$set": data})
    return {"ok": True}

@api.post("/inventory/items/bulk-update")
async def bulk_update_items(body: dict, user=Depends(require_roles("admin", "manager"))):
    """Update many items at once (price / stock / reorder / hsn / gst). Body: {updates:[{id, ...fields}]}."""
    updates = body.get("updates") or []
    allowed = {"name", "category", "uom", "unit_cost", "sale_price", "purchase_price", "reorder_level", "hsn", "gst_rate", "qty_on_hand"}
    n = 0
    for u in updates:
        iid = u.get("id")
        if not iid:
            continue
        data = {k: v for k, v in u.items() if k in allowed}
        if not data:
            continue
        await db.items.update_one({"id": iid}, {"$set": data})
        n += 1
    return {"updated": n}

@api.delete("/inventory/items/{iid}")
async def del_item(iid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.items.find_one({"id": iid}, {"_id": 0})
    await _recycle("items", "Inventory Item", doc, user)
    await db.items.delete_one({"id": iid})
    return {"ok": True}

@api.get("/inventory/items/{iid}/ledger")
async def item_ledger(iid: str, user=Depends(get_current_user)):
    """Document-centric detail view for one item, backing the Items dual-pane drill-down (mirrors
    the /parties/{pid}/transactions pattern): merges StockMovement + StockAdjustment (both keyed on
    item_id) into one Type/Number/Date/Quantity/Price-per-Unit/Status list, newest first, plus a
    read-only "used in BOM(s)" summary. BOM linkage is discovered by querying db.boms where any
    line's item_id matches (there is no reverse-reference field stored on the item itself) — an item
    with no BOM matches simply gets an empty list back, never an error, since most raw/consumable
    items are never used in an assembly."""
    item = await db.items.find_one({"id": iid}, {"_id": 0})
    if not item:
        raise HTTPException(404, "Item not found")

    rows = []
    for m in await db.movements.find({"item_id": iid}, {"_id": 0}).to_list(5000):
        qty = float(m.get("qty", 0) or 0)
        mtype = m.get("type")
        label = {"in": "Stock In", "out": "Stock Out", "adjust": "Adjustment",
                  "in_process": "In Process", "transfer": "Transfer"}.get(mtype, mtype or "Movement")
        rows.append({"type": label, "number": m.get("ref") or "", "date": m.get("created_at"),
                     "quantity": qty, "price_per_unit": float(item.get("unit_cost", 0) or 0),
                     "status": m.get("notes") or "", "doc_type": "movement", "id": m.get("id"),
                     "location": m.get("location"), "to_location": m.get("to_location")})
    for a in await db.stock_adjustments.find({"item_id": iid}, {"_id": 0}).to_list(5000):
        qty = float(a.get("qty", 0) or 0)
        rows.append({"type": "Adjustment", "number": a.get("code") or "", "date": a.get("created_at"),
                     "quantity": qty, "price_per_unit": float(a.get("at_price", item.get("unit_cost", 0)) or 0),
                     "status": a.get("reason") or "", "doc_type": "stock_adjustment", "id": a.get("id"),
                     "location": a.get("location")})
    rows.sort(key=lambda r: str(r.get("date") or ""), reverse=True)

    used_in_boms = []
    async for b in db.boms.find({"lines.item_id": iid}, {"_id": 0}):
        used_in_boms.append({"id": b.get("id"), "code": b.get("code"), "product_name": b.get("product_name"),
                              "revision": b.get("revision"), "bom_type": b.get("bom_type"),
                              "is_active": b.get("is_active")})

    return {
        "item": {"id": item.get("id"), "sku": item.get("sku"), "name": item.get("name"),
                  "category": item.get("category"), "uom": item.get("uom"),
                  "qty_on_hand": item.get("qty_on_hand", 0), "qty_in_process": item.get("qty_in_process", 0),
                  "qty_by_location": item.get("qty_by_location", {}), "reorder_level": item.get("reorder_level", 0),
                  "unit_cost": item.get("unit_cost", 0), "sale_price": item.get("sale_price", 0),
                  "purchase_price": item.get("purchase_price", 0), "hsn": item.get("hsn"),
                  "gst_rate": item.get("gst_rate"), "location": item.get("location")},
        "transactions": rows,
        "used_in_boms": used_in_boms,
    }

# ---------------------------------------------------------------------------
# Standard Parts Library endpoints
# ---------------------------------------------------------------------------
async def _ensure_part_lib_seeded():
    if await db.part_library.count_documents({}) == 0:
        await db.part_library.insert_many([dict(p) for p in STANDARD_PARTS_SEED])

@api.get("/part-library")
async def list_part_library(category: Optional[str] = None, q: Optional[str] = None, user=Depends(get_current_user)):
    await _ensure_part_lib_seeded()
    query = {}
    if category and category != "All":
        query["category"] = category
    rows = await list_collection(db.part_library, query=query, sort_key="name")
    if q:
        ql = q.lower()
        rows = [r for r in rows if ql in (r.get("name", "") + " " + r.get("standard", "") + " " + r.get("size", "") + " " + r.get("material", "")).lower()]
    rows.sort(key=lambda r: (r.get("category", ""), r.get("name", "")))
    return rows

@api.get("/part-library/categories")
async def part_library_categories(user=Depends(get_current_user)):
    return PART_LIB_CATEGORIES

@api.post("/part-library")
async def create_part_library(p: PartLibraryItem, user=Depends(get_current_user)):
    p.seeded = False
    doc = p.model_dump()
    await db.part_library.insert_one(doc)
    return serialize(doc)

@api.put("/part-library/{pid}")
async def update_part_library(pid: str, p: PartLibraryItem, user=Depends(get_current_user)):
    data = p.model_dump(); data.pop("id", None); data.pop("created_at", None); data.pop("seeded", None)
    await db.part_library.update_one({"id": pid}, {"$set": data})
    return {"ok": True}

@api.delete("/part-library/{pid}")
async def del_part_library(pid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.part_library.find_one({"id": pid}, {"_id": 0})
    if doc:
        await _recycle("part_library", "Part Library Item", doc, user)
    await db.part_library.delete_one({"id": pid})
    return {"ok": True}

@api.post("/part-library/{pid}/to-inventory")
async def part_library_to_inventory(pid: str, body: Optional[dict] = None, user=Depends(get_current_user)):
    """Create an inventory item from a library part. Optional body: {sku, qty_on_hand, location}."""
    p = await db.part_library.find_one({"id": pid}, {"_id": 0})
    if not p:
        raise HTTPException(404, "Part not found")
    body = body or {}
    sku = (body.get("sku") or "").strip()
    if not sku:
        base = "".join(ch for ch in (p.get("size") or p.get("name") or "PRT").upper() if ch.isalnum())[:10] or "PRT"
        sku = f"{p.get('category','STD')[:3].upper()}-{base}"
        n = 1
        while await db.items.find_one({"sku": sku}):
            n += 1; sku = f"{p.get('category','STD')[:3].upper()}-{base}-{n}"
    elif await db.items.find_one({"sku": sku}):
        raise HTTPException(400, "SKU already exists")
    qty = float(body.get("qty_on_hand") or 0)
    loc = (body.get("location") or "").strip()
    item = InventoryItem(
        sku=sku, name=p["name"], category="consumable" if p.get("category") == "Fastener" else "raw",
        uom=p.get("uom", "Nos"), qty_on_hand=qty, unit_cost=float(p.get("unit_cost") or 0),
        hsn=p.get("hsn", ""), gst_rate=float(p.get("gst_rate") or 18.0), location=loc,
        qty_by_location=({loc: qty} if loc and qty else {}),
    )
    doc = item.model_dump()
    await db.items.insert_one(doc)
    return serialize(doc)

# ---------------------------------------------------------------------------
# Admin: Reset Trial Data — hard-purge selected groups (for clean re-import).
# Protects system collections (users, settings, counters, ISO docs, registers,
# email accounts, etc.) — they can never be touched by this tool.
# ---------------------------------------------------------------------------
RESET_GROUPS = {
    "parties":    {"label": "Parties (Customers, Suppliers, Leads)",
                   "cols": ["customers", "suppliers", "leads"]},
    "inventory":  {"label": "Inventory items & stock movements",
                   "cols": ["items", "movements", "material_state_movements"]},
    "sales":      {"label": "Sales (Quotations, Sale Orders, Invoices, Proforma, Sale Returns, Payments-In, Credit Notes, Delivery Challans)",
                   "cols": ["quotations", "sale_orders", "invoices", "proforma_invoices", "sale_returns", "payments_in", "credit_notes", "delivery_challans"]},
    "purchase":   {"label": "Purchase (POs, Vendor Bills, Purchase Returns, Payments-Out, Expenses, Debit Notes)",
                   "cols": ["purchase_orders", "vendor_bills", "purchase_returns", "payments_out", "expenses", "debit_notes"]},
    "production": {"label": "Production (Work Orders, Operations, Job Cards, BOMs, Part Master, QC, NCR, CAPA, Job-Work-Out)",
                   "cols": ["work_orders", "wo_operations", "job_cards", "boms", "parts", "qc_inspections", "qc_reports", "ncrs", "capas", "job_work_out"]},
}
RESET_PROTECTED = {"users", "settings", "counters", "audit_logs", "email_accounts", "iso_documents",
                   "register_templates", "register_entries", "trial_requests", "recycle_bin",
                   "webhook_events", "machines", "instruments", "employees", "attendance"}

@api.get("/admin/reset-data/preview")
async def reset_data_preview(user=Depends(require_roles("admin"))):
    out = []
    for key, g in RESET_GROUPS.items():
        counts = {}
        total = 0
        for c in g["cols"]:
            n = await db[c].count_documents({})
            counts[c] = n
            total += n
        out.append({"key": key, "label": g["label"], "cols": g["cols"], "counts": counts, "total": total})
    return {"groups": out}

@api.post("/admin/reset-data")
async def reset_data(body: dict, user=Depends(require_roles("admin"))):
    """Hard-purge chosen groups. Body: {groups:[...], confirm:"RESET"}. Irreversible."""
    if (body or {}).get("confirm") != "RESET":
        raise HTTPException(400, "Confirmation text must be exactly RESET")
    groups = (body or {}).get("groups") or []
    if not groups:
        raise HTTPException(400, "No groups selected")
    deleted = {}
    cols_done = set()
    for key in groups:
        g = RESET_GROUPS.get(key)
        if not g:
            continue
        for c in g["cols"]:
            if c in RESET_PROTECTED or c in cols_done:
                continue
            res = await db[c].delete_many({})
            deleted[c] = res.deleted_count
            cols_done.add(c)
    # also clear recycle-bin entries for the wiped collections so nothing lingers
    if cols_done:
        await db.recycle_bin.delete_many({"collection": {"$in": list(cols_done)}})
    try:
        await db.audit_logs.insert_one({
            "id": new_id(), "action": "reset_trial_data", "by_user": (user or {}).get("email", ""),
            "groups": groups, "deleted": deleted, "created_at": now_iso(),
        })
    except Exception:
        pass
    return {"ok": True, "deleted": deleted, "total": sum(deleted.values())}

@api.post("/admin/seed-sample")
async def seed_sample_data(body: Optional[dict] = None, user=Depends(require_roles("admin"))):
    """Insert a coherent SAMPLE dataset (realistic items/HSN, fictional parties) across
    sales, purchase, payments, orders & expenses so every screen can be verified. Idempotent
    unless force=true. All docs carry seeded_demo=True."""
    body = body or {}
    if not body.get("force") and await db.invoices.count_documents({}) > 0:
        raise HTTPException(400, "ERP already has invoices. Pass force=true to seed anyway, or Reset first.")
    from datetime import timedelta as _td
    D = lambda days: (datetime.utcnow().date() - _td(days=days)).isoformat()
    tag = {"seeded_demo": True}

    def line(desc, code, hsn, qty, unit, rate, gst):
        return {"description": desc, "item_code": code, "hsn": hsn, "qty": qty, "unit": unit,
                "rate": rate, "discount_pct": 0, "discount_amount": 0, "gst_rate": gst}

    def amt(l):
        a = float(l["qty"]) * float(l["rate"]); a -= a * float(l.get("discount_pct", 0)) / 100; a -= float(l.get("discount_amount", 0))
        return a if a > 0 else 0

    def totals(lines, itype, inter, round_off=0):
        sub = sum(amt(l) for l in lines)
        tax = sum(amt(l) * float(l["gst_rate"]) / 100 for l in lines) if itype == "gst" else 0.0
        out = {"subtotal": round(sub, 2), "cgst": 0.0, "sgst": 0.0, "igst": 0.0, "gst_total": round(tax, 2), "total": round(sub + tax + round_off, 2)}
        if itype == "gst":
            if inter: out["igst"] = round(tax, 2)
            else: out["cgst"] = round(tax / 2, 2); out["sgst"] = round(tax / 2, 2)
        return out

    # ---- Items (real catalogue) ----
    items_spec = [
        ("ALU-ROD-50", "ROD ALU Ø50.8MM", "7604", "Kg", 533, 420, 18),
        ("HYD-ROD-80", "HYD CYL PISTON ROD 80MM", "7222", "Nos", 8870, 7000, 18),
        ("CAGE-SS304", "CAGE ARM - SS304", "8487", "Nos", 12760, 9500, 5),
        ("MS-PLT-16", "MS PLATE 600X600X16MM", "72082630", "Kg", 90, 70, 18),
        ("PEEK-FLAT", "PEEK FLAT 1240X62X25MM", "8483", "Nos", 38500, 31000, 5),
        ("BRASS-P8", "BRASS PIPE Ø8MM", "7412", "Mtr", 650, 480, 12),
        ("EN24-45", "ROUND BAR EN24 Ø45MM", "7228", "Kg", 225, 180, 18),
        ("MS-CH-100", "MS CHANNEL 100X50X6MM", "72163100", "Kg", 66, 52, 18),
        ("JOBWORK", "VMC Machining Job Work", "9988", "Nos", 5000, 0, 18),
        ("FREIGHT", "Freight", "9965", "Nos", 2000, 0, 18),
    ]
    items = []
    for sku, name, hsn, uom, sale, cost, gst in items_spec:
        it = {"id": new_id(), "sku": sku, "name": name, "category": "raw", "uom": uom,
              "qty_on_hand": 100, "qty_in_process": 0, "qty_by_location": {"Vatva": 60, "Santej": 40},
              "reorder_level": 10, "unit_cost": cost, "hsn": hsn, "gst_rate": gst, "location": "Vatva",
              "created_at": now_iso(), **tag}
        items.append(it)
    await db.items.insert_many([dict(i) for i in items])

    # ---- Customers (mix intra Gujarat / interstate / export) ----
    cust_spec = [
        ("Shakti Auto Components", "24ABCFS1234A1Z5", "Gujarat", False, "gst"),
        ("Anand Engineering Works", "24AAECA5678B1Z2", "Gujarat", False, "gst"),
        ("Deccan Hydraulics Pvt Ltd", "36AADCD9012C1Z8", "Telangana", True, "gst"),
        ("Northern Tooling Co", "06AAFCN3456D1Z1", "Haryana", True, "gst"),
        ("Surya Fabricators", "", "Gujarat", False, "non_gst"),
        ("Global Exports FZE", "", "Other Territory", True, "export"),
    ]
    customers = []
    for nm, gstin, state, inter, itype in cust_spec:
        customers.append({"id": new_id(), "name": nm, "gstin": gstin, "address": state,
                          "phone": "90000000" + str(10 + len(customers)), "email": "", "customer_type": "repeat",
                          "orders_count": 0, "created_at": now_iso(), "_state": state, "_inter": inter, "_itype": itype, **tag})
    await db.customers.insert_many([{k: v for k, v in c.items() if not k.startswith("_")} for c in customers])

    # ---- Suppliers ----
    sup_spec = [
        ("Maruti Steel Traders", "24AAACM1111E1Z3", "Gujarat", False),
        ("Bombay Alloys & Metals", "27AABCB2222F1Z9", "Maharashtra", True),
        ("Sanghvi Hardware", "24AAGFS3333G1Z7", "Gujarat", False),
    ]
    suppliers = []
    for nm, gstin, state, inter in sup_spec:
        suppliers.append({"id": new_id(), "name": nm, "gstin": gstin, "address": state,
                          "phone": "98000000" + str(10 + len(suppliers)), "email": "", "created_at": now_iso(),
                          "_state": state, "_inter": inter, **tag})
    await db.suppliers.insert_many([{k: v for k, v in s.items() if not k.startswith("_")} for s in suppliers])

    I = {i["sku"]: i for i in items}
    sale_price = {sku: sp for (sku, _n, _h, _u, sp, _c, _g) in items_spec}
    def SL(sku, qty):
        it = I[sku]; return line(it["name"], sku, it["hsn"], qty, it["uom"], sale_price[sku], it["gst_rate"])

    # ---- Sale Invoices ----
    inv_plan = [
        ("DEMO/26-0001", 0, [("HYD-ROD-80", 4), ("FREIGHT", 1)], 22),
        ("DEMO/26-0002", 1, [("MS-PLT-16", 120), ("MS-CH-100", 80)], 18),
        ("DEMO/26-0003", 2, [("CAGE-SS304", 6)], 14),     # interstate, 5%
        ("DEMO/26-0004", 3, [("EN24-45", 200), ("JOBWORK", 3)], 10),  # interstate 18%
        ("DEMO/26-0005", 0, [("BRASS-P8", 50)], 7),
        ("DEMO/26-0006", 1, [("ALU-ROD-50", 90)], 30),    # older -> overdue
        ("DEMO/26-0007", 4, [("MS-PLT-16", 60)], 5),      # non-GST customer
        ("DEMO/26-0008", 5, [("PEEK-FLAT", 2)], 12),      # export
    ]
    invoices = []
    for code, ci, lns, ago in inv_plan:
        c = customers[ci]; itype = c["_itype"]; inter = c["_inter"]
        lines = [SL(sku, q) for sku, q in lns]
        t = totals(lines, itype, inter)
        inv = {"id": new_id(), "code": code, "customer_id": c["id"], "customer_name": c["name"],
               "customer_gstin": c["gstin"], "place_of_supply": c["_state"], "is_interstate": inter,
               "invoice_type": itype, "date": D(ago), "due_date": D(ago - 30), "lines": lines,
               "round_off": 0, "tds": 0, "status": "sent", "godown": "Vatva", "created_at": now_iso(), **t, **tag}
        invoices.append(inv)
    await db.invoices.insert_many([dict(i) for i in invoices])

    # ---- Purchase Bills ----
    bill_plan = [
        ("BILL-MS-7781", 0, [("MS-PLT-16", 300), ("MS-CH-100", 150)], 20),
        ("BILL-ALY-552", 1, [("HYD-ROD-80", 5)], 16),       # interstate
        ("BILL-HW-1043", 2, [("BRASS-P8", 100), ("EN24-45", 250)], 9),
        ("BILL-MS-7802", 0, [("ALU-ROD-50", 120)], 6),
        ("BILL-ALY-560", 1, [("PEEK-FLAT", 2)], 3),         # interstate 5%
    ]
    bills = []
    for code, si, lns, ago in bill_plan:
        s = suppliers[si]; inter = s["_inter"]
        lines = [line(I[sku]["name"], sku, I[sku]["hsn"], q, I[sku]["uom"], I[sku]["unit_cost"], I[sku]["gst_rate"]) for sku, q in lns]
        t = totals(lines, "gst", inter)
        bills.append({"id": new_id(), "code": code, "supplier_id": s["id"], "supplier_name": s["name"],
                      "supplier_gstin": s["gstin"], "place_of_supply": s["_state"], "is_interstate": inter,
                      "date": D(ago), "due_date": D(ago - 30), "lines": lines, "round_off": 0,
                      "status": "unpaid", "created_at": now_iso(), **t, **tag})
    await db.vendor_bills.insert_many([dict(b) for b in bills])

    # ---- Payments In (with allocations -> bill-to-bill) ----
    def alloc(doc, amount, dtype="invoice"):
        return {"document_id": doc["id"], "document_code": doc["code"], "document_type": dtype, "amount": round(amount, 2)}
    pays_in = []
    # full settle inv1; partial inv2; full inv3
    plan_in = [(invoices[0], invoices[0]["total"], 20), (invoices[1], round(invoices[1]["total"] * 0.5, 2), 15), (invoices[2], invoices[2]["total"], 12)]
    for inv, amount, ago in plan_in:
        c = next(x for x in customers if x["id"] == inv["customer_id"])
        pays_in.append({"id": new_id(), "code": f"PMT-IN-{len(pays_in)+1:04d}", "party_id": c["id"], "party_name": c["name"],
                        "date": D(ago), "amount": amount, "allocated_amount": amount, "payment_type": "Bank Transfer",
                        "bank_name": "IndusInd", "allocations": [alloc(inv, amount)],
                        "status": "Used" if abs(amount - inv["total"]) < 0.01 else "Partially Used", "created_at": now_iso(), **tag})
    await db.payments_in.insert_many([dict(p) for p in pays_in])

    # ---- Payments Out ----
    pays_out = []
    plan_out = [(bills[0], bills[0]["total"], 18), (bills[2], round(bills[2]["total"] * 0.6, 2), 7)]
    for b, amount, ago in plan_out:
        s = next(x for x in suppliers if x["id"] == b["supplier_id"])
        pays_out.append({"id": new_id(), "code": f"PMT-OUT-{len(pays_out)+1:04d}", "party_id": s["id"], "party_name": s["name"],
                         "date": D(ago), "amount": amount, "allocated_amount": amount, "payment_type": "Bank Transfer",
                         "bank_name": "IndusInd", "allocations": [alloc(b, amount, "vendor_bill")],
                         "status": "Used" if abs(amount - b["total"]) < 0.01 else "Partially Used", "created_at": now_iso(), **tag})
    await db.payments_out.insert_many([dict(p) for p in pays_out])

    # ---- A Quotation, Proforma, Sale Order, Purchase Order, Expenses ----
    c0, c1 = customers[0], customers[1]; s0 = suppliers[0]
    q_lines = [SL("HYD-ROD-80", 4)]; qt = totals(q_lines, "gst", False)
    await db.quotations.insert_one({"id": new_id(), "code": "QT-26-0001", "customer_id": c0["id"], "customer_name": c0["name"],
        "customer_gstin": c0["gstin"], "place_of_supply": "Gujarat", "is_interstate": False, "date": D(25),
        "lines": q_lines, "status": "sent", "created_at": now_iso(), **qt, **tag})
    pf_lines = [SL("CAGE-SS304", 6)]; pft = totals(pf_lines, "gst", True)
    await db.proforma_invoices.insert_one({"id": new_id(), "code": "PI-26-0001", "customer_id": customers[2]["id"], "customer_name": customers[2]["name"],
        "customer_gstin": customers[2]["gstin"], "place_of_supply": "Telangana", "is_interstate": True, "date": D(20),
        "lines": pf_lines, "status": "sent", "created_at": now_iso(), **pft, **tag})
    so_lines = [SL("EN24-45", 150)]; sot = totals(so_lines, "gst", False)
    await db.sale_orders.insert_one({"id": new_id(), "code": "SO-26-0001", "customer_id": c1["id"], "customer_name": c1["name"],
        "customer_gstin": c1["gstin"], "place_of_supply": "Gujarat", "is_interstate": False, "date": D(18),
        "lines": so_lines, "status": "confirmed", "created_at": now_iso(), **sot, **tag})
    po_lines = [line(I["MS-PLT-16"]["name"], "MS-PLT-16", I["MS-PLT-16"]["hsn"], 500, "Kg", I["MS-PLT-16"]["unit_cost"], 18)]
    pot = totals(po_lines, "gst", False)
    await db.purchase_orders.insert_one({"id": new_id(), "code": "PO-26-0001", "supplier_id": s0["id"], "supplier_name": s0["name"],
        "supplier_gstin": s0["gstin"], "place_of_supply": "Gujarat", "is_interstate": False, "date": D(15),
        "lines": po_lines, "status": "sent", "created_at": now_iso(), **pot, **tag})
    await db.expenses.insert_many([
        {"id": new_id(), "code": "EXP-0001", "category": "Power & Fuel", "party_name": "Torrent Power", "date": D(12),
         "amount": 28500, "total": 28500, "notes": "Electricity - Vatva", "created_at": now_iso(), **tag},
        {"id": new_id(), "code": "EXP-0002", "category": "Consumables", "party_name": "Local Hardware", "date": D(6),
         "amount": 7400, "total": 7400, "notes": "Cutting tools & inserts", "created_at": now_iso(), **tag},
    ])

    return {"ok": True, "seeded": {
        "items": len(items), "customers": len(customers), "suppliers": len(suppliers),
        "invoices": len(invoices), "vendor_bills": len(bills), "payments_in": len(pays_in),
        "payments_out": len(pays_out), "quotations": 1, "proforma": 1, "sale_orders": 1,
        "purchase_orders": 1, "expenses": 2,
    }}

# ===========================================================================
# Bank & Cash Accounts  +  Document Masters  (Vyapar full-replacement modules)
# ===========================================================================
class FinAccount(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    name: str
    type: Literal["bank", "cash"] = "bank"
    opening_balance: float = 0
    opening_date: Optional[str] = ""
    bank_name: Optional[str] = ""
    account_no: Optional[str] = ""
    ifsc: Optional[str] = ""
    upi: Optional[str] = ""
    is_default: bool = False
    created_at: str = Field(default_factory=now_iso)

async def _ensure_accounts_seeded():
    if await db.fin_accounts.count_documents({}) == 0:
        await db.fin_accounts.insert_one(FinAccount(name="Cash", type="cash", is_default=True).model_dump())

def _acct_matches(p, acct):
    return p.get("account_id") == acct["id"] or (p.get("bank_name") and p.get("bank_name") == acct.get("name")) \
        or (acct.get("type") == "cash" and (p.get("payment_type") == "Cash") and not p.get("account_id") and not p.get("bank_name"))

async def _account_balance(acct):
    pin = await db.payments_in.find({}, {"_id": 0}).to_list(50000)
    pout = await db.payments_out.find({}, {"_id": 0}).to_list(50000)
    bal = float(acct.get("opening_balance", 0))
    for p in pin:
        if _acct_matches(p, acct): bal += float(p.get("amount", 0))
    for p in pout:
        if _acct_matches(p, acct): bal -= float(p.get("amount", 0))
    return round(bal, 2)

@api.get("/accounts")
async def list_accounts(user=Depends(get_current_user)):
    await _ensure_accounts_seeded()
    accts = await list_collection(db.fin_accounts, sort_key="name")
    for a in accts:
        a["balance"] = await _account_balance(a)
    return accts

@api.post("/accounts")
async def create_account(a: FinAccount, user=Depends(get_current_user)):
    doc = a.model_dump()
    if doc.get("is_default"):
        await db.fin_accounts.update_many({}, {"$set": {"is_default": False}})
    await db.fin_accounts.insert_one(doc)
    return serialize(doc)

@api.put("/accounts/{aid}")
async def update_account(aid: str, a: FinAccount, user=Depends(get_current_user)):
    data = a.model_dump(); data.pop("id", None); data.pop("created_at", None)
    if data.get("is_default"):
        await db.fin_accounts.update_many({}, {"$set": {"is_default": False}})
    await db.fin_accounts.update_one({"id": aid}, {"$set": data})
    return {"ok": True}

@api.delete("/accounts/{aid}")
async def delete_account(aid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.fin_accounts.find_one({"id": aid}, {"_id": 0})
    if doc:
        await _recycle("fin_accounts", "Account", doc, user)
    await db.fin_accounts.delete_one({"id": aid})
    return {"ok": True}

@api.get("/accounts/{aid}/ledger")
async def account_ledger(aid: str, date_from: str = "", date_to: str = "", user=Depends(get_current_user)):
    acct = await db.fin_accounts.find_one({"id": aid}, {"_id": 0})
    if not acct:
        raise HTTPException(404, "Account not found")
    pin = await db.payments_in.find({}, {"_id": 0}).to_list(50000)
    pout = await db.payments_out.find({}, {"_id": 0}).to_list(50000)
    rows = []
    for p in pin:
        if _acct_matches(p, acct):
            rows.append({"date": _date10(p.get("date")), "ref": p.get("code", ""), "particulars": p.get("party_name", ""),
                         "mode": p.get("payment_type", ""), "credit": float(p.get("amount", 0)), "debit": 0})
    for p in pout:
        if _acct_matches(p, acct):
            rows.append({"date": _date10(p.get("date")), "ref": p.get("code", ""), "particulars": p.get("party_name", ""),
                         "mode": p.get("payment_type", ""), "credit": 0, "debit": float(p.get("amount", 0))})
    if date_from: rows = [r for r in rows if r["date"] >= date_from]
    if date_to: rows = [r for r in rows if r["date"] <= date_to]
    rows.sort(key=lambda r: r["date"])
    bal = float(acct.get("opening_balance", 0))
    for r in rows:
        bal += r["credit"] - r["debit"]; r["balance"] = round(bal, 2)
    return {"account": acct, "opening_balance": float(acct.get("opening_balance", 0)), "rows": rows, "closing_balance": round(bal, 2)}

@api.get("/cheques")
async def cheque_register(user=Depends(get_current_user)):
    out = []
    for coll, kind in (("payments_in", "in"), ("payments_out", "out")):
        for p in await db[coll].find({"payment_type": "Cheque"}, {"_id": 0}).to_list(20000):
            out.append({"id": p.get("id"), "coll": coll, "direction": kind, "code": p.get("code", ""),
                        "date": _date10(p.get("date")), "party": p.get("party_name", ""), "amount": float(p.get("amount", 0)),
                        "ref_no": p.get("ref_no", ""), "bank_name": p.get("bank_name", ""),
                        "cheque_status": p.get("cheque_status", "Pending")})
    out.sort(key=lambda r: r["date"], reverse=True)
    return out

@api.put("/cheques/{coll}/{pid}/status")
async def update_cheque_status(coll: str, pid: str, body: dict, user=Depends(get_current_user)):
    if coll not in ("payments_in", "payments_out"):
        raise HTTPException(400, "Bad collection")
    status = (body or {}).get("cheque_status")
    if status not in ("Pending", "Cleared", "Bounced"):
        raise HTTPException(400, "Bad status")
    await db[coll].update_one({"id": pid}, {"$set": {"cheque_status": status}})
    return {"ok": True}

# ---- Document Masters (Terms & Conditions, prefixes, payment terms, bank) ----
async def _get_setting(key, default):
    doc = await db.settings.find_one({"_id": key}, {"_id": 0})
    return doc.get("value", default) if doc else default

async def _put_setting(key, value):
    await db.settings.update_one({"_id": key}, {"$set": {"value": value}}, upsert=True)

MASTER_KEYS = {"doc_terms": "masters_doc_terms", "payment_terms": "masters_payment_terms",
               "prefixes": "masters_prefixes", "company_bank": "masters_company_bank",
               "tds_sections": "masters_tds_sections", "doc_custom_fields": "masters_doc_custom_fields"}

DEFAULT_DOC_CUSTOM_FIELDS = [
    {"name": "Transport Name", "enabled": True, "type": "text"},
    {"name": "Vehicle Number", "enabled": True, "type": "text"},
    {"name": "Delivery Date", "enabled": True, "type": "date"},
    {"name": "Delivery Location", "enabled": True, "type": "text"},
]

DEFAULT_TDS_SECTIONS = [
    {"section": "192", "name": "Payment of salary", "rate": 1.0},
    {"section": "194C", "name": "Contractors — HUF/Individual", "rate": 1.0},
    {"section": "194C", "name": "Contractors — Others", "rate": 2.0},
    {"section": "194J", "name": "Technical services", "rate": 2.0},
    {"section": "194J", "name": "Professional services", "rate": 10.0},
]

@api.get("/masters")
async def get_masters(user=Depends(get_current_user)):
    return {
        "doc_terms": await _get_setting(MASTER_KEYS["doc_terms"], {}),
        "payment_terms": await _get_setting(MASTER_KEYS["payment_terms"], []),
        "prefixes": await _get_setting(MASTER_KEYS["prefixes"], {}),
        "company_bank": await _get_setting(MASTER_KEYS["company_bank"], {}),
        "tds_sections": await _get_setting(MASTER_KEYS["tds_sections"], DEFAULT_TDS_SECTIONS),
        "doc_custom_fields": await _get_setting(MASTER_KEYS["doc_custom_fields"], DEFAULT_DOC_CUSTOM_FIELDS),
    }

@api.put("/masters/{section}")
async def put_masters(section: str, body: dict, user=Depends(get_current_user)):
    if section not in MASTER_KEYS:
        raise HTTPException(400, "Unknown masters section")
    await _put_setting(MASTER_KEYS[section], (body or {}).get("value"))
    return {"ok": True}

@api.get("/settings/opening-balances")
async def get_opening_balances(user=Depends(get_current_user)):
    return await _get_setting("opening_balances", {"capital": 0, "opening_stock": 0, "as_of": ""})

@api.put("/settings/opening-balances")
async def put_opening_balances(body: dict, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    b = body or {}
    await _put_setting("opening_balances", {
        "capital": float(b.get("capital", 0) or 0),
        "opening_stock": float(b.get("opening_stock", 0) or 0),
        "as_of": b.get("as_of", "") or "",
    })
    return {"ok": True}

@api.post("/admin/seed-masters")
async def seed_masters(user=Depends(require_roles("admin", "manager"))):
    """Load Denplex's real Vyapar masters (T&C library, FY prefixes, payment terms, bank) as ERP defaults."""
    doc_terms = {
        "Sale Invoice": "*Subject to Ahmedabad jurisdiction only\n1) The bill must be paid within due date otherwise interest @18% will be charged extra\n2) Goods once sold will not be taken back.\n3) Our responsibility ceases on delivery the goods to the carrier.\n4) Payment requested by CASH/CHEQUE/Bank Transfer only\n5) If any rejection or rework occurs please notify within 10 days of material receipt; after that it won't be accepted.",
        "Sale Order": "Thanks for doing business with us!",
        "Delivery Challan": "1. Handle the equipment with care and take all necessary precautions to prevent damage.\n2. Keep the equipment in good working condition and perform routine cleaning and maintenance.\n3. Do not modify, tamper with, or repair the equipment without prior written consent.\n4. If the equipment is damaged, a damage fee will be charged based on the extent of damage.",
        "Estimate Quotation": "•\tPayment Terms : 50% advance, 50% against delivery\n•\tDelivery Time : 15-20 working days\n•\tFreight : Extra at actual\n•\tStock : Ex-works Ahmedabad",
        "Proforma Invoice": "Thanks for doing business with us!",
        "Purchase Bill": "Thanks for doing business with us!",
        "Purchase Order": "1. All consignments must carry proper documents, like invoices, test certificates etc.\n2. Invoice must have PO No., Material No., HSN/SAC Code, Vendor No., full address with our PAN & GSTIN.\n3. Test certificate does not relieve responsibility for product quality; results from our QA dept. shall be final.\n4. If goods are not delivered as per schedule, Denplex may buy elsewhere and supplier shall be liable for actual costs and damages.\n5. On any non-compliance of GST provisions (Returns/E-way/E-invoice) resulting in ITC loss, the supplier shall indemnify such loss with interest & penalty.",
    }
    payment_terms = [
        {"name": "Due on Receipt", "days": 0, "is_default": False},
        {"name": "Net 15", "days": 15, "is_default": False},
        {"name": "Net 30", "days": 30, "is_default": True},
        {"name": "Net 45", "days": 45, "is_default": False},
        {"name": "Net 60", "days": 60, "is_default": False},
    ]
    prefixes = {"invoice": "2627/", "purchase_order": "2025-26/", "proforma": "202526",
                "sale_order": "SO-26-", "credit_note": "CN-26-", "delivery_challan": "DC-26-"}
    company_bank = {"account_name": "Denplex Engineering Company", "bank_name": "IndusInd Bank, Bodakdev",
                    "account_no": "259033338999", "ifsc": "INDB0000232", "upi": "denplexengineering-1@okicici"}
    await _put_setting(MASTER_KEYS["doc_terms"], doc_terms)
    await _put_setting(MASTER_KEYS["payment_terms"], payment_terms)
    await _put_setting(MASTER_KEYS["prefixes"], prefixes)
    await _put_setting(MASTER_KEYS["company_bank"], company_bank)
    await _put_setting(MASTER_KEYS["tds_sections"], DEFAULT_TDS_SECTIONS)
    await _put_setting(MASTER_KEYS["doc_custom_fields"], DEFAULT_DOC_CUSTOM_FIELDS)
    # also create a Bank account from these details if none exists yet
    await _ensure_accounts_seeded()
    if not await db.fin_accounts.find_one({"name": company_bank["account_name"]}):
        await db.fin_accounts.insert_one(FinAccount(
            name=company_bank["account_name"], type="bank", opening_balance=1066594.5,
            bank_name=company_bank["bank_name"], account_no=company_bank["account_no"],
            ifsc=company_bank["ifsc"], upi=company_bank["upi"], is_default=True).model_dump())
    return {"ok": True, "doc_terms": len(doc_terms), "payment_terms": len(payment_terms), "prefixes": len(prefixes)}

@api.post("/inventory/movements")
async def create_movement(m: StockMovement, user=Depends(get_current_user)):
    item = await db.items.find_one({"id": m.item_id}, {"_id": 0})
    if not item:
        raise HTTPException(404, "Item not found")
    qty = float(m.qty)
    new_oh = item["qty_on_hand"]
    new_ip = item.get("qty_in_process", 0)
    loc = (m.location or "").strip()
    by_loc = dict(item.get("qty_by_location") or {})
    cur_loc = float(by_loc.get(loc, 0)) if loc else 0
    if m.type == "in":
        new_oh += qty
        if loc: by_loc[loc] = cur_loc + qty
    elif m.type == "out":
        new_oh -= qty
        if loc: by_loc[loc] = cur_loc - qty
    elif m.type == "adjust":
        # set the chosen location to qty (or the grand total if no location given)
        if loc:
            by_loc[loc] = qty
            new_oh = sum(float(v) for v in by_loc.values())
        else:
            new_oh = qty
    elif m.type == "in_process":
        new_ip += qty
        new_oh -= qty
        if loc: by_loc[loc] = cur_loc - qty
    await db.items.update_one({"id": m.item_id}, {"$set": {"qty_on_hand": new_oh, "qty_in_process": new_ip, "qty_by_location": by_loc}})
    doc = m.model_dump()
    doc["item_sku"] = item["sku"]
    doc["item_name"] = item["name"]
    doc["by_user"] = user["name"]
    await db.movements.insert_one(doc)

    # ===== M.4b hook: auto-record material state movement =====
    try:
        state_from, state_to = "", ""
        if m.type == "in":
            state_to = "raw"
        elif m.type == "out":
            state_from = "raw"
        elif m.type == "in_process":
            state_from = "raw"
            state_to = "wip"
        # "adjust" sets qty directly — no state movement recorded
        if state_from or state_to:
            await record_state_movement(
                item_id=m.item_id,
                item_sku=item["sku"],
                item_name=item["name"],
                qty=qty,
                from_state=state_from,
                to_state=state_to,
                ref_type="StockMovement",
                ref_id=doc.get("id", ""),
                ref_code=(m.ref or "") if hasattr(m, "ref") else "",
                note=(m.notes or "") if hasattr(m, "notes") else "",
                user_email=user.get("email", "") if isinstance(user, dict) else "",
            )
    except Exception as _e:
        # Non-fatal — never fail the main movement because of state tracking
        try: logger.warning(f"Material state hook failed (non-fatal): {_e}")
        except Exception: pass
    # ===== end M.4b hook =====

    return serialize(doc)

@api.get("/inventory/movements")
async def list_movements(user=Depends(get_current_user)):
    return await list_collection(db.movements)

@api.get("/inventory/locations")
async def get_inventory_locations(user=Depends(get_current_user)):
    s = await db.settings.find_one({"_id": "inventory_locations"})
    return {"locations": (s or {}).get("locations", DEFAULT_LOCATIONS)}

@api.put("/inventory/locations")
async def set_inventory_locations(body: dict, user=Depends(require_roles("admin", "manager"))):
    locs = [str(x).strip() for x in (body.get("locations") or []) if str(x).strip()]
    if not locs:
        locs = list(DEFAULT_LOCATIONS)
    await db.settings.update_one({"_id": "inventory_locations"}, {"$set": {"locations": locs}}, upsert=True)
    return {"locations": locs}

@api.post("/inventory/transfer")
async def transfer_stock(t: StockTransfer, user=Depends(get_current_user)):
    """Move qty of an item from one location to another. Grand total qty_on_hand
    is unchanged; only the per-location split shifts. Records a 'transfer' movement."""
    item = await db.items.find_one({"id": t.item_id}, {"_id": 0})
    if not item:
        raise HTTPException(404, "Item not found")
    frm = (t.from_location or "").strip(); to = (t.to_location or "").strip()
    qty = float(t.qty)
    if not frm or not to:
        raise HTTPException(400, "from_location and to_location are required")
    if frm == to:
        raise HTTPException(400, "Source and destination must differ")
    if qty <= 0:
        raise HTTPException(400, "Quantity must be greater than zero")
    by_loc = dict(item.get("qty_by_location") or {})
    if float(by_loc.get(frm, 0)) < qty:
        raise HTTPException(400, f"Not enough stock at {frm} (have {by_loc.get(frm, 0)}, need {qty})")
    by_loc[frm] = float(by_loc.get(frm, 0)) - qty
    by_loc[to] = float(by_loc.get(to, 0)) + qty
    await db.items.update_one({"id": t.item_id}, {"$set": {"qty_by_location": by_loc}})
    doc = StockMovement(item_id=t.item_id, item_sku=item["sku"], item_name=item["name"],
                        type="transfer", qty=qty, location=frm, to_location=to,
                        notes=t.notes or "", by_user=user.get("name", "")).model_dump()
    await db.movements.insert_one(doc)
    return {"ok": True, "qty_by_location": by_loc}

# ---------------------------------------------------------------------------
# Stock Adjustments — add/reduce stock WITH A REASON (Vyapar parity: opening
# stock, damage, correction, scrap...). Auditable register + movement trail.
# ---------------------------------------------------------------------------
STOCK_ADJ_REASONS = ["Opening Stock", "Physical Count Correction", "Damaged", "Lost / Theft",
                     "Scrap", "Sample / Testing", "Internal Use", "Other"]

class StockAdjustmentIn(BaseModel):
    model_config = ConfigDict(extra="ignore")
    item_id: str
    adj_type: Literal["add", "reduce"]
    qty: float
    reason: str = "Other"
    at_price: float = 0            # optional value per unit for the adjustment
    location: Optional[str] = ""
    date: Optional[str] = ""
    notes: Optional[str] = ""

@api.get("/inventory/adjustment-reasons")
async def get_adjustment_reasons(user=Depends(get_current_user)):
    return {"reasons": STOCK_ADJ_REASONS}

@api.post("/inventory/adjustments")
async def create_stock_adjustment(a: StockAdjustmentIn, user=Depends(get_current_user)):
    item = await db.items.find_one({"id": a.item_id}, {"_id": 0})
    if not item:
        raise HTTPException(404, "Item not found")
    qty = float(a.qty)
    if qty <= 0:
        raise HTTPException(400, "Quantity must be greater than zero")
    signed = qty if a.adj_type == "add" else -qty
    new_oh = float(item.get("qty_on_hand") or 0) + signed
    by_loc = dict(item.get("qty_by_location") or {})
    loc = (a.location or "").strip()
    if loc:
        by_loc[loc] = float(by_loc.get(loc, 0)) + signed
    code = await gen_code("ADJ", "stock_adjustment")
    doc = {"id": new_id(), "code": code, "item_id": a.item_id, "item_sku": item["sku"],
           "item_name": item["name"], "adj_type": a.adj_type, "qty": qty,
           "reason": (a.reason or "Other").strip(), "at_price": float(a.at_price or 0),
           "value": round(qty * float(a.at_price or 0), 2), "location": loc,
           "date": (a.date or now_iso()[:10]), "notes": (a.notes or "").strip(),
           "qty_before": float(item.get("qty_on_hand") or 0), "qty_after": new_oh,
           "by_user": user.get("name", ""), "created_at": now_iso()}
    await db.items.update_one({"id": a.item_id}, {"$set": {"qty_on_hand": new_oh, "qty_by_location": by_loc}})
    await db.stock_adjustments.insert_one(doc)
    # movement trail entry
    mv = StockMovement(item_id=a.item_id, item_sku=item["sku"], item_name=item["name"],
                       type=("in" if a.adj_type == "add" else "out"), qty=qty, location=loc,
                       ref=code, notes=f"Stock adjustment: {doc['reason']}",
                       by_user=user.get("name", "")).model_dump()
    await db.movements.insert_one(mv)
    await write_audit(user.get("name", ""), "stock_adjustment", "inventory", doc["id"],
                      {"item": item["name"], "type": a.adj_type, "qty": qty, "reason": doc["reason"]})
    return serialize(doc)

@api.get("/inventory/adjustments")
async def list_stock_adjustments(q: Optional[str] = None, user=Depends(get_current_user)):
    docs = await db.stock_adjustments.find({}, {"_id": 0}).sort("created_at", -1).to_list(5000)
    if q:
        ql = q.lower()
        docs = [d for d in docs if ql in (d.get("item_name", "") + d.get("item_sku", "") + d.get("reason", "") + d.get("code", "")).lower()]
    return docs

@api.delete("/inventory/adjustments/{aid}")
async def delete_stock_adjustment(aid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.stock_adjustments.find_one({"id": aid}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Adjustment not found")
    # reverse the stock effect so delete = undo
    item = await db.items.find_one({"id": doc["item_id"]}, {"_id": 0})
    if item:
        signed = -float(doc["qty"]) if doc["adj_type"] == "add" else float(doc["qty"])
        by_loc = dict(item.get("qty_by_location") or {})
        if doc.get("location"):
            by_loc[doc["location"]] = float(by_loc.get(doc["location"], 0)) + signed
        await db.items.update_one({"id": doc["item_id"]},
                                  {"$set": {"qty_on_hand": float(item.get("qty_on_hand") or 0) + signed,
                                            "qty_by_location": by_loc}})
    _ = await _recycle("stock_adjustments", "Stock Adjustment", doc, user) if "_recycle" in globals() else None
    await db.stock_adjustments.delete_one({"id": aid})
    await write_audit(user.get("name", ""), "stock_adjustment_deleted", "inventory", aid,
                      {"item": doc.get("item_name"), "qty": doc.get("qty"), "reversed": True})
    return {"ok": True, "reversed": True}

BILL_EXTRACT_PROMPT = (
    "You are an accounts clerk for a precision-machining company. Read this purchase bill / "
    "tax invoice and extract its data. Return ONLY a JSON object, no prose: "
    "{\"supplier_name\":str, \"gstin\":str, \"bill_number\":str, \"bill_date\":str, "
    "\"items\":[{\"description\":str, \"hsn\":str, \"qty\":number, \"uom\":str, "
    "\"rate\":number, \"gst_rate\":number, \"amount\":number}], "
    "\"sub_total\":number, \"tax\":number, \"total\":number}. "
    "rate = price per unit before tax; gst_rate = the GST percent (e.g. 18); amount = line total. "
    "Use empty string or 0 when a field is missing. Capture every line item."
)

def _parse_bill_json(text: str):
    import json
    t = (text or "").strip()
    if "```" in t:
        for part in t.split("```"):
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            if part.startswith("{") or part.startswith("["):
                t = part; break
    if not t.startswith("{") and "{" in t and "}" in t:
        t = t[t.index("{"): t.rindex("}") + 1]
    try:
        raw = json.loads(t)
    except Exception:
        return {"supplier_name": "", "bill_number": "", "items": []}
    items = []
    for it in (raw.get("items") or []):
        if not isinstance(it, dict):
            continue
        desc = str(it.get("description") or it.get("name") or "").strip()
        if not desc:
            continue
        items.append({
            "description": desc,
            "hsn": str(it.get("hsn") or "").strip(),
            "qty": it.get("qty") or it.get("quantity") or 0,
            "uom": it.get("uom") or "pcs",
            "rate": it.get("rate") or it.get("price") or 0,
            "gst_rate": it.get("gst_rate") or it.get("gst") or 18,
            "amount": it.get("amount") or 0,
        })
    return {
        "supplier_name": str(raw.get("supplier_name") or "").strip(),
        "gstin": str(raw.get("gstin") or "").strip(),
        "bill_number": str(raw.get("bill_number") or raw.get("invoice_no") or "").strip(),
        "bill_date": str(raw.get("bill_date") or "").strip(),
        "items": items,
        "sub_total": raw.get("sub_total") or 0,
        "tax": raw.get("tax") or 0,
        "total": raw.get("total") or 0,
    }

@api.post("/inventory/scan-bill")
async def scan_bill(payload: BillScanIn, user=Depends(get_current_user)):
    """AI bill OCR via Claude vision. Returns {extracted:{supplier_name, bill_number, items:[...]}}."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "AI bill scanning is not configured. Set ANTHROPIC_API_KEY in the backend environment.")
    raw = (payload.image_base64 or "").split(",")[-1].strip()
    if not raw:
        raise HTTPException(400, "Empty image")
    mime = (payload.mime or "image/jpeg").lower()
    if "pdf" in mime:
        block = {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": raw}}
    else:
        media = mime if mime.startswith("image/") else "image/jpeg"
        block = {"type": "image", "source": {"type": "base64", "media_type": media, "data": raw}}
    body = {"model": QC_VISION_MODEL, "max_tokens": 3000,
            "messages": [{"role": "user", "content": [block, {"type": "text", "text": BILL_EXTRACT_PROMPT}]}]}
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=120) as cx:
            r = await cx.post(f"{ANTHROPIC_BASE_URL}/v1/messages", json=body, headers=headers)
    except Exception as e:
        raise HTTPException(502, f"Could not reach the vision API: {e}")
    if r.status_code >= 400:
        raise HTTPException(502, f"Vision API error {r.status_code}: {r.text[:300]}")
    try:
        data = r.json()
        text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
    except Exception as e:
        raise HTTPException(502, f"Unexpected vision API response: {e}")
    return {"extracted": _parse_bill_json(text)}

# ---------------- BOM ----------------
@api.post("/bom")
async def create_bom(b: BOM, user=Depends(get_current_user)):
    doc = b.model_dump()
    doc["code"] = await gen_code("BOM", "bom")
    # Back-fill parent part name + lines from PartMaster references
    if doc.get("parent_part_id") and not doc.get("parent_part_number"):
        p = await db.parts.find_one({"id": doc["parent_part_id"]}, {"_id": 0, "part_number": 1, "name": 1})
        if p:
            doc["parent_part_number"] = p.get("part_number", "")
            if not doc.get("product_name"): doc["product_name"] = p.get("name", "")
    for line in doc.get("lines", []):
        if line.get("component_part_id") and not line.get("component_part_number"):
            p = await db.parts.find_one({"id": line["component_part_id"]}, {"_id": 0, "part_number": 1, "name": 1, "sourcing": 1})
            if p:
                line["component_part_number"] = p.get("part_number", "")
                line["component_part_name"] = p.get("name", "")
                if not line.get("sourcing"): line["sourcing"] = p.get("sourcing", "")
    if not doc.get("design_code"):
        doc["design_code"] = f"DSGN-{datetime.now(timezone.utc).strftime('%y%m')}-{await next_seq('design'):04d}"
    await db.boms.insert_one(doc)
    return serialize(doc)

@api.get("/bom")
async def list_bom(user=Depends(get_current_user)):
    return await list_collection(db.boms)

@api.put("/bom/{bid}")
async def update_bom(bid: str, b: BOM, user=Depends(get_current_user)):
    data = b.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.boms.update_one({"id": bid}, {"$set": data})
    return {"ok": True}

@api.delete("/bom/{bid}")
async def del_bom(bid: str, user=Depends(get_current_user)):
    await db.boms.delete_one({"id": bid})
    return {"ok": True}

# ---------------- Work Orders ----------------
@api.post("/work-orders")
async def create_wo(w: WorkOrder, user=Depends(get_current_user)):
    doc = w.model_dump()
    doc["code"] = await gen_code("WO", "wo")
    await db.work_orders.insert_one(doc)
    if doc.get("customer_id"):
        await db.customers.update_one({"id": doc["customer_id"]}, {"$inc": {"orders_count": 1}})
        cust = await db.customers.find_one({"id": doc["customer_id"]}, {"_id": 0})
        if cust and cust.get("orders_count", 0) >= 2:
            await db.customers.update_one({"id": doc["customer_id"]}, {"$set": {"customer_type": "repeat"}})
    return serialize(doc)

@api.get("/work-orders")
async def list_wo(user=Depends(get_current_user)):
    return await list_collection(db.work_orders)

async def _wo_status_hook(wid, prev, data, user):
    try:
        old_status = (prev or {}).get("status", "")
        new_status = data.get("status", "")
        TRANSITIONS = {
            ("planned", "in_progress"): ("raw", "wip"),
            ("in_progress", "qc"):       ("wip", "inspection_hold"),
        }
        if old_status == new_status:
            return
        trans = (old_status, new_status)
        if trans not in TRANSITIONS:
            return
        from_state, to_state = TRANSITIONS[trans]
        qty = float(data.get("qty") or (prev or {}).get("qty", 0) or 0)
        if qty <= 0:
            return
        await record_state_movement(
            item_name=(data.get("product") or (prev or {}).get("product", "")),
            part_number=(data.get("part_number") or (prev or {}).get("part_number", "")),
            qty=qty, from_state=from_state, to_state=to_state,
            ref_type="WO", ref_id=wid,
            ref_code=(prev or {}).get("code", "") or wid,
            note=f"WO status: {old_status} -> {new_status}",
            user_email=user.get("email", "") if isinstance(user, dict) else "",
        )
    except Exception as _e:
        try: logger.warning(f"WO hook failed: {_e}")
        except Exception: pass

@api.put("/work-orders/{wid}")
async def update_wo(wid: str, w: WorkOrder, user=Depends(get_current_user)):
    prev = await db.work_orders.find_one({"id": wid}, {"_id": 0})
    data = w.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.work_orders.update_one({"id": wid}, {"$set": data})
    await _wo_status_hook(wid, prev, data, user)
    return {"ok": True}

@api.delete("/work-orders/{wid}")
async def del_wo(wid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.work_orders.find_one({"id": wid}, {"_id": 0})
    await _recycle("work_orders", "Work Order", doc, user)
    await db.work_orders.delete_one({"id": wid})
    return {"ok": True}

# ---------------- Job Cards ----------------
@api.post("/job-cards")
async def create_jc(j: JobCard, user=Depends(get_current_user)):
    doc = j.model_dump()
    doc["code"] = await gen_code("JC", "jc")
    wo = await db.work_orders.find_one({"id": j.work_order_id}, {"_id": 0})
    if wo:
        doc["work_order_code"] = wo.get("code", "")
    await db.job_cards.insert_one(doc)
    return serialize(doc)

@api.get("/job-cards")
async def list_jc(user=Depends(get_current_user)):
    return await list_collection(db.job_cards)

@api.put("/job-cards/{jid}")
async def update_jc(jid: str, j: JobCard, user=Depends(get_current_user)):
    data = j.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.job_cards.update_one({"id": jid}, {"$set": data})
    return {"ok": True}

@api.delete("/job-cards/{jid}")
async def del_jc(jid: str, user=Depends(get_current_user)):
    await db.job_cards.delete_one({"id": jid})
    return {"ok": True}

# ---------------- Machines (master) ----------------
class Machine(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    name: str
    machine_type: Optional[str] = ""        # CNC Turning, VMC, Surface Grinder, Bandsaw ...
    group: Optional[str] = ""               # work-center group (for future capacity planning)
    status: Literal["available", "running", "maintenance", "idle"] = "available"
    hourly_rate: float = 0                  # machine-hour cost (for future job costing)
    location: Optional[str] = ""
    is_active: bool = True
    notes: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

@api.post("/machines")
async def create_machine(m: Machine, user=Depends(get_current_user)):
    doc = m.model_dump()
    if not doc.get("code"):
        doc["code"] = await gen_code("MC", "machine")
    await db.machines.insert_one(doc)
    return serialize(doc)

@api.get("/machines")
async def list_machines(user=Depends(get_current_user)):
    return await list_collection(db.machines)

@api.put("/machines/{mid}")
async def update_machine(mid: str, m: Machine, user=Depends(get_current_user)):
    data = m.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.machines.update_one({"id": mid}, {"$set": data})
    return {"ok": True}

@api.delete("/machines/{mid}")
async def del_machine(mid: str, user=Depends(require_roles("admin", "manager"))):
    await db.machines.delete_one({"id": mid})
    return {"ok": True}


# ---------------- Work Order Operations (MES routing) ----------------
class WOOperation(BaseModel):
    """A single routing operation on a Work Order — the MES layer.
    Operation -> Machine -> Operator -> Status -> Time."""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    work_order_id: str = ""
    work_order_code: Optional[str] = ""
    seq: int = 0                              # routing sequence (10, 20, 30 ...)
    operation: str                            # Cutting, CNC Turning, Grinding, QC ...
    machine: Optional[str] = ""               # machine name / code
    machine_id: Optional[str] = ""
    operator: Optional[str] = ""              # operator name
    operator_id: Optional[str] = ""
    status: Literal["pending", "running", "done", "hold"] = "pending"
    planned_minutes: float = 0
    actual_minutes: float = 0
    qty_done: float = 0
    started_at: Optional[str] = ""
    finished_at: Optional[str] = ""
    notes: Optional[str] = ""
    photos: List[str] = []                    # base64 inspection photos from shopfloor
    created_at: str = Field(default_factory=now_iso)

def _op_minutes(start_iso: str, end_iso: str) -> float:
    try:
        s = datetime.fromisoformat((start_iso or "").replace("Z", "+00:00"))
        e = datetime.fromisoformat((end_iso or "").replace("Z", "+00:00"))
        return round(max(0.0, (e - s).total_seconds() / 60.0), 1)
    except Exception:
        return 0.0

async def _wo_or_404(wid: str):
    wo = await db.work_orders.find_one({"id": wid}, {"_id": 0})
    if not wo:
        raise HTTPException(404, "Work order not found")
    return wo

@api.get("/work-orders/{wid}/operations")
async def list_wo_operations(wid: str, user=Depends(get_current_user)):
    ops = await db.wo_operations.find({"work_order_id": wid}, {"_id": 0}).to_list(500)
    ops.sort(key=lambda o: (o.get("seq", 0), o.get("created_at", "")))
    return ops

@api.post("/work-orders/{wid}/operations")
async def add_wo_operation(wid: str, op: WOOperation, user=Depends(get_current_user)):
    wo = await _wo_or_404(wid)
    doc = op.model_dump()
    doc["work_order_id"] = wid
    doc["work_order_code"] = wo.get("code", "")
    if not doc.get("seq"):
        last = await db.wo_operations.find(
            {"work_order_id": wid}, {"_id": 0, "seq": 1}).sort("seq", -1).to_list(1)
        doc["seq"] = (last[0]["seq"] + 10) if last else 10
    await db.wo_operations.insert_one(doc)
    return serialize(doc)

@api.put("/work-orders/{wid}/operations/{op_id}")
async def update_wo_operation(wid: str, op_id: str, op: WOOperation, user=Depends(get_current_user)):
    data = op.model_dump(); data.pop("id", None); data.pop("created_at", None)
    data.pop("work_order_id", None); data.pop("work_order_code", None)
    if data.get("started_at") and data.get("finished_at"):
        data["actual_minutes"] = _op_minutes(data["started_at"], data["finished_at"])
    await db.wo_operations.update_one({"id": op_id, "work_order_id": wid}, {"$set": data})
    return {"ok": True}

@api.delete("/work-orders/{wid}/operations/{op_id}")
async def del_wo_operation(wid: str, op_id: str, user=Depends(get_current_user)):
    await db.wo_operations.delete_one({"id": op_id, "work_order_id": wid})
    return {"ok": True}

@api.post("/work-orders/{wid}/operations/{op_id}/start")
async def start_wo_operation(wid: str, op_id: str, user=Depends(get_current_user)):
    now = now_iso()
    await db.wo_operations.update_one(
        {"id": op_id, "work_order_id": wid},
        {"$set": {"status": "running", "started_at": now}},
    )
    wo = await db.work_orders.find_one({"id": wid}, {"_id": 0})
    if wo and wo.get("status") == "planned":
        await db.work_orders.update_one({"id": wid}, {"$set": {"status": "in_progress"}})
    return {"ok": True, "started_at": now}

@api.post("/work-orders/{wid}/operations/{op_id}/complete")
async def complete_wo_operation(wid: str, op_id: str, qty_done: Optional[float] = None,
                                user=Depends(get_current_user)):
    op = await db.wo_operations.find_one({"id": op_id, "work_order_id": wid}, {"_id": 0})
    if not op:
        raise HTTPException(404, "Operation not found")
    now = now_iso()
    started = op.get("started_at") or now
    upd = {"status": "done", "finished_at": now, "started_at": started,
           "actual_minutes": _op_minutes(started, now)}
    if qty_done is not None:
        upd["qty_done"] = qty_done
    await db.wo_operations.update_one({"id": op_id, "work_order_id": wid}, {"$set": upd})
    remaining = await db.wo_operations.count_documents(
        {"work_order_id": wid, "status": {"$ne": "done"}})
    if remaining == 0:
        wo = await db.work_orders.find_one({"id": wid}, {"_id": 0})
        if wo and wo.get("status") in ("planned", "in_progress"):
            await db.work_orders.update_one({"id": wid}, {"$set": {"status": "qc"}})
    return {"ok": True, "finished_at": now, "actual_minutes": upd["actual_minutes"]}

@api.post("/work-orders/{wid}/operations/{op_id}/hold")
async def hold_wo_operation(wid: str, op_id: str, user=Depends(get_current_user)):
    """Put an operation on hold (status only — preserves photos/timings)."""
    res = await db.wo_operations.update_one(
        {"id": op_id, "work_order_id": wid}, {"$set": {"status": "hold"}})
    if res.matched_count == 0:
        raise HTTPException(404, "Operation not found")
    return {"ok": True}

class OpPhotoIn(BaseModel):
    photo: str  # base64 (data URL or raw)

@api.post("/work-orders/{wid}/operations/{op_id}/photo")
async def add_op_photo(wid: str, op_id: str, body: OpPhotoIn, user=Depends(get_current_user)):
    """Shopfloor: attach an inspection photo to an operation (mobile)."""
    photo = (body.photo or "").split(",")[-1].strip()  # strip data-URL prefix if present
    if not photo:
        raise HTTPException(400, "Empty photo")
    op = await db.wo_operations.find_one({"id": op_id, "work_order_id": wid}, {"_id": 0, "photos": 1})
    if op is None:
        raise HTTPException(404, "Operation not found")
    photos = (op.get("photos") or [])[:19]  # cap at 20 total
    photos.append(await _drive_offload(photo, f"shopfloor-{wid}-{op_id}-{len(photos)}.jpg", "image/jpeg", "Shopfloor Photos"))
    await db.wo_operations.update_one({"id": op_id, "work_order_id": wid}, {"$set": {"photos": photos}})
    return {"ok": True, "photo_count": len(photos)}

@api.get("/work-orders/{wid}/operations/{op_id}/photos")
async def get_op_photos(wid: str, op_id: str, user=Depends(get_current_user)):
    op = await db.wo_operations.find_one({"id": op_id, "work_order_id": wid}, {"_id": 0, "photos": 1})
    if op is None:
        raise HTTPException(404, "Operation not found")
    out = []
    for ph in (op.get("photos") or []):
        if isinstance(ph, str) and ph.startswith("gdrive:"):
            try:
                data = await _resolve_b64_or_drive(ph)
                out.append("data:image/jpeg;base64," + base64.b64encode(data).decode())
            except Exception:
                out.append("")
        elif isinstance(ph, str) and ph.startswith("data:"):
            out.append(ph)
        else:
            out.append("data:image/jpeg;base64," + str(ph))
    return {"photos": out}

@api.post("/work-orders/{wid}/operations/seed-from-part")
async def seed_operations_from_part(wid: str, user=Depends(get_current_user)):
    """Generate routing operations from the WO part's Part Master process list."""
    wo = await _wo_or_404(wid)
    part = None
    pn = wo.get("part_number") or ""
    if pn:
        part = await db.parts.find_one({"part_number": pn}, {"_id": 0})
    if not part:
        raise HTTPException(404, "No Part Master matches this work order's part number. "
                                 "Create the part (with a process list) first.")
    processes = part.get("process") or []
    if not processes:
        raise HTTPException(400, "The matched Part Master has no process list to seed from.")
    existing = await db.wo_operations.count_documents({"work_order_id": wid})
    if existing:
        raise HTTPException(400, "This work order already has operations. Delete them to re-seed.")
    created = []
    seq = 10
    cyc = float(part.get("cycle_time_minutes") or 0)
    qty = float(wo.get("qty") or 0)
    for p in processes:
        doc = WOOperation(
            work_order_id=wid, work_order_code=wo.get("code", ""),
            seq=seq, operation=str(p),
            planned_minutes=round(cyc * qty, 1) if (cyc and qty) else 0,
        ).model_dump()
        await db.wo_operations.insert_one(doc)
        created.append(serialize(doc))
        seq += 10
    return {"ok": True, "created": created, "count": len(created)}


# ---------------- Helper: totals ----------------
def compute_totals(lines: List[Dict[str, Any]], round_off: float = 0) -> Dict[str, float]:
    subtotal = 0.0; gst_total = 0.0
    for l in lines:
        amt = float(l.get("qty", 0)) * float(l.get("rate", 0))
        amt -= amt * float(l.get("discount_pct", 0) or 0) / 100.0
        amt -= float(l.get("discount_amount", 0) or 0)
        if amt < 0:
            amt = 0.0
        gst = amt * float(l.get("gst_rate", 0)) / 100.0
        subtotal += amt; gst_total += gst
    return {"subtotal": round(subtotal, 2), "gst_total": round(gst_total, 2),
            "total": round(subtotal + gst_total + float(round_off or 0), 2)}

# ---------------- Quotations ----------------
class QuoteEstimateIn(BaseModel):
    image_base64: str = ""
    mime: str = "image/png"
    material: str = ""
    qty: float = 1
    part_name: str = ""

QUOTE_ESTIMATE_PROMPT = (
    "You are a senior estimator at a precision-machining / jigs & fixtures company. From this "
    "engineering drawing, estimate a manufacturing quote AND draft techno-commercial content. "
    "Return ONLY JSON, no prose: {\"part_name\":str, \"process_sequence\":[str], "
    "\"machining_minutes_per_pc\":number, \"material_cost_per_pc\":number, \"machining_cost_per_pc\":number, "
    "\"suggested_unit_price\":number, \"assumptions\":str, \"key_highlights\":[str], "
    "\"technical_specifications\":[str], \"cycle_of_operation\":[str], \"inspection_criteria\":[str], "
    "\"scope_of_buyer\":[str]}. All money in INR. suggested_unit_price = (material + machining) per piece "
    "plus a reasonable margin (~25-35%). key_highlights: 3-5 selling points. technical_specifications: key "
    "material/tolerances/specs. cycle_of_operation: ordered manufacturing or usage steps. inspection_criteria: "
    "what is checked. scope_of_buyer: what the customer must provide. Be realistic for an Indian SME shop."
)

def _loads_tolerant(t):
    """Parse JSON, salvaging truncated output by closing open strings/arrays/objects."""
    import json, re
    t = (t or "").strip()
    if "```" in t:
        for part in t.split("```"):
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            if part.startswith("{"):
                t = part; break
    if not t.startswith("{") and "{" in t:
        t = t[t.index("{"):]
    try:
        return json.loads(t)
    except Exception:
        pass
    instr = False; esc = False; depth_obj = 0; depth_arr = 0
    for ch in t:
        if esc: esc = False; continue
        if ch == "\\": esc = True; continue
        if ch == '"': instr = not instr; continue
        if instr: continue
        if ch == '{': depth_obj += 1
        elif ch == '}': depth_obj -= 1
        elif ch == '[': depth_arr += 1
        elif ch == ']': depth_arr -= 1
    s2 = t + ('"' if instr else "")
    s2 = re.sub(r',\s*$', '', s2.rstrip())
    s2 = s2 + ("]" * max(0, depth_arr)) + ("}" * max(0, depth_obj))
    try:
        return json.loads(s2)
    except Exception:
        return {}

def _parse_quote_json(text: str):
    raw = _loads_tolerant(text)
    if not isinstance(raw, dict) or not raw:
        return {}
    return {
        "part_name": str(raw.get("part_name") or "").strip(),
        "process_sequence": [str(x) for x in (raw.get("process_sequence") or []) if str(x).strip()],
        "machining_minutes_per_pc": raw.get("machining_minutes_per_pc") or 0,
        "material_cost_per_pc": raw.get("material_cost_per_pc") or 0,
        "machining_cost_per_pc": raw.get("machining_cost_per_pc") or 0,
        "suggested_unit_price": raw.get("suggested_unit_price") or 0,
        "assumptions": str(raw.get("assumptions") or "").strip(),
        "key_highlights": [str(x) for x in (raw.get("key_highlights") or []) if str(x).strip()],
        "technical_specifications": [str(x) for x in (raw.get("technical_specifications") or []) if str(x).strip()],
        "cycle_of_operation": [str(x) for x in (raw.get("cycle_of_operation") or []) if str(x).strip()],
        "inspection_criteria": [str(x) for x in (raw.get("inspection_criteria") or []) if str(x).strip()],
        "scope_of_buyer": [str(x) for x in (raw.get("scope_of_buyer") or []) if str(x).strip()],
    }

@api.post("/quotations/ai-estimate")
async def ai_quote_estimate(body: QuoteEstimateIn, user=Depends(get_current_user)):
    """AI Quote Generator: drawing + material -> machining time, costs, process, suggested price."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "AI quote estimation is not configured. Set ANTHROPIC_API_KEY in the backend environment.")
    raw = (body.image_base64 or "").split(",")[-1].strip()
    if not raw:
        raise HTTPException(400, "Please attach a drawing (PDF or image).")
    mime = (body.mime or "image/png").lower()
    if "pdf" in mime:
        block = {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": raw}}
    else:
        media = mime if mime.startswith("image/") else "image/png"
        block = {"type": "image", "source": {"type": "base64", "media_type": media, "data": raw}}
    ctx = f"Material: {body.material or 'unspecified'}. Quantity: {body.qty or 1} pcs. Part name: {body.part_name or '(read from drawing)'}.\n"
    payload = {"model": QC_VISION_MODEL, "max_tokens": 4000,
               "messages": [{"role": "user", "content": [block, {"type": "text", "text": ctx + QUOTE_ESTIMATE_PROMPT}]}]}
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=120) as cx:
            r = await cx.post(f"{ANTHROPIC_BASE_URL}/v1/messages", json=payload, headers=headers)
    except Exception as e:
        raise HTTPException(502, f"Could not reach the vision API: {e}")
    if r.status_code >= 400:
        raise HTTPException(502, f"Vision API error {r.status_code}: {r.text[:300]}")
    try:
        data = r.json()
        text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
    except Exception as e:
        raise HTTPException(502, f"Unexpected vision API response: {e}")
    est = _parse_quote_json(text)
    if not est:
        raise HTTPException(502, "Could not parse an estimate from the drawing.")
    return {"estimate": est}


# ---------------- Quotation Word (.docx) generator ----------------
_QD_BANK = [("GST No.", "24AALFD1671P1Z2"), ("PAN Card No.", "AALFD1671P"),
            ("MSME UDYAM No.", "UDYAM-GJ-09-0005351"), ("Account Name", "DENPLEX ENGINEERING COMPANY"),
            ("Bank Name", "INDUSIND BANK"), ("Branch", "BODAKDEV BRANCH"),
            ("IFSC Code", "INDB0000232"), ("A/C Number", "259033338999")]
_QD_TERMS = [("Payment Terms", "{payment}"),
             ("Delivery Timeline", "Within 35-45 working days from date of PO and receipt of advance payment."),
             ("Installation & Commissioning (I&C)", "{ic}"),
             ("Standard Warranty", "1 year / standard warranty for electronic parts provided by manufacturer."),
             ("Packaging & Forwarding (P&F)", "Charges applicable at actuals."),
             ("Freight", "To be borne by the buyer at actuals."),
             ("GST", "18% applicable as per prevailing tax laws."),
             ("Offer Validity", "{validity} days from the date of quotation."),
             ("Site Visit", "{sitevisit}"),
             ("Jurisdiction", "All disputes are subject to Ahmedabad jurisdiction only.")]

def build_quotation_docx(d, fmt="general"):
    """Render a Denplex-letterhead quotation as editable .docx (general | techno)."""
    import io as _io
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    RED = RGBColor(0xCC, 0x00, 0x00); BLACK = RGBColor(0x1A, 0x1A, 0x1A)
    GREY = RGBColor(0x55, 0x55, 0x55); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    LETTERHEAD = str(ROOT_DIR / "letterhead.docx")

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc); tcPr.append(sh)
    def borders(table, color="999999", sz="4"):
        tblPr = table._tbl.tblPr; b = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
            e.set(qn('w:space'), '0'); e.set(qn('w:color'), color); b.append(e)
        tblPr.append(b)
    def run(par, text, size=10, bold=False, color=BLACK):
        r = par.add_run(text); r.font.size = Pt(size); r.font.bold = bold
        r.font.color.rgb = color; r.font.name = "Arial"; return r
    def li(doc, marker="•  "):
        par = doc.add_paragraph(); par.paragraph_format.left_indent = Mm(5)
        par.paragraph_format.space_after = Pt(2); run(par, marker, 9); return par
    def kv(doc, k, v):
        par = li(doc); run(par, f"{k}: ", 9, True); run(par, v, 9)
    def heading(doc, text):
        par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(6)
        par.paragraph_format.space_after = Pt(2); run(par, text, 12, True, RED); return par
    def inr(n):
        try: return f"₹ {float(n):,.0f}/-"
        except Exception: return str(n)

    try:
        doc = Document(LETTERHEAD)
    except Exception:
        doc = Document()
    doc._body.clear_content()
    try:
        nrm = doc.styles["Normal"].font; nrm.name = "Arial"; nrm.size = Pt(10)
    except Exception:
        pass

    t = doc.add_table(rows=3, cols=2); borders(t)
    run(t.cell(0, 0).paragraphs[0], "TO,", 9, True); run(t.cell(0, 1).paragraphs[0], f"Date: {d['date']}", 9, True)
    run(t.cell(1, 0).paragraphs[0], d.get('attn', ''), 9); run(t.cell(1, 1).paragraphs[0], f"Qtn. No.: {d.get('qtn_no','')}", 9)
    run(t.cell(2, 0).paragraphs[0], d.get('customer', ''), 9, True); run(t.cell(2, 1).paragraphs[0], d.get('customer_addr', ''), 8)
    doc.add_paragraph()

    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; run(par, d.get('title', 'Quotation'), 14, True, BLACK)
    if d.get('subtitle'):
        ps = doc.add_paragraph(); ps.alignment = WD_ALIGN_PARAGRAPH.CENTER; run(ps, d['subtitle'], 10, False, GREY)
    doc.add_paragraph()

    def pricing():
        tb = doc.add_table(rows=1, cols=5); borders(tb)
        for i, cn in enumerate(["Sr.", "Description", "Qty", "Rate (INR)", "Total (INR)"]):
            run(tb.rows[0].cells[i].paragraphs[0], cn, 9, True, WHITE); shade(tb.rows[0].cells[i], "1A1A1A")
        sub = 0
        for i, line in enumerate(d['lines'], 1):
            row = tb.add_row().cells; tot = float(line['qty']) * float(line['rate']); sub += tot
            run(row[0].paragraphs[0], str(i), 9); run(row[1].paragraphs[0], line['description'], 9)
            run(row[2].paragraphs[0], str(line['qty']), 9)
            run(row[3].paragraphs[0], inr(line['rate']) if float(line['rate']) else "Included", 9)
            run(row[4].paragraphs[0], inr(tot) if tot else "Included", 9)
        gst = sub * 0.18
        for label, val, hexc, white in [("Sub Total", sub, "EFEFEF", False), ("GST @ 18%", gst, "EFEFEF", False), ("GRAND TOTAL (INR)", sub + gst, "CC0000", True)]:
            row = tb.add_row().cells; m = row[0].merge(row[1]).merge(row[2]).merge(row[3])
            run(m.paragraphs[0], label, 9, True, WHITE if white else BLACK); m.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run(row[4].paragraphs[0], inr(val), 9, True, WHITE if white else BLACK)
            for cc in row: shade(cc, hexc)

    def photos():
        if d.get('photos'):
            heading(doc, "Concept Image (For Reference)")
            for ph in d['photos'][:3]:
                try: doc.add_paragraph().add_run().add_picture(_io.BytesIO(ph), width=Mm(120))
                except Exception: pass

    def terms(num=""):
        heading(doc, f"{num}Terms & Conditions")
        tv = {"payment": d.get('payment', "50% advance, balance before dispatch"), "ic": d.get('ic', "Not applicable"),
              "validity": str(d.get('validity', 7)), "sitevisit": d.get('sitevisit', "Not applicable")}
        for k, v in _QD_TERMS: kv(doc, k, v.format(**tv))

    def bank(num=""):
        heading(doc, f"{num}Company Details & Bank Details")
        for k, v in _QD_BANK: kv(doc, k, v)

    if fmt == "techno":
        if d.get('highlights'):
            box = doc.add_table(rows=1, cols=1); borders(box, "CC0000"); shade(box.cell(0, 0), "FBE9E9")
            run(box.cell(0, 0).paragraphs[0], "KEY HIGHLIGHTS", 10, True, RED)
            for hl in d['highlights']: run(box.cell(0, 0).add_paragraph(), f"•  {hl}", 9)
            doc.add_paragraph()
        heading(doc, "1.  Pricing Summary"); pricing(); doc.add_paragraph()
        heading(doc, "2.  Technical Specifications")
        for x in d.get('specs', []): run(li(doc), x, 9)
        heading(doc, "3.  Standard Proposed Cycle of Operation")
        for i, x in enumerate(d.get('cycle', []), 1): run(li(doc, f"{i}.  "), x, 9)
        heading(doc, "4.  Inspection Criteria")
        for x in d.get('inspection', []): run(li(doc), x, 9)
        heading(doc, "5.  Scope of Buyer")
        for x in d.get('scope', []): run(li(doc), x, 9)
        photos(); terms("6.  "); bank("7.  ")
    else:
        photos(); pricing(); doc.add_paragraph(); terms(); bank()

    doc.add_paragraph(); run(doc.add_paragraph(), "Looking forward to your kind acknowledgement.", 9, True)
    run(doc.add_paragraph(), "\nFor DENPLEX ENGINEERING COMPANY\n\n\nAuthorised Signatory", 9, True)
    out = _io.BytesIO(); doc.save(out); return out.getvalue()


class QuoteDocLine(BaseModel):
    description: str = ""
    qty: float = 1
    rate: float = 0

class QuoteDocIn(BaseModel):
    format: str = "general"
    date: str = ""
    qtn_no: str = ""
    attn: str = ""
    customer: str = ""
    customer_addr: str = ""
    title: str = "Quotation"
    subtitle: str = ""
    lines: List[QuoteDocLine] = []
    highlights: List[str] = []
    specs: List[str] = []
    cycle: List[str] = []
    inspection: List[str] = []
    scope: List[str] = []
    photos: List[str] = []
    payment: str = ""
    ic: str = ""
    validity: int = 7
    sitevisit: str = ""

@api.post("/quotations/docx")
async def quotation_docx(body: QuoteDocIn, user=Depends(get_current_user)):
    """Generate an editable Word quotation (general | techno) on the Denplex letterhead."""
    photos = []
    for ph in (body.photos or []):
        try: photos.append(base64.b64decode((ph or "").split(",")[-1]))
        except Exception: pass
    data = {
        "date": body.date or datetime.now(timezone.utc).strftime("%d-%m-%Y"),
        "qtn_no": body.qtn_no or "", "attn": body.attn or "",
        "customer": body.customer or "", "customer_addr": body.customer_addr or "",
        "title": body.title or "Quotation", "subtitle": body.subtitle or "",
        "lines": [{"description": l.description, "qty": l.qty, "rate": l.rate} for l in body.lines] or [{"description": "Item", "qty": 1, "rate": 0}],
        "highlights": body.highlights, "specs": body.specs, "cycle": body.cycle,
        "inspection": body.inspection, "scope": body.scope, "photos": photos,
        "payment": body.payment or "50% advance with PO, balance before dispatch",
        "ic": body.ic or "As applicable", "validity": body.validity or 7,
        "sitevisit": body.sitevisit or "As applicable",
    }
    fmt = "techno" if (body.format or "").lower().startswith("techno") else "general"
    try:
        blob = build_quotation_docx(data, fmt)
    except Exception as e:
        raise HTTPException(500, f"Could not generate the Word document: {e}")
    fname = f"Quotation-{(body.qtn_no or 'draft').replace('/', '-')}.docx"
    return Response(content=blob,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={fname}"})

@api.post("/quotations")
async def create_quote(q: Quotation, user=Depends(get_current_user)):
    doc = q.model_dump()
    doc["code"] = await gen_code("QT", "quote")
    t = compute_totals([l for l in doc["lines"]])
    doc.update(t)
    await db.quotations.insert_one(doc)
    return serialize(doc)

@api.get("/quotations")
async def list_quotes(user=Depends(get_current_user)):
    return await list_collection(db.quotations)

@api.put("/quotations/{qid}")
async def update_quote(qid: str, q: Quotation, user=Depends(get_current_user)):
    data = q.model_dump(); data.pop("id", None); data.pop("created_at", None)
    t = compute_totals(data["lines"]); data.update(t)
    await db.quotations.update_one({"id": qid}, {"$set": data})
    return {"ok": True}

@api.delete("/quotations/{qid}")
async def del_quote(qid: str, user=Depends(get_current_user)):
    doc = await db.quotations.find_one({"id": qid}, {"_id": 0})
    await _recycle("quotations", "Quotation", doc, user)
    await db.quotations.delete_one({"id": qid})
    return {"ok": True}

# ---------------- Purchase Orders ----------------
@api.post("/purchase-orders")
async def create_po(p: PurchaseOrder, user=Depends(get_current_user)):
    doc = p.model_dump()
    doc["code"] = (p.code or "").strip() or await gen_code("PO", "po")
    t = compute_totals(doc["lines"], doc.get("round_off", 0)); doc.update(t)
    await db.purchase_orders.insert_one(doc)
    return serialize(doc)

@api.get("/purchase-orders")
async def list_po(user=Depends(get_current_user)):
    return await list_collection(db.purchase_orders)

@api.put("/purchase-orders/{pid}")
async def update_po(pid: str, p: PurchaseOrder, user=Depends(get_current_user)):
    data = p.model_dump(); data.pop("id", None); data.pop("created_at", None)
    t = compute_totals(data["lines"]); data.update(t)
    await db.purchase_orders.update_one({"id": pid}, {"$set": data})
    return {"ok": True}

@api.delete("/purchase-orders/{pid}")
async def del_po(pid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.purchase_orders.find_one({"id": pid}, {"_id": 0})
    await _recycle("purchase_orders", "Purchase Order", doc, user)
    await db.purchase_orders.delete_one({"id": pid})
    return {"ok": True}

# ---------------- Invoices ----------------
def compute_invoice_totals(lines: List[Dict[str, Any]], interstate: bool, round_off: float = 0, tds: float = 0,
                           extra_charges: Optional[List[Dict[str, Any]]] = None, tcs: float = 0) -> Dict[str, float]:
    subtotal = 0.0; gst = 0.0
    for l in lines:
        amt = float(l.get("qty", 0)) * float(l.get("rate", 0))
        # per-line discount: percentage first, then a flat amount
        amt -= amt * float(l.get("discount_pct", 0) or 0) / 100.0
        amt -= float(l.get("discount_amount", 0) or 0)
        if amt < 0:
            amt = 0.0
        g = amt * float(l.get("gst_rate", 0)) / 100.0
        subtotal += amt; gst += g
    charges_total = sum(float(c.get("amount", 0) or 0) for c in (extra_charges or []))
    total = subtotal + gst + charges_total + float(round_off or 0) - float(tds or 0) + float(tcs or 0)
    extra = {"charges_total": round(charges_total, 2), "tds": round(float(tds or 0), 2), "tcs": round(float(tcs or 0), 2)}
    if interstate:
        return {"subtotal": round(subtotal, 2), "cgst": 0.0, "sgst": 0.0, "igst": round(gst, 2), "total": round(total, 2), **extra}
    return {"subtotal": round(subtotal, 2), "cgst": round(gst/2, 2), "sgst": round(gst/2, 2), "igst": 0.0, "total": round(total, 2), **extra}

@api.post("/invoices")
async def create_invoice(inv: Invoice, user=Depends(get_current_user)):
    from datetime import timedelta
    doc = inv.model_dump()
    doc["code"] = (inv.code or "").strip() or await gen_code("INV", "invoice")
    doc.update(compute_invoice_totals(doc["lines"], doc.get("is_interstate", False), doc.get("round_off", 0), doc.get("tds", 0), doc.get("extra_charges"), doc.get("tcs", 0)))
    # If an e-way bill number is entered at creation, compute its validity window.
    if doc.get("eway_bill_no") and not doc.get("eway_valid_until"):
        gen = (doc.get("eway_generated_at") or str(doc.get("date", ""))[:10] or datetime.utcnow().date().isoformat())[:10]
        try:
            gen_d = datetime.strptime(gen, "%Y-%m-%d").date()
        except Exception:
            gen_d = datetime.utcnow().date(); gen = gen_d.isoformat()
        days = _eway_validity_days(doc.get("eway_distance_km", 0), doc.get("eway_over_dimensional", False))
        doc["eway_generated_at"] = gen
        doc["eway_valid_until"] = (gen_d + timedelta(days=days)).isoformat()
    await db.invoices.insert_one(doc)
    return serialize(doc)

@api.get("/invoices")
async def list_invoices(user=Depends(get_current_user)):
    return await list_collection(db.invoices)

def _eway_validity_days(distance_km: float, over_dimensional: bool = False) -> int:
    """E-way bill validity: 1 day per 200 km (normal cargo), 1 day per 20 km (over-dimensional). Min 1 day."""
    import math
    per = 20 if over_dimensional else 200
    dist = float(distance_km or 0)
    return max(1, math.ceil(dist / per)) if dist > 0 else 1

def _eway_edit_locked(inv: Dict[str, Any]) -> bool:
    """True when an invoice carries an e-way bill whose validity period has ended → editing is locked."""
    vu = (inv or {}).get("eway_valid_until")
    if (inv or {}).get("eway_bill_no") and vu:
        try:
            return datetime.strptime(str(vu)[:10], "%Y-%m-%d").date() < datetime.utcnow().date()
        except Exception:
            return False
    return False

class EwayBillIn(BaseModel):
    eway_bill_no: Optional[str] = ""
    distance_km: float = 0
    generated_at: Optional[str] = ""           # ISO date; defaults to today
    over_dimensional: bool = False
    details: dict = Field(default_factory=dict)   # full NIC form: dispatch/ship/transport/Part-B/value

@api.post("/invoices/{iid}/eway-bill")
async def record_eway_bill(iid: str, body: EwayBillIn, user=Depends(get_current_user)):
    """Record an e-way bill (and its full NIC form data) against an invoice + auto-compute validity.
    NOTE: actually filing to the government NIC system needs a GSP API connection; until then this
    captures the prepared payload and the number you generate on the NIC portal."""
    from datetime import timedelta
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv:
        raise HTTPException(404, "Invoice not found")
    gen = (body.generated_at or datetime.utcnow().date().isoformat())[:10]
    try:
        gen_d = datetime.strptime(gen, "%Y-%m-%d").date()
    except Exception:
        gen_d = datetime.utcnow().date(); gen = gen_d.isoformat()
    days = _eway_validity_days(body.distance_km, body.over_dimensional)
    valid_until = (gen_d + timedelta(days=days)).isoformat()
    upd = {"eway_bill_no": (body.eway_bill_no or "").strip(), "eway_generated_at": gen,
           "eway_distance_km": float(body.distance_km or 0), "eway_over_dimensional": bool(body.over_dimensional),
           "eway_valid_until": valid_until, "eway_details": body.details or {}}
    await db.invoices.update_one({"id": iid}, {"$set": upd})
    return {"ok": True, "valid_days": days, **upd}

@api.get("/invoices/settled-summary")
async def invoices_settled_summary(user=Depends(get_current_user)):
    """Bulk invoice_id -> settled amount map (reuses _settled_per_invoice with no filter) — powers
    the Sales report page's Total/Received/Balance summary cards + per-row Balance column without
    an N+1 call per invoice. MUST be registered before /invoices/{iid} — FastAPI matches routes in
    registration order, so a literal path like this needs to come first or the {iid} pattern below
    swallows it (iid="settled-summary") and returns a false 404. Hit this bug once already; if you
    add another literal /invoices/<word> route, put it up here too."""
    return await _settled_per_invoice()

@api.get("/invoices/{iid}")
async def get_invoice(iid: str, user=Depends(get_current_user)):
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv:
        raise HTTPException(404, "Invoice not found")
    return inv

@api.get("/invoices/{iid}/payments")
async def invoice_payments(iid: str, user=Depends(get_current_user)):
    """Payment-In records allocated against this invoice, plus settled/balance — the data behind the
    Invoice dual-pane detail's payment history (mirrors the party-transactions merge pattern, scoped
    to a single invoice instead of a whole party). Reuses the same allocations shape as
    _settled_per_invoice / the payments-in module."""
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv:
        raise HTTPException(404, "Invoice not found")
    pays = await db.payments_in.find({"allocations.document_id": iid}, {"_id": 0}).to_list(2000)
    rows = []
    settled = 0.0
    for p in pays:
        for a in (p.get("allocations") or []):
            if a.get("document_id") == iid and a.get("document_type") == "invoice":
                amt = float(a.get("amount") or 0) + float(a.get("tds_amount") or 0)
                settled += amt
                rows.append({"payment_id": p.get("id"), "payment_code": p.get("code"), "date": p.get("date"),
                             "amount": round(amt, 2), "payment_type": p.get("payment_type")})
    rows.sort(key=lambda r: str(r.get("date") or ""), reverse=True)
    total = float(inv.get("total", 0) or 0)
    return {"total": round(total, 2), "settled": round(settled, 2), "balance": round(total - settled, 2), "payments": rows}

# ---------------- E-way bill / e-Invoice — Adaequare GSP (Enriched APIs) ----------------
# Adaequare Info Pvt Ltd is our chosen GSP (confirmed July 2026). Enriched APIs handle NIC/IRP
# encryption+session management for us — we call Adaequare's own gateway with plain JSON.
# SECRETS are read from Railway env vars (never stored in code/db):
#   EWAY_GSP_CLIENT_ID      -> Adaequare "gspappid"      (from the App created on gsp.adaequare.com)
#   EWAY_GSP_CLIENT_SECRET  -> Adaequare "gspappsecret"
#   EWAY_GSP_USERNAME       -> the API username created on einvoice1.gst.gov.in / ewaybillgst.gov.in
#                              (Registration -> For GSP -> select "Adaequare Info Private Limited")
#   EWAY_GSP_PASSWORD       -> that API user's password
#   EWAY_GSP_ENV            -> "staging" (default, sandbox) or "live" (production)
#   COMPANY_GSTIN           -> falls back to Settings > company_gstin if unset
# Non-secret overrides (rarely needed) may be tuned in db.settings id="eway_gsp".
ADQ_BASE_URL = "https://gsp.adaequare.com"

async def _eway_gsp_config() -> Dict[str, Any]:
    s = await get_setting("eway_gsp") or {}
    company = await get_setting("integrations") or {}
    env = os.getenv("EWAY_GSP_ENV", s.get("env", "staging")).strip().lower()
    return {
        "base_url": os.getenv("EWAY_GSP_BASE_URL", s.get("base_url", ADQ_BASE_URL)).rstrip("/"),
        "env": "live" if env == "live" else "staging",
        "prefix": "/enriched" if env == "live" else "/test/enriched",
        "client_id": os.getenv("EWAY_GSP_CLIENT_ID", s.get("client_id", "")),        # gspappid
        "client_secret": os.getenv("EWAY_GSP_CLIENT_SECRET", s.get("client_secret", "")),  # gspappsecret
        "username": os.getenv("EWAY_GSP_USERNAME", s.get("username", "")),           # API username on govt portal
        "password": os.getenv("EWAY_GSP_PASSWORD", s.get("password", "")),
        "gstin": os.getenv("COMPANY_GSTIN", company.get("company_gstin", "")),
        "provider": "Adaequare",
    }

def _gsp_ready(c: Dict[str, Any]) -> bool:
    return bool(c.get("base_url") and c.get("client_id") and c.get("client_secret")
                and c.get("username") and c.get("password") and c.get("gstin"))

@api.get("/eway/gsp-status")
async def eway_gsp_status(user=Depends(get_current_user)):
    c = await _eway_gsp_config()
    return {"configured": _gsp_ready(c), "provider": c.get("provider", ""), "gstin": c.get("gstin", ""), "env": c.get("env", "")}

async def _adq_token(c: Dict[str, Any]) -> str:
    """Get an Adaequare GSP access token, cached in db.settings until near expiry.
    Per their docs the token is valid ~6h for Enriched API calls; we refresh 5 min early."""
    cached = await get_setting("eway_gsp_token") or {}
    if cached.get("token") and cached.get("expires_at", "") > now_iso():
        return cached["token"]
    async with httpx.AsyncClient(timeout=30) as cl:
        r = await cl.post(f"{c['base_url']}/gsp/authenticate?action=GSP&grant_type=token",
                          headers={"gspappid": c["client_id"], "gspappsecret": c["client_secret"]})
        r.raise_for_status()
        j = r.json()
    token = j.get("access_token", "")
    if not token:
        raise HTTPException(502, f"Adaequare auth did not return a token: {str(j)[:300]}")
    expires_in = int(j.get("expires_in") or 21600)  # default 6h
    expires_at = (datetime.utcnow() + timedelta(seconds=max(expires_in - 300, 60))).isoformat()
    await set_setting("eway_gsp_token", {"token": token, "expires_at": expires_at})
    return token

def _adq_ewb_headers(c: Dict[str, Any], token: str) -> Dict[str, str]:
    return {"Content-Type": "application/json", "username": c["username"], "password": c["password"],
            "gstin": c["gstin"], "requestid": new_id(), "Authorization": f"Bearer {token}"}

def _adq_ei_headers(c: Dict[str, Any], token: str) -> Dict[str, str]:
    # E-Invoice APIs use "user_name" (underscore) per Adaequare docs, unlike EWB's "username"
    return {"Content-Type": "application/json", "user_name": c["username"], "password": c["password"],
            "gstin": c["gstin"], "requestid": new_id(), "Authorization": f"Bearer {token}"}

def _ddmmyyyy(s: str) -> str:
    try:
        return datetime.strptime(str(s)[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return datetime.utcnow().strftime("%d/%m/%Y")

def _gstin_state_code(gstin: str) -> int:
    """First 2 digits of a GSTIN are the numeric state code (e.g. 24 = Gujarat)."""
    g = (gstin or "").strip()
    try:
        return int(g[:2])
    except Exception:
        return 0

def _build_nic_ewb_payload(inv: Dict[str, Any], d: Dict[str, Any], company: Dict[str, Any]) -> Dict[str, Any]:
    """Adaequare enRiched EWB 'GENEWAYBILL' JSON — field names verified against their API doc sample."""
    inter = bool(inv.get("is_interstate"))
    from_gstin = company.get("company_gstin", "")
    to_gstin = inv.get("customer_gstin", "") or d.get("ship_gstin", "")
    from_state = _gstin_state_code(from_gstin)
    to_state = _gstin_state_code(to_gstin) or from_state
    items = [{
        "productName": (l.get("description", "") or "")[:100], "productDesc": (l.get("description", "") or "")[:100],
        "hsnCode": int(re.sub(r"\D", "", str(l.get("hsn", "") or "0")) or 0),
        "quantity": float(l.get("qty", 0)), "qtyUnit": (l.get("unit", "NOS") or "NOS").upper()[:3],
        "taxableAmount": round(float(l.get("qty", 0)) * float(l.get("rate", 0)), 2),
        "sgstRate": (0 if inter else float(l.get("gst_rate", 0)) / 2),
        "cgstRate": (0 if inter else float(l.get("gst_rate", 0)) / 2),
        "igstRate": (float(l.get("gst_rate", 0)) if inter else 0),
        "cessRate": 0, "cessAdvol": 0,
    } for l in (inv.get("lines", []) or [])]
    return {
        "supplyType": "O", "subSupplyType": "1", "subSupplyDesc": "",
        "docType": "INV", "docNo": inv.get("code", ""), "docDate": _ddmmyyyy(inv.get("date", "")),
        "fromGstin": from_gstin, "fromTrdName": company.get("company_name", ""),
        "fromAddr1": (d.get("dispatch_address", "") or company.get("company_address", ""))[:120], "fromAddr2": "",
        "fromPlace": d.get("dispatch_place", "") or d.get("dispatch_state", ""),
        "fromPincode": int(re.sub(r"\D", "", str(d.get("dispatch_pin", "") or "0")) or 0),
        "actFromStateCode": from_state, "fromStateCode": from_state,
        "toGstin": to_gstin, "toTrdName": inv.get("customer_name", ""),
        "toAddr1": (d.get("ship_address", "") or "")[:120], "toAddr2": "",
        "toPlace": d.get("ship_place", "") or d.get("ship_state", ""),
        "toPincode": int(re.sub(r"\D", "", str(d.get("ship_pin", "") or "0")) or 0),
        "actToStateCode": to_state, "toStateCode": to_state,
        "totalValue": inv.get("subtotal", 0), "cgstValue": inv.get("cgst", 0), "sgstValue": inv.get("sgst", 0),
        "igstValue": inv.get("igst", 0), "cessValue": 0, "totInvValue": inv.get("total", 0),
        "transporterId": d.get("transporter_id", ""), "transporterName": d.get("transporter_name", ""),
        "transMode": {"Road": "1", "Rail": "2", "Air": "3", "Ship": "4"}.get(d.get("mode", "Road"), "1"),
        "transDistance": str(int(float(d.get("distance_km", 0) or 0))),
        "vehicleNo": d.get("vehicle_no", ""),
        "vehicleType": "O" if d.get("vehicle_type") == "Over Dimensional Cargo" else "R",
        "transDocNo": d.get("trans_doc_no", ""),
        "transDocDate": _ddmmyyyy(d.get("trans_doc_date", "")) if d.get("trans_doc_date") else "",
        "itemList": items,
    }

@api.post("/invoices/{iid}/eway-bill/file")
async def file_eway_to_nic(iid: str, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    """One-click file an e-way bill via Adaequare GSP (enRiched EWB API, action=GENEWAYBILL)."""
    c = await _eway_gsp_config()
    if not _gsp_ready(c):
        raise HTTPException(400, "GSP not configured. Add Adaequare credentials (client id/secret + API "
                                 "username/password + GSTIN) in Railway env vars, then retry.")
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv:
        raise HTTPException(404, "Invoice not found")
    if inv.get("eway_bill_no"):
        raise HTTPException(400, "This invoice already has an e-way bill number. Cancel it first to regenerate.")
    company = await get_setting("integrations") or {}
    details = inv.get("eway_details", {}) or {}
    payload = _build_nic_ewb_payload(inv, details, company)
    try:
        token = await _adq_token(c)
        async with httpx.AsyncClient(timeout=30) as cl:
            gen = await cl.post(f"{c['base_url']}{c['prefix']}/ewb/ewayapi?action=GENEWAYBILL",
                                headers=_adq_ewb_headers(c, token), json=payload)
            gd = gen.json() if gen.content else {}
            gen.raise_for_status()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Adaequare EWB call failed: {e}. Response: {str(locals().get('gd', ''))[:300]}")
    if not gd.get("success"):
        raise HTTPException(502, f"Adaequare rejected the e-way bill: {gd.get('message') or gd}")
    result = gd.get("result") or {}
    ewb_no = str(result.get("ewayBillNo") or "")
    valid_upto = str(result.get("validUpto") or "")
    if not ewb_no:
        raise HTTPException(502, f"GSP did not return an e-way bill number. Response: {str(gd)[:300]}")
    upd = {"eway_bill_no": ewb_no, "eway_generated_at": datetime.utcnow().date().isoformat(),
           "eway_valid_until": valid_upto[:10] if valid_upto else "", "eway_details": {**details, "gsp_response": gd}}
    await db.invoices.update_one({"id": iid}, {"$set": upd})
    await write_audit(user.get("name", ""), "eway_bill_filed", "invoice", iid, {"eway_bill_no": ewb_no, "via": "Adaequare"})
    return {"ok": True, "eway_bill_no": ewb_no, "valid_upto": valid_upto, "filed_via": "Adaequare"}

class EwayCancelIn(BaseModel):
    reason_code: str = "1"   # 1=Duplicate, 2=Data Entry Mistake, 3=Order Cancelled, 4=Other
    reason_remark: str = ""

@api.post("/invoices/{iid}/eway-bill/cancel")
async def cancel_eway_via_gsp(iid: str, body: EwayCancelIn, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    """Cancel an already-filed e-way bill via Adaequare GSP (action=CANEWB). NIC only allows
    cancellation within 24 hours of generation and only if not yet verified in transit."""
    c = await _eway_gsp_config()
    if not _gsp_ready(c):
        raise HTTPException(400, "GSP not configured.")
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv or not inv.get("eway_bill_no"):
        raise HTTPException(404, "No e-way bill on file for this invoice.")
    try:
        token = await _adq_token(c)
        async with httpx.AsyncClient(timeout=30) as cl:
            r = await cl.post(f"{c['base_url']}{c['prefix']}/ewb/ewayapi?action=CANEWB",
                              headers=_adq_ewb_headers(c, token),
                              json={"ewbNo": int(inv["eway_bill_no"]), "cancelRsnCode": int(body.reason_code),
                                    "cancelRmrk": body.reason_remark or "Cancelled via ERP"})
            gd = r.json() if r.content else {}
            r.raise_for_status()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Adaequare cancel call failed: {e}")
    if not gd.get("success"):
        raise HTTPException(502, f"Adaequare rejected the cancellation: {gd.get('message') or gd}")
    await db.invoices.update_one({"id": iid}, {"$set": {"eway_bill_no": "", "eway_valid_until": "",
                                                        "eway_details": {**(inv.get("eway_details") or {}), "cancel_response": gd}}})
    await write_audit(user.get("name", ""), "eway_bill_cancelled", "invoice", iid, {"via": "Adaequare"})
    return {"ok": True}

# ---------------- e-Invoice (IRN + signed QR) ----------------
def _gstin_state(gstin: str) -> str:
    return (gstin or "")[:2] if gstin else ""

def _build_irp_einvoice_payload(inv: Dict[str, Any], company: Dict[str, Any], d: Dict[str, Any]) -> Dict[str, Any]:
    """IRP e-invoice schema 1.1 JSON for Adaequare's 'EI - Generate IRN' API — field names/casing
    verified against their sample request (SlNo/PrdDesc/HsnCd/Qty/UnitPrice/AssAmt/GstRt.. camel-caps)."""
    inter = bool(inv.get("is_interstate"))
    seller_gstin = company.get("company_gstin", ""); buyer_gstin = inv.get("customer_gstin", "")
    items = []
    for i, l in enumerate(inv.get("lines", []) or [], 1):
        qty = float(l.get("qty", 0)); rate = float(l.get("rate", 0)); gross = round(qty * rate, 2)
        ass = round(gross - (gross * float(l.get("discount_pct", 0) or 0) / 100) - float(l.get("discount_amount", 0) or 0), 2)
        if ass < 0: ass = 0.0
        gstrt = float(l.get("gst_rate", 0)); gst = round(ass * gstrt / 100, 2)
        items.append({"SlNo": str(i), "PrdDesc": (l.get("description", "") or "")[:300], "IsServc": "N",
            "HsnCd": str(l.get("hsn", "") or ""), "Qty": qty, "Unit": (l.get("unit", "NOS") or "NOS").upper()[:3],
            "UnitPrice": rate, "TotAmt": gross, "Discount": 0, "AssAmt": ass, "GstRt": gstrt,
            "IgstAmt": (gst if inter else 0), "CgstAmt": (0 if inter else round(gst / 2, 2)),
            "SgstAmt": (0 if inter else round(gst / 2, 2)), "CesRt": 0, "CesAmt": 0, "CesNonAdvlAmt": 0,
            "StateCesRt": 0, "StateCesAmt": 0, "OthChrg": 0, "TotItemVal": round(ass + gst, 2)})
    pin = lambda v: int(re.sub(r"\D", "", str(v or "0")) or 0)
    return {
        "Version": "1.1",
        "TranDtls": {"TaxSch": "GST", "SupTyp": "B2B", "RegRev": "N", "EcmGstin": None, "IgstOnIntra": "N"},
        "DocDtls": {"Typ": "INV", "No": inv.get("code", ""), "Dt": _ddmmyyyy(inv.get("date", ""))},
        "SellerDtls": {"Gstin": seller_gstin, "LglNm": company.get("company_name", ""), "TrdNm": company.get("company_name", ""),
            "Addr1": (d.get("dispatch_address", "") or company.get("company_address", ""))[:100],
            "Loc": d.get("dispatch_place", "") or d.get("dispatch_state", "") or "",
            "Pin": pin(d.get("dispatch_pin")), "Stcd": _gstin_state(seller_gstin),
            "Ph": (company.get("company_phone", "") or "")[:12], "Em": company.get("company_email", "")},
        "BuyerDtls": {"Gstin": buyer_gstin, "LglNm": inv.get("customer_name", ""), "TrdNm": inv.get("customer_name", ""),
            "Pos": _gstin_state(buyer_gstin) or _gstin_state(seller_gstin),
            "Addr1": (d.get("ship_address", "") or "")[:100],
            "Loc": d.get("ship_place", "") or d.get("ship_state", "") or "",
            "Pin": pin(d.get("ship_pin")), "Stcd": _gstin_state(buyer_gstin)},
        "ItemList": items,
        "ValDtls": {"AssVal": inv.get("subtotal", 0), "CgstVal": inv.get("cgst", 0), "SgstVal": inv.get("sgst", 0),
            "IgstVal": inv.get("igst", 0), "CesVal": 0, "StCesVal": 0, "Discount": 0, "OthChrg": 0,
            "RndOffAmt": inv.get("round_off", 0) or 0, "TotInvVal": inv.get("total", 0)},
    }

class EInvoiceRecordIn(BaseModel):
    irn: Optional[str] = ""
    ack_no: Optional[str] = ""
    ack_date: Optional[str] = ""
    signed_qr: Optional[str] = ""

@api.post("/invoices/{iid}/e-invoice/record")
async def record_einvoice(iid: str, body: EInvoiceRecordIn, user=Depends(get_current_user)):
    """Manually record an IRN / Ack / signed-QR generated on the IRP portal (used until the GSP API is wired)."""
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv:
        raise HTTPException(404, "Invoice not found")
    upd = {"irn": (body.irn or "").strip(), "ack_no": (body.ack_no or "").strip(),
           "ack_date": (body.ack_date or "").strip(), "signed_qr": (body.signed_qr or "").strip(),
           "einvoice_status": "generated" if (body.irn or "").strip() else ""}
    await db.invoices.update_one({"id": iid}, {"$set": upd})
    return {"ok": True, **upd}

@api.post("/invoices/{iid}/e-invoice/file")
async def file_einvoice(iid: str, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    """One-click generate an e-invoice IRN via Adaequare GSP (EI - Generate IRN API)."""
    c = await _eway_gsp_config()
    if not _gsp_ready(c):
        raise HTTPException(400, "GSP not configured. Add Adaequare credentials in Railway env vars to enable e-invoice filing.")
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv:
        raise HTTPException(404, "Invoice not found")
    if inv.get("invoice_type", "gst") != "gst":
        raise HTTPException(400, "e-Invoice applies to GST invoices only.")
    if inv.get("irn"):
        raise HTTPException(400, "This invoice already has an IRN. Cancel it first to regenerate.")
    company = await get_setting("integrations") or {}
    payload = _build_irp_einvoice_payload(inv, company, inv.get("eway_details", {}) or {})
    try:
        token = await _adq_token(c)
        async with httpx.AsyncClient(timeout=30) as cl:
            gen = await cl.post(f"{c['base_url']}{c['prefix']}/ei/api/invoice",
                                headers=_adq_ei_headers(c, token), json=payload)
            gd = gen.json() if gen.content else {}
            gen.raise_for_status()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Adaequare e-invoice call failed: {e}. Response: {str(locals().get('gd', ''))[:300]}")
    if not gd.get("success"):
        raise HTTPException(502, f"Adaequare rejected the e-invoice: {gd.get('message') or gd}")
    result = gd.get("result") or {}
    irn = str(result.get("Irn") or "")
    if not irn:
        raise HTTPException(502, f"IRP did not return an IRN. Response: {str(gd)[:300]}")
    upd = {"irn": irn, "ack_no": str(result.get("AckNo") or ""), "ack_date": str(result.get("AckDt") or ""),
           "signed_qr": str(result.get("SignedQRCode") or ""),
           "einvoice_status": "generated", "einvoice_details": gd}
    # If NIC also returned an e-way bill as part of IRN generation (EwbDtls was supplied), capture it
    if result.get("EwbNo"):
        upd["eway_bill_no"] = str(result.get("EwbNo"))
        upd["eway_valid_until"] = str(result.get("EwbValidTill") or "")[:10]
    await db.invoices.update_one({"id": iid}, {"$set": upd})
    await write_audit(user.get("name", ""), "einvoice_filed", "invoice", iid, {"irn": irn, "via": "Adaequare"})
    return {"ok": True, "irn": irn, "ack_no": upd["ack_no"], "ack_date": upd["ack_date"]}

class EInvoiceCancelIn(BaseModel):
    reason_code: str = "1"   # 1=Duplicate, 2=Data Entry Mistake, 3=Order Cancelled, 4=Other
    reason_remark: str = ""

@api.post("/invoices/{iid}/e-invoice/cancel")
async def cancel_einvoice_via_gsp(iid: str, body: EInvoiceCancelIn, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    """Cancel an IRN via Adaequare GSP. NIC only allows cancellation within 24 hours of generation."""
    c = await _eway_gsp_config()
    if not _gsp_ready(c):
        raise HTTPException(400, "GSP not configured.")
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv or not inv.get("irn"):
        raise HTTPException(404, "No IRN on file for this invoice.")
    try:
        token = await _adq_token(c)
        async with httpx.AsyncClient(timeout=30) as cl:
            r = await cl.post(f"{c['base_url']}{c['prefix']}/ei/api/invoice/cancel",
                              headers=_adq_ei_headers(c, token),
                              json={"irn": inv["irn"], "cnlrsn": body.reason_code,
                                    "cnlrem": body.reason_remark or "Cancelled via ERP"})
            gd = r.json() if r.content else {}
            r.raise_for_status()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Adaequare cancel call failed: {e}")
    if not gd.get("success"):
        raise HTTPException(502, f"Adaequare rejected the cancellation: {gd.get('message') or gd}")
    await db.invoices.update_one({"id": iid}, {"$set": {"irn": "", "ack_no": "", "ack_date": "", "signed_qr": "",
                                                        "einvoice_status": "cancelled",
                                                        "einvoice_details": {**(inv.get("einvoice_details") or {}), "cancel_response": gd}}})
    await write_audit(user.get("name", ""), "einvoice_cancelled", "invoice", iid, {"via": "Adaequare"})
    return {"ok": True}

# ---------------- Inbound webhooks (HR/attendance e.g. Mewurk) ----------------
WEBHOOK_BASE = os.getenv("API_PUBLIC_URL", "https://denplex-erp-production.up.railway.app")

async def _process_webhook(source: str, body: Any) -> str:
    """Best-effort map an inbound HR/attendance payload into the attendance collection.
    Tolerant of field names — refined once a real Mewurk payload is seen in the event log."""
    if isinstance(body, list):
        records = body
    elif isinstance(body, dict):
        records = body.get("data") or body.get("records") or body.get("attendance") or body.get("items") or [body]
    else:
        records = []
    if isinstance(records, dict):
        records = [records]
    n = 0
    for r in records:
        if not isinstance(r, dict):
            continue
        emp_code = str(r.get("employee_code") or r.get("emp_code") or r.get("employee_id") or r.get("empId")
                       or r.get("emp_id") or r.get("code") or "").strip()
        date = str(r.get("date") or r.get("attendance_date") or r.get("day") or "")[:10]
        if not (emp_code and date):
            continue
        emp = await db.employees.find_one({"code": emp_code}) or await db.employees.find_one({"id": emp_code})
        sraw = str(r.get("status") or r.get("attendance_status") or "present").lower()
        if "absent" in sraw or sraw == "a":
            st = "absent"
        elif "half" in sraw:
            st = "half_day"
        elif "leave" in sraw:
            st = "leave"
        else:
            st = "present"
        try:
            hours = float(r.get("hours") or r.get("working_hours") or (8 if st == "present" else 0))
        except Exception:
            hours = 8 if st == "present" else 0
        doc = {"id": new_id(), "employee_id": (emp.get("id") if emp else emp_code),
               "employee_name": (emp.get("name") if emp else str(r.get("employee_name") or r.get("name") or "")),
               "date": date, "status": st, "hours": hours, "notes": f"via {source}",
               "created_at": now_iso(), "source": source}
        await db.attendance.update_one({"employee_id": doc["employee_id"], "date": date}, {"$set": doc}, upsert=True)
        n += 1
    return f"attendance upserted: {n}"

@api.get("/webhooks/config")
async def get_webhook_config(user=Depends(require_roles("admin", "manager"))):
    cfg = await get_setting("webhooks") or {}
    src = cfg.get("mewurk") or {}
    if not src.get("secret"):
        src = {"secret": secrets.token_urlsafe(24), "enabled": True}
        await db.settings.update_one({"_id": "webhooks"}, {"$set": {"mewurk": src}}, upsert=True)
    return {"mewurk": {"enabled": src.get("enabled", True),
                       "url": f"{WEBHOOK_BASE}/api/webhooks/mewurk/{src['secret']}",
                       "secret": src["secret"]}}

@api.post("/webhooks/config")
async def set_webhook_config(body: dict, user=Depends(require_roles("admin", "manager"))):
    cfg = await get_setting("webhooks") or {}
    src = cfg.get("mewurk") or {}
    if body.get("rotate") or not src.get("secret"):
        src["secret"] = secrets.token_urlsafe(24)
    if "enabled" in body:
        src["enabled"] = bool(body["enabled"])
    if "secret" not in src:
        src["secret"] = secrets.token_urlsafe(24)
    await db.settings.update_one({"_id": "webhooks"}, {"$set": {"mewurk": src}}, upsert=True)
    return {"ok": True, "url": f"{WEBHOOK_BASE}/api/webhooks/mewurk/{src['secret']}", "secret": src["secret"]}

@api.get("/webhooks/events")
async def list_webhook_events(user=Depends(require_roles("admin", "manager"))):
    return await db.webhook_events.find({}, {"_id": 0}).sort("received_at", -1).to_list(100)

@api.post("/webhooks/{source}/{token}")
async def receive_webhook(source: str, token: str, request: Request):
    """Public inbound webhook (no JWT) — authenticated by the secret token in the URL.
    Logs every event raw so the real payload format can be inspected, then best-effort maps it."""
    cfg = await get_setting("webhooks") or {}
    src = cfg.get(source) or {}
    if not src or not src.get("enabled", True):
        raise HTTPException(404, "Webhook source not enabled")
    if not token or token != src.get("secret"):
        raise HTTPException(401, "Invalid webhook token")
    try:
        body = await request.json()
    except Exception:
        body = {"_raw": (await request.body()).decode("utf-8", "ignore")[:5000]}
    ev = {"id": new_id(), "source": source, "received_at": now_iso(), "body": body, "processed": False, "result": ""}
    try:
        ev["result"] = await _process_webhook(source, body)
        ev["processed"] = True
    except Exception as e:
        ev["result"] = f"map-skipped: {e}"
    await db.webhook_events.insert_one(dict(ev))
    return {"ok": True, "received": True, "result": ev["result"]}

@api.put("/invoices/{iid}")
async def update_invoice(iid: str, inv: Invoice, user=Depends(get_current_user)):
    existing = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if existing and _eway_edit_locked(existing):
        raise HTTPException(423, "Invoice locked: its e-way bill validity period has ended, so it can no longer be edited.")
    data = inv.model_dump(); data.pop("id", None); data.pop("created_at", None)
    data.update(compute_invoice_totals(data["lines"], data.get("is_interstate", False), data.get("round_off", 0), data.get("tds", 0), data.get("extra_charges"), data.get("tcs", 0)))
    await db.invoices.update_one({"id": iid}, {"$set": data})
    return {"ok": True}

@api.delete("/invoices/{iid}")
async def del_invoice(iid: str, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    doc = await db.invoices.find_one({"id": iid}, {"_id": 0})
    await _recycle("invoices", "Invoice", doc, user)
    await db.invoices.delete_one({"id": iid})
    return {"ok": True}

# ---------------- QC Reports ----------------
@api.post("/qc-reports")
async def create_qc(q: QCReport, user=Depends(get_current_user)):
    doc = q.model_dump()
    doc["code"] = await gen_code("QC", "qc")
    if q.work_order_id:
        wo = await db.work_orders.find_one({"id": q.work_order_id}, {"_id": 0})
        if wo:
            doc["work_order_code"] = wo.get("code", "")
            doc["customer_id"] = doc.get("customer_id") or wo.get("customer_id", "")
            doc["customer_name"] = doc.get("customer_name") or wo.get("customer_name", "")
    if doc.get("photos"):
        doc["photos"] = [await _drive_offload(ph, f"qc-{doc.get('code','')}-{i}.jpg", "image/jpeg", "QC Photos") for i, ph in enumerate(doc["photos"])]
    await db.qc_reports.insert_one(doc)
    return serialize(doc)

@api.get("/qc-reports")
async def list_qc(user=Depends(get_current_user)):
    return await list_collection(db.qc_reports)

@api.delete("/qc-reports/{qid}")
async def del_qc(qid: str, user=Depends(require_roles("admin", "manager", "qc"))):
    await db.qc_reports.delete_one({"id": qid})
    return {"ok": True}

# ---------------- Documents ----------------
def _doc_mime(name):
    import mimetypes
    return mimetypes.guess_type(name or "")[0] or "application/octet-stream"

@api.post("/documents")
async def upload_doc(d: DocumentMeta, user=Depends(get_current_user)):
    doc = d.model_dump()
    doc["uploaded_by"] = user["name"]
    doc["file_base64"] = await _drive_offload(doc.get("file_base64"), doc.get("name") or "document", _doc_mime(doc.get("name")), "Documents")
    await db.documents.insert_one(doc)
    return serialize(doc)

@api.get("/documents")
async def list_docs(linked_to: Optional[str] = None, user=Depends(get_current_user)):
    q = {"linked_to": linked_to} if linked_to else {}
    docs = await db.documents.find(q, {"_id": 0, "file_base64": 0}).sort("created_at", -1).to_list(500)
    for d in docs:
        d["has_file"] = True
    return docs

@api.get("/documents/{did}/download")
async def download_doc(did: str, user=Depends(get_current_user)):
    doc = await db.documents.find_one({"id": did}, {"_id": 0})
    if not doc or not doc.get("file_base64"):
        raise HTTPException(404, "No file on this document")
    data = await _resolve_b64_or_drive(doc["file_base64"])
    return Response(content=data, media_type=_doc_mime(doc.get("name")),
        headers={"Content-Disposition": f'attachment; filename="{doc.get("name","document")}"'})

@api.get("/documents/{did}/revisions/{rev_no}/download")
async def download_doc_rev(did: str, rev_no: int, user=Depends(get_current_user)):
    doc = await db.documents.find_one({"id": did}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Document not found")
    rev = next((r for r in doc.get("revisions", []) if int(r.get("rev_no", -1)) == rev_no), None)
    if not rev or not rev.get("file_base64"):
        raise HTTPException(404, "No file for this revision")
    data = await _resolve_b64_or_drive(rev["file_base64"])
    return Response(content=data, media_type=_doc_mime(doc.get("name")),
        headers={"Content-Disposition": f'attachment; filename="{doc.get("name","document")}-rev{rev_no}"'})

@api.delete("/documents/{did}")
async def del_doc(did: str, user=Depends(get_current_user)):
    await db.documents.delete_one({"id": did})
    return {"ok": True}

# ---------------- Dashboard ----------------
@api.get("/dashboard/stats")
async def dashboard_stats(user=Depends(get_current_user)):
    # Projection-limited fetches to keep dashboard fast even with thousands of records
    items = await db.items.find(
        {}, {"_id": 0, "id": 1, "sku": 1, "name": 1, "qty_on_hand": 1, "reorder_level": 1, "uom": 1}
    ).to_list(5000)
    low_stock = [i for i in items if i.get("qty_on_hand", 0) <= i.get("reorder_level", 0)]
    open_wo = await db.work_orders.count_documents({"status": {"$in": ["planned", "in_progress"]}})
    qc_pending = await db.work_orders.count_documents({"status": "qc"})
    leads_open = await db.leads.count_documents({"status": {"$in": ["new", "contacted", "qualified"]}})
    customers = await db.customers.count_documents({})
    # Aggregate revenue at the DB layer instead of pulling 2000 docs into Python
    rev_pipeline = [
        {"$match": {"status": {"$in": ["paid", "sent"]}}},
        {"$group": {"_id": None, "total": {"$sum": "$total"}}},
    ]
    rev_cursor = db.invoices.aggregate(rev_pipeline)
    rev_doc = await rev_cursor.to_list(1)
    revenue = float(rev_doc[0]["total"]) if rev_doc else 0.0
    items_count = await db.items.count_documents({})
    repeat_customers = await db.customers.count_documents({"customer_type": "repeat"})
    recent = await db.work_orders.find({}, {"_id": 0}).sort("created_at", -1).to_list(8)
    return {
        "open_wo": open_wo,
        "qc_pending": qc_pending,
        "low_stock_count": len(low_stock),
        "low_stock_items": low_stock[:10],
        "leads_open": leads_open,
        "customers": customers,
        "repeat_customers": repeat_customers,
        "revenue": round(revenue, 2),
        "items_count": items_count,
        "recent_wo": recent,
    }

# ---------------- P1: Accounting (Expenses + GST Report) ----------------
@api.post("/expenses")
async def create_expense(e: Expense, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    doc = e.model_dump()
    amt = float(doc.get("amount") or 0)
    rate = float(doc.get("gst_rate") or 0)
    doc["gst_amount"] = round(amt * rate / 100.0, 2)
    doc["total"] = round(amt + doc["gst_amount"], 2)
    await db.expenses.insert_one(doc)
    return serialize(doc)

@api.get("/expenses")
async def list_expenses(user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    return await list_collection(db.expenses)

@api.delete("/expenses/{eid}")
async def del_expense(eid: str, user=Depends(require_roles("admin", "accountant", "ca"))):
    await db.expenses.delete_one({"id": eid})
    return {"ok": True}

@api.get("/accounting/gst-report")
async def gst_report(period_from: Optional[str] = None, period_to: Optional[str] = None,
                     user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    invq = {}
    if period_from or period_to:
        invq["date"] = {}
        if period_from: invq["date"]["$gte"] = period_from
        if period_to: invq["date"]["$lte"] = period_to
    invs = await db.invoices.find(invq, {"_id": 0}).to_list(5000)
    output_cgst = sum(i.get("cgst", 0) for i in invs)
    output_sgst = sum(i.get("sgst", 0) for i in invs)
    output_igst = sum(i.get("igst", 0) for i in invs)
    output_taxable = sum(i.get("subtotal", 0) for i in invs)

    exq = {}
    if period_from or period_to:
        exq["date"] = {}
        if period_from: exq["date"]["$gte"] = period_from
        if period_to: exq["date"]["$lte"] = period_to
    expenses = await db.expenses.find(exq, {"_id": 0}).to_list(5000)
    input_gst = sum(e.get("gst_amount", 0) for e in expenses)
    input_taxable = sum(e.get("amount", 0) for e in expenses)

    return {
        "period_from": period_from or "",
        "period_to": period_to or "",
        "output": {
            "taxable": round(output_taxable, 2),
            "cgst": round(output_cgst, 2),
            "sgst": round(output_sgst, 2),
            "igst": round(output_igst, 2),
            "total_gst": round(output_cgst + output_sgst + output_igst, 2),
            "invoice_count": len(invs),
        },
        "input": {
            "taxable": round(input_taxable, 2),
            "total_gst": round(input_gst, 2),
            "expense_count": len(expenses),
        },
        "net_liability": round((output_cgst + output_sgst + output_igst) - input_gst, 2),
    }

# ---------------- P1: HR (Employees + Attendance) ----------------
@api.post("/employees")
async def create_emp(e: Employee, user=Depends(require_roles("admin", "manager"))):
    doc = e.model_dump()
    doc["code"] = await gen_code("EMP", "employee")
    await db.employees.insert_one(doc)
    return serialize(doc)

@api.get("/employees")
async def list_emp(user=Depends(get_current_user)):
    return await list_collection(db.employees)

@api.put("/employees/{eid}")
async def upd_emp(eid: str, e: Employee, user=Depends(require_roles("admin", "manager"))):
    data = e.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.employees.update_one({"id": eid}, {"$set": data})
    return {"ok": True}

@api.delete("/employees/{eid}")
async def del_emp(eid: str, user=Depends(require_roles("admin"))):
    await db.employees.delete_one({"id": eid})
    return {"ok": True}

@api.post("/attendance")
async def create_att(a: Attendance, user=Depends(require_roles("admin", "manager"))):
    emp = await db.employees.find_one({"id": a.employee_id}, {"_id": 0})
    doc = a.model_dump()
    if emp: doc["employee_name"] = emp.get("name", "")
    await db.attendance.insert_one(doc)
    return serialize(doc)

@api.get("/attendance")
async def list_att(user=Depends(get_current_user)):
    return await list_collection(db.attendance, sort_key="date")

@api.delete("/attendance/{aid}")
async def del_att(aid: str, user=Depends(require_roles("admin", "manager"))):
    await db.attendance.delete_one({"id": aid})
    return {"ok": True}

# ---------------- P1: Marketing / Social Campaigns ----------------
@api.post("/campaigns")
async def create_camp(c: Campaign, user=Depends(get_current_user)):
    doc = c.model_dump()
    await db.campaigns.insert_one(doc)
    return serialize(doc)

@api.get("/campaigns")
async def list_camp(user=Depends(get_current_user)):
    return await list_collection(db.campaigns)

@api.put("/campaigns/{cid}")
async def upd_camp(cid: str, c: Campaign, user=Depends(get_current_user)):
    data = c.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.campaigns.update_one({"id": cid}, {"$set": data})
    return {"ok": True}

@api.delete("/campaigns/{cid}")
async def del_camp(cid: str, user=Depends(get_current_user)):
    await db.campaigns.delete_one({"id": cid})
    return {"ok": True}

# ---------------- P1: Document Revisions (ISO 9001 etc.) ----------------
@api.post("/documents/{did}/revisions")
async def add_revision(did: str, payload: DocRevIn, user=Depends(get_current_user)):
    doc = await db.documents.find_one({"id": did}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Document not found")
    revs = doc.get("revisions", [])
    if not revs:
        revs.append({
            "rev_no": 0,
            "file_base64": doc.get("file_base64", ""),
            "notes": "Initial version",
            "by": doc.get("uploaded_by", ""),
            "created_at": doc.get("created_at", now_iso()),
        })
    new_rev_no = max([r.get("rev_no", 0) for r in revs]) + 1
    new_file = await _drive_offload(payload.file_base64, f"{doc.get('name','document')}-rev{new_rev_no}", _doc_mime(doc.get("name")), "Documents")
    revs.append({
        "rev_no": new_rev_no,
        "file_base64": new_file,
        "notes": payload.notes or "",
        "by": user["name"],
        "created_at": now_iso(),
    })
    await db.documents.update_one(
        {"id": did},
        {"$set": {"revisions": revs, "file_base64": new_file, "current_revision": new_rev_no}},
    )
    return {"ok": True, "current_revision": new_rev_no}

@api.get("/documents/{did}/revisions")
async def get_revisions(did: str, user=Depends(get_current_user)):
    doc = await db.documents.find_one({"id": did}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Document not found")
    return {"current_revision": doc.get("current_revision", 0), "revisions": doc.get("revisions", [])}



# ---------------- Customer Portal (public) ----------------
@api.get("/portal/track")
async def portal_track(ref: str = Query(..., min_length=2)):
    ref_u = ref.strip()
    wo = await db.work_orders.find_one({"$or": [{"code": ref_u}, {"po_ref": ref_u}]}, {"_id": 0})
    if not wo:
        raise HTTPException(404, "Reference not found")
    jcs = await db.job_cards.find({"work_order_id": wo["id"]}, {"_id": 0}).to_list(200)
    qcs = await db.qc_reports.find({"work_order_id": wo["id"]}, {"_id": 0, "photos": 0}).to_list(200)
    return {
        "work_order": {k: wo.get(k) for k in ["code", "product", "qty", "status", "priority", "progress", "start_date", "due_date", "customer_name", "po_ref", "created_at"]},
        "job_cards": jcs,
        "qc_reports": qcs,
    }

# ---------------- Health ----------------
@api.get("/")
async def root():
    return {"ok": True, "service": "precision-erp"}


# ================ P2: Settings, Twilio WhatsApp, Resend, Indiamart, TradeIndia, GSTR, 2FA, Audit ================

async def get_setting(key: str) -> Dict[str, Any]:
    doc = await db.settings.find_one({"_id": key}, {"_id": 0})
    return doc or {}

async def set_setting(key: str, data: Dict[str, Any]):
    await db.settings.replace_one({"_id": key}, {"_id": key, **data}, upsert=True)

# ---------------------------------------------------------------------------
# Public marketing site config — makes the Landing page (pre-login homepage)
# config-driven instead of hardcoded, so the SAME codebase can be redeployed
# for a future client by editing this settings doc (company name/tagline/
# features/modules/branding), and Denplex can turn the "Free Trial / sandbox"
# sales pitch off (trial_enabled=false) while using the ERP internally.
# PUBLIC endpoint — never put secrets (GSTIN, bank, API keys) in this doc.
# ---------------------------------------------------------------------------
DEFAULT_SITE_CONFIG: Dict[str, Any] = {
    "brand_name": "DENPLEX ERP",
    "logo_url": "/denplex-logo.png",
    "overline": "Denplex Engineering Company · Jig & Fixtures · Precision Job Work",
    "heading_prefix": "Run your workshop with",
    "heading_highlight": "±0.001 mm",
    "heading_suffix": "cut.",
    "subheading": "One operating system for your inventory, BOM, work orders, job cards, QC, quotations, GST invoices, leads, and customer portal — built for small-scale precision engineering businesses.",
    "hero_image": "https://static.prod-images.emergentagent.com/jobs/7f514505-bc8d-48ed-954b-5815c5f6170a/images/0475821f5ac9696210216b4d7b8fd14a0f3b01c44be9da7e3b33a4e81d154066.png",
    "feature_image": "https://images.unsplash.com/photo-1666634157070-6fd830fb5672?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA1OTN8MHwxfHNlYXJjaHwxfHxwcmVjaXNpb24lMjBjbmMlMjBtYWNoaW5pbmclMjBtZXRhbHxlbnwwfHx8fDE3NzkxMjY0NDR8MA&ixlib=rb-4.1.0&q=85",
    "hero_badge_line1": "Live · Workshop floor",
    "hero_badge_line2": "Open WO #WO-26-0042",
    "hero_badge_line3": "SS-316 Drill Jig · Qty 24 · Stage: Milling",
    "stats": [{"k": "16+", "v": "Modules"}, {"k": "ISO", "v": "9001 ready"}, {"k": "GST", "v": "CGST/SGST/IGST"}],
    "features": [
        {"icon": "Boxes", "title": "Inventory & AI bill scan", "desc": "Stock in/out, in-process tracking, adjustments. Scan a paper or digital bill and Claude reads it into your system."},
        {"icon": "Layers", "title": "Work Orders & Job Cards", "desc": "Plan production, assign machines and operators, track progress in real time."},
        {"icon": "ShieldCheck", "title": "QC reports with photos", "desc": "Capture parameters, measurements, pass/fail. Link to work orders & customers forever."},
        {"icon": "BarChart3", "title": "GST invoices & quotations", "desc": "CGST / SGST / IGST automatically. Quotations, POs, invoices in one place."},
        {"icon": "Users", "title": "CRM + WhatsApp", "desc": "Leads, repeat vs one-time customers, click-to-WhatsApp to send POs, quotes, and follow-ups."},
        {"icon": "Cog", "title": "Customer portal", "desc": "Customers track their order with a PO or reference number. Public, branded, mobile-friendly."},
    ],
    "modules": ["Inventory", "BOM", "Work Orders", "Job Cards", "Quotations", "Purchase Orders", "Invoices",
                "QC Reports", "CRM / Leads", "Suppliers", "Customers", "Documents", "Dashboard",
                "Customer Portal", "User Roles", "GST"],
    "footer_cta_heading": "Ready to run a tighter shop?",
    "footer_copyright": "© 2026 Denplex Engineering Company",
    "footer_version": "Denplex ERP v0.3 · Built for MSMEs",
    # --- Trial/sales pitch toggle: OFF by default while Denplex is the only user of its own ERP.
    # Flip to true (via PUT /settings/marketing-site, admin only) when onboarding an external client.
    "trial_enabled": False,
    "footer_cta_sub": "Start a 30-day free trial — no card needed. We verify and email you in 24 hours.",
    "sandbox_note": "",   # e.g. "Or try the sandbox: admin@erp.com · Admin@123" — leave blank in production
}

@api.get("/public/site-config")
async def get_public_site_config():
    """Unauthenticated — powers the pre-login marketing Landing page. Never returns secrets."""
    s = await get_setting("marketing_site")
    return {**DEFAULT_SITE_CONFIG, **s}

@api.put("/settings/marketing-site")
async def update_marketing_site(body: dict, user=Depends(require_roles("admin"))):
    clean = {k: v for k, v in (body or {}).items() if k in DEFAULT_SITE_CONFIG}
    await set_setting("marketing_site", clean)
    return {**DEFAULT_SITE_CONFIG, **clean}

class IntegrationSettingsIn(BaseModel):
    twilio_account_sid: Optional[str] = ""
    twilio_auth_token: Optional[str] = ""
    twilio_whatsapp_from: Optional[str] = ""
    indiamart_crm_key: Optional[str] = ""
    tradeindia_webhook_secret: Optional[str] = ""
    company_name: Optional[str] = "Denplex Engineering Company"
    company_gstin: Optional[str] = ""
    company_state: Optional[str] = ""
    company_address: Optional[str] = ""
    # Multi-unit support: list of {name, address} dicts. Renders one block per unit
    # in the PDF header. If empty, falls back to company_address.
    company_units: Optional[List[Dict[str, str]]] = []
    company_tagline: Optional[str] = "Precision Engineered Solutions"
    company_phone: Optional[str] = ""
    company_email: Optional[str] = ""
    company_udyam: Optional[str] = ""  # UDYAM / MSME registration shown on letterhead
    # Bank / UPI block (printed on every invoice per standard tax-invoice layout)
    bank_name: Optional[str] = ""
    bank_account_no: Optional[str] = ""
    bank_ifsc: Optional[str] = ""
    bank_branch: Optional[str] = ""
    upi_id: Optional[str] = ""  # e.g. denplex@axisbank — used to auto-generate QR
    # Signatory image (base64 PNG/JPG, optional)
    signatory_image_b64: Optional[str] = ""
    signatory_label: Optional[str] = "Authorised Signatory"
    # Terms & default sale description
    invoice_terms: Optional[str] = ("*Subject to Ahmedabad jurisdiction only\n"
                                    "1) The bill must be paid within due date otherwise interest @18% will be charged extra\n"
                                    "2) Goods once sold will not be taken back.\n"
                                    "3) Our responsibility ceases on delivery the goods to the carries.\n"
                                    "4) Payment requested by CASH/CHEQUE/Bank Transfer only\n"
                                    "5) If any rejection or rework occurs please notify withing 10 days of material receipt after that it won't be accepted.")
    invoice_description: Optional[str] = ""

@api.get("/settings/integrations")
async def get_integrations(user=Depends(require_roles("admin"))):
    return await get_setting("integrations")

@api.put("/settings/integrations")
async def update_integrations(payload: IntegrationSettingsIn, user=Depends(require_roles("admin"))):
    data = payload.model_dump()
    await set_setting("integrations", data)
    return data

# ---------- Invoice Template (per-section visibility toggles) ----------
class InvoiceTemplateIn(BaseModel):
    """Per-section visibility flags for the printed PDF."""
    show_company_logo: bool = True
    show_company_address: bool = True
    show_company_gstin: bool = True
    show_company_email: bool = True
    show_company_phone: bool = True
    show_company_udyam: bool = True
    show_ship_to: bool = True
    show_bill_from: bool = False             # Off by default; auto-on when invoice has explicit bill_from
    show_ship_from: bool = False             # Off by default; auto-on when invoice has explicit ship_from
    show_due_date: bool = True
    show_place_of_supply: bool = True
    show_hsn_column: bool = True
    show_item_code_column: bool = True       # Item Code column
    show_po_meta: bool = True                # PO Date / PO No / Purchaser Name in meta box
    show_discount_column: bool = True
    show_tax_summary: bool = True            # HSN-wise CGST/SGST/IGST breakup
    show_totals_sidebar: bool = True         # Sub Total / Discount / Tax / TCS / Total
    show_amount_in_words: bool = True
    show_payment_mode: bool = True
    show_description: bool = True
    show_terms: bool = True
    show_bank_details: bool = True
    show_upi_qr: bool = True
    show_signatory_image: bool = True
    show_bank_on_new_page: bool = True       # Move bank details + signature to a new page
    show_unit_column: bool = True            # Unit column in items table (Mtr/Nos/Kg)
    show_inline_gst_column: bool = False     # Off by default — GST shown only in Tax Summary
    show_split_tax_in_sidebar: bool = False  # Off = single "Tax (X%)" line; on = CGST + SGST split
    print_original_duplicate: bool = True
    paper_size: Literal["A4", "A5"] = "A4"
    orientation: Literal["portrait", "landscape"] = "portrait"
    amount_in_words_locale: Literal["en_IN", "en"] = "en_IN"
    # Style preset: "standard" = full tax-invoice layout, "compact" = single-page minimal,
    # "modern" = clean serif with accent lines and more whitespace.
    template_style: Literal["standard", "compact", "modern"] = "standard"

@api.get("/settings/invoice-template")
async def get_invoice_template(doc_type: Optional[str] = None, user=Depends(get_current_user)):
    """Return template settings. If `doc_type` is supplied, returns the merged
    {default + doc_type override} flags. Otherwise returns the full map keyed by doc_type."""
    s = await get_setting("invoice_template") or {}
    defaults = InvoiceTemplateIn().model_dump()
    base = {**defaults, **(s.get("default") or {})}
    if doc_type:
        override = (s.get(doc_type) or {})
        return {**base, **override}
    # Return full map for the UI
    KNOWN = ["default", "invoice", "quotation", "purchase_order", "sale_order",
            "delivery_challan", "job_work_out", "credit_note", "vendor_bill"]
    out: Dict[str, Any] = {}
    for k in KNOWN:
        out[k] = {**base, **((s.get(k) or {}))}
    return out

@api.put("/settings/invoice-template")
async def update_invoice_template(payload: Dict[str, Any], user=Depends(require_roles("admin"))):
    """Accepts either a flat `InvoiceTemplateIn` (treated as `default`) or a map
    {doc_type: {flags...}}. Storing only the overrides keeps each doc_type small."""
    # Validate by passing through model where possible
    allowed = set(InvoiceTemplateIn().model_dump().keys())
    s = await get_setting("invoice_template") or {}
    if any(k in payload for k in allowed):
        # Flat payload → save as `default`
        clean = {k: v for k, v in payload.items() if k in allowed}
        s["default"] = clean
    else:
        # Map payload → merge each doc_type
        for dt, flags in payload.items():
            if not isinstance(flags, dict): continue
            s[dt] = {k: v for k, v in flags.items() if k in allowed}
    await set_setting("invoice_template", s)
    return s

# ---------- Audit log ----------
def get_client_ip(request: Optional[Request]) -> str:
    if not request:
        return ""
    fwd = request.headers.get("x-forwarded-for", "")
    if fwd:
        return fwd.split(",")[0].strip()
    return request.client.host if request.client else ""

async def write_audit(user_name: str, action: str, entity: str, entity_id: Optional[str] = None, details: Optional[Dict] = None, request: Optional[Request] = None):
    try:
        ip = get_client_ip(request)
        ua = (request.headers.get("user-agent", "") if request else "")[:300]
        await db.audit_logs.insert_one({
            "id": new_id(),
            "user": user_name,
            "action": action,
            "entity": entity,
            "entity_id": entity_id or "",
            "details": details or {},
            "ip": ip,
            "user_agent": ua,
            "created_at": now_iso(),
        })
    except Exception as e:
        logger.exception("audit failed: %s", e)

@api.get("/audit-logs")
async def list_audit(limit: int = 200, user=Depends(require_roles("admin"))):
    rows = await db.audit_logs.find({}, {"_id": 0}).sort("created_at", -1).to_list(min(limit, 1000))
    return rows

# ---------- Twilio WhatsApp ----------
def _format_in_whatsapp(raw: str) -> str:
    raw = (raw or "").strip()
    if not raw:
        raise ValueError("Empty phone number")
    if raw.lower().startswith("whatsapp:"):
        return raw
    try:
        if raw.startswith("+"):
            parsed = phonenumbers.parse(raw, None)
        else:
            parsed = phonenumbers.parse(raw, "IN")
        if not phonenumbers.is_valid_number(parsed):
            raise ValueError("Invalid phone number")
        e164 = phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
    except phonenumbers.NumberParseException as ex:
        raise ValueError(f"Phone parse error: {ex}")
    return f"whatsapp:{e164}"

class WhatsAppSendIn(BaseModel):
    to_phone: str
    body: str
    media_url: Optional[str] = ""

@api.post("/whatsapp/send")
async def whatsapp_send(payload: WhatsAppSendIn, user=Depends(get_current_user)):
    cfg = await get_setting("integrations")
    sid = cfg.get("twilio_account_sid"); tok = cfg.get("twilio_auth_token"); frm = cfg.get("twilio_whatsapp_from")
    if not (sid and tok and frm):
        raise HTTPException(400, "Twilio not configured. Set credentials in Settings → Integrations.")
    try:
        to = _format_in_whatsapp(payload.to_phone)
    except ValueError as e:
        raise HTTPException(422, str(e))
    def _send():
        client = TwilioClient(sid, tok)
        kwargs = {"body": payload.body, "from_": frm, "to": to}
        if payload.media_url:
            kwargs["media_url"] = [payload.media_url]
        return client.messages.create(**kwargs)
    try:
        msg = await asyncio.to_thread(_send)
    except TwilioRestException as e:
        raise HTTPException(502, f"Twilio: {e.msg}")
    await write_audit(user["name"], "whatsapp_send", "message", msg.sid, {"to": to})
    return {"sid": msg.sid, "status": msg.status}

# ---------- Resend deprecated — using Gmail / Outlook OAuth instead ----------

# ---------- PDF builders (Denplex Red/Black branding) ----------
def _money(n) -> str:
    try:
        return f"Rs. {float(n):,.2f}"
    except Exception:
        return f"Rs. {n}"

def _amount_in_words(n: float, locale: str = "en_IN") -> str:
    try:
        from num2words import num2words
        rupees = int(n)
        paise = round((float(n) - rupees) * 100)
        words_r = num2words(rupees, lang="en_IN" if locale == "en_IN" else "en").title()
        out = f"{words_r} Rupees"
        if paise:
            out += f" and {num2words(paise, lang='en').title()} Paise"
        return out + " only"
    except Exception:
        return ""

def _upi_qr_png(upi_id: str, payee_name: str, amount: float = 0.0, note: str = "") -> Optional[bytes]:
    """Generate UPI QR (BHIM/PhonePe/GPay scannable) as PNG bytes."""
    if not upi_id:
        return None
    try:
        import qrcode
        from urllib.parse import urlencode, quote
        params = {"pa": upi_id, "pn": payee_name or "Denplex", "cu": "INR"}
        if amount and amount > 0:
            params["am"] = f"{amount:.2f}"
        if note:
            params["tn"] = note[:50]
        # Use quote_via to keep & encoded correctly
        upi_url = "upi://pay?" + urlencode(params, quote_via=quote)
        qr = qrcode.QRCode(box_size=3, border=1)
        qr.add_data(upi_url); qr.make(fit=True)
        img = qr.make_image(fill_color="#0A0A0A", back_color="#FFFFFF")
        out = io.BytesIO(); img.save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


# ---------------- QR codes (scan to open WO / Part / Inventory) ----------------
QR_BASE_URL = os.environ.get("FRONTEND_URL", "https://erp.denplex.co").rstrip("/")
# entity -> (collection name, code field, title field)
_QR_ENTITIES = {
    "work-order": ("work_orders", "code", "product"),
    "part":       ("parts", "part_number", "name"),
    "inventory":  ("items", "sku", "name"),
}

def _entity_qr_png(data: str) -> bytes:
    import qrcode
    qr = qrcode.QRCode(box_size=8, border=2,
                       error_correction=qrcode.constants.ERROR_CORRECT_M)
    qr.add_data(data); qr.make(fit=True)
    img = qr.make_image(fill_color="#0A0A0A", back_color="#FFFFFF")
    out = io.BytesIO(); img.save(out, format="PNG"); return out.getvalue()

def _qr_spec(entity: str):
    spec = _QR_ENTITIES.get(entity)
    if not spec:
        raise HTTPException(404, "Unknown QR entity")
    return spec

@api.get("/qr/{entity}/{eid}.png")
async def entity_qr_png(entity: str, eid: str, user=Depends(get_current_user)):
    _qr_spec(entity)
    url = f"{QR_BASE_URL}/app/scan/{entity}/{eid}"
    return Response(content=_entity_qr_png(url), media_type="image/png",
                    headers={"Cache-Control": "public, max-age=86400"})

@api.get("/scan/{entity}/{eid}")
async def scan_resolve(entity: str, eid: str, user=Depends(get_current_user)):
    cname, code_field, title_field = _qr_spec(entity)
    coll = db[cname]
    doc = await coll.find_one({"id": eid}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Record not found for this QR code")
    out = {
        "entity": entity,
        "id": eid,
        "code": doc.get(code_field) or doc.get("code") or eid,
        "title": doc.get(title_field) or "",
        "record": doc,
    }
    if entity == "work-order":
        ops = await db.wo_operations.find(
            {"work_order_id": eid}, {"_id": 0}).to_list(200)
        ops.sort(key=lambda o: (o.get("seq", 0), o.get("created_at", "")))
        light = []
        for o in ops:
            pc = len(o.get("photos") or [])
            o2 = {k: v for k, v in o.items() if k != "photos"}
            o2["photo_count"] = pc
            light.append(o2)
        out["operations"] = light
    return out


# ---------------- Planning & Scheduling ----------------
@api.get("/planning/overview")
async def planning_overview(user=Depends(get_current_user)):
    """Machine loading (from open operation minutes) + due-date schedule buckets."""
    from datetime import date as _date, timedelta as _td
    today = _date.today()

    # Machine load: sum planned minutes of non-done operations, grouped by machine
    pipeline = [
        {"$match": {"status": {"$in": ["pending", "running", "hold"]}}},
        {"$group": {
            "_id": {"$ifNull": ["$machine", ""]},
            "ops": {"$sum": 1},
            "minutes": {"$sum": {"$ifNull": ["$planned_minutes", 0]}},
            "running": {"$sum": {"$cond": [{"$eq": ["$status", "running"]}, 1, 0]}},
        }},
    ]
    rows = await db.wo_operations.aggregate(pipeline).to_list(300)
    machine_load = [{
        "machine": (r["_id"] or "Unassigned"),
        "ops": r["ops"],
        "minutes": round(r.get("minutes", 0) or 0, 1),
        "hours": round((r.get("minutes", 0) or 0) / 60.0, 1),
        "running": r.get("running", 0),
    } for r in rows]
    machine_load.sort(key=lambda x: x["minutes"], reverse=True)

    # Known machines with zero load still listed (so planners see idle capacity)
    loaded_names = {m["machine"] for m in machine_load}
    masters = await db.machines.find({"is_active": {"$ne": False}},
                                     {"_id": 0, "name": 1, "group": 1, "status": 1}).to_list(500)
    for m in masters:
        nm = m.get("name") or ""
        if nm and nm not in loaded_names:
            machine_load.append({"machine": nm, "ops": 0, "minutes": 0, "hours": 0, "running": 0})

    # Due-date buckets for active work orders
    wos = await db.work_orders.find(
        {"status": {"$nin": ["completed", "cancelled"]}},
        {"_id": 0, "code": 1, "product": 1, "part_number": 1, "customer_name": 1,
         "due_date": 1, "status": 1, "priority": 1, "id": 1, "qty": 1},
    ).to_list(3000)

    def which(d):
        if not d:
            return "no_date"
        try:
            dd = _date.fromisoformat(str(d)[:10])
        except Exception:
            return "no_date"
        if dd < today:
            return "overdue"
        if dd == today:
            return "today"
        if dd <= today + _td(days=7):
            return "this_week"
        return "later"

    buckets = {"overdue": [], "today": [], "this_week": [], "later": [], "no_date": []}
    for w in wos:
        buckets[which(w.get("due_date"))].append(w)
    for k in buckets:
        buckets[k].sort(key=lambda w: (str(w.get("due_date") or "9999"),
                                       {"high": 0, "medium": 1, "low": 2}.get(w.get("priority"), 1)))

    return {
        "today": today.isoformat(),
        "machine_load": machine_load,
        "buckets": {k: {"count": len(v), "items": v[:60]} for k, v in buckets.items()},
        "active_wo": len(wos),
    }


# ---------------- Costing & Profitability ----------------
COSTING_DEFAULT_MACHINE_RATE = float(os.environ.get("COSTING_MACHINE_RATE", "400"))
COSTING_DEFAULT_LABOUR_RATE = float(os.environ.get("COSTING_LABOUR_RATE", "120"))

async def _costing_rates():
    doc = await db.settings.find_one({"_id": "costing"}, {"_id": 0}) or {}
    mr = float(doc.get("default_machine_rate") or 0) or COSTING_DEFAULT_MACHINE_RATE
    lr = float(doc.get("default_labour_rate") or 0) or COSTING_DEFAULT_LABOUR_RATE
    return mr, lr

class CostingRatesIn(BaseModel):
    default_machine_rate: float = 400
    default_labour_rate: float = 120

@api.put("/costing/rates")
async def set_costing_rates(body: CostingRatesIn, user=Depends(get_current_user)):
    await db.settings.replace_one({"_id": "costing"},
        {"_id": "costing", "default_machine_rate": body.default_machine_rate,
         "default_labour_rate": body.default_labour_rate}, upsert=True)
    return {"ok": True}

@api.get("/costing/overview")
async def costing_overview(user=Depends(get_current_user)):
    """Job costing per WO + machine/operator/customer profitability, from operation minutes x rates."""
    mrate_default, lrate_default = await _costing_rates()

    machines = await db.machines.find({}, {"_id": 0, "name": 1, "hourly_rate": 1}).to_list(500)
    mrate = {m["name"]: (float(m.get("hourly_rate") or 0) or mrate_default) for m in machines if m.get("name")}
    emps = await db.employees.find({}, {"_id": 0, "name": 1, "monthly_salary": 1}).to_list(2000)
    erate = {}
    for e in emps:
        nm = (e.get("name") or "").strip().lower()
        if nm:
            erate[nm] = (float(e.get("monthly_salary") or 0) / 208.0) if e.get("monthly_salary") else lrate_default
    items = await db.items.find({}, {"_id": 0, "name": 1, "unit_cost": 1}).to_list(8000)
    icost = {(i.get("name") or "").strip().lower(): float(i.get("unit_cost") or 0) for i in items}

    def op_minutes(o):
        a = float(o.get("actual_minutes") or 0)
        return a if a > 0 else float(o.get("planned_minutes") or 0)
    def mrate_for(name):
        return mrate.get(name, mrate_default) if name else mrate_default
    def lrate_for(name):
        return erate.get((name or "").strip().lower(), lrate_default) if name else lrate_default

    ops = await db.wo_operations.find({}, {"_id": 0}).to_list(30000)
    per_machine = {}; per_operator = {}; by_wo = {}
    for o in ops:
        mins = op_minutes(o)
        mc = mins * mrate_for(o.get("machine")) / 60.0
        lc = mins * lrate_for(o.get("operator")) / 60.0
        mn = o.get("machine") or "Unassigned"
        pm = per_machine.setdefault(mn, {"machine": mn, "minutes": 0, "cost": 0, "ops": 0})
        pm["minutes"] += mins; pm["cost"] += mc; pm["ops"] += 1
        on = o.get("operator") or "Unassigned"
        po = per_operator.setdefault(on, {"operator": on, "minutes": 0, "ops": 0, "done": 0})
        po["minutes"] += mins; po["ops"] += 1; po["done"] += 1 if o.get("status") == "done" else 0
        w = by_wo.setdefault(o.get("work_order_id"), {"machining": 0, "labour": 0, "minutes": 0, "ops": 0})
        w["machining"] += mc; w["labour"] += lc; w["minutes"] += mins; w["ops"] += 1

    movs = await db.material_state_movements.find({"ref_type": "WO"}, {"_id": 0, "ref_id": 1, "item_name": 1, "qty": 1}).to_list(30000)
    wo_material = {}
    for mv in movs:
        c = icost.get((mv.get("item_name") or "").strip().lower(), 0) * float(mv.get("qty") or 0)
        rid = mv.get("ref_id")
        wo_material[rid] = wo_material.get(rid, 0) + c

    wos = await db.work_orders.find({}, {"_id": 0, "id": 1, "code": 1, "product": 1, "customer_name": 1, "qty": 1, "status": 1}).to_list(5000)
    per_wo = []
    for w in wos:
        c = by_wo.get(w["id"], {"machining": 0, "labour": 0, "minutes": 0, "ops": 0})
        mat = wo_material.get(w["id"], 0)
        total = c["machining"] + c["labour"] + mat
        qty = float(w.get("qty") or 0)
        per_wo.append({"id": w["id"], "code": w.get("code"), "product": w.get("product"),
            "customer": w.get("customer_name"), "qty": w.get("qty"), "status": w.get("status"),
            "machining_cost": round(c["machining"], 1), "labour_cost": round(c["labour"], 1),
            "material_cost": round(mat, 1), "total_cost": round(total, 1),
            "cost_per_pc": round(total / qty, 2) if qty else 0, "ops": c["ops"]})
    per_wo.sort(key=lambda x: x["total_cost"], reverse=True)

    for d in per_machine.values():
        d["hours"] = round(d["minutes"] / 60, 1); d["cost"] = round(d["cost"], 0)
    for d in per_operator.values():
        d["hours"] = round(d["minutes"] / 60, 1)
    pm = sorted(per_machine.values(), key=lambda x: x["cost"], reverse=True)
    po = sorted(per_operator.values(), key=lambda x: x["minutes"], reverse=True)

    invs = await db.invoices.find({}, {"_id": 0, "customer_name": 1, "total": 1}).to_list(15000)
    rev = {}
    for iv in invs:
        nm = iv.get("customer_name") or "—"; rev[nm] = rev.get(nm, 0) + float(iv.get("total") or 0)
    wcost = {}
    for w in per_wo:
        nm = w.get("customer") or "—"; wcost[nm] = wcost.get(nm, 0) + w["total_cost"]
    per_customer = []
    for nm in set(list(rev.keys()) + list(wcost.keys())):
        r = round(rev.get(nm, 0), 0); c = round(wcost.get(nm, 0), 0)
        per_customer.append({"customer": nm, "revenue": r, "work_cost": c, "margin": round(r - c, 0)})
    per_customer.sort(key=lambda x: x["revenue"], reverse=True)

    return {
        "rates": {"machine": mrate_default, "labour": lrate_default},
        "totals": {"wip_cost": round(sum(w["total_cost"] for w in per_wo), 0), "wo_count": len(per_wo)},
        "per_wo": per_wo[:100], "per_machine": pm, "per_operator": po, "per_customer": per_customer[:50],
    }


# ---------------- Google Drive (OAuth: file store + data backup) ----------------
GOOGLE_OAUTH_CLIENT_ID = os.environ.get("GOOGLE_OAUTH_CLIENT_ID", "")
GOOGLE_OAUTH_CLIENT_SECRET = os.environ.get("GOOGLE_OAUTH_CLIENT_SECRET", "")
GOOGLE_OAUTH_REDIRECT_URI = os.environ.get(
    "GOOGLE_OAUTH_REDIRECT_URI",
    "https://denplex-erp-production.up.railway.app/api/google/oauth/callback")
GOOGLE_DRIVE_SCOPE = ("https://www.googleapis.com/auth/drive.file "
                      "https://www.googleapis.com/auth/userinfo.email")
_GOOGLE_TOKEN_URL = "https://oauth2.googleapis.com/token"

async def _gdrive_cfg():
    return await db.settings.find_one({"_id": "google_drive"}, {"_id": 0}) or {}

async def _gdrive_access_token():
    cfg = await _gdrive_cfg()
    rt = cfg.get("refresh_token")
    if not rt:
        raise HTTPException(400, "Google Drive is not connected.")
    if not (GOOGLE_OAUTH_CLIENT_ID and GOOGLE_OAUTH_CLIENT_SECRET):
        raise HTTPException(503, "Google OAuth not configured (missing client id/secret).")
    async with httpx.AsyncClient(timeout=30) as cx:
        r = await cx.post(_GOOGLE_TOKEN_URL, data={
            "client_id": GOOGLE_OAUTH_CLIENT_ID, "client_secret": GOOGLE_OAUTH_CLIENT_SECRET,
            "refresh_token": rt, "grant_type": "refresh_token"})
    if r.status_code >= 400:
        raise HTTPException(502, f"Google token refresh failed: {r.text[:200]}")
    return r.json().get("access_token")

async def _drive_ensure_folder(name, parent=None):
    token = await _gdrive_access_token()
    q = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent:
        q += f" and '{parent}' in parents"
    async with httpx.AsyncClient(timeout=30) as cx:
        r = await cx.get("https://www.googleapis.com/drive/v3/files",
            params={"q": q, "fields": "files(id,name)", "spaces": "drive"},
            headers={"Authorization": f"Bearer {token}"})
        files = r.json().get("files", []) if r.status_code < 400 else []
        if files:
            return files[0]["id"]
        meta = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
        if parent:
            meta["parents"] = [parent]
        c = await cx.post("https://www.googleapis.com/drive/v3/files",
            json=meta, headers={"Authorization": f"Bearer {token}"})
        if c.status_code >= 400:
            raise HTTPException(502, f"Drive folder create failed: {c.text[:200]}")
        return c.json()["id"]

async def _drive_upload(name, mime, data: bytes, parent=None):
    import json as _json
    token = await _gdrive_access_token()
    meta = {"name": name}
    if parent:
        meta["parents"] = [parent]
    boundary = "denplexerpboundary"
    body = (f"--{boundary}\r\nContent-Type: application/json; charset=UTF-8\r\n\r\n"
            f"{_json.dumps(meta)}\r\n--{boundary}\r\nContent-Type: {mime}\r\n\r\n").encode() + data + f"\r\n--{boundary}--".encode()
    async with httpx.AsyncClient(timeout=180) as cx:
        r = await cx.post("https://www.googleapis.com/upload/drive/v3/files",
            params={"uploadType": "multipart", "fields": "id,name,webViewLink"},
            content=body, headers={"Authorization": f"Bearer {token}",
                "Content-Type": f"multipart/related; boundary={boundary}"})
    if r.status_code >= 400:
        raise HTTPException(502, f"Drive upload failed: {r.text[:200]}")
    return r.json()

@api.get("/google/status")
async def google_status(user=Depends(get_current_user)):
    cfg = await _gdrive_cfg()
    return {"connected": bool(cfg.get("refresh_token")), "email": cfg.get("email", ""),
            "configured": bool(GOOGLE_OAUTH_CLIENT_ID and GOOGLE_OAUTH_CLIENT_SECRET),
            "last_backup": cfg.get("last_backup", ""), "auto_backup": cfg.get("auto_backup", True)}

@api.get("/google/oauth/start")
async def google_oauth_start(user=Depends(get_current_user)):
    if not GOOGLE_OAUTH_CLIENT_ID:
        raise HTTPException(503, "Google OAuth not configured. Set GOOGLE_OAUTH_CLIENT_ID / SECRET in the backend.")
    from urllib.parse import urlencode
    state = secrets.token_urlsafe(16)
    await db.settings.update_one({"_id": "google_drive"}, {"$set": {"oauth_state": state}}, upsert=True)
    params = {"client_id": GOOGLE_OAUTH_CLIENT_ID, "redirect_uri": GOOGLE_OAUTH_REDIRECT_URI,
              "response_type": "code", "scope": GOOGLE_DRIVE_SCOPE,
              "access_type": "offline", "prompt": "consent", "state": state}
    return {"auth_url": "https://accounts.google.com/o/oauth2/v2/auth?" + urlencode(params)}

@api.get("/google/oauth/callback")
async def google_oauth_callback(code: str = "", state: str = "", error: str = ""):
    from fastapi.responses import RedirectResponse
    front = QR_BASE_URL
    if error or not code:
        return RedirectResponse(f"{front}/app/settings?gdrive=error")
    cfg = await _gdrive_cfg()
    if state and cfg.get("oauth_state") and state != cfg.get("oauth_state"):
        return RedirectResponse(f"{front}/app/settings?gdrive=error")
    async with httpx.AsyncClient(timeout=30) as cx:
        r = await cx.post(_GOOGLE_TOKEN_URL, data={
            "client_id": GOOGLE_OAUTH_CLIENT_ID, "client_secret": GOOGLE_OAUTH_CLIENT_SECRET,
            "code": code, "redirect_uri": GOOGLE_OAUTH_REDIRECT_URI, "grant_type": "authorization_code"})
    if r.status_code >= 400:
        return RedirectResponse(f"{front}/app/settings?gdrive=error")
    tok = r.json(); rt = tok.get("refresh_token"); email = ""
    try:
        async with httpx.AsyncClient(timeout=20) as cx:
            ui = await cx.get("https://www.googleapis.com/oauth2/v2/userinfo",
                headers={"Authorization": f"Bearer {tok.get('access_token')}"})
            if ui.status_code < 400:
                email = ui.json().get("email", "")
    except Exception:
        pass
    upd = {"connected": True, "email": email}
    if rt:
        upd["refresh_token"] = rt
    await db.settings.update_one({"_id": "google_drive"}, {"$set": upd, "$unset": {"oauth_state": ""}}, upsert=True)
    return RedirectResponse(f"{front}/app/settings?gdrive=connected")

@api.post("/google/disconnect")
async def google_disconnect(user=Depends(require_roles("admin"))):
    await db.settings.update_one({"_id": "google_drive"},
        {"$unset": {"refresh_token": "", "email": "", "connected": ""}}, upsert=True)
    return {"ok": True}

async def _run_backup():
    import json as _json
    root = await _drive_ensure_folder("Denplex ERP")
    backups = await _drive_ensure_folder("Backups", root)
    cols = ["work_orders", "wo_operations", "machines", "parts", "boms", "items",
            "material_state_movements", "customers", "suppliers", "leads", "quotations",
            "invoices", "purchase_orders", "qc_inspections", "qc_reports", "employees",
            "expenses", "payments_in", "payments_out", "documents", "campaigns"]
    dump = {}
    for c in cols:
        try:
            dump[c] = await db[c].find({}, {"_id": 0}).to_list(100000)
        except Exception:
            dump[c] = []
    try:
        dump["users"] = await db.users.find({}, {"_id": 0, "password": 0, "totp_secret": 0}).to_list(5000)
    except Exception:
        dump["users"] = []
    payload = _json.dumps(dump, default=str).encode("utf-8")
    fname = f"denplex-erp-backup-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.json"
    res = await _drive_upload(fname, "application/json", payload, backups)
    await db.settings.update_one({"_id": "google_drive"}, {"$set": {"last_backup": now_iso()}}, upsert=True)
    return {"ok": True, "file": res.get("name"), "link": res.get("webViewLink"),
            "size_kb": round(len(payload) / 1024, 1), "collections": len(dump)}

@api.post("/google/backup")
async def google_backup(user=Depends(get_current_user)):
    """Export ERP data to JSON and upload to Google Drive (Denplex ERP / Backups)."""
    return await _run_backup()

class AutoBackupIn(BaseModel):
    enabled: bool = True
    interval_hours: int = 24

@api.put("/google/auto-backup")
async def google_auto_backup(body: AutoBackupIn, user=Depends(get_current_user)):
    await db.settings.update_one({"_id": "google_drive"},
        {"$set": {"auto_backup": bool(body.enabled),
                  "auto_backup_interval_hours": max(1, int(body.interval_hours or 24))}}, upsert=True)
    return {"ok": True}

# ---- Drive file storage: transparent base64 <-> Drive offload ----
async def _drive_offload(b64, filename, mime, category):
    """If Drive is connected, upload a base64 blob to Drive and return 'gdrive:<id>'.
    Returns the original value unchanged on any failure / if not connected / already offloaded."""
    if not b64 or (isinstance(b64, str) and b64.startswith("gdrive:")):
        return b64
    try:
        cfg = await _gdrive_cfg()
        if not cfg.get("refresh_token"):
            return b64
        raw = b64.split(",", 1)[1] if isinstance(b64, str) and b64.startswith("data:") else b64
        data = base64.b64decode(raw)
        root = await _drive_ensure_folder("Denplex ERP")
        folder = await _drive_ensure_folder(category, root)
        res = await _drive_upload(filename or "file", mime, data, folder)
        return "gdrive:" + res.get("id")
    except Exception as e:
        try: logger.warning("drive offload failed: %s", e)
        except Exception: pass
        return b64

async def _resolve_b64_or_drive(value):
    """Return raw bytes for a stored field that is either base64 or 'gdrive:<id>'."""
    if isinstance(value, str) and value.startswith("gdrive:"):
        token = await _gdrive_access_token()
        async with httpx.AsyncClient(timeout=120) as cx:
            r = await cx.get(f"https://www.googleapis.com/drive/v3/files/{value[len('gdrive:'):]}",
                params={"alt": "media"}, headers={"Authorization": f"Bearer {token}"})
        if r.status_code >= 400:
            raise HTTPException(502, f"Drive download failed: {r.text[:120]}")
        return r.content
    raw = value.split(",", 1)[1] if isinstance(value, str) and value.startswith("data:") else value
    return base64.b64decode(raw)

# ---- Automatic backup background loop ----
async def _backup_loop():
    await asyncio.sleep(90)
    while True:
        try:
            cfg = await _gdrive_cfg()
            if cfg.get("refresh_token") and cfg.get("auto_backup", True):
                last = cfg.get("last_backup")
                interval = float(cfg.get("auto_backup_interval_hours", 24) or 24)
                due = True
                if last:
                    try:
                        lt = datetime.fromisoformat(str(last).replace("Z", "+00:00"))
                        due = (datetime.now(timezone.utc) - lt).total_seconds() >= interval * 3600
                    except Exception:
                        due = True
                if due:
                    await _run_backup()
                    logger.info("Auto Drive backup completed")
        except Exception as e:
            try: logger.warning("auto backup loop error: %s", e)
            except Exception: pass
        await asyncio.sleep(1800)

@app.on_event("startup")
async def _start_backup_loop():
    try:
        asyncio.create_task(_backup_loop())
    except Exception:
        pass


# ---------------- ISO QMS: NCR + CAPA registers ----------------
class NCR(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    date: str = Field(default_factory=now_iso)
    source: str = "production"      # production | internal_audit | customer_complaint | supplier | other
    work_order_id: Optional[str] = ""
    work_order_code: Optional[str] = ""
    process_name: Optional[str] = ""
    product: Optional[str] = ""
    part_number: Optional[str] = ""
    customer_name: Optional[str] = ""
    supplier_name: Optional[str] = ""
    qty: float = 0
    description: str = ""           # non-conformity description (F/PRD/03)
    root_cause: Optional[str] = ""
    correction: Optional[str] = ""  # immediate correction
    disposition: str = "rework"     # rework | repair | regrade | scrap | use_as_is | return_to_supplier
    capa_id: Optional[str] = ""
    capa_code: Optional[str] = ""
    status: Literal["open", "closed"] = "open"
    raised_by: Optional[str] = ""
    closed_by: Optional[str] = ""
    closed_date: Optional[str] = ""
    remarks: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class CAPA(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = None
    date: str = Field(default_factory=now_iso)
    source: str = "ncr"            # ncr | internal_audit | customer_complaint | management_review | other
    ncr_id: Optional[str] = ""
    ncr_code: Optional[str] = ""
    nonconformity: str = ""
    root_cause: Optional[str] = ""
    corrective_action: Optional[str] = ""
    preventive_action: Optional[str] = ""
    responsibility: Optional[str] = ""
    target_date: Optional[str] = ""
    risk_assessment: Optional[str] = ""
    effectiveness: Optional[str] = ""   # verification of effectiveness
    verified_by: Optional[str] = ""
    verified_date: Optional[str] = ""
    iso_clause: Optional[str] = "8.7, 10.2"
    status: Literal["open", "in_progress", "verified", "closed"] = "open"
    raised_by: Optional[str] = ""
    remarks: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

@api.post("/ncrs")
async def create_ncr(n: NCR, user=Depends(get_current_user)):
    doc = n.model_dump(); doc["code"] = await gen_code("NCR", "ncr")
    if not doc.get("raised_by"): doc["raised_by"] = user.get("name", "")
    if doc.get("work_order_id"):
        wo = await db.work_orders.find_one({"id": doc["work_order_id"]}, {"_id": 0})
        if wo: doc["work_order_code"] = wo.get("code", "")
    await db.ncrs.insert_one(doc)
    return serialize(doc)

@api.get("/ncrs")
async def list_ncrs(user=Depends(get_current_user)):
    return await list_collection(db.ncrs)

@api.put("/ncrs/{nid}")
async def update_ncr(nid: str, n: NCR, user=Depends(get_current_user)):
    data = n.model_dump(); data.pop("id", None); data.pop("created_at", None)
    if data.get("status") == "closed" and not data.get("closed_date"):
        data["closed_date"] = now_iso(); data["closed_by"] = user.get("name", "")
    await db.ncrs.update_one({"id": nid}, {"$set": data})
    return {"ok": True}

@api.delete("/ncrs/{nid}")
async def del_ncr(nid: str, user=Depends(require_roles("admin", "manager", "qc"))):
    await db.ncrs.delete_one({"id": nid})
    return {"ok": True}

@api.post("/capas")
async def create_capa(c: CAPA, user=Depends(get_current_user)):
    doc = c.model_dump(); doc["code"] = await gen_code("CAPA", "capa")
    if not doc.get("raised_by"): doc["raised_by"] = user.get("name", "")
    if doc.get("ncr_id"):
        ncr = await db.ncrs.find_one({"id": doc["ncr_id"]}, {"_id": 0})
        if ncr:
            doc["ncr_code"] = ncr.get("code", "")
            await db.ncrs.update_one({"id": doc["ncr_id"]}, {"$set": {"capa_id": doc["id"], "capa_code": doc["code"]}})
    await db.capas.insert_one(doc)
    return serialize(doc)

@api.get("/capas")
async def list_capas(user=Depends(get_current_user)):
    return await list_collection(db.capas)

@api.put("/capas/{cid}")
async def update_capa(cid: str, c: CAPA, user=Depends(get_current_user)):
    data = c.model_dump(); data.pop("id", None); data.pop("created_at", None)
    if data.get("status") in ("verified", "closed") and not data.get("verified_date"):
        data["verified_date"] = now_iso(); data["verified_by"] = user.get("name", "")
    await db.capas.update_one({"id": cid}, {"$set": data})
    return {"ok": True}

@api.delete("/capas/{cid}")
async def del_capa(cid: str, user=Depends(require_roles("admin", "manager", "qc"))):
    await db.capas.delete_one({"id": cid})
    return {"ok": True}


# ---------------- ISO: Calibration + Supplier Quality + ISO form PDFs ----------------
class Instrument(BaseModel):                       # F/QCD/03 Calibration
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    instrument_name: str
    make: Optional[str] = ""
    range: Optional[str] = ""
    identification_no: Optional[str] = ""
    location: Optional[str] = ""                   # Vatva / Santej
    calibration_date: Optional[str] = ""
    due_date: Optional[str] = ""
    calibrated_by: Optional[str] = ""
    frequency_months: int = 12
    remarks: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class ApprovedSupplier(BaseModel):                 # F/PUR/03
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    name: str
    address: Optional[str] = ""
    material_service: Optional[str] = ""
    supplier_type: str = "trader"                  # manufacturer | trader | job_work | service
    approval_criteria: List[str] = []              # A..G codes
    approval_date: Optional[str] = ""
    status: Literal["approved", "on_hold", "removed"] = "approved"
    remarks: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

class SupplierEvaluation(BaseModel):               # F/PUR/04
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    supplier_name: str
    period: Optional[str] = ""
    quality_score: float = 0
    delivery_score: float = 0
    cost_score: float = 0
    responsiveness_score: float = 0
    system_score: float = 0
    total_pct: float = 0
    rating: Optional[str] = ""                     # A | B | C
    evaluated_by: Optional[str] = ""
    date: str = Field(default_factory=now_iso)
    remarks: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

# --- Calibration / Instruments ---
@api.post("/instruments")
async def create_instrument(i: Instrument, user=Depends(get_current_user)):
    doc = i.model_dump(); await db.instruments.insert_one(doc); return serialize(doc)
@api.get("/instruments")
async def list_instruments(user=Depends(get_current_user)):
    return await list_collection(db.instruments, sort_key="due_date")
@api.put("/instruments/{iid}")
async def update_instrument(iid: str, i: Instrument, user=Depends(get_current_user)):
    data = i.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.instruments.update_one({"id": iid}, {"$set": data}); return {"ok": True}
@api.delete("/instruments/{iid}")
async def del_instrument(iid: str, user=Depends(require_roles("admin","manager","qc"))):
    await db.instruments.delete_one({"id": iid}); return {"ok": True}

# --- Approved suppliers ---
@api.post("/approved-suppliers")
async def create_apsup(a: ApprovedSupplier, user=Depends(get_current_user)):
    doc = a.model_dump(); await db.approved_suppliers.insert_one(doc); return serialize(doc)
@api.get("/approved-suppliers")
async def list_apsup(user=Depends(get_current_user)):
    return await list_collection(db.approved_suppliers, sort_key="name")
@api.put("/approved-suppliers/{sid}")
async def update_apsup(sid: str, a: ApprovedSupplier, user=Depends(get_current_user)):
    data = a.model_dump(); data.pop("id", None); data.pop("created_at", None)
    await db.approved_suppliers.update_one({"id": sid}, {"$set": data}); return {"ok": True}
@api.delete("/approved-suppliers/{sid}")
async def del_apsup(sid: str, user=Depends(require_roles("admin","manager"))):
    await db.approved_suppliers.delete_one({"id": sid}); return {"ok": True}

# --- Supplier evaluations ---
@api.post("/supplier-evaluations")
async def create_supeval(e: SupplierEvaluation, user=Depends(get_current_user)):
    doc = e.model_dump()
    scores = [doc["quality_score"], doc["delivery_score"], doc["cost_score"], doc["responsiveness_score"], doc["system_score"]]
    if not doc.get("total_pct"):
        doc["total_pct"] = round(sum(float(x or 0) for x in scores) / (5 * 10) * 100, 1) if any(scores) else 0
    if not doc.get("rating"):
        t = doc["total_pct"]; doc["rating"] = "A" if t >= 85 else "B" if t >= 60 else "C"
    if not doc.get("evaluated_by"): doc["evaluated_by"] = user.get("name", "")
    await db.supplier_evaluations.insert_one(doc); return serialize(doc)
@api.get("/supplier-evaluations")
async def list_supeval(user=Depends(get_current_user)):
    return await list_collection(db.supplier_evaluations)
@api.put("/supplier-evaluations/{eid}")
async def update_supeval(eid: str, e: SupplierEvaluation, user=Depends(get_current_user)):
    data = e.model_dump(); data.pop("id", None); data.pop("created_at", None)
    scores = [data["quality_score"], data["delivery_score"], data["cost_score"], data["responsiveness_score"], data["system_score"]]
    data["total_pct"] = round(sum(float(x or 0) for x in scores) / (5 * 10) * 100, 1) if any(scores) else 0
    data["rating"] = "A" if data["total_pct"] >= 85 else "B" if data["total_pct"] >= 60 else "C"
    await db.supplier_evaluations.update_one({"id": eid}, {"$set": data}); return {"ok": True}
@api.delete("/supplier-evaluations/{eid}")
async def del_supeval(eid: str, user=Depends(require_roles("admin","manager"))):
    await db.supplier_evaluations.delete_one({"id": eid}); return {"ok": True}

# --- ISO record PDF (Denplex letterhead) for NCR / CAPA ---
def _iso_record_pdf(title, code, doc_no, sections):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=27*mm, bottomMargin=22*mm, leftMargin=14*mm, rightMargin=14*mm, title=title)
    styles = getSampleStyleSheet()
    BLACK = colors.HexColor("#1A1A1A"); RED = colors.HexColor("#CC0000")
    story = []
    story.append(Paragraph(title, ParagraphStyle("t", parent=styles["Title"], fontSize=14, fontName=_PDF_FONT_BOLD, textColor=BLACK, alignment=1)))
    head = Table([[f"Doc No: {doc_no}", f"No: {code or '-'}"]], colWidths=[(doc.width)/2.0]*2)
    head.setStyle(TableStyle([("FONT",(0,0),(-1,-1),_PDF_FONT_BOLD,9),("TEXTCOLOR",(0,0),(-1,-1),RED),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story.append(head); story.append(Spacer(1, 3*mm))
    for heading, rows in sections:
        if heading:
            story.append(Paragraph(heading, ParagraphStyle("h", parent=styles["Normal"], fontSize=10, fontName=_PDF_FONT_BOLD, textColor=RED, spaceBefore=6, spaceAfter=2)))
        data = [[Paragraph(f"<b>{k}</b>", ParagraphStyle("k", parent=styles["Normal"], fontSize=9, fontName=_PDF_FONT_REGULAR)),
                 Paragraph(str(v or "-"), ParagraphStyle("v", parent=styles["Normal"], fontSize=9, fontName=_PDF_FONT_REGULAR))] for k, v in rows]
        t = Table(data, colWidths=[55*mm, doc.width - 55*mm])
        t.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#CCCCCC")),
            ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#F5F5F5")),("VALIGN",(0,0),(-1,-1),"TOP"),
            ("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
        story.append(t)
    story.append(Spacer(1, 8*mm))
    story.append(Table([["Prepared By", "Approved By"]], colWidths=[doc.width/2.0]*2,
        style=TableStyle([("FONT",(0,0),(-1,-1),_PDF_FONT_BOLD,9),("TOPPADDING",(0,0),(-1,-1),18),("ALIGN",(0,0),(-1,-1),"CENTER")])))
    doc.build(story, onFirstPage=_qc_header_footer, onLaterPages=_qc_header_footer)
    return buf.getvalue()

@api.get("/ncrs/{nid}/pdf")
async def ncr_pdf(nid: str, user=Depends(get_current_user)):
    n = await db.ncrs.find_one({"id": nid}, {"_id": 0})
    if not n: raise HTTPException(404, "NCR not found")
    sections = [("", [("Date", (n.get("date") or "")[:10]), ("Source", (n.get("source") or "").replace("_", " ").title()),
        ("Process", n.get("process_name")), ("Product / Part", f"{n.get('product','')}  {n.get('part_number','')}"),
        ("Qty", n.get("qty")), ("Customer / Supplier", n.get("customer_name") or n.get("supplier_name"))]),
        ("Non-Conformity", [("Description", n.get("description")), ("Root Cause", n.get("root_cause")),
        ("Correction", n.get("correction")), ("Disposition", (n.get("disposition") or "").replace("_", " ").title()),
        ("Linked CAPA", n.get("capa_code")), ("Status", (n.get("status") or "").title()), ("Remarks", n.get("remarks"))])]
    pdf = _iso_record_pdf("Non-Conformance Report", n.get("code"), "F/PRD/03", sections)
    return Response(content=pdf, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{n.get("code","NCR")}.pdf"'})

@api.get("/capas/{cid}/pdf")
async def capa_pdf(cid: str, user=Depends(get_current_user)):
    c = await db.capas.find_one({"id": cid}, {"_id": 0})
    if not c: raise HTTPException(404, "CAPA not found")
    sections = [("", [("Date", (c.get("date") or "")[:10]), ("Source", (c.get("source") or "").replace("_", " ").title()),
        ("Linked NCR", c.get("ncr_code")), ("ISO Clause", c.get("iso_clause"))]),
        ("Analysis & Action", [("Non-Conformity", c.get("nonconformity")), ("Root Cause", c.get("root_cause")),
        ("Corrective Action", c.get("corrective_action")), ("Preventive Action", c.get("preventive_action")),
        ("Responsibility", c.get("responsibility")), ("Target Date", (c.get("target_date") or "")[:10]),
        ("Risk Assessment", c.get("risk_assessment")), ("Effectiveness Verification", c.get("effectiveness")),
        ("Status", (c.get("status") or "").replace("_", " ").title())])]
    pdf = _iso_record_pdf("Corrective & Preventive Action", c.get("code"), "F/QMS/10", sections)
    return Response(content=pdf, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{c.get("code","CAPA")}.pdf"'})


# ==================== ISO Documents Library (F/QMS docs, procedures, policies) ====================
# Stores QMS documents imported from the company's ISO Google-Drive folder so they can be
# browsed, edited in-app (rich text) and downloaded as Denplex-letterhead DOCX / PDF.
# - doc_type "text"  -> editable html_content, exports to docx/pdf
# - doc_type "file"  -> binary original (register/annexure/EHS pdf) streamed from Drive by source_drive_id
class ISODocument(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = ""            # e.g. PR/QAC/01, F/QMS/05
    title: str
    category: str = "General"           # Manual, QMS Procedure, Department Procedure, Work Instruction, Policy, Quality Objective, Master List, Annexure, Register, EHS, Inspection
    scope: Literal["master", "fy26-27"] = "master"
    doc_type: Literal["text", "file"] = "text"
    html_content: Optional[str] = ""
    source_drive_id: Optional[str] = ""
    source_url: Optional[str] = ""
    file_name: Optional[str] = ""
    mime: Optional[str] = ""
    revision: int = 0
    department: Optional[str] = ""       # Production, QC, Design, Purchase, Store, HR, Marketing, Maintenance, Management
    status: Literal["approved", "pending_approval"] = "approved"
    pending_html: Optional[str] = None   # proposed edit awaiting approval
    submitted_by: Optional[str] = ""
    submitted_at: Optional[str] = ""
    approved_by: Optional[str] = ""
    approved_at: Optional[str] = ""
    history: List[Dict[str, Any]] = []   # [{revision, html_content, updated_by, updated_at}]
    updated_at: str = Field(default_factory=now_iso)
    created_at: str = Field(default_factory=now_iso)

class ISODocBulkIn(BaseModel):
    documents: List[ISODocument] = []
    replace_all: bool = False           # if true, wipe the collection first (clean re-seed)

import html as _htmlmod
from html.parser import HTMLParser as _HTMLParser

_ISO_INLINE = {"b": "b", "strong": "b", "i": "i", "em": "i", "u": "u"}
_ISO_BLOCK = {"p", "div", "h1", "h2", "h3", "h4", "h5", "li", "tr", "br"}

class _ISOBlockParser(_HTMLParser):
    """Convert a subset of HTML into an ordered list of blocks:
       ('h', level, inline_html) | ('p', inline_html) | ('li', ordered_bool, inline_html) | ('table', [[cell_html,...],...])"""
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self.buf = ""
        self.pending = None          # ('h', level) or ('li', ordered)
        self.list_stack = []         # 'ul' | 'ol'
        self.in_table = False
        self.table_rows = []
        self.cur_row = None
        self.cur_cell = None
    def _flush(self):
        txt = self.buf.strip()
        if txt:
            if self.pending and self.pending[0] == "h":
                self.blocks.append(("h", self.pending[1], txt))
            elif self.pending and self.pending[0] == "li":
                self.blocks.append(("li", self.pending[1], txt))
            else:
                self.blocks.append(("p", txt))
        self.buf = ""
        self.pending = None
    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag == "table":
            self._flush(); self.in_table = True; self.table_rows = []; return
        if self.in_table:
            if tag == "tr": self.cur_row = []
            elif tag in ("td", "th"): self.cur_cell = ""
            elif tag in _ISO_INLINE and self.cur_cell is not None: self.cur_cell += f"<{_ISO_INLINE[tag]}>"
            return
        if tag in ("ul", "ol"): self._flush(); self.list_stack.append(tag)
        elif tag == "li": self._flush(); self.pending = ("li", (self.list_stack[-1] == "ol") if self.list_stack else False)
        elif tag in ("h1", "h2", "h3", "h4", "h5"): self._flush(); self.pending = ("h", min(int(tag[1]), 3))
        elif tag in ("p", "div"): self._flush()
        elif tag == "br": self.buf += "\n"
        elif tag in _ISO_INLINE: self.buf += f"<{_ISO_INLINE[tag]}>"
    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag == "table":
            if self.cur_row: self.table_rows.append(self.cur_row)
            if self.table_rows: self.blocks.append(("table", self.table_rows))
            self.in_table = False; self.table_rows = []; self.cur_row = None; self.cur_cell = None; return
        if self.in_table:
            if tag == "tr":
                if self.cur_row is not None: self.table_rows.append(self.cur_row); self.cur_row = None
            elif tag in ("td", "th"):
                if self.cur_row is not None: self.cur_row.append((self.cur_cell or "").strip()); self.cur_cell = None
            elif tag in _ISO_INLINE and self.cur_cell is not None: self.cur_cell += f"</{_ISO_INLINE[tag]}>"
            return
        if tag in ("ul", "ol"):
            self._flush()
            if self.list_stack: self.list_stack.pop()
        elif tag in ("li", "p", "div", "h1", "h2", "h3", "h4", "h5"): self._flush()
        elif tag in _ISO_INLINE: self.buf += f"</{_ISO_INLINE[tag]}>"
    def handle_data(self, data):
        if self.in_table:
            if self.cur_cell is not None: self.cur_cell += _htmlmod.escape(data)
        else:
            self.buf += _htmlmod.escape(data)
    def close(self):
        super().close(); self._flush()
        return self.blocks

def _iso_blocks(html_content: str):
    p = _ISOBlockParser()
    try:
        p.feed(html_content or "")
        return p.close()
    except Exception:
        # fall back to plain text paragraphs
        txt = re.sub("<[^>]+>", "", html_content or "")
        return [("p", _htmlmod.escape(line)) for line in txt.splitlines() if line.strip()]

_ISO_INLINE_RE = re.compile(r"(</?[biu]>)")
def _iso_inline_runs(inline_html: str):
    """Yield (text, bold, italic, underline) runs from inline html with <b>/<i>/<u>."""
    bold = ital = und = 0
    for tok in _ISO_INLINE_RE.split(inline_html or ""):
        if not tok: continue
        if tok == "<b>": bold += 1
        elif tok == "</b>": bold = max(0, bold - 1)
        elif tok == "<i>": ital += 1
        elif tok == "</i>": ital = max(0, ital - 1)
        elif tok == "<u>": und += 1
        elif tok == "</u>": und = max(0, und - 1)
        else:
            yield (_htmlmod.unescape(tok), bold > 0, ital > 0, und > 0)

def _iso_inline_to_rl(inline_html: str) -> str:
    """Sanitize inline html for a reportlab Paragraph (keeps b/i/u, escapes the rest, \\n -> <br/>)."""
    out = []
    for text, b, i, u in _iso_inline_runs(inline_html):
        t = _htmlmod.escape(text).replace("\n", "<br/>")
        if u: t = f"<u>{t}</u>"
        if i: t = f"<i>{t}</i>"
        if b: t = f"<b>{t}</b>"
        out.append(t)
    return "".join(out) or "&nbsp;"

def build_iso_doc_docx(d: Dict[str, Any]) -> bytes:
    """Render an ISO text document as an editable Denplex-letterhead .docx."""
    import io as _io
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    RED = RGBColor(0xCC, 0x00, 0x00); BLACK = RGBColor(0x1A, 0x1A, 0x1A)
    LETTERHEAD = str(ROOT_DIR / "letterhead.docx")
    try: doc = Document(LETTERHEAD)
    except Exception: doc = Document()
    doc._body.clear_content()
    try:
        nrm = doc.styles["Normal"].font; nrm.name = "Arial"; nrm.size = Pt(10)
    except Exception: pass

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc); tcPr.append(sh)
    def borders(table, color="999999"):
        tblPr = table._tbl.tblPr; b = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            e = OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
            e.set(qn('w:space'),'0'); e.set(qn('w:color'),color); b.append(e)
        tblPr.append(b)
    def add_inline(par, inline_html, size=10, base_bold=False, color=BLACK):
        any_run = False
        for text, b, i, u in _iso_inline_runs(inline_html):
            for j, seg in enumerate(text.split("\n")):
                if j > 0: par.add_run().add_break()
                if not seg: continue
                r = par.add_run(seg); r.font.size = Pt(size); r.font.name = "Arial"
                r.font.bold = bool(b or base_bold); r.font.italic = bool(i); r.font.underline = bool(u)
                r.font.color.rgb = color; any_run = True
        return any_run

    # Title header block
    head = doc.add_table(rows=1, cols=2); borders(head, "CC0000")
    c0 = head.cell(0, 0).paragraphs[0]; r = c0.add_run(d.get("title", "Document")); r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLACK; r.font.name = "Arial"
    c1 = head.cell(0, 1).paragraphs[0]; c1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta = []
    if d.get("code"): meta.append(f"Doc No: {d['code']}")
    if d.get("category"): meta.append(d["category"])
    meta.append(f"Rev: {d.get('revision', 0)}")
    rr = c1.add_run("\n".join(meta)); rr.font.size = Pt(8); rr.font.color.rgb = RED; rr.font.bold = True; rr.font.name = "Arial"
    doc.add_paragraph()

    for blk in _iso_blocks(d.get("html_content", "")):
        if blk[0] == "h":
            par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(6); par.paragraph_format.space_after = Pt(2)
            sizes = {1: 13, 2: 11.5, 3: 10.5}
            add_inline(par, blk[2], size=sizes.get(blk[1], 11), base_bold=True, color=RED)
        elif blk[0] == "p":
            par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(3); add_inline(par, blk[1], 10)
        elif blk[0] == "li":
            par = doc.add_paragraph(style=None); par.paragraph_format.left_indent = Mm(6); par.paragraph_format.space_after = Pt(2)
            par.add_run("•  ").font.name = "Arial"; add_inline(par, blk[2], 10)
        elif blk[0] == "table":
            rows = blk[1]; cols = max((len(r) for r in rows), default=1)
            if cols and rows:
                tb = doc.add_table(rows=0, cols=cols); borders(tb)
                for ri, row in enumerate(rows):
                    cells = tb.add_row().cells
                    for ci in range(cols):
                        cell_html = row[ci] if ci < len(row) else ""
                        add_inline(cells[ci].paragraphs[0], cell_html, 9, base_bold=(ri == 0))
                        if ri == 0: shade(cells[ci], "F0F0F0")
            doc.add_paragraph()
    out = _io.BytesIO(); doc.save(out); return out.getvalue()

def build_iso_doc_pdf(d: Dict[str, Any]) -> bytes:
    """Render an ISO text document as a Denplex-letterhead PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=27*mm, bottomMargin=22*mm, leftMargin=14*mm, rightMargin=14*mm, title=d.get("title", "Document"))
    styles = getSampleStyleSheet()
    BLACK = colors.HexColor("#1A1A1A"); RED = colors.HexColor("#CC0000")
    body = ParagraphStyle("b", parent=styles["Normal"], fontName=_PDF_FONT_REGULAR, fontSize=9.5, leading=13, textColor=BLACK, spaceAfter=3)
    h_styles = {1: ParagraphStyle("h1", parent=body, fontName=_PDF_FONT_BOLD, fontSize=12, textColor=RED, spaceBefore=8, spaceAfter=3),
                2: ParagraphStyle("h2", parent=body, fontName=_PDF_FONT_BOLD, fontSize=10.5, textColor=RED, spaceBefore=6, spaceAfter=2),
                3: ParagraphStyle("h3", parent=body, fontName=_PDF_FONT_BOLD, fontSize=10, textColor=BLACK, spaceBefore=5, spaceAfter=2)}
    bullet = ParagraphStyle("bul", parent=body, leftIndent=10, bulletIndent=0)
    story = []
    story.append(Paragraph(d.get("title", "Document"), ParagraphStyle("t", parent=styles["Title"], fontSize=14, fontName=_PDF_FONT_BOLD, textColor=BLACK, alignment=1)))
    meta = " · ".join([x for x in [f"Doc No: {d.get('code')}" if d.get('code') else "", d.get("category", ""), f"Rev: {d.get('revision', 0)}"] if x])
    head = Table([[meta]], colWidths=[doc.width])
    head.setStyle(TableStyle([("FONT", (0,0), (-1,-1), _PDF_FONT_BOLD, 8), ("TEXTCOLOR", (0,0), (-1,-1), RED), ("BOTTOMPADDING", (0,0), (-1,-1), 6)]))
    story.append(head)
    for blk in _iso_blocks(d.get("html_content", "")):
        try:
            if blk[0] == "h":
                story.append(Paragraph(_iso_inline_to_rl(blk[2]), h_styles.get(blk[1], h_styles[2])))
            elif blk[0] == "p":
                story.append(Paragraph(_iso_inline_to_rl(blk[1]), body))
            elif blk[0] == "li":
                story.append(Paragraph("•&nbsp;&nbsp;" + _iso_inline_to_rl(blk[2]), bullet))
            elif blk[0] == "table":
                rows = blk[1]; cols = max((len(r) for r in rows), default=1)
                data = [[Paragraph(_iso_inline_to_rl(c), body) for c in (r + [""] * (cols - len(r)))] for r in rows]
                t = Table(data, colWidths=[doc.width / cols] * cols)
                t.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
                    ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F0F0F0")), ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4), ("TOPPADDING", (0,0), (-1,-1), 3), ("BOTTOMPADDING", (0,0), (-1,-1), 3)]))
                story.append(Spacer(1, 2*mm)); story.append(t); story.append(Spacer(1, 2*mm))
        except Exception:
            continue
    doc.build(story, onFirstPage=_qc_header_footer, onLaterPages=_qc_header_footer)
    return buf.getvalue()

@api.post("/iso-documents")
async def create_iso_document(d: ISODocument, user=Depends(get_current_user)):
    doc = d.model_dump(); doc["updated_at"] = now_iso()
    await db.iso_documents.insert_one(doc)
    return serialize(doc)

@api.post("/iso-documents/bulk")
async def bulk_iso_documents(body: ISODocBulkIn, user=Depends(get_current_user)):
    if user.get("role") not in ("admin", "manager"):
        raise HTTPException(403, "Only admin/manager can bulk import")
    if body.replace_all:
        await db.iso_documents.delete_many({})
    n = 0
    for d in body.documents:
        doc = d.model_dump(); doc["updated_at"] = now_iso()
        # upsert by (code, scope, title) so re-seeding is idempotent
        key = {"scope": doc["scope"], "title": doc["title"]}
        if doc.get("code"): key = {"scope": doc["scope"], "code": doc["code"], "title": doc["title"]}
        existing = await db.iso_documents.find_one(key)
        if existing:
            doc["id"] = existing["id"]; doc["created_at"] = existing.get("created_at", doc["created_at"])
            await db.iso_documents.replace_one({"id": existing["id"]}, doc)
        else:
            await db.iso_documents.insert_one(doc)
        n += 1
    return {"imported": n}

@api.get("/iso-documents")
async def list_iso_documents(scope: Optional[str] = None, category: Optional[str] = None,
                             department: Optional[str] = None, status: Optional[str] = None,
                             user=Depends(get_current_user)):
    q = {}
    if scope: q["scope"] = scope
    if category: q["category"] = category
    if department: q["department"] = department
    if status: q["status"] = status
    cursor = db.iso_documents.find(q, {"_id": 0, "html_content": 0, "pending_html": 0, "history": 0}).sort([("category", 1), ("code", 1), ("title", 1)])
    return await cursor.to_list(5000)

@api.get("/iso-documents/{did}")
async def get_iso_document(did: str, user=Depends(get_current_user)):
    d = await db.iso_documents.find_one({"id": did}, {"_id": 0})
    if not d: raise HTTPException(404, "Document not found")
    return d

_ISO_APPROVERS = ("admin", "manager")

@api.put("/iso-documents/{did}")
async def update_iso_document(did: str, payload: Dict[str, Any], user=Depends(get_current_user)):
    """Metadata edits apply immediately. Content (html_content) edits by admin/manager apply immediately
    and bump the revision; content edits by other roles are queued as pending_approval."""
    d = await db.iso_documents.find_one({"id": did})
    if not d: raise HTTPException(404, "Document not found")
    set_fields = {k: payload[k] for k in ("title", "code", "category", "scope", "department") if k in payload}
    set_fields["updated_at"] = now_iso()
    is_approver = user.get("role") in _ISO_APPROVERS
    if "html_content" in payload and payload["html_content"] != d.get("html_content"):
        if is_approver:
            hist = d.get("history", [])
            hist.append({"revision": d.get("revision", 0), "html_content": d.get("html_content", ""),
                         "updated_by": d.get("approved_by") or d.get("submitted_by") or "", "updated_at": d.get("updated_at", "")})
            set_fields.update({"html_content": payload["html_content"], "revision": int(d.get("revision", 0)) + 1,
                               "status": "approved", "approved_by": user.get("name") or user.get("email", ""),
                               "approved_at": now_iso(), "pending_html": None, "history": hist[-30:]})
        else:
            set_fields.update({"pending_html": payload["html_content"], "status": "pending_approval",
                               "submitted_by": user.get("name") or user.get("email", ""), "submitted_at": now_iso()})
    await db.iso_documents.update_one({"id": did}, {"$set": set_fields})
    return await db.iso_documents.find_one({"id": did}, {"_id": 0})

@api.post("/iso-documents/{did}/approve")
async def approve_iso_document(did: str, user=Depends(require_roles("admin", "manager"))):
    d = await db.iso_documents.find_one({"id": did})
    if not d: raise HTTPException(404, "Document not found")
    if not d.get("pending_html"): raise HTTPException(400, "No pending change to approve")
    hist = d.get("history", [])
    hist.append({"revision": d.get("revision", 0), "html_content": d.get("html_content", ""),
                 "updated_by": d.get("approved_by") or "", "updated_at": d.get("updated_at", "")})
    await db.iso_documents.update_one({"id": did}, {"$set": {
        "html_content": d["pending_html"], "revision": int(d.get("revision", 0)) + 1, "status": "approved",
        "pending_html": None, "approved_by": user.get("name") or user.get("email", ""),
        "approved_at": now_iso(), "updated_at": now_iso(), "history": hist[-30:]}})
    return await db.iso_documents.find_one({"id": did}, {"_id": 0})

@api.post("/iso-documents/{did}/reject")
async def reject_iso_document(did: str, user=Depends(require_roles("admin", "manager"))):
    d = await db.iso_documents.find_one({"id": did})
    if not d: raise HTTPException(404, "Document not found")
    await db.iso_documents.update_one({"id": did}, {"$set": {
        "pending_html": None, "status": "approved", "updated_at": now_iso()}})
    return {"ok": True}

@api.delete("/iso-documents/{did}")
async def delete_iso_document(did: str, user=Depends(get_current_user)):
    doc = await db.iso_documents.find_one({"id": did}, {"_id": 0})
    await _recycle("iso_documents", "ISO Document", doc, user)
    await db.iso_documents.delete_one({"id": did})
    return {"ok": True}

@api.get("/iso-documents/{did}/docx")
async def iso_document_docx(did: str, user=Depends(get_current_user)):
    d = await db.iso_documents.find_one({"id": did}, {"_id": 0})
    if not d: raise HTTPException(404, "Document not found")
    data = build_iso_doc_docx(d)
    fn = (d.get("code") or d.get("title") or "document").replace("/", "-")
    return Response(content=data, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fn}.docx"'})

@api.get("/iso-documents/{did}/pdf")
async def iso_document_pdf(did: str, user=Depends(get_current_user)):
    d = await db.iso_documents.find_one({"id": did}, {"_id": 0})
    if not d: raise HTTPException(404, "Document not found")
    data = build_iso_doc_pdf(d)
    fn = (d.get("code") or d.get("title") or "document").replace("/", "-")
    return Response(content=data, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fn}.pdf"'})

@api.get("/iso-documents/{did}/file")
async def iso_document_file(did: str, user=Depends(get_current_user)):
    """Stream the original binary (register/annexure/EHS pdf) from Google Drive by source_drive_id."""
    d = await db.iso_documents.find_one({"id": did}, {"_id": 0})
    if not d: raise HTTPException(404, "Document not found")
    fid = d.get("source_drive_id")
    if not fid: raise HTTPException(404, "No original file linked")
    token = await _gdrive_access_token()
    async with httpx.AsyncClient(timeout=120) as cx:
        r = await cx.get(f"https://www.googleapis.com/drive/v3/files/{fid}",
            params={"alt": "media"}, headers={"Authorization": f"Bearer {token}"})
    if r.status_code >= 400:
        raise HTTPException(502, f"Drive download failed: {r.text[:120]}")
    fn = d.get("file_name") or (d.get("title") or "file")
    return Response(content=r.content, media_type=d.get("mime") or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{fn}"'})


# ==================== Registers (periodic data-entry forms: daily/weekly/monthly) ====================
# A generic, configurable register engine. Each RegisterTemplate defines columns; users add dated
# RegisterEntry rows (free data entry), exportable to Excel and Denplex-letterhead PDF.
class RegisterColumn(BaseModel):
    model_config = ConfigDict(extra="ignore")
    key: str
    label: str
    type: Literal["text", "number", "date", "select", "textarea"] = "text"
    options: List[str] = []

class RegisterTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    code: Optional[str] = ""             # e.g. F/PRD/02
    name: str
    department: str = "General"
    frequency: Literal["daily", "weekly", "monthly", "quarterly", "yearly", "as_required"] = "as_required"
    description: Optional[str] = ""
    columns: List[RegisterColumn] = []
    active: bool = True
    created_at: str = Field(default_factory=now_iso)

class RegisterEntry(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    template_id: Optional[str] = ""      # taken from the URL path, not required in the body
    date: str = Field(default_factory=lambda: now_iso()[:10])
    data: Dict[str, Any] = {}
    created_by: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)
    updated_at: str = Field(default_factory=now_iso)

class RegisterBulkIn(BaseModel):
    templates: List[RegisterTemplate] = []
    replace_all: bool = False

class RegisterEntryBulkIn(BaseModel):
    entries: List[RegisterEntry] = []
    replace_all: bool = False           # wipe this register's existing entries first

@api.post("/registers")
async def create_register(t: RegisterTemplate, user=Depends(require_roles("admin", "manager"))):
    doc = t.model_dump(); await db.register_templates.insert_one(doc); return serialize(doc)

@api.post("/registers/bulk")
async def bulk_registers(body: RegisterBulkIn, user=Depends(require_roles("admin", "manager"))):
    if body.replace_all:
        await db.register_templates.delete_many({})
    n = 0
    for t in body.templates:
        doc = t.model_dump()
        existing = await db.register_templates.find_one({"name": doc["name"], "department": doc["department"]})
        if existing:
            doc["id"] = existing["id"]; doc["created_at"] = existing.get("created_at", doc["created_at"])
            await db.register_templates.replace_one({"id": existing["id"]}, doc)
        else:
            await db.register_templates.insert_one(doc)
        n += 1
    return {"imported": n}

@api.get("/registers")
async def list_registers(department: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if department: q["department"] = department
    cur = db.register_templates.find(q, {"_id": 0}).sort([("department", 1), ("name", 1)])
    return await cur.to_list(2000)

@api.get("/registers/{tid}")
async def get_register(tid: str, user=Depends(get_current_user)):
    t = await db.register_templates.find_one({"id": tid}, {"_id": 0})
    if not t: raise HTTPException(404, "Register not found")
    return t

@api.put("/registers/{tid}")
async def update_register(tid: str, payload: Dict[str, Any], user=Depends(require_roles("admin", "manager"))):
    payload.pop("id", None); payload.pop("_id", None)
    await db.register_templates.update_one({"id": tid}, {"$set": payload})
    return await db.register_templates.find_one({"id": tid}, {"_id": 0})

@api.delete("/registers/{tid}")
async def delete_register(tid: str, user=Depends(require_roles("admin", "manager"))):
    doc = await db.register_templates.find_one({"id": tid}, {"_id": 0})
    if doc:
        entries = await db.register_entries.find({"template_id": tid}, {"_id": 0}).to_list(20000)
        doc["_entries_backup"] = entries   # keep entries with the snapshot so a restore brings data back
        await _recycle("register_templates", "Register", doc, user)
    await db.register_templates.delete_one({"id": tid})
    await db.register_entries.delete_many({"template_id": tid})
    return {"ok": True}

@api.get("/registers/{tid}/entries")
async def list_register_entries(tid: str, user=Depends(get_current_user)):
    cur = db.register_entries.find({"template_id": tid}, {"_id": 0}).sort([("date", -1), ("created_at", -1)])
    return await cur.to_list(20000)

@api.post("/registers/{tid}/entries")
async def create_register_entry(tid: str, e: RegisterEntry, user=Depends(get_current_user)):
    t = await db.register_templates.find_one({"id": tid})
    if not t: raise HTTPException(404, "Register not found")
    doc = e.model_dump(); doc["template_id"] = tid
    doc["created_by"] = user.get("name") or user.get("email", "")
    await db.register_entries.insert_one(doc); return serialize(doc)

@api.post("/registers/{tid}/entries/bulk")
async def bulk_register_entries(tid: str, body: RegisterEntryBulkIn, user=Depends(get_current_user)):
    """Insert many entries in one request (used for historical data import)."""
    t = await db.register_templates.find_one({"id": tid})
    if not t: raise HTTPException(404, "Register not found")
    if body.replace_all:
        await db.register_entries.delete_many({"template_id": tid})
    by = user.get("name") or user.get("email", "")
    docs = []
    for e in body.entries:
        d = e.model_dump(); d["template_id"] = tid; d["created_by"] = by
        docs.append(d)
    if docs:
        await db.register_entries.insert_many(docs)
    return {"imported": len(docs)}

@api.put("/registers/{tid}/entries/{eid}")
async def update_register_entry(tid: str, eid: str, payload: Dict[str, Any], user=Depends(get_current_user)):
    upd = {k: payload[k] for k in ("date", "data") if k in payload}
    upd["updated_at"] = now_iso()
    await db.register_entries.update_one({"id": eid, "template_id": tid}, {"$set": upd})
    return await db.register_entries.find_one({"id": eid}, {"_id": 0})

@api.delete("/registers/{tid}/entries/{eid}")
async def delete_register_entry(tid: str, eid: str, user=Depends(get_current_user)):
    await db.register_entries.delete_one({"id": eid, "template_id": tid})
    return {"ok": True}

async def _register_rows(tid: str):
    t = await db.register_templates.find_one({"id": tid}, {"_id": 0})
    if not t: raise HTTPException(404, "Register not found")
    entries = await db.register_entries.find({"template_id": tid}, {"_id": 0}).sort([("date", 1), ("created_at", 1)]).to_list(20000)
    cols = t.get("columns", [])
    return t, cols, entries

def _filter_register_entries(cols, entries, location: str = "", month: str = "", q: str = ""):
    """Apply the same filters the UI shows (location chip / month / search) so an
    export reflects exactly what the user is looking at, not the whole register."""
    loc_col = next((c for c in cols if c.get("key") == "location" or str(c.get("label", "")).lower() == "location"), None)
    location = (location or "").strip(); month = (month or "").strip(); q = (q or "").strip().lower()
    out = []
    for e in entries:
        data = e.get("data", {}) or {}
        if loc_col and location and str(data.get(loc_col.get("key"), "")).strip() != location:
            continue
        if month and str(e.get("date", ""))[:7] != month:
            continue
        if q:
            hay = (str(e.get("date", "")) + " " + " ".join(str(data.get(c.get("key"), "")) for c in cols)).lower()
            if q not in hay:
                continue
        out.append(e)
    return out

@api.get("/registers/{tid}/export/xlsx")
async def register_export_xlsx(tid: str, location: str = "", month: str = "", q: str = "", user=Depends(get_current_user)):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    t, cols, entries = await _register_rows(tid)
    entries = _filter_register_entries(cols, entries, location, month, q)
    wb = Workbook(); ws = wb.active
    ws.title = re.sub(r"[\\/*?:\[\]]", "-", (t.get("code") or t.get("name") or "Register"))[:31]
    RED = "CC0000"; thin = Side(style="thin", color="CCCCCC"); border = Border(thin, thin, thin, thin)
    ncol = max(1, len(cols) + 1)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncol)
    c = ws.cell(1, 1, "DENPLEX ENGINEERING COMPANY"); c.font = Font(bold=True, size=14, color=RED); c.alignment = Alignment(horizontal="center")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncol)
    c = ws.cell(2, 1, f"{t.get('name','')}" + (f"   ·   Doc No: {t.get('code')}" if t.get("code") else "") + f"   ·   Frequency: {t.get('frequency','')}")
    c.font = Font(bold=True, size=10); c.alignment = Alignment(horizontal="center")
    hdr_row = 4
    headers = ["Date"] + [col.get("label", col.get("key", "")) for col in cols]
    for j, h in enumerate(headers, 1):
        cell = ws.cell(hdr_row, j, h); cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1A1A1A"); cell.alignment = Alignment(horizontal="center", wrap_text=True); cell.border = border
    r = hdr_row + 1
    for e in entries:
        ws.cell(r, 1, (e.get("date") or "")[:10]).border = border
        for j, col in enumerate(cols, 2):
            ws.cell(r, j, e.get("data", {}).get(col.get("key"), "")).border = border
        r += 1
    for j in range(1, len(headers) + 1):
        ws.column_dimensions[ws.cell(hdr_row, j).column_letter].width = 20
    buf = BytesIO(); wb.save(buf)
    fn = (t.get("code") or t.get("name") or "register").replace("/", "-")
    return Response(content=buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fn}.xlsx"'})

@api.get("/registers/{tid}/export/pdf")
async def register_export_pdf(tid: str, location: str = "", month: str = "", q: str = "", user=Depends(get_current_user)):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    t, cols, entries = await _register_rows(tid)
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=27*mm, bottomMargin=18*mm, leftMargin=10*mm, rightMargin=10*mm, title=t.get("name", "Register"))
    styles = getSampleStyleSheet()
    BLACK = colors.HexColor("#1A1A1A"); RED = colors.HexColor("#CC0000")
    cell = ParagraphStyle("c", parent=styles["Normal"], fontName=_PDF_FONT_REGULAR, fontSize=7.5, leading=9)
    hcell = ParagraphStyle("hc", parent=cell, fontName=_PDF_FONT_BOLD, textColor=colors.white)
    story = [Paragraph(t.get("name", "Register"), ParagraphStyle("t", parent=styles["Title"], fontSize=13, fontName=_PDF_FONT_BOLD, textColor=BLACK, alignment=1))]
    meta = " · ".join([x for x in [f"Doc No: {t.get('code')}" if t.get("code") else "", f"Dept: {t.get('department','')}", f"Frequency: {t.get('frequency','')}"] if x])
    story.append(Paragraph(meta, ParagraphStyle("m", parent=cell, fontName=_PDF_FONT_BOLD, textColor=RED, alignment=1, spaceAfter=4)))
    headers = ["Date"] + [c.get("label", c.get("key", "")) for c in cols]
    data = [[Paragraph(h, hcell) for h in headers]]
    for e in entries:
        row = [Paragraph((e.get("date") or "")[:10], cell)]
        for col in cols:
            row.append(Paragraph(str(e.get("data", {}).get(col.get("key"), "") or ""), cell))
        data.append(row)
    if len(data) == 1:
        data.append([Paragraph("No entries yet", cell)] + [Paragraph("", cell) for _ in cols])
    tbl = Table(data, colWidths=[doc.width / len(headers)] * len(headers), repeatRows=1)
    tbl.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1A1A1A")), ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 3), ("RIGHTPADDING", (0,0), (-1,-1), 3), ("TOPPADDING", (0,0), (-1,-1), 2), ("BOTTOMPADDING", (0,0), (-1,-1), 2),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F7F7F7")])]))
    story.append(tbl)
    doc.build(story, onFirstPage=_qc_header_footer, onLaterPages=_qc_header_footer)
    fn = (t.get("code") or t.get("name") or "register").replace("/", "-")
    return Response(content=buf.getvalue(), media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fn}.pdf"'})


def _hsn_tax_summary(lines: List[Dict[str, Any]], is_interstate: bool) -> List[Dict[str, Any]]:
    """Aggregate per-HSN tax breakup for the Tax Summary block."""
    bucket: Dict[str, Dict[str, float]] = {}
    for l in lines or []:
        hsn = str(l.get("hsn") or "")
        qty = float(l.get("qty", 0) or 0)
        rate = float(l.get("rate", 0) or 0)
        gst_rate = float(l.get("gst_rate", 0) or 0)
        taxable = qty * rate
        b = bucket.setdefault(hsn, {"taxable": 0.0, "cgst_amt": 0.0, "sgst_amt": 0.0, "igst_amt": 0.0, "cgst_rate": 0.0, "sgst_rate": 0.0, "igst_rate": 0.0})
        b["taxable"] += taxable
        if is_interstate:
            b["igst_rate"] = gst_rate
            b["igst_amt"] += taxable * gst_rate / 100
        else:
            b["cgst_rate"] = gst_rate / 2
            b["sgst_rate"] = gst_rate / 2
            b["cgst_amt"] += taxable * gst_rate / 200
            b["sgst_amt"] += taxable * gst_rate / 200
    rows = []
    for hsn, v in bucket.items():
        rows.append({"hsn": hsn, **v, "total_tax": v["cgst_amt"] + v["sgst_amt"] + v["igst_amt"]})
    return rows

def _build_doc_pdf(title: str, code: str, party_label: str, party_name: str, date_s: str,
                   lines: List[Dict[str, Any]], totals: Dict[str, float], gst_breakup: Optional[Dict[str, float]] = None,
                   company: Optional[Dict[str, Any]] = None, notes: str = "",
                   tpl: Optional[Dict[str, Any]] = None,
                   party_extra: Optional[Dict[str, Any]] = None,
                   ship_to: Optional[Dict[str, Any]] = None,
                   doc_meta: Optional[Dict[str, Any]] = None,
                   copy_label: str = "ORIGINAL FOR RECIPIENT",
                   bill_from: Optional[Dict[str, Any]] = None,
                   ship_from: Optional[Dict[str, Any]] = None) -> bytes:
    """Tax-invoice/quotation/PO PDF — Denplex Red+Black branded.
    `tpl` overrides which sections are visible (defaults to all-on).
    `party_extra` carries gstin/phone/state/address for the Bill-To party.
    `ship_to` is an optional shipping address dict.
    `doc_meta` may contain {due_date, place_of_supply, payment_mode, time}.
    """
    from reportlab.platypus import Image as RLImage
    tpl = tpl or {}
    # Compact preset turns off heavy blocks by default; user toggles still win.
    _COMPACT_OFF_BY_DEFAULT = {"show_tax_summary", "show_bank_details", "show_upi_qr",
                                "show_signatory_image", "show_description", "show_terms"}
    _style_default = (tpl.get("template_style") or "standard").lower()
    def show(k: str, default: bool = True) -> bool:
        v = tpl.get(k)
        if v is None and _style_default == "compact" and k in _COMPACT_OFF_BY_DEFAULT:
            return False
        return default if v is None else bool(v)
    party_extra = party_extra or {}
    doc_meta = doc_meta or {}
    company = company or {}
    is_interstate = bool((doc_meta or {}).get("is_interstate")) or bool(gst_breakup and gst_breakup.get("igst"))
    _taxable_doc = ((doc_meta or {}).get("invoice_type", "gst") == "gst")   # non_gst / export → no GST blocks

    RED = colors.HexColor("#DC2626")
    BLACK = colors.HexColor("#0A0A0A")
    INK = colors.HexColor("#334155")     # softer body ink (slate-700) — lighter, less heavy than pure black
    GREY = colors.HexColor("#475569")
    LIGHTGREY = colors.HexColor("#F4F6F8")
    BORDER = colors.HexColor("#D7DEE7")

    # ---- Style preset (standard / compact / modern) ----
    style_preset = (tpl.get("template_style") or "standard").lower()
    if style_preset == "compact":
        _margin = 6*mm; _body = 7.5; _title_sz = 16; _company_sz = 12; _border_w = 0.4
        _accent = colors.HexColor("#475569")  # slate grey accent
    elif style_preset == "modern":
        _margin = 14*mm; _body = 9.0; _title_sz = 22; _company_sz = 14; _border_w = 0.0  # no full borders
        _accent = colors.HexColor("#1E293B")  # near-black accent
    else:  # standard — kept deliberately light/subtle
        _margin = 10*mm; _body = 7.5; _title_sz = 13.5; _company_sz = 11; _border_w = 0.4
        _accent = RED

    # Pick fonts: use registered TTF (e.g. DejaVuSans) for proper ₹ rendering
    _FONT = _PDF_FONT_REGULAR
    _FONT_B = _PDF_FONT_BOLD
    _TITLE_FONT = _FONT_B if style_preset != "modern" else (_FONT_B)  # serif if available later

    buf = io.BytesIO()
    page_size = A4
    doc = SimpleDocTemplate(buf, pagesize=page_size,
                            rightMargin=_margin, leftMargin=_margin,
                            topMargin=max(_margin - 2*mm, 6*mm),
                            bottomMargin=_margin)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ttl", parent=styles["Heading1"], fontName=_TITLE_FONT,
                                 fontSize=_title_sz, textColor=BLACK, leading=_title_sz+2,
                                 alignment=(0 if style_preset == "modern" else 1), spaceAfter=0)
    h2_style = ParagraphStyle("h2", parent=styles["Heading2"], fontName=_FONT_B,
                              fontSize=14, textColor=BLACK, leading=16, spaceAfter=0)
    small = ParagraphStyle("sm", parent=styles["Normal"], fontName=_FONT,
                           fontSize=_body, textColor=INK, leading=_body+2.5)
    smallb = ParagraphStyle("smb", parent=small, fontName=_FONT_B)
    tiny = ParagraphStyle("ti", parent=styles["Normal"], fontName=_FONT,
                          fontSize=max(_body-1, 6.5), textColor=GREY, leading=max(_body, 8.5))
    box_label = ParagraphStyle("bl", parent=smallb, fontSize=_body+0.5, textColor=BLACK)
    copy_lbl = ParagraphStyle("cl", parent=small, fontSize=_body-0.5, textColor=GREY, alignment=2)
    flow = []

    # Modern style: render an accent line under title later
    _is_modern = (style_preset == "modern")
    _is_compact = (style_preset == "compact")

    # ---------- Top right "Original for Recipient" ----------
    if show("print_original_duplicate") and copy_label:
        flow.append(Paragraph(copy_label.upper(), copy_lbl))

    # ---------- Title ----------
    flow.append(Paragraph(f"<b>{title}</b>", title_style))
    if _is_modern:
        # Accent line beneath the title
        from reportlab.platypus import HRFlowable
        flow.append(HRFlowable(width="100%", thickness=1.2, color=_accent, spaceBefore=2, spaceAfter=4))
    else:
        flow.append(Spacer(1, 3*mm))

    # ---------- Company header card ----------
    logo_cell = ""
    if show("show_company_logo"):
        logo_path = str(ROOT_DIR / "logo.png")
        try:
            logo_cell = RLImage(logo_path, width=22*mm, height=22*mm)
        except Exception:
            logo_cell = Paragraph("<b>DENPLEX</b>", h2_style)
    company_lines = [Paragraph(f"<font size=11><b>{company.get('company_name','Denplex Engineering Company')}</b></font>", smallb)]
    if show("show_company_udyam") and company.get("company_udyam"):
        company_lines.append(Spacer(1, 1.5*mm))
        company_lines.append(Paragraph(f"<font size=8 color='#475569'>™ UDYAM REGISTRATION NUMBER - <b>{company['company_udyam']}</b></font>", tiny))
        company_lines.append(Spacer(1, 1*mm))
    # Multi-unit address takes priority; fall back to single company_address
    _units = company.get("company_units") or []
    if show("show_company_address"):
        if _units and isinstance(_units, list):
            for u in _units:
                if not isinstance(u, dict): continue
                u_name = (u.get("name") or "").strip()
                u_addr = (u.get("address") or "").strip()
                if not (u_name or u_addr):
                    continue
                if u_name:
                    company_lines.append(Paragraph(f"<b>{u_name}:</b> {u_addr}", tiny))
                else:
                    company_lines.append(Paragraph(u_addr, tiny))
        elif company.get("company_address"):
            company_lines.append(Paragraph(company["company_address"], tiny))
    # Phone + Email + State row
    contact_bits = []
    if show("show_company_phone") and company.get("company_phone"):
        contact_bits.append(f"<b>Phone:</b> {company['company_phone']}")
    if show("show_company_gstin") and company.get("company_gstin"):
        contact_bits.append(f"<b>GSTIN:</b> {company['company_gstin']}")
    if contact_bits:
        company_lines.append(Paragraph(" &nbsp;&nbsp; ".join(contact_bits), tiny))
    contact2 = []
    if show("show_company_email") and company.get("company_email"):
        contact2.append(f"<b>Email:</b> {company['company_email']}")
    if company.get("company_state"):
        contact2.append(f"<b>State:</b> {company['company_state']}")
    if contact2:
        company_lines.append(Paragraph(" &nbsp;&nbsp; ".join(contact2), tiny))

    header_tbl = Table([[logo_cell, company_lines]], colWidths=[28*mm, 162*mm])
    header_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOX", (0, 0), (-1, -1), _border_w, BORDER) if _border_w > 0 else
        ("LINEBELOW", (0, 0), (-1, -1), 0.5, BORDER),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    flow.append(header_tbl)

    # ---------- e-Invoice IRN + signed QR (shown at top when present) ----------
    if doc_meta.get("irn"):
        qr_cell = Paragraph("", tiny)
        try:
            import qrcode as _qr
            _qimg = _qr.make(doc_meta.get("signed_qr") or doc_meta.get("irn"))
            _bio = io.BytesIO(); _qimg.save(_bio, format="PNG"); _bio.seek(0)
            qr_cell = RLImage(_bio, width=20*mm, height=20*mm)
        except Exception:
            qr_cell = Paragraph("", tiny)
        einv_txt = [Paragraph("<b>e-Invoice</b>", smallb), Paragraph(f"IRN: {doc_meta['irn']}", tiny)]
        if doc_meta.get("ack_no"):
            einv_txt.append(Paragraph(f"Ack No: {doc_meta['ack_no']} &nbsp; Ack Date: {doc_meta.get('ack_date','')}", tiny))
        einv_tbl = Table([[einv_txt, qr_cell]], colWidths=[160*mm, 30*mm])
        einv_tbl.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("BOX", (0,0), (-1,-1), 0.5, BORDER),
            ("ALIGN", (1,0), (1,-1), "RIGHT"), ("LEFTPADDING", (0,0), (-1,-1), 6), ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 4), ("BOTTOMPADDING", (0,0), (-1,-1), 4)]))
        flow.append(einv_tbl)

    # ---------- Bill To / Invoice Details (two-column box) ----------
    bill_lines = [Paragraph(f"<b>{party_label}:</b>", box_label),
                  Paragraph(f"<font size=10><b>{party_name}</b></font>", smallb)]
    if party_extra.get("address"): bill_lines.append(Paragraph(party_extra["address"], tiny))
    if party_extra.get("phone"):   bill_lines.append(Paragraph(f"Contact No.: <b>{party_extra['phone']}</b>", tiny))
    if party_extra.get("gstin"):   bill_lines.append(Paragraph(f"GSTIN: <b>{party_extra['gstin']}</b>", tiny))
    if party_extra.get("state"):   bill_lines.append(Paragraph(f"State: <b>{party_extra['state']}</b>", tiny))

    # 2-column meta: main fields on left, PO/purchaser on right
    meta_main = [Paragraph(f"<b>{title.split()[0]} Details:</b>", box_label),
                 Paragraph(f"{title.split()[0]} No.: <b>{code}</b>", smallb),
                 Paragraph(f"Date: <b>{date_s}</b>", smallb)]
    if show("show_due_date") and doc_meta.get("due_date"):
        meta_main.append(Paragraph(f"Due Date: <b>{doc_meta['due_date']}</b>", smallb))
    if show("show_place_of_supply") and doc_meta.get("place_of_supply"):
        meta_main.append(Paragraph(f"Place of Supply: <b>{doc_meta['place_of_supply']}</b>", smallb))
    if doc_meta.get("eway_bill_no"):
        meta_main.append(Paragraph(f"E-Way Bill No: <b>{doc_meta['eway_bill_no']}</b>", smallb))
    for _cfn, _cfv in (doc_meta.get("custom_fields") or {}).items():
        if str(_cfv or "").strip():
            meta_main.append(Paragraph(f"{_cfn}: <b>{_cfv}</b>", smallb))
    meta_po = []
    if show("show_po_meta"):
        if doc_meta.get("po_date"):
            meta_po.append(Paragraph(f"PO Date: <b>{doc_meta['po_date']}</b>", smallb))
        if doc_meta.get("po_number") or doc_meta.get("po_no"):
            meta_po.append(Paragraph(f"PO No: <b>{doc_meta.get('po_number') or doc_meta.get('po_no')}</b>", smallb))
        if doc_meta.get("purchaser_name"):
            meta_po.append(Paragraph(f"Purchaser Name: <b>{doc_meta['purchaser_name']}</b>", smallb))
    # Compose meta_lines: either flat list or nested 2-column table when PO meta is present
    if meta_po:
        meta_inner = Table([[meta_main, meta_po]], colWidths=[46*mm, 47*mm])
        meta_inner.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        meta_lines = [meta_inner]
    else:
        meta_lines = meta_main

    bd_tbl = Table([[bill_lines, meta_lines]], colWidths=[95*mm, 95*mm])
    bd_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOX", (0, 0), (-1, -1), 0.75, BORDER),
        ("LINEAFTER", (0, 0), (0, -1), 0.5, BORDER),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    flow.append(bd_tbl)

    # ---------- Ship To ----------
    if show("show_ship_to") and ship_to and (ship_to.get("name") or ship_to.get("address")):
        ship_block = [Paragraph("<b>Ship To:</b>", box_label),
                      Paragraph(ship_to.get("name", ""), smallb)]
        if ship_to.get("address"):
            ship_block.append(Paragraph(ship_to["address"], tiny))
        ship_tbl = Table([[ship_block]], colWidths=[190*mm])
        ship_tbl.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOX", (0, 0), (-1, -1), 0.75, BORDER),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        flow.append(ship_tbl)

    # ---------- Bill From / Ship From (purchase docs or multi-unit shipments) ----------
    def _addr_block(label, payload):
        block = [Paragraph(f"<b>{label}:</b>", box_label),
                 Paragraph(payload.get("name", ""), smallb)]
        if payload.get("address"):
            block.append(Paragraph(payload["address"], tiny))
        tbl = Table([[block]], colWidths=[190*mm])
        tbl.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOX", (0, 0), (-1, -1), 0.75, BORDER),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return tbl
    if show("show_bill_from", False) and bill_from and (bill_from.get("name") or bill_from.get("address")):
        flow.append(_addr_block("Bill From", bill_from))
    if show("show_ship_from", False) and ship_from and (ship_from.get("name") or ship_from.get("address")):
        flow.append(_addr_block("Ship From", ship_from))

    # ---------- Line items ----------
    show_hsn = show("show_hsn_column")
    # Smart-show columns: explicit toggle wins; else auto-detect data presence
    def _resolve_show(toggle_name, predicate, default_when_no_data=False):
        explicit = tpl.get(toggle_name)
        if explicit is not None:
            return bool(explicit)
        any_data = any(predicate(l) for l in (lines or []))
        return any_data if any_data else default_when_no_data
    show_code = _resolve_show("show_item_code_column", lambda l: l.get("item_code") or l.get("code"))
    show_unit = _resolve_show("show_unit_column", lambda l: l.get("unit"), default_when_no_data=True)
    show_disc = _resolve_show("show_discount_column", lambda l: float(l.get("discount_amount") or 0) > 0 or float(l.get("discount_pct") or 0) > 0)
    show_inline_gst = bool(tpl.get("show_inline_gst_column"))  # default off — GST shown only in Tax Summary
    # Default columns: # | Item name | Item Code | HSN/SAC | Qty | Unit | Price/Unit | Amount
    cols = ["#", "Item name"]
    widths = [7*mm, 56*mm]
    if show_code:
        cols.append("Item Code"); widths.append(22*mm)
    if show_hsn:
        cols.append("HSN/SAC"); widths.append(18*mm)
    cols.append("Qty"); widths.append(14*mm)
    if show_unit:
        cols.append("Unit"); widths.append(14*mm)
    cols.append("Price/Unit"); widths.append(22*mm)
    if show_disc:
        cols.append("Discount"); widths.append(22*mm)
    if show_inline_gst:
        cols.append("GST"); widths.append(22*mm)
    cols.append("Amount"); widths.append(24*mm)

    data = [cols]
    subtotal = 0.0; total_discount = 0.0; total_gst = 0.0; total_amount = 0.0
    for i, l in enumerate(lines or [], 1):
        qty = float(l.get("qty", 0) or 0)
        rate = float(l.get("rate", 0) or 0)
        disc_amt = float(l.get("discount_amount") or 0)
        disc_pct = float(l.get("discount_pct") or 0)
        # If percent provided, derive amount; else, percent for display
        gross = qty * rate
        if disc_amt == 0 and disc_pct:
            disc_amt = gross * disc_pct / 100
        net = max(gross - disc_amt, 0)
        gst_rate = float(l.get("gst_rate", 0) or 0)
        gst_amt = net * gst_rate / 100
        # Amount = net (qty*rate - disc). GST captured separately in Tax Summary.
        # When inline GST column is on, Amount = net + GST (optional behavior).
        amt_display = (net + gst_amt) if tpl.get("show_inline_gst_column") else net
        subtotal += gross
        total_discount += disc_amt
        total_gst += gst_amt
        total_amount += amt_display
        row = [str(i), Paragraph(l.get("description", ""), small)]
        if show_code:
            row.append(str(l.get("item_code") or l.get("code") or ""))
        if show_hsn:
            row.append(str(l.get("hsn") or ""))
        row.append(f"{qty:g}")
        if show_unit:
            row.append(str(l.get("unit") or "Nos"))
        row.append(f"₹ {rate:,.2f}")
        if show_disc:
            row.append(f"₹ {disc_amt:,.2f} ({disc_pct:g}%)" if disc_amt else "—")
        if show_inline_gst:
            row.append(f"₹ {gst_amt:,.2f} ({gst_rate:g}%)" if gst_amt else "—")
        row.append(f"₹ {amt_display:,.2f}")
        data.append(row)
    # Total row
    tot_row = ["", Paragraph("<b>Total</b>", smallb)]
    if show_code: tot_row.append("")
    if show_hsn: tot_row.append("")
    tot_row.append(f"{sum(float(l.get('qty',0) or 0) for l in (lines or [])):g}")
    if show_unit: tot_row.append("")
    tot_row.append("")  # Price/Unit total cell — left blank
    if show_disc: tot_row.append(f"₹ {total_discount:,.2f}" if total_discount else "")
    if show_inline_gst: tot_row.append(f"₹ {total_gst:,.2f}")
    tot_row.append(f"₹ {total_amount:,.2f}")
    data.append(tot_row)

    tbl = Table(data, colWidths=widths, repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), LIGHTGREY),
        ("FONTNAME", (0, 0), (-1, 0), _PDF_FONT_BOLD),
        ("FONTSIZE", (0, 0), (-1, -1), 8.0),
        ("TEXTCOLOR", (0, 0), (-1, -1), INK),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (2 if show_hsn else 2, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.4, BORDER),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        # Total row
        ("BACKGROUND", (0, -1), (-1, -1), LIGHTGREY),
        ("FONTNAME", (0, -1), (-1, -1), _PDF_FONT_BOLD),
        ("LINEABOVE", (0, -1), (-1, -1), 0.8, BLACK),
    ]
    tbl.setStyle(TableStyle(style))
    flow.append(tbl)

    # ---------- Tax Summary + Totals sidebar ----------
    bottom_left_blocks = []
    if show("show_tax_summary") and (lines or []) and _taxable_doc:
        rows = _hsn_tax_summary(lines or [], is_interstate)
        if rows:
            if is_interstate:
                hdr = ["HSN/SAC", "Taxable Amount", "IGST Rate", "IGST Amount", "Total Tax Amount"]
                t_data = [hdr]
                ct = 0; gt_taxable = 0; gt_total_tax = 0
                for r in rows:
                    gt_taxable += r["taxable"]; ct += r["igst_amt"]; gt_total_tax += r["total_tax"]
                    t_data.append([r["hsn"] or "—", f"₹ {r['taxable']:,.2f}", f"{r['igst_rate']:g}%", f"₹ {r['igst_amt']:,.2f}", f"₹ {r['total_tax']:,.2f}"])
                t_data.append(["Total", f"₹ {gt_taxable:,.2f}", "", f"₹ {ct:,.2f}", f"₹ {gt_total_tax:,.2f}"])
                widths_ts = [22*mm, 28*mm, 18*mm, 28*mm, 30*mm]
            else:
                # Nested header: HSN/SAC | Taxable | CGST(Rate|Amt) | SGST(Rate|Amt) | Total Tax
                hdr1 = ["HSN/SAC", "Taxable Amount (₹)", "CGST", "", "SGST", "", "Total Tax (₹)"]
                hdr2 = ["", "", "Rate (%)", "Amt (₹)", "Rate (%)", "Amt (₹)", ""]
                t_data = [hdr1, hdr2]
                gt_taxable=0; gt_c=0; gt_s=0; gt_total_tax=0
                for r in rows:
                    gt_taxable += r["taxable"]; gt_c += r["cgst_amt"]; gt_s += r["sgst_amt"]; gt_total_tax += r["total_tax"]
                    t_data.append([r["hsn"] or "—", f"{r['taxable']:,.2f}", f"{r['cgst_rate']:g}", f"{r['cgst_amt']:,.2f}", f"{r['sgst_rate']:g}", f"{r['sgst_amt']:,.2f}", f"{r['total_tax']:,.2f}"])
                t_data.append(["Total", f"{gt_taxable:,.2f}", "", f"{gt_c:,.2f}", "", f"{gt_s:,.2f}", f"{gt_total_tax:,.2f}"])
                widths_ts = [18*mm, 24*mm, 12*mm, 18*mm, 12*mm, 18*mm, 20*mm]
            ts = Table(t_data, colWidths=widths_ts)
            ts_style = [
                ("BACKGROUND", (0,0), (-1,1), LIGHTGREY) if not is_interstate else ("BACKGROUND", (0,0), (-1,0), LIGHTGREY),
                ("FONTNAME", (0,0), (-1, 1 if not is_interstate else 0), _PDF_FONT_BOLD),
                ("FONTNAME", (0,0), (-1,-1), _PDF_FONT_REGULAR),
                ("FONTSIZE", (0,0), (-1,-1), 7.5),
                ("ALIGN", (1,0), (-1,-1), "CENTER"),
                ("ALIGN", (0,0), (0,-1), "LEFT"),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                ("GRID", (0,0), (-1,-1), 0.4, BORDER),
                ("BACKGROUND", (0,-1), (-1,-1), LIGHTGREY),
                ("FONTNAME", (0,-1), (-1,-1), _PDF_FONT_BOLD),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ]
            # For intrastate, span the nested headers
            if not is_interstate:
                ts_style += [
                    ("SPAN", (0,0), (0,1)),    # HSN/SAC
                    ("SPAN", (1,0), (1,1)),    # Taxable Amount
                    ("SPAN", (2,0), (3,0)),    # CGST (Rate+Amt)
                    ("SPAN", (4,0), (5,0)),    # SGST (Rate+Amt)
                    ("SPAN", (6,0), (6,1)),    # Total Tax
                ]
            ts.setStyle(TableStyle(ts_style))
            bottom_left_blocks.append(Paragraph("<b>Tax Summary:</b>", smallb))
            bottom_left_blocks.append(ts)

    # Totals sidebar — single "Tax (X%)" line by default; can split via toggle
    sidebar = []
    if show("show_totals_sidebar"):
        sd = []
        # Sub Total = sum of nets — same as total_amount when inline GST off
        st_value = (subtotal - total_discount) if not tpl.get("show_inline_gst_column") else (subtotal - total_discount + total_gst)
        sd.append(["Sub Total", f"₹ {st_value:,.2f}"])
        if total_discount:
            sd.append(["Discount", f"₹ {total_discount:,.2f}"])
        _split_tax = bool(tpl.get("show_split_tax_in_sidebar"))  # default off (combined Tax line)
        if not _taxable_doc:
            pass  # non-GST / export invoice: no tax line in totals
        elif _split_tax:
            if is_interstate:
                sd.append(["IGST", f"₹ {(gst_breakup or {}).get('igst', total_gst):,.2f}"])
            else:
                cg = (gst_breakup or {}).get("cgst", total_gst/2)
                sg = (gst_breakup or {}).get("sgst", total_gst/2)
                sd.append(["CGST", f"₹ {cg:,.2f}"])
                sd.append(["SGST", f"₹ {sg:,.2f}"])
        else:
            # Combined "Tax (X%)" — derive avg rate from line items
            net_taxable = sum(float(l.get("qty",0) or 0) * float(l.get("rate",0) or 0) for l in (lines or []))
            if net_taxable > 0:
                avg_rate = (total_gst / net_taxable) * 100
                # Round to common GST rates if close enough
                for cand in (5, 12, 18, 28):
                    if abs(avg_rate - cand) < 0.5:
                        avg_rate = cand; break
                label = f"Tax ({avg_rate:g}%)"
            else:
                label = "Tax"
            sd.append([label, f"₹ {total_gst:,.2f}"])
        # Additional charges (lump sum) + TDS (deducted) + TCS (collected)
        _charges = totals.get("extra_charges") or []
        _ct = float(totals.get("charges_total") or 0)
        if _charges:
            for _c in _charges:
                if float(_c.get("amount", 0) or 0):
                    sd.append([str(_c.get("name") or "Charges"), f"₹ {float(_c.get('amount', 0) or 0):,.2f}"])
        elif _ct:
            sd.append(["Additional Charges", f"₹ {_ct:,.2f}"])
        _tds = float(totals.get("tds") or 0)
        if _tds:
            _tr = float(totals.get("tds_rate") or 0)
            sd.append([f"TDS{f' ({_tr:g}%)' if _tr else ''}", f"- ₹ {_tds:,.2f}"])
        _tcs = float(totals.get("tcs") or 0)
        if _tcs:
            _cr = float(totals.get("tcs_rate") or 0)
            sd.append([f"TCS{f' ({_cr:g}%)' if _cr else ''}", f"+ ₹ {_tcs:,.2f}"])
        # Round-off: round total to whole rupees.
        # The grand total = net + tax; total_amount may be net-only when inline GST is off.
        grand_total_raw = (total_amount + total_gst) if not tpl.get("show_inline_gst_column") else total_amount
        raw_total = float(totals.get('total') or grand_total_raw)
        rounded_total = round(raw_total)
        round_off = rounded_total - raw_total
        if abs(round_off) > 0.005:
            sd.append(["Round Off", f"{'+' if round_off > 0 else '-'} ₹ {abs(round_off):,.2f}"])
        sd.append(["Total", f"₹ {rounded_total:,.2f}"])
        # Stash for amount-in-words
        totals["__rounded_total"] = rounded_total
        side = Table(sd, colWidths=[28*mm, 30*mm])
        side.setStyle(TableStyle([
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("ALIGN", (1,0), (1,-1), "RIGHT"),
            ("GRID", (0,0), (-1,-1), 0.4, BORDER),
            ("FONTNAME", (0,-1), (-1,-1), _PDF_FONT_BOLD),
            ("BACKGROUND", (0,-1), (-1,-1), RED),
            ("TEXTCOLOR", (0,-1), (-1,-1), colors.white),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
        ]))
        sidebar.append(side)
    # Amount in words (under totals)
    if show("show_amount_in_words"):
        # Prefer the rounded total if available (set above by sidebar logic)
        words_value = float(totals.get("__rounded_total") or totals.get("total", total_amount) or 0)
        words = _amount_in_words(round(words_value), tpl.get("amount_in_words_locale", "en_IN"))
        if words:
            sidebar.append(Spacer(1, 2*mm))
            sidebar.append(Paragraph("<b>Invoice Amount In Words:</b>", tiny))
            sidebar.append(Paragraph(words, smallb))

    # Lay out [tax summary | sidebar]
    if bottom_left_blocks or sidebar:
        left_cell = bottom_left_blocks or [Spacer(1, 1)]
        right_cell = sidebar or [Spacer(1, 1)]
        ts_tbl = Table([[left_cell, right_cell]], colWidths=[125*mm, 65*mm])
        ts_tbl.setStyle(TableStyle([
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 0),
            ("RIGHTPADDING", (0,0), (-1,-1), 0),
            ("TOPPADDING", (0,0), (-1,-1), 4),
        ]))
        flow.append(ts_tbl)

    # ---------- Payment Mode ----------
    if show("show_payment_mode") and doc_meta.get("payment_mode"):
        flow.append(Spacer(1, 2*mm))
        pm = Table([[Paragraph("<b>Payment Mode:</b>", smallb), Paragraph(doc_meta["payment_mode"], small)]],
                   colWidths=[35*mm, 155*mm])
        pm.setStyle(TableStyle([("BOX",(0,0),(-1,-1),0.5,BORDER),("INNERGRID",(0,0),(-1,-1),0.4,BORDER),
                                ("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),
                                ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
        flow.append(pm)

    # ---------- Description / Terms ----------
    if show("show_description") or show("show_terms"):
        desc_text = company.get("invoice_description", "") if show("show_description") else ""
        terms_text = company.get("invoice_terms", "") if show("show_terms") else ""
        if notes:
            desc_text = (desc_text + "\n" + notes).strip() if desc_text else notes
        # Build single-cell or two-cell layout depending on what content exists
        has_desc = bool(desc_text)
        has_terms = bool(terms_text)
        if has_desc and has_terms:
            cells_l = [Paragraph("<b>Description:</b>", smallb), Paragraph(desc_text.replace("\n", "<br/>"), small)]
            cells_r = [Paragraph("<b>Terms &amp; Conditions:</b>", smallb), Paragraph(terms_text.replace("\n", "<br/>"), small)]
        elif has_terms:
            cells_l = [Paragraph("<b>Terms &amp; Conditions:</b>", smallb), Paragraph(terms_text.replace("\n", "<br/>"), small)]
            cells_r = []
        elif has_desc:
            cells_l = [Paragraph("<b>Description:</b>", smallb), Paragraph(desc_text.replace("\n", "<br/>"), small)]
            cells_r = []
        else:
            cells_l = None
            cells_r = None
        if cells_l is not None:
            if cells_r:
                dt_tbl = Table([[cells_l, cells_r]], colWidths=[95*mm, 95*mm])
                dt_tbl.setStyle(TableStyle([
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ("BOX", (0,0), (-1,-1), 0.5, BORDER),
                    ("LINEAFTER", (0,0), (0,-1), 0.4, BORDER),
                    ("LEFTPADDING", (0,0), (-1,-1), 6),
                    ("RIGHTPADDING", (0,0), (-1,-1), 6),
                    ("TOPPADDING", (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                ]))
            else:
                dt_tbl = Table([[cells_l]], colWidths=[190*mm])
                dt_tbl.setStyle(TableStyle([
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ("BOX", (0,0), (-1,-1), 0.5, BORDER),
                    ("LEFTPADDING", (0,0), (-1,-1), 6),
                    ("RIGHTPADDING", (0,0), (-1,-1), 6),
                    ("TOPPADDING", (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                ]))
            flow.append(Spacer(1, 2*mm))
            flow.append(dt_tbl)

    # ---------- Bank details + QR + Signatory (optional page break) ----------
    has_bank = show("show_bank_details") and any(company.get(k) for k in ("bank_name","bank_account_no","bank_ifsc","upi_id"))
    has_sig = show("show_signatory_image")
    if has_bank or has_sig:
        # Bank/signatory flow inline right after Terms — single-page invoice like the
        # reference layout. (Previous page-break behaviour removed per user request.)
        # Bank cell with optional QR
        bank_block = [Paragraph("<b>Bank Details:</b>", smallb)]
        # QR
        qr_img = None
        if show("show_upi_qr") and company.get("upi_id"):
            qr_png = _upi_qr_png(company.get("upi_id", ""), company.get("company_name", "Denplex"),
                                 amount=float(totals.get("total", 0) or 0), note=code)
            if qr_png:
                qr_img = RLImage(io.BytesIO(qr_png), width=22*mm, height=22*mm)
        bank_lines = []
        if company.get("bank_name"): bank_lines.append(Paragraph(f"Bank Name: <b>{company['bank_name']}</b>", small))
        if company.get("bank_account_no"): bank_lines.append(Paragraph(f"Bank Account No.: <b>{company['bank_account_no']}</b>", small))
        if company.get("bank_ifsc"): bank_lines.append(Paragraph(f"Bank IFSC code: <b>{company['bank_ifsc']}</b>", small))
        if company.get("bank_branch"): bank_lines.append(Paragraph(f"Branch: <b>{company['bank_branch']}</b>", small))
        if company.get("upi_id"): bank_lines.append(Paragraph(f"UPI: <b>{company['upi_id']}</b>", small))
        if qr_img:
            left_bank = Table([[qr_img, bank_lines]], colWidths=[24*mm, 71*mm])
            left_bank.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),0),
                                           ("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),0),
                                           ("BOTTOMPADDING",(0,0),(-1,-1),0)]))
            bank_block.append(left_bank)
        else:
            bank_block.extend(bank_lines)

        # Signatory cell
        sig_block = [Paragraph(f"<b>For: {company.get('company_name','Denplex Engineering Company')}</b>", smallb)]
        sig_b64 = company.get("signatory_image_b64") or ""
        if has_sig and sig_b64:
            try:
                b = sig_b64.split(",",1)[1] if sig_b64.startswith("data:") else sig_b64
                raw = base64.b64decode(b)
                sig_img = RLImage(io.BytesIO(raw), width=40*mm, height=18*mm, kind="proportional")
                sig_block.append(Spacer(1, 1*mm))
                sig_block.append(sig_img)
            except Exception:
                sig_block.append(Spacer(1, 12*mm))
        else:
            sig_block.append(Spacer(1, 12*mm))
        sig_block.append(Paragraph(f"<font color='#475569'>{company.get('signatory_label','Authorised Signatory')}</font>", tiny))

        bs_tbl = Table([[bank_block, sig_block]], colWidths=[95*mm, 95*mm])
        bs_tbl.setStyle(TableStyle([
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("BOX", (0,0), (-1,-1), 0.5, BORDER),
            ("LINEAFTER", (0,0), (0,-1), 0.4, BORDER),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("ALIGN", (1,0), (1,-1), "CENTER"),
        ]))
        flow.append(Spacer(1, 2*mm))
        flow.append(bs_tbl)

    doc.build(flow)
    return buf.getvalue()

async def _resolve_doc(coll, doc_id: str, party_id_key: str, party_name_key: str):
    doc = await coll.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Not found")
    return doc

async def _tpl_for(doc_type: str) -> Dict[str, Any]:
    s = await get_setting("invoice_template") or {}
    defaults = InvoiceTemplateIn().model_dump()
    base = {**defaults, **(s.get("default") or {})}
    return {**base, **((s.get(doc_type) or {}))}

@api.get("/invoices/{iid}/pdf")
async def invoice_pdf(iid: str, copy: Optional[str] = "ORIGINAL FOR RECIPIENT", user=Depends(get_current_user)):
    inv = await db.invoices.find_one({"id": iid}, {"_id": 0})
    if not inv: raise HTTPException(404, "Not found")
    company = await get_setting("integrations")
    tpl = await _tpl_for("invoice")
    party_extra = {}
    if inv.get("customer_id"):
        c = await db.customers.find_one({"id": inv["customer_id"]}, {"_id": 0}) or {}
        party_extra = {"address": c.get("address",""), "phone": c.get("phone",""), "gstin": c.get("gstin",""), "state": c.get("state","")}
    # Build optional Ship To: explicit on invoice wins; else fall back to customer address
    ship_to = None
    if inv.get("ship_to_address") or inv.get("ship_to_name"):
        ship_to = {"name": inv.get("ship_to_name") or inv.get("customer_name", ""),
                   "address": inv.get("ship_to_address", ""),
                   "gstin": inv.get("ship_to_gstin", "")}
    elif party_extra.get("address"):
        ship_to = {"name": inv.get("customer_name", ""), "address": party_extra.get("address", ""),
                   "gstin": party_extra.get("gstin", "")}
    bill_from = None
    if inv.get("bill_from_name") or inv.get("bill_from_address"):
        bill_from = {"name": inv.get("bill_from_name", ""), "address": inv.get("bill_from_address", "")}
    ship_from = None
    if inv.get("ship_from_name") or inv.get("ship_from_address"):
        ship_from = {"name": inv.get("ship_from_name", ""), "address": inv.get("ship_from_address", "")}
    # Auto-enable bill_from / ship_from blocks when invoice has those fields
    if bill_from: tpl = {**tpl, "show_bill_from": True}
    if ship_from: tpl = {**tpl, "show_ship_from": True}
    doc_meta = {
        "due_date": str(inv.get("due_date",""))[:10],
        "place_of_supply": inv.get("place_of_supply",""),
        "payment_mode": inv.get("payment_mode",""),
        "po_number": inv.get("po_number",""),
        "po_date": inv.get("po_date",""),
        "purchaser_name": inv.get("purchaser_name",""),
        "eway_bill_no": inv.get("eway_bill_no",""),
        "custom_fields": inv.get("custom_fields", {}),
        "is_interstate": bool(inv.get("is_interstate")),
        "invoice_type": inv.get("invoice_type", "gst"),
        "irn": inv.get("irn",""), "ack_no": inv.get("ack_no",""),
        "ack_date": inv.get("ack_date",""), "signed_qr": inv.get("signed_qr",""),
    }
    _itype = inv.get("invoice_type", "gst")
    _doc_title = "Export Invoice" if _itype == "export" else ("Bill of Supply" if _itype == "non_gst" else "Tax Invoice")
    pdf = _build_doc_pdf(_doc_title, inv.get("code", ""), "Bill To", inv.get("customer_name", ""), str(inv.get("date", ""))[:10],
                         inv.get("lines", []),
                         {"subtotal": inv.get("subtotal", 0), "total": inv.get("total", 0), "gst_total": inv.get("cgst",0)+inv.get("sgst",0)+inv.get("igst",0),
                          "charges_total": inv.get("charges_total", 0), "extra_charges": inv.get("extra_charges", []),
                          "tds": inv.get("tds", 0), "tds_rate": inv.get("tds_rate", 0), "tcs": inv.get("tcs", 0), "tcs_rate": inv.get("tcs_rate", 0)},
                         gst_breakup={"cgst": inv.get("cgst", 0), "sgst": inv.get("sgst", 0), "igst": inv.get("igst", 0)},
                         company=company, notes=(inv.get("terms_text") or inv.get("notes", "")), tpl=tpl,
                         party_extra=party_extra, ship_to=ship_to, doc_meta=doc_meta,
                         copy_label=copy or "ORIGINAL FOR RECIPIENT",
                         bill_from=bill_from, ship_from=ship_from)
    return Response(content=pdf, media_type="application/pdf",
                    headers={"Content-Disposition": f'inline; filename="{inv.get("code","invoice")}.pdf"'})

@api.get("/quotations/{qid}/pdf")
async def quote_pdf(qid: str, user=Depends(get_current_user)):
    q = await db.quotations.find_one({"id": qid}, {"_id": 0})
    if not q: raise HTTPException(404, "Not found")
    company = await get_setting("integrations")
    tpl = await _tpl_for("quotation")
    party_extra = {}
    if q.get("customer_id"):
        c = await db.customers.find_one({"id": q["customer_id"]}, {"_id": 0}) or {}
        party_extra = {"address": c.get("address",""), "phone": c.get("phone",""), "gstin": c.get("gstin",""), "state": c.get("state","")}
    pdf = _build_doc_pdf("Quotation", q.get("code", ""), "To", q.get("customer_name", ""), str(q.get("date", ""))[:10],
                         q.get("lines", []), {"subtotal": q.get("subtotal", 0), "gst_total": q.get("gst_total", 0), "total": q.get("total", 0)},
                         company=company, notes=q.get("notes", ""), tpl=tpl, party_extra=party_extra,
                         copy_label="")
    return Response(content=pdf, media_type="application/pdf",
                    headers={"Content-Disposition": f'inline; filename="{q.get("code","quote")}.pdf"'})

@api.get("/purchase-orders/{pid}/pdf")
async def po_pdf(pid: str, user=Depends(get_current_user)):
    p = await db.purchase_orders.find_one({"id": pid}, {"_id": 0})
    if not p: raise HTTPException(404, "Not found")
    company = await get_setting("integrations")
    tpl = await _tpl_for("purchase_order")
    party_extra = {}
    if p.get("supplier_id"):
        c = await db.suppliers.find_one({"id": p["supplier_id"]}, {"_id": 0}) or {}
        party_extra = {"address": c.get("address",""), "phone": c.get("phone",""), "gstin": c.get("gstin",""), "state": c.get("state","")}
    pdf = _build_doc_pdf("Purchase Order", p.get("code", ""), "Supplier", p.get("supplier_name", ""), str(p.get("date", ""))[:10],
                         p.get("lines", []), {"subtotal": p.get("subtotal", 0), "gst_total": p.get("gst_total", 0), "total": p.get("total", 0)},
                         company=company, notes=p.get("notes", ""), tpl=tpl, party_extra=party_extra,
                         copy_label="")
    return Response(content=pdf, media_type="application/pdf",
                    headers={"Content-Disposition": f'inline; filename="{p.get("code","po")}.pdf"'})

# ---------- New document types (imported from Vyapar) ----------
def _party_extra_from(c: Dict[str, Any]) -> Dict[str, Any]:
    return {"address": (c or {}).get("address",""), "phone": (c or {}).get("phone",""),
            "gstin": (c or {}).get("gstin",""), "state": (c or {}).get("state","")}

async def _generic_doc_pdf(coll, did: str, title: str, party_label: str, party_field: str, party_coll, doc_type: str, copy_label: str = ""):
    d = await coll.find_one({"id": did}, {"_id": 0})
    if not d: raise HTTPException(404, "Not found")
    company = await get_setting("integrations")
    tpl = await _tpl_for(doc_type)
    party_id_key = f"{party_field}_id"
    party_name_key = f"{party_field}_name"
    party_extra = _party_extra_from(await party_coll.find_one({"id": d.get(party_id_key)}, {"_id": 0}) if d.get(party_id_key) else {})
    totals = {"subtotal": d.get("subtotal", 0), "gst_total": d.get("gst_total", 0), "total": d.get("total", 0)}
    pdf = _build_doc_pdf(title, d.get("code",""), party_label, d.get(party_name_key,""), str(d.get("date",""))[:10],
                         d.get("lines", []), totals,
                         gst_breakup={"cgst": d.get("cgst",0), "sgst": d.get("sgst",0), "igst": d.get("igst",0)},
                         company=company, notes=d.get("notes",""), tpl=tpl,
                         party_extra=party_extra,
                         doc_meta={"due_date": str(d.get("due_date",""))[:10],
                                   "place_of_supply": d.get("place_of_supply",""),
                                   "payment_mode": d.get("payment_mode",""),
                                   "is_interstate": bool(d.get("is_interstate"))},
                         copy_label=copy_label)
    return Response(content=pdf, media_type="application/pdf",
                    headers={"Content-Disposition": f'inline; filename="{d.get("code",doc_type)}.pdf"'})

def _doc_totals_with_gst(lines, interstate, round_off=0, tds=0, extra_charges=None, tcs=0):
    t = compute_invoice_totals(lines, interstate, round_off, tds, extra_charges, tcs)
    t["gst_total"] = round(t.get("cgst", 0) + t.get("sgst", 0) + t.get("igst", 0), 2)
    return t

@api.post("/sale-orders")
async def create_sale_order(so: SaleOrder, user=Depends(get_current_user)):
    doc = so.model_dump()
    doc["code"] = (so.code or "").strip() or await gen_code("SO", "sale_order")
    doc.update(_doc_totals_with_gst(doc["lines"], doc.get("is_interstate", False), doc.get("round_off", 0)))
    await db.sale_orders.insert_one(doc)
    return serialize(doc)

@api.post("/vendor-bills")
async def create_vendor_bill(vb: VendorBill, user=Depends(get_current_user)):
    doc = vb.model_dump()
    doc["code"] = (vb.code or "").strip() or await gen_code("PB", "vendor_bill")
    doc.update(_doc_totals_with_gst(doc["lines"], doc.get("is_interstate", False), doc.get("round_off", 0), doc.get("tds", 0), doc.get("extra_charges"), doc.get("tcs", 0)))
    await db.vendor_bills.insert_one(doc)
    return serialize(doc)

@api.get("/vendor-bills")
async def list_vendor_bills(user=Depends(get_current_user)):
    return await list_collection(db.vendor_bills, sort_key="date")

@api.get("/vendor-bills/settled-summary")
async def vendor_bills_settled_summary(user=Depends(get_current_user)):
    """Bulk vendor_bill id -> settled amount map (reuses _settled_per_bill with no filter) — powers
    the Purchase Bills report page's Paid/Unpaid/Overdue/Total summary cards + per-row Balance Due
    column without an N+1 call per bill."""
    return await _settled_per_bill()

@api.delete("/vendor-bills/{did}")
async def del_vendor_bill(did: str, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    """Was missing — vendor_bills previously had no delete route at all (only list/pdf/create),
    unlike every other document type which already gets recycle-bin-backed delete. Mirrors
    del_invoice's pattern exactly."""
    doc = await db.vendor_bills.find_one({"id": did}, {"_id": 0})
    await _recycle("vendor_bills", "Purchase Bill", doc, user)
    await db.vendor_bills.delete_one({"id": did})
    return {"ok": True}

@api.get("/vendor-bills/{did}/pdf")
async def vendor_bill_pdf(did: str, user=Depends(get_current_user)):
    return await _generic_doc_pdf(db.vendor_bills, did, "Purchase Bill", "Supplier", "supplier", db.suppliers, "vendor_bill")

@api.get("/sale-orders")
async def list_sale_orders(user=Depends(get_current_user)):
    return await list_collection(db.sale_orders, sort_key="date")

@api.get("/sale-orders/{did}/pdf")
async def sale_order_pdf(did: str, user=Depends(get_current_user)):
    return await _generic_doc_pdf(db.sale_orders, did, "Sale Order", "Customer", "customer", db.customers, "sale_order")

@api.get("/delivery-challans")
async def list_delivery_challans(user=Depends(get_current_user)):
    return await list_collection(db.delivery_challans, sort_key="date")

@api.get("/delivery-challans/{did}/pdf")
async def delivery_challan_pdf(did: str, user=Depends(get_current_user)):
    return await _generic_doc_pdf(db.delivery_challans, did, "Delivery Challan", "Ship To", "customer", db.customers, "delivery_challan")

@api.get("/job-work-out")
async def list_job_work_out(user=Depends(get_current_user)):
    return await list_collection(db.job_work_out, sort_key="date")

@api.get("/job-work-out/{did}/pdf")
async def job_work_out_pdf(did: str, user=Depends(get_current_user)):
    return await _generic_doc_pdf(db.job_work_out, did, "Job Work Out Challan", "Job Worker", "customer", db.customers, "job_work_out")

@api.get("/credit-notes")
async def list_credit_notes(user=Depends(get_current_user)):
    return await list_collection(db.credit_notes, sort_key="date")

@api.get("/credit-notes/{did}/pdf")
async def credit_note_pdf(did: str, user=Depends(get_current_user)):
    return await _generic_doc_pdf(db.credit_notes, did, "Credit Note", "Bill To", "customer", db.customers, "credit_note")

# ---------- Indiamart pull leads ----------
@api.post("/integrations/indiamart/sync")
async def indiamart_sync(user=Depends(require_roles("admin", "manager", "sales"))):
    cfg = await get_setting("integrations")
    key = cfg.get("indiamart_crm_key")
    if not key:
        raise HTTPException(400, "Indiamart CRM key not set in Settings")
    # Indiamart Lead Manager: pulls last 7 days by default
    url = f"https://mapi.indiamart.com/wservce/crm/crmListing/v2/?glusr_crm_key={key}"
    try:
        async with httpx.AsyncClient(timeout=30) as cli:
            r = await cli.get(url)
        data = r.json()
    except Exception as e:
        raise HTTPException(502, f"Indiamart fetch failed: {e}")
    if data.get("CODE") not in (200, "200"):
        raise HTTPException(400, f"Indiamart: {data.get('MESSAGE', 'error')}")
    rows = data.get("RESPONSE", []) or []
    added = 0
    for row in rows:
        ext_id = str(row.get("UNIQUE_QUERY_ID") or row.get("QUERY_ID") or "")
        if ext_id and await db.leads.find_one({"external_id": ext_id}):
            continue
        lead = {
            "id": new_id(),
            "external_id": ext_id,
            "name": row.get("SENDER_NAME", "Unknown"),
            "company": row.get("SENDER_COMPANY", ""),
            "phone": row.get("SENDER_MOBILE", "") or row.get("SENDER_PHONE", ""),
            "email": row.get("SENDER_EMAIL", ""),
            "source": "indiamart",
            "requirement": row.get("QUERY_PRODUCT_NAME", "") or row.get("QUERY_MESSAGE", ""),
            "status": "new",
            "notes": f"City: {row.get('SENDER_CITY','')}, State: {row.get('SENDER_STATE','')}",
            "created_at": now_iso(),
        }
        await db.leads.insert_one(lead)
        added += 1
    await write_audit(user["name"], "indiamart_sync", "leads", None, {"added": added, "fetched": len(rows)})
    return {"added": added, "fetched": len(rows)}

# ---------- TradeIndia webhook ----------
class TradeIndiaLead(BaseModel):
    name: Optional[str] = ""
    company: Optional[str] = ""
    phone: Optional[str] = ""
    mobile: Optional[str] = ""
    email: Optional[str] = ""
    message: Optional[str] = ""
    product: Optional[str] = ""
    city: Optional[str] = ""
    state: Optional[str] = ""
    external_id: Optional[str] = ""

@api.post("/integrations/tradeindia/webhook")
async def tradeindia_webhook(payload: TradeIndiaLead, token: str = Query(...)):
    cfg = await get_setting("integrations")
    expected = cfg.get("tradeindia_webhook_secret") or ""
    if not expected or token != expected:
        raise HTTPException(401, "Invalid webhook token")
    name = payload.name or "Unknown"
    phone = payload.mobile or payload.phone or ""
    if payload.external_id and await db.leads.find_one({"external_id": payload.external_id}):
        return {"ok": True, "skipped": "duplicate"}
    lead = {
        "id": new_id(),
        "external_id": payload.external_id or "",
        "name": name, "company": payload.company or "",
        "phone": phone, "email": payload.email or "",
        "source": "tradeindia",
        "requirement": payload.product or payload.message or "",
        "status": "new",
        "notes": f"City: {payload.city or ''}, State: {payload.state or ''}",
        "created_at": now_iso(),
    }
    await db.leads.insert_one(lead)
    return {"ok": True}

# ---------- GSTR-1 / GSTR-3B CSV exports ----------
def _csv_response(rows: List[List[Any]], filename: str) -> Response:
    buf = io.StringIO()
    w = csv.writer(buf)
    for r in rows:
        w.writerow(r)
    return Response(content=buf.getvalue(), media_type="text/csv",
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'})

@api.get("/accounting/gstr1.csv")
async def gstr1_csv(period_from: Optional[str] = None, period_to: Optional[str] = None,
                    user=Depends(require_roles("admin", "accountant", "ca"))):
    q = {}
    if period_from or period_to:
        q["date"] = {}
        if period_from: q["date"]["$gte"] = period_from
        if period_to: q["date"]["$lte"] = period_to
    invs = await db.invoices.find(q, {"_id": 0}).to_list(5000)
    rows: List[List[Any]] = [["GSTIN/UIN of Recipient", "Receiver Name", "Invoice Number", "Invoice Date",
                              "Invoice Value", "Place of Supply", "Reverse Charge", "Applicable % of Tax Rate",
                              "Invoice Type", "E-Commerce GSTIN", "Rate", "Taxable Value", "Cess Amount"]]
    for i in invs:
        # group lines by rate
        by_rate: Dict[float, float] = {}
        for ln in i.get("lines", []):
            r = float(ln.get("gst_rate", 0))
            amt = float(ln.get("qty", 0)) * float(ln.get("rate", 0))
            by_rate[r] = by_rate.get(r, 0) + amt
        for rate, taxable in by_rate.items():
            rows.append([
                i.get("customer_gstin", ""), i.get("customer_name", ""), i.get("code", ""),
                str(i.get("date", ""))[:10], round(i.get("total", 0), 2), i.get("place_of_supply", ""),
                "N", "", "Regular", "", rate, round(taxable, 2), 0,
            ])
    return _csv_response(rows, f"GSTR1_{period_from or 'all'}_{period_to or 'all'}.csv")

@api.get("/accounting/gstr3b.csv")
async def gstr3b_csv(period_from: Optional[str] = None, period_to: Optional[str] = None,
                     user=Depends(require_roles("admin", "accountant", "ca"))):
    q = {}
    if period_from or period_to:
        q["date"] = {}
        if period_from: q["date"]["$gte"] = period_from
        if period_to: q["date"]["$lte"] = period_to
    invs = await db.invoices.find(q, {"_id": 0}).to_list(5000)
    exq = q.copy()
    expenses = await db.expenses.find(exq, {"_id": 0}).to_list(5000)
    out_tax = sum(i.get("subtotal", 0) for i in invs)
    out_cgst = sum(i.get("cgst", 0) for i in invs)
    out_sgst = sum(i.get("sgst", 0) for i in invs)
    out_igst = sum(i.get("igst", 0) for i in invs)
    in_tax = sum(e.get("amount", 0) for e in expenses)
    in_gst = sum(e.get("gst_amount", 0) for e in expenses)
    rows = [
        ["GSTR-3B Summary", f"{period_from or ''} to {period_to or ''}"],
        [],
        ["3.1 Outward Supplies (Taxable)"],
        ["Nature", "Taxable Value", "IGST", "CGST", "SGST", "Cess"],
        ["(a) Outward taxable supplies (other than zero rated, nil rated and exempted)",
         round(out_tax, 2), round(out_igst, 2), round(out_cgst, 2), round(out_sgst, 2), 0],
        [],
        ["4. Eligible ITC"],
        ["Source", "IGST", "CGST", "SGST", "Cess"],
        ["(A)(5) All other ITC", round(in_gst/2, 2) if not any(i.get("is_interstate") for i in invs) else 0,
         round(in_gst/2, 2), round(in_gst/2, 2), 0],
        [],
        ["Net GST Payable", round(out_cgst + out_sgst + out_igst - in_gst, 2)],
        ["Output taxable total", round(out_tax, 2)],
        ["Input taxable total", round(in_tax, 2)],
    ]
    return _csv_response(rows, f"GSTR3B_{period_from or 'all'}_{period_to or 'all'}.csv")

# ---------- 2FA TOTP ----------
class TotpVerifyIn(BaseModel):
    code: str

@api.post("/auth/2fa/setup")
async def totp_setup(user=Depends(get_current_user)):
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    if u.get("totp_enabled"):
        raise HTTPException(400, "2FA already enabled")
    secret = pyotp.random_base32()
    await db.users.update_one({"id": user["id"]}, {"$set": {"totp_secret": secret, "totp_enabled": False}})
    uri = pyotp.totp.TOTP(secret).provisioning_uri(name=u["email"], issuer_name="Precision ERP")
    return {"secret": secret, "otpauth_url": uri}

@api.post("/auth/2fa/enable")
async def totp_enable(payload: TotpVerifyIn, request: Request, user=Depends(get_current_user)):
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    secret = u.get("totp_secret")
    if not secret:
        raise HTTPException(400, "Run setup first")
    if not pyotp.TOTP(secret).verify(payload.code, valid_window=1):
        raise HTTPException(400, "Invalid code")
    await db.users.update_one({"id": user["id"]}, {"$set": {"totp_enabled": True}})
    await write_audit(user["name"], "2fa_enable", "user", user["id"], request=request)
    return {"ok": True}

@api.post("/auth/2fa/disable")
async def totp_disable(payload: TotpVerifyIn, request: Request, user=Depends(get_current_user)):
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    if not u.get("totp_enabled"): return {"ok": True}
    if not pyotp.TOTP(u["totp_secret"]).verify(payload.code, valid_window=1):
        raise HTTPException(400, "Invalid code")
    await db.users.update_one({"id": user["id"]}, {"$set": {"totp_enabled": False, "totp_secret": ""}})
    await write_audit(user["name"], "2fa_disable", "user", user["id"], request=request)
    return {"ok": True}

@api.get("/auth/2fa/status")
async def totp_status(user=Depends(get_current_user)):
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    return {"enabled": bool(u.get("totp_enabled"))}



# ================ Email Accounts (Gmail App Password — SMTP + IMAP) ================
# No OAuth: users connect each Gmail/Workspace mailbox by providing the email + a
# Google "App Password" (16 chars). Same flow works for Outlook/Yahoo via SMTP/IMAP.

# Per-user-encrypted secrets use a Fernet key derived from JWT_SECRET so they survive
# restarts without an extra env var. If you ever need to rotate, change JWT_SECRET and
# users will be asked to reconnect their mailboxes.
def _email_fernet() -> Fernet:
    digest = hashlib.sha256((JWT_SECRET or "denplex-erp").encode("utf-8")).digest()
    return Fernet(base64.urlsafe_b64encode(digest))

def _enc(plain: str) -> str:
    if plain is None:
        plain = ""
    return _email_fernet().encrypt(plain.encode("utf-8")).decode("utf-8")

def _dec(token: str) -> str:
    if not token:
        return ""
    try:
        return _email_fernet().decrypt(token.encode("utf-8")).decode("utf-8")
    except (InvalidToken, ValueError):
        return ""

# Common provider presets (autodetected from email domain)
EMAIL_PROVIDERS = {
    "gmail.com":      {"smtp_host": "smtp.gmail.com",      "smtp_port": 465, "imap_host": "imap.gmail.com",      "imap_port": 993, "label": "Gmail"},
    "googlemail.com": {"smtp_host": "smtp.gmail.com",      "smtp_port": 465, "imap_host": "imap.gmail.com",      "imap_port": 993, "label": "Gmail"},
    "outlook.com":    {"smtp_host": "smtp.office365.com",  "smtp_port": 587, "imap_host": "outlook.office365.com","imap_port": 993, "label": "Outlook"},
    "hotmail.com":    {"smtp_host": "smtp.office365.com",  "smtp_port": 587, "imap_host": "outlook.office365.com","imap_port": 993, "label": "Outlook"},
    "live.com":       {"smtp_host": "smtp.office365.com",  "smtp_port": 587, "imap_host": "outlook.office365.com","imap_port": 993, "label": "Outlook"},
    "yahoo.com":      {"smtp_host": "smtp.mail.yahoo.com", "smtp_port": 465, "imap_host": "imap.mail.yahoo.com",  "imap_port": 993, "label": "Yahoo"},
    "zoho.com":       {"smtp_host": "smtp.zoho.com",       "smtp_port": 465, "imap_host": "imap.zoho.com",       "imap_port": 993, "label": "Zoho"},
}
DEFAULT_PROVIDER = {"smtp_host": "smtp.gmail.com", "smtp_port": 465, "imap_host": "imap.gmail.com", "imap_port": 993, "label": "Gmail"}

def _autodetect(email_addr: str) -> Dict[str, Any]:
    dom = (email_addr or "").lower().split("@")[-1]
    return EMAIL_PROVIDERS.get(dom, DEFAULT_PROVIDER)

def _app_pw_hint(label: str) -> str:
    """Provider-specific hint shown in 4xx/5xx auth errors."""
    if label == "Outlook":
        return "Open https://account.microsoft.com/security → Advanced security → App passwords and paste the generated password here (2-Step Verification must be ON)."
    if label == "Yahoo":
        return "Open Yahoo → Account Info → Account security → Generate app password and paste it here."
    if label == "Zoho":
        return "Open Zoho Mail → Settings → Mail Accounts → App Specific Passwords and paste it here."
    return "Make sure 2-Step Verification is ON in your Google Account, then create an App Password at https://myaccount.google.com/apppasswords and paste it here (spaces ok, we strip them)."

class EmailAccountIn(BaseModel):
    email: EmailStr
    app_password: str
    label: Optional[str] = ""
    is_default: bool = False
    # Optional manual overrides (rarely needed; autodetected from domain)
    smtp_host: Optional[str] = ""
    smtp_port: Optional[int] = 0
    imap_host: Optional[str] = ""
    imap_port: Optional[int] = 0

class EmailSendIn(BaseModel):
    to: List[EmailStr]
    subject: str
    html: str
    cc: Optional[List[EmailStr]] = None
    bcc: Optional[List[EmailStr]] = None
    attachment_base64: Optional[str] = ""
    attachment_filename: Optional[str] = ""
    attachment_mime: Optional[str] = "application/pdf"
    account_id: Optional[str] = ""  # if blank uses default account

def _account_public(a: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": a.get("id"),
        "email": a.get("email"),
        "label": a.get("label") or _autodetect(a.get("email", "")).get("label"),
        "smtp_host": a.get("smtp_host"),
        "smtp_port": a.get("smtp_port"),
        "imap_host": a.get("imap_host"),
        "imap_port": a.get("imap_port"),
        "is_default": bool(a.get("is_default")),
        "last_test_at": a.get("last_test_at", ""),
        "last_test_ok": bool(a.get("last_test_ok", False)),
        "last_test_error": a.get("last_test_error", ""),
        "created_at": a.get("created_at", ""),
    }

def _smtp_test(smtp_host: str, smtp_port: int, email_addr: str, password: str) -> None:
    """Raise on any failure. Uses SSL on 465, STARTTLS otherwise."""
    ctx = ssl.create_default_context()
    if int(smtp_port) == 465:
        with smtplib.SMTP_SSL(smtp_host, int(smtp_port), context=ctx, timeout=20) as s:
            s.login(email_addr, password)
    else:
        with smtplib.SMTP(smtp_host, int(smtp_port), timeout=20) as s:
            s.ehlo()
            s.starttls(context=ctx)
            s.ehlo()
            s.login(email_addr, password)

def _imap_open(imap_host: str, imap_port: int, email_addr: str, password: str) -> imaplib.IMAP4_SSL:
    m = imaplib.IMAP4_SSL(imap_host, int(imap_port))
    m.login(email_addr, password)
    return m

async def _record_test(user_id: str, acct_id: str, ok: bool, err: str = ""):
    await db.email_accounts.update_one(
        {"id": acct_id, "user_id": user_id},
        {"$set": {"last_test_at": now_iso(), "last_test_ok": bool(ok), "last_test_error": (err or "")[:300]}},
    )

@api.post("/email/accounts")
async def add_email_account(payload: EmailAccountIn, user=Depends(get_current_user)):
    preset = _autodetect(payload.email)
    smtp_host = (payload.smtp_host or preset["smtp_host"]).strip()
    smtp_port = int(payload.smtp_port or preset["smtp_port"]) 
    imap_host = (payload.imap_host or preset["imap_host"]).strip()
    imap_port = int(payload.imap_port or preset["imap_port"]) 
    pw = (payload.app_password or "").strip().replace(" ", "")  # Google shows spaces in groups of 4
    if not pw:
        raise HTTPException(400, "App password required")
    label = preset["label"]
    # Test SMTP login synchronously before persisting (fail fast with a helpful message)
    try:
        await asyncio.to_thread(_smtp_test, smtp_host, smtp_port, str(payload.email), pw)
    except smtplib.SMTPAuthenticationError as e:
        raise HTTPException(400, f"SMTP login failed ({e.smtp_code}). {_app_pw_hint(label)}")
    except Exception as e:
        raise HTTPException(400, f"SMTP test failed: {e}")
    # Quick IMAP test (non-fatal — if IMAP is blocked we'll still allow sending)
    imap_err = ""
    try:
        m = await asyncio.to_thread(_imap_open, imap_host, imap_port, str(payload.email), pw)
        await asyncio.to_thread(m.logout)
    except Exception as e:
        imap_err = f"IMAP (inbox read) failed: {e}"
    # If this is the user's first account or marked default, clear others
    if payload.is_default:
        await db.email_accounts.update_many({"user_id": user["id"]}, {"$set": {"is_default": False}})
    has_any = await db.email_accounts.find_one({"user_id": user["id"]})
    doc = {
        "id": new_id(),
        "user_id": user["id"],
        "email": str(payload.email).lower(),
        "label": payload.label or _autodetect(str(payload.email))["label"],
        "encrypted_password": _enc(pw),
        "smtp_host": smtp_host, "smtp_port": smtp_port,
        "imap_host": imap_host, "imap_port": imap_port,
        "is_default": bool(payload.is_default) or (has_any is None),
        "created_at": now_iso(),
        "last_test_at": now_iso(),
        "last_test_ok": True,
        "last_test_error": imap_err,
    }
    await db.email_accounts.insert_one(doc)
    await write_audit(user["name"], "email_account_added", "email_account", doc["id"], {"email": doc["email"]})
    return {**_account_public(doc), "imap_warning": (imap_err or None)}

@api.get("/email/accounts")
async def list_email_accounts(user=Depends(get_current_user)):
    rows = await db.email_accounts.find({"user_id": user["id"]}, {"_id": 0}).sort("created_at", 1).to_list(50)
    return [_account_public(r) for r in rows]

@api.delete("/email/accounts/{acct_id}")
async def delete_email_account(acct_id: str, user=Depends(get_current_user)):
    a = await db.email_accounts.find_one({"id": acct_id, "user_id": user["id"]})
    if not a:
        raise HTTPException(404, "Not found")
    await db.email_accounts.delete_one({"id": acct_id, "user_id": user["id"]})
    # If we removed the default, promote the next one
    if a.get("is_default"):
        nxt = await db.email_accounts.find_one({"user_id": user["id"]})
        if nxt:
            await db.email_accounts.update_one({"id": nxt["id"]}, {"$set": {"is_default": True}})
    await write_audit(user["name"], "email_account_removed", "email_account", acct_id, {"email": a.get("email")})
    return {"ok": True}

@api.post("/email/accounts/{acct_id}/default")
async def set_default_email(acct_id: str, user=Depends(get_current_user)):
    a = await db.email_accounts.find_one({"id": acct_id, "user_id": user["id"]})
    if not a:
        raise HTTPException(404, "Not found")
    await db.email_accounts.update_many({"user_id": user["id"]}, {"$set": {"is_default": False}})
    await db.email_accounts.update_one({"id": acct_id, "user_id": user["id"]}, {"$set": {"is_default": True}})
    return {"ok": True}

@api.post("/email/accounts/{acct_id}/test")
async def test_email_account(acct_id: str, user=Depends(get_current_user)):
    a = await db.email_accounts.find_one({"id": acct_id, "user_id": user["id"]})
    if not a:
        raise HTTPException(404, "Not found")
    pw = _dec(a.get("encrypted_password", ""))
    err = ""
    ok = True
    try:
        await asyncio.to_thread(_smtp_test, a["smtp_host"], a["smtp_port"], a["email"], pw)
    except Exception as e:
        ok = False
        err = f"SMTP: {e}"
    if ok:
        try:
            m = await asyncio.to_thread(_imap_open, a["imap_host"], a["imap_port"], a["email"], pw)
            await asyncio.to_thread(m.logout)
        except Exception as e:
            err = f"IMAP: {e}"  # don't flip ok; send still works
    await _record_test(user["id"], acct_id, ok, err)
    return {"ok": ok, "error": err}

async def _pick_account(user_id: str, account_id: str = "") -> Dict[str, Any]:
    if account_id:
        a = await db.email_accounts.find_one({"id": account_id, "user_id": user_id})
        if not a:
            raise HTTPException(404, "Email account not found")
        return a
    a = await db.email_accounts.find_one({"user_id": user_id, "is_default": True})
    if a:
        return a
    a = await db.email_accounts.find_one({"user_id": user_id})
    if not a:
        raise HTTPException(400, "No email account connected. Open Settings → Email Accounts and add one.")
    return a

@api.post("/email/send")
async def email_send(payload: EmailSendIn, user=Depends(get_current_user)):
    a = await _pick_account(user["id"], payload.account_id or "")
    pw = _dec(a.get("encrypted_password", ""))
    msg = EmailMessage()
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0}) or {}
    from_name = u.get("name") or "Denplex ERP"
    msg["From"] = f'{from_name} <{a["email"]}>'
    msg["To"] = ", ".join([str(e) for e in payload.to])
    if payload.cc: msg["Cc"] = ", ".join([str(e) for e in payload.cc])
    if payload.bcc: msg["Bcc"] = ", ".join([str(e) for e in payload.bcc])
    msg["Subject"] = payload.subject
    msg.set_content("This email contains HTML content. Please use an HTML-capable client.")
    msg.add_alternative(payload.html, subtype="html")
    if payload.attachment_base64 and payload.attachment_filename:
        b64 = payload.attachment_base64
        if b64.startswith("data:") and "," in b64:
            b64 = b64.split(",", 1)[1]
        raw = base64.b64decode(b64)
        maintype, _, subtype = (payload.attachment_mime or "application/octet-stream").partition("/")
        msg.add_attachment(raw, maintype=maintype, subtype=subtype, filename=payload.attachment_filename)
    rcpts: List[str] = [str(e) for e in payload.to] + [str(e) for e in (payload.cc or [])] + [str(e) for e in (payload.bcc or [])]
    def _do_send():
        ctx = ssl.create_default_context()
        if int(a["smtp_port"]) == 465:
            with smtplib.SMTP_SSL(a["smtp_host"], int(a["smtp_port"]), context=ctx, timeout=30) as s:
                s.login(a["email"], pw)
                s.send_message(msg, from_addr=a["email"], to_addrs=rcpts)
        else:
            with smtplib.SMTP(a["smtp_host"], int(a["smtp_port"]), timeout=30) as s:
                s.ehlo(); s.starttls(context=ctx); s.ehlo()
                s.login(a["email"], pw)
                s.send_message(msg, from_addr=a["email"], to_addrs=rcpts)
    try:
        await asyncio.to_thread(_do_send)
    except smtplib.SMTPAuthenticationError as e:
        hint = _app_pw_hint(_autodetect(a["email"])["label"])
        raise HTTPException(502, f"Login rejected by mail server ({e.smtp_code}). Your App Password may have expired. {hint}")
    except Exception as e:
        raise HTTPException(502, f"Email send failed: {e}")
    await write_audit(user["name"], "email_send", "email", None, {"from": a["email"], "to": payload.to, "subject": payload.subject})
    return {"ok": True, "from": a["email"]}

def _decode_header_str(value: str) -> str:
    if not value:
        return ""
    parts = decode_header(value)
    out = []
    for txt, enc in parts:
        if isinstance(txt, bytes):
            try:
                out.append(txt.decode(enc or "utf-8", errors="replace"))
            except Exception:
                out.append(txt.decode("utf-8", errors="replace"))
        else:
            out.append(txt)
    return "".join(out)

def _fetch_inbox(host: str, port: int, email_addr: str, password: str, max_n: int) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    m = imaplib.IMAP4_SSL(host, int(port))
    try:
        m.login(email_addr, password)
        m.select("INBOX", readonly=True)
        typ, data = m.search(None, "ALL")
        if typ != "OK" or not data or not data[0]:
            return out
        ids = data[0].split()
        recent = ids[-int(max_n):][::-1]
        for mid in recent:
            typ, msg_data = m.fetch(mid, "(BODY.PEEK[HEADER.FIELDS (FROM SUBJECT DATE)] BODY.PEEK[TEXT]<0.500>)")
            if typ != "OK" or not msg_data:
                continue
            headers = ""
            snippet = ""
            for part in msg_data:
                if isinstance(part, tuple) and len(part) >= 2:
                    chunk = part[1].decode("utf-8", errors="replace") if isinstance(part[1], (bytes, bytearray)) else str(part[1])
                    if "FROM" in (part[0].decode(errors="replace") if isinstance(part[0], (bytes, bytearray)) else str(part[0])).upper():
                        headers = chunk
                    else:
                        snippet = chunk
            msg = emaillib.message_from_string(headers or "")
            frm_raw = _decode_header_str(msg.get("From", ""))
            subj = _decode_header_str(msg.get("Subject", ""))
            date_s = msg.get("Date", "")
            try:
                date_iso = parsedate_to_datetime(date_s).isoformat() if date_s else ""
            except Exception:
                date_iso = ""
            name, em_addr = parseaddr(frm_raw)
            # Clean snippet: strip HTML tags, collapse whitespace
            snippet_text = re.sub(r"<[^>]+>", " ", snippet)
            snippet_text = re.sub(r"&nbsp;|&zwnj;|&amp;|&lt;|&gt;|&quot;", " ", snippet_text)
            snippet_clean = re.sub(r"\s+", " ", snippet_text).strip()[:240]
            out.append({
                "uid": mid.decode("utf-8", errors="replace") if isinstance(mid, (bytes, bytearray)) else str(mid),
                "from_name": name or (em_addr.split("@")[0] if em_addr else ""),
                "from_email": em_addr or "",
                "subject": subj,
                "date": date_iso,
                "snippet": snippet_clean,
            })
    finally:
        try:
            m.close()
        except Exception:
            pass
        try:
            m.logout()
        except Exception:
            pass
    return out

@api.get("/email/accounts/{acct_id}/inbox")
async def email_inbox(acct_id: str, max: int = 25, user=Depends(get_current_user)):
    a = await db.email_accounts.find_one({"id": acct_id, "user_id": user["id"]})
    if not a:
        raise HTTPException(404, "Not found")
    pw = _dec(a.get("encrypted_password", ""))
    try:
        msgs = await asyncio.to_thread(_fetch_inbox, a["imap_host"], a["imap_port"], a["email"], pw, max)
    except Exception as e:
        raise HTTPException(502, f"Inbox fetch failed: {e}")
    return {"account": a["email"], "messages": msgs}

async def _sync_one_account(user_id: str, a: Dict[str, Any], max_n: int = 25) -> Dict[str, int]:
    pw = _dec(a.get("encrypted_password", ""))
    try:
        msgs = await asyncio.to_thread(_fetch_inbox, a["imap_host"], a["imap_port"], a["email"], pw, max_n)
    except Exception:
        return {"added": 0, "scanned": 0}
    added = 0
    for m in msgs:
        em_addr = (m.get("from_email") or "").lower()
        if not em_addr:
            continue
        if "noreply" in em_addr or "no-reply" in em_addr or "mailer-daemon" in em_addr:
            continue
        ext = f"{a['email']}::{m.get('uid')}"
        if await db.leads.find_one({"external_id": ext}):
            continue
        await db.leads.insert_one({
            "id": new_id(),
            "external_id": ext,
            "name": m.get("from_name") or em_addr.split("@")[0],
            "company": "",
            "phone": "",
            "email": em_addr,
            "source": "email",
            "requirement": m.get("subject", "") or m.get("snippet", "")[:160],
            "status": "new",
            "notes": (m.get("snippet", "") or "")[:500],
            "created_at": now_iso(),
        })
        added += 1
    return {"added": added, "scanned": len(msgs)}

@api.post("/email/accounts/{acct_id}/sync-leads")
async def email_sync_one(acct_id: str, max: int = 25, user=Depends(get_current_user)):
    a = await db.email_accounts.find_one({"id": acct_id, "user_id": user["id"]})
    if not a:
        raise HTTPException(404, "Not found")
    res = await _sync_one_account(user["id"], a, max)
    await write_audit(user["name"], "email_sync_leads", "leads", None, {"account": a["email"], **res})
    return {"account": a["email"], **res}

@api.post("/email/sync-leads")
async def email_sync_all(max: int = 25, user=Depends(get_current_user)):
    accts = await db.email_accounts.find({"user_id": user["id"]}, {"_id": 0}).to_list(50)
    if not accts:
        raise HTTPException(400, "No email accounts connected.")
    total_added = 0; total_scanned = 0; per = []
    for a in accts:
        res = await _sync_one_account(user["id"], a, max)
        per.append({"account": a["email"], **res})
        total_added += res["added"]; total_scanned += res["scanned"]
    await write_audit(user["name"], "email_sync_leads_all", "leads", None, {"added": total_added, "scanned": total_scanned})
    return {"added": total_added, "scanned": total_scanned, "per_account": per}




# ================ Vyapar Import (.vyb / .xlsx / .csv) ================
# Vyapar has no public API; we accept (a) Excel/CSV exports from Vyapar's
# Reports menu (most reliable), or (b) a raw `.vyb` backup which we attempt
# to identify (SQLite / ZIP / encrypted). Encrypted backups fall back to
# instructions on how to do the Excel export.

VYAPAR_UPLOAD_DIR = Path("/tmp/vyapar_uploads")
VYAPAR_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

def _vyb_inspect(path: Path) -> Dict[str, Any]:
    """Identify what's inside an uploaded file."""
    import sqlite3, zipfile
    out: Dict[str, Any] = {"kind": "unknown", "tables": [], "counts": {}, "notes": ""}
    with open(path, "rb") as f:
        head = f.read(64)
    if head.startswith(b"SQLite format 3"):
        out["kind"] = "sqlite"
        try:
            con = sqlite3.connect(str(path)); cur = con.cursor()
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
            tables = [r[0] for r in cur.fetchall()]
            out["tables"] = tables
            counts = {}
            for t in tables:
                try:
                    cur.execute(f'SELECT COUNT(*) FROM "{t}"')
                    counts[t] = cur.fetchone()[0]
                except Exception:
                    pass
            out["counts"] = counts
            con.close()
            out["notes"] = "Plain SQLite database — Vyapar tables can be extracted directly."
        except Exception as e:
            out["notes"] = f"SQLite open failed: {e}"
        return out
    if head[:2] == b"PK":
        # could be xlsx or zip — sniff by extension first
        if path.suffix.lower() in (".xlsx", ".xls"):
            out["kind"] = "xlsx"
            try:
                from openpyxl import load_workbook
                wb = load_workbook(filename=str(path), read_only=True, data_only=True)
                out["tables"] = wb.sheetnames
                counts = {}
                for s in wb.sheetnames:
                    ws = wb[s]
                    counts[s] = max(0, (ws.max_row or 1) - 1)
                out["counts"] = counts
                wb.close()
                out["notes"] = "Excel workbook — sheets will be mapped to Parties / Items / Sales / Purchases by column heuristics."
            except Exception as e:
                out["notes"] = f"Excel read failed: {e}"
            return out
        out["kind"] = "zip"
        try:
            with zipfile.ZipFile(path) as z:
                names = z.namelist()
                out["tables"] = names[:50]
                inner_sqlite = [n for n in names if n.endswith(".db") or n.endswith(".sqlite") or n.endswith(".vyp")]
                if inner_sqlite:
                    out["notes"] = f"Vyapar backup container — extracting {inner_sqlite[0]} for inspection."
                    # Extract the inner DB and recurse so we can preview the real tables
                    import tempfile
                    tmpd = Path(tempfile.mkdtemp())
                    z.extract(inner_sqlite[0], tmpd)
                    inner_path = tmpd / inner_sqlite[0]
                    inner_info = _vyb_inspect(inner_path)
                    if inner_info.get("kind") == "sqlite":
                        # Replace the inspect result with the inner SQLite analysis but keep the
                        # outer kind=zip so import logic knows it must unzip first.
                        out["kind"] = "zip_sqlite"
                        out["inner_sqlite"] = inner_sqlite[0]
                        out["tables"] = inner_info.get("tables", [])
                        out["counts"] = inner_info.get("counts", {})
                        out["notes"] = ("Vyapar SQLite backup — " + (
                            f"{out['counts'].get('kb_names', 0)} parties, "
                            f"{out['counts'].get('kb_items', 0)} items, "
                            f"{out['counts'].get('kb_transactions', 0)} transactions."))
                else:
                    out["notes"] = "Archive contains: " + ", ".join(names[:10])
        except Exception as e:
            out["notes"] = f"ZIP read failed: {e}"
        return out
    if path.suffix.lower() == ".csv":
        out["kind"] = "csv"
        try:
            import csv as _csv
            with open(path, "r", encoding="utf-8-sig", errors="replace", newline="") as f:
                rdr = _csv.reader(f); rows = list(rdr)
            if rows:
                out["tables"] = [",".join(rows[0][:10])]
                out["counts"] = {"rows": len(rows) - 1}
            out["notes"] = "CSV file."
        except Exception as e:
            out["notes"] = f"CSV read failed: {e}"
        return out
    out["kind"] = "unsupported"
    out["notes"] = ("This file is encrypted or in a proprietary Vyapar format we cannot decrypt without their internal keys. "
                    "The reliable workaround is to export Excel files from Vyapar's Reports menu (Sale, Party, Item, Purchase). "
                    "Re-upload the .xlsx file here.")
    return out

def _norm(s: Any) -> str:
    return re.sub(r"[^a-z0-9]", "", str(s or "").strip().lower())

PARTY_COLS = {"name": ["name", "partyname", "customername", "suppliername"],
              "phone": ["phone", "mobile", "phoneno", "mobileno", "contact", "contactno"],
              "email": ["email", "emailid"],
              "gstin": ["gstin", "gstno", "gst"],
              "address": ["address", "billingaddress"],
              "state": ["state", "billingstate"]}
ITEM_COLS  = {"name": ["itemname", "name", "productname"],
              "hsn": ["hsnsac", "hsn", "hsncode", "sac"],
              "sale_price": ["saleprice", "price", "rate", "mrp"],
              "purchase_price": ["purchaseprice", "cost"],
              "stock": ["openingstock", "stock", "currentstock", "quantity"],
              "unit": ["unit", "uom"]}
SALE_COLS  = {"code": ["invoiceno", "saleinvoiceno", "billno", "refno"],
              "date": ["date", "invoicedate"],
              "customer_name": ["partyname", "customername"],
              "total": ["totalamount", "amount", "total", "grandtotal"],
              "balance": ["balance", "balancedue"],
              "payment_type": ["paymentmode", "paymenttype"]}
PURCHASE_COLS = {"code": ["billno", "purchaseno", "refno", "invoiceno"],
                 "date": ["date", "billdate"],
                 "supplier_name": ["partyname", "suppliername"],
                 "total": ["totalamount", "amount", "total", "grandtotal"]}

def _classify_sheet(headers: List[str]) -> str:
    h = [_norm(x) for x in headers]
    score = {"sales": 0, "purchases": 0, "items": 0, "parties": 0}
    for col in h:
        if col in {"invoiceno", "saleinvoiceno", "billno"}: score["sales"] += 2
        if col in {"purchaseno"}: score["purchases"] += 2
        if col in {"itemname", "hsnsac", "saleprice", "purchaseprice", "openingstock"}: score["items"] += 1
        if col in {"partyname", "customername", "suppliername", "gstin", "mobileno"}: score["parties"] += 1
        if col in {"totalamount", "grandtotal"} and "partyname" in h: score["sales"] += 1
    best = max(score, key=score.get)
    return best if score[best] > 0 else "unknown"

def _map_row(row: Dict[str, Any], schema: Dict[str, List[str]]) -> Dict[str, Any]:
    norm = {_norm(k): v for k, v in row.items()}
    out: Dict[str, Any] = {}
    for canonical, synonyms in schema.items():
        for s in synonyms:
            if s in norm and norm[s] not in (None, "", "—"):
                out[canonical] = norm[s]
                break
    return out

class VyaparImportIn(BaseModel):
    token: str
    parties: bool = True
    items: bool = True
    sales: bool = True
    purchases: bool = True
    expenses: bool = True
    dry_run: bool = False

@api.post("/integrations/vyapar/inspect")
async def vyapar_inspect(file: UploadFile = File(...), user=Depends(require_roles("admin"))):
    if not file.filename:
        raise HTTPException(400, "No file")
    token = new_id()
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", file.filename)[-80:]
    path = VYAPAR_UPLOAD_DIR / f"{token}_{safe_name}"
    raw = await file.read()
    if len(raw) > 50 * 1024 * 1024:
        raise HTTPException(400, "File too large (>50 MB)")
    path.write_bytes(raw)
    info = await asyncio.to_thread(_vyb_inspect, path)
    info["token"] = token
    info["filename"] = safe_name
    info["size_bytes"] = len(raw)
    await db.vyapar_uploads.insert_one({
        "id": token, "user_id": user["id"], "filename": safe_name,
        "path": str(path), "kind": info.get("kind"), "created_at": now_iso(),
    })
    return info

async def _do_import_sqlite(path: Path, opts: VyaparImportIn, auto_seed_company: bool = True) -> Dict[str, Any]:
    """Import a Vyapar SQLite DB. Recognises kb_* tables (Vyapar's real schema)."""
    import sqlite3
    res = {"parties": 0, "items": 0, "sales": 0, "purchases": 0,
           "quotations": 0, "sale_orders": 0, "purchase_orders": 0,
           "delivery_challans": 0, "job_work_out": 0, "sale_returns": 0,
           "company_seeded": False}
    con = sqlite3.connect(str(path)); con.row_factory = sqlite3.Row
    cur = con.cursor()

    # --- Detect schema: native Vyapar (kb_*) vs generic ---
    cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
    tables = {r[0].lower(): r[0] for r in cur.fetchall()}
    is_vyapar = "kb_names" in tables and "kb_transactions" in tables and "kb_items" in tables

    if not is_vyapar:
        # Fall back to generic best-effort (unchanged behavior for non-Vyapar SQLite)
        return await _do_import_generic_sqlite(path, opts)

    # --- Auto-seed company from kb_firms (only if user hasn't already saved) ---
    if auto_seed_company and "kb_firms" in tables:
        try:
            cur.execute('SELECT * FROM "kb_firms" LIMIT 1')
            f = cur.fetchone()
            if f:
                f = {k: f[k] for k in f.keys()}
                existing = await get_setting("integrations") or {}
                # Only fill blanks; never overwrite something the user already set
                payload: Dict[str, Any] = {**existing}
                for src_key, dst_key in [
                    ("firm_name", "company_name"),
                    ("firm_phone", "company_phone"),
                    ("firm_email", "company_email"),
                    ("firm_address", "company_address"),
                    ("firm_gstin_number", "company_gstin"),
                    ("firm_state", "company_state"),
                    ("firm_bank_name", "bank_name"),
                    ("firm_bank_account_number", "bank_account_no"),
                    ("firm_bank_ifsc_code", "bank_ifsc"),
                ]:
                    if not payload.get(dst_key) and f.get(src_key):
                        payload[dst_key] = str(f[src_key])
                # UDYAM is embedded in firm_address; extract if present and not set
                if not payload.get("company_udyam") and f.get("firm_address"):
                    m = re.search(r"UDYAM-[A-Z]+-\d+-\d+", str(f["firm_address"]))
                    if m:
                        payload["company_udyam"] = m.group(0)
                if not opts.dry_run:
                    await set_setting("integrations", payload)
                    res["company_seeded"] = True
        except Exception:
            pass

    # --- HSN tax-id → rate lookup (kb_tax_code) ---
    tax_rate: Dict[int, float] = {}
    if "kb_tax_code" in tables:
        try:
            cur.execute('SELECT tax_code_id, tax_rate FROM "kb_tax_code"')
            for r in cur.fetchall():
                tax_rate[int(r["tax_code_id"])] = float(r["tax_rate"] or 0)
        except Exception:
            pass

    # --- Determine each party's role from how it's used in transactions ---
    cust_ids = set(); supp_ids = set()
    try:
        cur.execute("SELECT txn_name_id, txn_type FROM kb_transactions WHERE txn_name_id IS NOT NULL")
        for r in cur.fetchall():
            nid = int(r["txn_name_id"]); tt = int(r["txn_type"] or 0)
            # Verified vs real Denplex .vyb: customer-side types 1,3,21,24,27,30,65; supplier-side 2,4,23,28
            if tt in (1, 3, 21, 24, 27, 30, 65):
                cust_ids.add(nid)
            elif tt in (2, 4, 23, 28):
                supp_ids.add(nid)
    except Exception:
        pass

    # --- Unit lookup (kb_item_units) for item UOM ---
    unit_lookup: Dict[int, str] = {}
    try:
        cur.execute('SELECT unit_id, unit_short_name FROM "kb_item_units"')
        for r in cur.fetchall():
            unit_lookup[int(r["unit_id"])] = (r["unit_short_name"] or "").strip()
    except Exception:
        pass

    # --- Parties (kb_names) ---
    if opts.parties:
        # name_type: 1 = party (customer/supplier); 2 = EXPENSE CATEGORY (imported separately below)
        cur.execute('SELECT * FROM "kb_names" WHERE name_type = 1')
        for r in cur.fetchall():
            full_name = (r["full_name"] or "").strip()
            if not full_name:
                continue
            doc = {
                "id": new_id(),
                "name": full_name[:160],
                "phone": str(r["phone_number"] or "").strip(),
                "email": str(r["email"] or "").strip(),
                "gstin": str(r["name_gstin_number"] or "").strip(),
                "address": str(r["address"] or "").strip(),
                "state": str(r["name_state"] or "").strip(),
                "vyapar_id": str(r["name_id"]),
                "source": "vyapar",
                "created_at": str(r["date_created"] or now_iso()),
            }
            nid = int(r["name_id"])
            is_supp = nid in supp_ids
            is_cust = (nid in cust_ids) or (not is_supp)   # default unknown parties to customer
            if not opts.dry_run:
                if is_cust:
                    await db.customers.update_one({"name": doc["name"]}, {"$setOnInsert": doc}, upsert=True)
                if is_supp:
                    await db.suppliers.update_one({"name": doc["name"]}, {"$setOnInsert": doc}, upsert=True)
            res["parties"] += 1

    # Build name_id -> party_name map for txn rows
    name_lookup: Dict[int, str] = {}
    try:
        cur.execute('SELECT name_id, full_name FROM "kb_names"')
        for r in cur.fetchall():
            if r["full_name"]:
                name_lookup[int(r["name_id"])] = r["full_name"].strip()
    except Exception:
        pass

    # --- Items (kb_items) ---
    if opts.items:
        cur.execute('SELECT * FROM "kb_items" WHERE item_is_active IS NULL OR item_is_active != 0')
        for r in cur.fetchall():
            name = (r["item_name"] or "").strip()
            if not name or name == "item 1":  # skip the default placeholder
                continue
            sku_raw = (r["item_code"] or f"VY-{r['item_id']}").strip()[:32] or f"VY-{r['item_id']}"
            doc = {
                "id": new_id(),
                "sku": sku_raw,
                "name": name[:200],
                "category": "raw",
                "uom": unit_lookup.get(int(r["base_unit_id"]), "Nos") if r["base_unit_id"] is not None else "Nos",
                "secondary_unit": unit_lookup.get(int(r["secondary_unit_id"]), "") if r["secondary_unit_id"] is not None else "",
                "conversion_factor": 0,
                "qty_on_hand": float(r["item_stock_quantity"] or 0),
                "qty_in_process": 0.0,
                "reorder_level": float(r["item_min_stock_quantity"] or 0),
                "unit_cost": float(r["item_purchase_unit_price"] or 0) or float(r["item_sale_unit_price"] or 0),
                "hsn": str(r["item_hsn_sac_code"] or "").strip(),
                "gst_rate": (tax_rate.get(int(r["item_tax_id"]), 18.0) if r["item_tax_id"] is not None else 18.0),
                "location": "",
                "sale_price": float(r["item_sale_unit_price"] or 0),
                "description": str(r["item_description"] or "").strip(),
                "vyapar_id": str(r["item_id"]),
                "source": "vyapar",
                "created_at": str(r["item_date_created"] or now_iso()),
            }
            if not opts.dry_run:
                await db.items.update_one({"name": doc["name"]}, {"$setOnInsert": doc}, upsert=True)
            res["items"] += 1

    # Build item_id -> (name, hsn, rate) for line items
    item_lookup: Dict[int, Dict[str, Any]] = {}
    try:
        cur.execute('SELECT item_id, item_name, item_hsn_sac_code, item_sale_unit_price FROM "kb_items"')
        for r in cur.fetchall():
            item_lookup[int(r["item_id"])] = {
                "name": (r["item_name"] or "").strip(),
                "hsn": (r["item_hsn_sac_code"] or "").strip(),
                "rate": float(r["item_sale_unit_price"] or 0),
            }
    except Exception:
        pass

    # Purchaser Name (Vyapar custom/UDF field) -> txn_id lookup, so it can print on the invoice
    # exactly like Vyapar does (kb_udf_fields defines the field, kb_udf_values holds per-txn values).
    purchaser_name_lookup: Dict[int, str] = {}
    try:
        cur.execute('SELECT udf_field_id FROM "kb_udf_fields" WHERE udf_field_name = ?', ("Purchaser Name",))
        _pf_ids = [int(r["udf_field_id"]) for r in cur.fetchall()]
        if _pf_ids:
            qmarks = ",".join("?" * len(_pf_ids))
            cur.execute(f'SELECT udf_ref_id, udf_value FROM "kb_udf_values" WHERE udf_value_field_id IN ({qmarks})', _pf_ids)
            for r in cur.fetchall():
                if r["udf_value"]:
                    purchaser_name_lookup[int(r["udf_ref_id"])] = str(r["udf_value"]).strip()
    except Exception:
        pass

    # Payment type names (moved up from the Payment-In/Out pass below so invoices/bills can also
    # print a Payment Mode — Vyapar shows "Credit" when no payment type is set on the doc itself).
    ptype_name: Dict[int, str] = {}
    try:
        cur.execute('SELECT paymentType_id, paymentType_name FROM "kb_paymentTypes"')
        for r in cur.fetchall():
            ptype_name[int(r["paymentType_id"])] = (r["paymentType_name"] or "")
    except Exception:
        pass

    def _map_ptype(pid):
        if pid is None:
            return "Credit"
        nm = (ptype_name.get(int(pid)) if pid is not None else "") or ""
        nmu = nm.upper()
        if "CASH" in nmu: return "Cash"
        if "CHEQUE" in nmu: return "Cheque"
        if "UPI" in nmu: return "UPI"
        if "CARD" in nmu: return "Card"
        return "Bank Transfer" if nm else "Credit"

    # --- Line items pre-grouped by txn_id ---
    lines_by_txn: Dict[int, List[Dict[str, Any]]] = {}
    try:
        cur.execute('SELECT * FROM "kb_lineitems"')
        for r in cur.fetchall():
            tx = int(r["lineitem_txn_id"])
            it = item_lookup.get(int(r["item_id"]) if r["item_id"] is not None else -1, {})
            qty = float(r["quantity"] or 0)
            rate = float(r["priceperunit"] or 0)
            disc_amt = float(r["lineitem_discount_amount"] or 0)
            tax_amt = float(r["lineitem_tax_amount"] or 0)
            taxable = max(qty * rate - disc_amt, 0.000001)
            implied_gst_rate = round(tax_amt / taxable * 100, 2) if tax_amt else 0
            lines_by_txn.setdefault(tx, []).append({
                "description": (r["lineitem_description"] or it.get("name", "")).strip() or it.get("name", ""),
                "hsn": it.get("hsn", ""),
                "qty": qty,
                "rate": rate,
                "discount_amount": disc_amt,
                "discount_pct": float(r["lineitem_discount_percent"] or 0),
                "gst_rate": implied_gst_rate,
                "gst_amount": tax_amt,
            })
    except Exception:
        pass

    # --- Transactions: route by txn_type ---
    # Map VERIFIED against the real Denplex .vyb (July 2026 backup):
    # 1=Sale, 2=Purchase, 3=Payment-In, 4=Payment-Out, 7=EXPENSE (not sale return!),
    # 21=Sale Return (Credit Note), 23=Purchase Return, 24=Sale Order,
    # 27=Quotation/Estimate, 28=Purchase Order, 30=Delivery Challan, 65=Proforma Invoice
    TYPE_ROUTES = {
        1:  ("invoices", "sales", "customer", "Tax Invoice"),
        2:  ("vendor_bills", "purchases", "supplier", "Purchase Bill"),
        21: ("credit_notes", "sale_returns", "customer", "Credit Note"),
        23: ("purchase_returns", "purchase_returns", "supplier", "Purchase Return"),
        24: ("sale_orders", "sale_orders", "customer", "Sale Order"),
        27: ("quotations", "quotations", "customer", "Quotation"),
        28: ("purchase_orders", "purchase_orders", "supplier", "Purchase Order"),
        30: ("delivery_challans", "delivery_challans", "customer", "Delivery Challan"),
        65: ("proforma_invoices", "proformas", "customer", "Proforma Invoice"),
    }
    # TDS section/rate lookup (tds_tax_rates)
    tds_lookup: Dict[int, Dict[str, Any]] = {}
    try:
        cur.execute('SELECT id, name, percentage FROM "tds_tax_rates"')
        for r in cur.fetchall():
            tds_lookup[int(r["id"])] = {"name": (r["name"] or ""), "rate": float(r["percentage"] or 0)}
    except Exception:
        pass

    res.update({"expenses": 0, "purchase_returns": 0, "proformas": 0,
                "cash_on_doc_payments": 0, "opening_balances": 0, "skipped_unknown": []})
    txn_to_doc: Dict[int, Dict[str, Any]] = {}
    cash_on_doc: List[Dict[str, Any]] = []            # cash received/paid on the doc itself -> synthetic payment below
    inv_open: Dict[int, List[Any]] = {}               # txn_id -> [party_kind, party_name, ERP-derived outstanding]
    cur.execute("SELECT * FROM kb_transactions")
    for r in cur.fetchall():
        t = int(r["txn_type"]) if r["txn_type"] is not None else 0
        if t in (3, 4, 7):
            continue                                   # payments + expenses are handled in their own passes
        route = TYPE_ROUTES.get(t)
        if not route:
            res["skipped_unknown"].append({"txn_id": r["txn_id"], "txn_type": t,
                                           "ref": (r["txn_ref_number_char"] or "").strip(),
                                           "amount": round(float(r["txn_cash_amount"] or 0) + float(r["txn_balance_amount"] or 0), 2)})
            continue
        collection_name, counter_key, party_kind, _title = route
        # Honour user toggles
        if party_kind == "customer" and not opts.sales: continue
        if party_kind == "supplier" and not opts.purchases: continue

        party_name = name_lookup.get(int(r["txn_name_id"]) if r["txn_name_id"] is not None else -1, "Unknown")
        ref = (r["txn_ref_number_char"] or "").strip()
        prefix = (r["txn_invoice_prefix"] or "").strip()
        code = (f"{prefix}{ref}" if ref else f"VY-{r['txn_id']}").strip()
        cash = float(r["txn_cash_amount"] or 0)
        bal = float(r["txn_balance_amount"] or 0)      # balance AT CREATION
        cur_bal = float(r["txn_current_balance"] or 0) # LIVE balance (net of later linked payments)
        total = cash + bal
        sub = max(total - float(r["txn_tax_amount"] or 0), 0)
        ln = lines_by_txn.get(int(r["txn_id"]), [])
        # Split GST based on place_of_supply vs firm_state (heuristic: if intra-state, half/half; if interstate, all IGST)
        is_interstate = bool(r["txn_place_of_supply"]) and "gujarat" not in str(r["txn_place_of_supply"]).lower()
        tax_total = float(r["txn_tax_amount"] or 0)
        cgst = 0.0; sgst = 0.0; igst = 0.0
        if is_interstate: igst = tax_total
        else: cgst = sgst = tax_total / 2
        tds_amt = float(r["txn_tds_tax_amount"] or 0)
        tds_info = tds_lookup.get(int(r["txn_tds_tax_id"])) if r["txn_tds_tax_id"] is not None else None

        # Vyapar's kb_lineitems.lineitem_tax_amount is often 0/unset even when the transaction as a
        # whole clearly carries tax (txn_tax_amount > 0) — this used to silently print "0%" GST on
        # every line of the PDF Tax Summary even though the invoice total was correct. When that
        # mismatch is detected, redistribute the doc-level tax proportionally across its lines and
        # snap to the nearest real GST slab, so the printed Tax Summary matches the invoice total.
        if ln and tax_total > 0.01 and not any(float(l.get("gst_rate", 0) or 0) > 0 for l in ln):
            _line_taxable_sum = sum(
                max(float(l.get("qty", 0) or 0) * float(l.get("rate", 0) or 0) - float(l.get("discount_amount", 0) or 0), 0)
                for l in ln
            )
            if _line_taxable_sum > 0:
                _raw_rate = tax_total / _line_taxable_sum * 100
                _rate = min([0, 5, 12, 18, 28], key=lambda s: abs(s - _raw_rate))
                for l in ln:
                    _lt = max(float(l.get("qty", 0) or 0) * float(l.get("rate", 0) or 0) - float(l.get("discount_amount", 0) or 0), 0)
                    l["gst_rate"] = _rate
                    l["gst_amount"] = round(_lt * _rate / 100, 2)

        doc: Dict[str, Any] = {
            "id": new_id(),
            "code": code,
            "date": str(r["txn_date"] or "")[:19] or now_iso(),
            "due_date": str(r["txn_due_date"] or "")[:10],
            "lines": ln,
            "subtotal": sub,
            "total": total,
            "round_off": float(r["txn_round_off_amount"] or 0),
            "notes": (r["txn_description"] or "").strip(),
            "place_of_supply": (r["txn_place_of_supply"] or "").strip(),
            "is_interstate": is_interstate,
            "vyapar_id": str(r["txn_id"]),
            "vyapar_txn_type": t,
            "source": "vyapar",
            "created_at": str(r["txn_date_created"] or now_iso()),
            "po_date": str(r["txn_po_date"] or "")[:10],
            "po_number": (r["txn_po_ref_number"] or "").strip(),
            "purchaser_name": purchaser_name_lookup.get(int(r["txn_id"]), ""),
            "payment_mode": _map_ptype(r["txn_payment_type_id"]),
            "ship_to_address": (r["txn_shipping_address"] or "").strip(),
        }
        if tds_amt:
            doc["tds"] = round(tds_amt, 2)
            if tds_info:
                doc["tds_rate"] = tds_info["rate"]
                doc["tds_section"] = tds_info["name"]
        if party_kind == "customer":
            doc["customer_id"] = ""
            doc["customer_name"] = party_name
            doc["cgst"] = cgst; doc["sgst"] = sgst; doc["igst"] = igst
            doc["gst_total"] = tax_total
            doc["status"] = "paid" if cur_bal <= 0.01 else "sent"
            doc["outstanding"] = round(max(cur_bal, 0.0), 2)  # Vyapar's own live balance = ground truth
        else:
            doc["supplier_id"] = ""
            doc["supplier_name"] = party_name
            doc["gst_total"] = tax_total
            doc["status"] = "received" if cur_bal <= 0.01 else "open"
            doc["outstanding"] = round(max(cur_bal, 0.0), 2)  # Vyapar's own live balance = ground truth

        if not opts.dry_run:
            target = getattr(db, collection_name)
            await target.update_one({"code": doc["code"], "vyapar_id": doc["vyapar_id"]},
                                    {"$setOnInsert": doc}, upsert=True)
        if t in (1, 2):
            txn_to_doc[int(r["txn_id"])] = {"id": doc["id"], "code": doc["code"], "dtype": "invoice" if t == 1 else "vendor_bill"}
            if cur_bal > 0.01:
                # ERP will show: total - (cash-on-doc payment) - later allocations = starts at bal
                inv_open[int(r["txn_id"])] = [party_kind, party_name, bal]
            if cash > 0.01 and bal > 0.01:
                # Cash taken on the doc itself must become a payment record, else ERP overstates outstanding
                cash_on_doc.append({"is_in": t == 1, "party": party_name,
                                    "date": str(r["txn_date"] or "")[:10] or now_iso()[:10],
                                    "amount": round(cash, 2), "ptype_id": r["txn_payment_type_id"],
                                    "doc_id": doc["id"], "doc_code": doc["code"], "txn_id": int(r["txn_id"])})
        # Increment counter
        if t == 1: res["sales"] += 1
        elif t == 2: res["purchases"] += 1
        elif t == 21: res["sale_returns"] += 1
        elif t == 23: res["purchase_returns"] += 1
        elif t == 24: res["sale_orders"] += 1
        elif t == 27: res["quotations"] += 1
        elif t == 28: res["purchase_orders"] += 1
        elif t == 30: res["delivery_challans"] += 1
        elif t == 65: res["proformas"] += 1

    # --- Payments In/Out (kb_transactions type 3/4) + bill-to-bill allocations (kb_txn_links) ---
    # (ptype_name / _map_ptype now built earlier, above the invoice/bill loop, so both passes share it —
    # for payments-in/out specifically, "Credit" doesn't make sense, so treat a missing type as Cash.)
    def _map_ptype_payment(pid):
        return "Cash" if pid is None else _map_ptype(pid)

    pay_allocs: Dict[int, List] = {}
    try:
        cur.execute('SELECT txn_links_txn_1_id, txn_links_txn_2_id, txn_links_amount FROM "kb_txn_links" WHERE txn_links_txn_1_type IN (3,4)')
        for r in cur.fetchall():
            pay_allocs.setdefault(int(r["txn_links_txn_1_id"]), []).append((int(r["txn_links_txn_2_id"]), float(r["txn_links_amount"] or 0)))
    except Exception:
        pass

    res["payments_in"] = 0; res["payments_out"] = 0
    cur.execute("SELECT * FROM kb_transactions WHERE txn_type IN (3,4)")
    for r in cur.fetchall():
        t = int(r["txn_type"]); is_in = (t == 3)
        if is_in and not opts.sales: continue
        if (not is_in) and not opts.purchases: continue
        party = name_lookup.get(int(r["txn_name_id"]) if r["txn_name_id"] is not None else -1, "Unknown")
        amt = float(r["txn_cash_amount"] or 0) + float(r["txn_balance_amount"] or 0)
        allocs = []
        for (doc_txn, a) in pay_allocs.get(int(r["txn_id"]), []):
            m = txn_to_doc.get(doc_txn)
            if m and a > 0:
                allocs.append({"document_id": m["id"], "document_code": m["code"],
                               "document_type": ("invoice" if is_in else "vendor_bill"), "amount": round(a, 2), "tds_amount": 0})
                if doc_txn in inv_open:
                    inv_open[doc_txn][2] -= a
        allocated = round(sum(x["amount"] for x in allocs), 2)
        code = (f"PMT-IN-VY-{r['txn_id']}" if is_in else f"PMT-OUT-VY-{r['txn_id']}")
        pdoc = {"id": new_id(), "code": code, "party_id": "", "party_name": party,
                "date": str(r["txn_date"] or "")[:10] or now_iso()[:10], "amount": round(amt, 2),
                "allocated_amount": allocated, "payment_type": _map_ptype_payment(r["txn_payment_type_id"]),
                "ref_no": (r["txn_ref_number_char"] or "").strip(), "bank_name": "", "allocations": allocs,
                "notes": (r["txn_description"] or "").strip(), "vyapar_id": str(r["txn_id"]), "source": "vyapar",
                "created_at": str(r["txn_date_created"] or now_iso())}
        pdoc["status"] = "Used" if (allocated >= amt - 0.01 and allocated > 0) else ("Partially Used" if allocated > 0 else "Unused")
        if not opts.dry_run:
            await (db.payments_in if is_in else db.payments_out).update_one(
                {"vyapar_id": pdoc["vyapar_id"], "code": pdoc["code"]}, {"$setOnInsert": pdoc}, upsert=True)
        if is_in: res["payments_in"] += 1
        else: res["payments_out"] += 1

    # --- Cash taken on the doc itself -> synthetic payment with allocation (keeps outstanding exact) ---
    for cp in cash_on_doc:
        pdoc = {"id": new_id(),
                "code": ("PMT-IN-VYCASH-" if cp["is_in"] else "PMT-OUT-VYCASH-") + str(cp["txn_id"]),
                "party_id": "", "party_name": cp["party"], "date": cp["date"], "amount": cp["amount"],
                "allocated_amount": cp["amount"], "payment_type": _map_ptype_payment(cp["ptype_id"]),
                "ref_no": "", "bank_name": "",
                "allocations": [{"document_id": cp["doc_id"], "document_code": cp["doc_code"],
                                 "document_type": "invoice" if cp["is_in"] else "vendor_bill",
                                 "amount": cp["amount"], "tds_amount": 0}],
                "notes": "Cash on document (Vyapar import)", "status": "Used",
                "vyapar_id": f"cash-{cp['txn_id']}", "source": "vyapar", "created_at": now_iso()}
        if not opts.dry_run:
            await (db.payments_in if cp["is_in"] else db.payments_out).update_one(
                {"vyapar_id": pdoc["vyapar_id"]}, {"$setOnInsert": pdoc}, upsert=True)
        res["cash_on_doc_payments"] += 1

    # --- Expenses (txn_type 7; category via txn_category_id -> kb_names name_type=2) ---
    if getattr(opts, "expenses", True):
        cat_map: Dict[int, Dict[str, str]] = {}
        try:
            cur.execute('SELECT name_id, full_name FROM "kb_names" WHERE name_type = 2')
            for r in cur.fetchall():
                nm = (r["full_name"] or "").strip()
                if not nm:
                    continue
                existing = await db.expense_categories.find_one({"name": nm}, {"_id": 0, "id": 1})
                if existing:
                    cid = existing["id"]
                else:
                    cid = new_id()
                    if not opts.dry_run:
                        await db.expense_categories.insert_one(
                            {"id": cid, "name": nm, "classification": "indirect", "created_at": now_iso()})
                cat_map[int(r["name_id"])] = {"id": cid, "name": nm}
        except Exception:
            pass
        cur.execute("SELECT * FROM kb_transactions WHERE txn_type = 7")
        for r in cur.fetchall():
            e_cash = float(r["txn_cash_amount"] or 0); e_bal = float(r["txn_balance_amount"] or 0)
            cat_id = int(r["txn_category_id"]) if r["txn_category_id"] is not None else -1
            cat = cat_map.get(cat_id, {"id": "", "name": "Other"})
            edoc = {"id": new_id(), "code": f"EXP-VY-{r['txn_id']}",
                    "category_id": cat["id"], "category_name": cat["name"],
                    "party_id": "", "party_name": "",
                    "date": str(r["txn_date"] or "")[:19] or now_iso(),
                    "amount": round(e_cash + e_bal, 2), "paid_amount": round(e_cash, 2),
                    "payment_type": _map_ptype_payment(r["txn_payment_type_id"]),
                    "ref_no": (r["txn_ref_number_char"] or "").strip(),
                    "notes": (r["txn_description"] or "").strip(),
                    "status": "Paid" if e_bal <= 0.01 else ("Partial" if e_cash > 0.01 else "Unpaid"),
                    "vyapar_id": str(r["txn_id"]), "source": "vyapar",
                    "created_at": str(r["txn_date_created"] or now_iso())}
            if not opts.dry_run:
                await db.expenses.update_one({"vyapar_id": edoc["vyapar_id"]}, {"$setOnInsert": edoc}, upsert=True)
            res["expenses"] += 1

    # --- Party opening balances: gap between Vyapar's live party balance and ERP-derived outstanding ---
    # (covers pre-history opening balances, advances, unlinked payments, credit-note effects)
    derived_open: Dict[str, Dict[str, float]] = {"customer": {}, "supplier": {}}
    for v in inv_open.values():
        if v[2] > 0.01:
            derived_open[v[0]][v[1]] = derived_open[v[0]].get(v[1], 0.0) + v[2]
    if opts.parties:
        try:
            cur.execute('SELECT full_name, amount FROM "kb_names" WHERE name_type = 1')
            for r in cur.fetchall():
                nm = (r["full_name"] or "").strip()
                if not nm:
                    continue
                vy = float(r["amount"] or 0)
                # Handle BOTH sides for every party (dual-role parties + zero-balance parties with open
                # docs + advances held with the "other" side, e.g. a supplier we overpaid = receivable).
                # Upsert so the record is created on whichever side must carry the balance.
                cust_diff = round(max(vy, 0.0) - derived_open["customer"].get(nm, 0.0), 2)
                if abs(cust_diff) > 0.01:
                    if not opts.dry_run:
                        await db.customers.update_one(
                            {"name": nm},
                            {"$set": {"opening_balance": cust_diff, "vyapar_balance": vy},
                             "$setOnInsert": {"id": new_id(), "source": "vyapar", "created_at": now_iso()}},
                            upsert=True)
                    res["opening_balances"] += 1
                supp_diff = round(max(-vy, 0.0) - derived_open["supplier"].get(nm, 0.0), 2)
                if abs(supp_diff) > 0.01:
                    if not opts.dry_run:
                        await db.suppliers.update_one(
                            {"name": nm},
                            {"$set": {"opening_balance": supp_diff, "vyapar_balance": vy},
                             "$setOnInsert": {"id": new_id(), "source": "vyapar", "created_at": now_iso()}},
                            upsert=True)
                    res["opening_balances"] += 1
        except Exception:
            pass

    # Refresh paid/unpaid status now that allocations exist
    if not opts.dry_run:
        try:
            await _refresh_invoice_paid_status([m["id"] for m in txn_to_doc.values() if m["dtype"] == "invoice"])
            await _refresh_bill_paid_status([m["id"] for m in txn_to_doc.values() if m["dtype"] == "vendor_bill"])
        except Exception:
            pass

    con.close()
    return res

async def _do_import_generic_sqlite(path: Path, opts: VyaparImportIn) -> Dict[str, Any]:
    """Fallback for non-Vyapar SQLite DBs (e.g. an exported generic db). Best-effort."""
    import sqlite3
    res = {"parties": 0, "items": 0, "sales": 0, "purchases": 0}
    con = sqlite3.connect(str(path)); con.row_factory = sqlite3.Row
    cur = con.cursor()
    cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
    tables = {r[0].lower(): r[0] for r in cur.fetchall()}
    def _try_select(candidates: List[str]):
        for c in candidates:
            if c.lower() in tables:
                try:
                    cur.execute(f'SELECT * FROM "{tables[c.lower()]}"'); return cur.fetchall()
                except Exception: pass
        return None
    if opts.parties:
        for r in _try_select(["parties", "party"]) or []:
            d = {k: r[k] for k in r.keys()}; p = _map_row(d, PARTY_COLS)
            if not p.get("name"): continue
            doc = {"id": new_id(), "name": str(p["name"])[:120], "phone": str(p.get("phone","") or ""),
                   "email": str(p.get("email","") or ""), "gstin": str(p.get("gstin","") or ""),
                   "address": str(p.get("address","") or ""), "state": str(p.get("state","") or ""),
                   "source": "vyapar", "created_at": now_iso()}
            if not opts.dry_run:
                await db.customers.update_one({"name": doc["name"]}, {"$setOnInsert": doc}, upsert=True)
            res["parties"] += 1
    con.close()
    return res

async def _do_import_xlsx(path: Path, opts: VyaparImportIn) -> Dict[str, Any]:
    from openpyxl import load_workbook
    res = {"parties": 0, "items": 0, "sales": 0, "purchases": 0}
    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        it = ws.iter_rows(values_only=True)
        try:
            headers_row = next(it)
        except StopIteration:
            continue
        headers = [str(h or "").strip() for h in headers_row]
        if not any(headers): continue
        kind = _classify_sheet(headers)
        rows = []
        for r in it:
            d = {headers[i]: (r[i] if i < len(r) else None) for i in range(len(headers))}
            rows.append(d)
        if kind == "parties" and opts.parties:
            for d in rows:
                p = _map_row(d, PARTY_COLS)
                if not p.get("name"): continue
                doc = {"id": new_id(), "name": str(p["name"])[:120],
                       "phone": str(p.get("phone","") or ""), "email": str(p.get("email","") or ""),
                       "gstin": str(p.get("gstin","") or ""), "address": str(p.get("address","") or ""),
                       "state": str(p.get("state","") or ""), "source": "vyapar", "created_at": now_iso()}
                if not opts.dry_run:
                    await db.customers.update_one({"name": doc["name"]}, {"$setOnInsert": doc}, upsert=True)
                res["parties"] += 1
        elif kind == "items" and opts.items:
            for d in rows:
                it_ = _map_row(d, ITEM_COLS)
                if not it_.get("name"): continue
                doc = {"id": new_id(), "item_code": "VY-" + new_id()[:8].upper(),
                       "name": str(it_["name"])[:160], "hsn": str(it_.get("hsn","") or ""),
                       "unit": str(it_.get("unit","NOS") or "NOS"),
                       "rate": float(it_.get("sale_price") or 0), "stock": float(it_.get("stock") or 0),
                       "reorder_level": 0, "source": "vyapar", "created_at": now_iso()}
                if not opts.dry_run:
                    await db.items.update_one({"name": doc["name"]}, {"$setOnInsert": doc}, upsert=True)
                res["items"] += 1
        elif kind == "sales" and opts.sales:
            for d in rows:
                s = _map_row(d, SALE_COLS)
                if not s.get("code"): continue
                doc = {"id": new_id(), "code": "VY-" + str(s["code"]),
                       "customer_id": "", "customer_name": str(s.get("customer_name","") or "Unknown"),
                       "date": str(s.get("date") or now_iso())[:19],
                       "lines": [], "subtotal": float(s.get("total") or 0),
                       "cgst": 0, "sgst": 0, "igst": 0,
                       "total": float(s.get("total") or 0),
                       "status": "paid" if (float(s.get("balance") or 0) == 0) else "sent",
                       "source": "vyapar", "created_at": now_iso()}
                if not opts.dry_run:
                    await db.invoices.update_one({"code": doc["code"]}, {"$setOnInsert": doc}, upsert=True)
                res["sales"] += 1
        elif kind == "purchases" and opts.purchases:
            for d in rows:
                p = _map_row(d, PURCHASE_COLS)
                if not p.get("code"): continue
                doc = {"id": new_id(), "code": "VY-" + str(p["code"]),
                       "supplier_id": "", "supplier_name": str(p.get("supplier_name","") or "Unknown"),
                       "date": str(p.get("date") or now_iso())[:19],
                       "lines": [], "subtotal": float(p.get("total") or 0),
                       "gst_total": 0, "total": float(p.get("total") or 0),
                       "status": "received", "source": "vyapar", "created_at": now_iso()}
                if not opts.dry_run:
                    await db.purchase_orders.update_one({"code": doc["code"]}, {"$setOnInsert": doc}, upsert=True)
                res["purchases"] += 1
    wb.close()
    return res

_VYAPAR_BG_TASKS: set = set()

async def _run_vyapar_import_job(job_id: str, path: Path, kind: str, payload: "VyaparImportIn", user: dict):
    """Runs the (potentially multi-minute) import off the request/response cycle so Railway's
    edge proxy never kills it with a 503 mid-import. Progress/result is polled via job_id."""
    try:
        p = path
        k = kind
        if k in ("zip", "zip_sqlite"):
            import zipfile, tempfile
            with zipfile.ZipFile(p) as z:
                inner = [n for n in z.namelist() if n.endswith(".db") or n.endswith(".sqlite") or n.endswith(".vyp")]
                if not inner:
                    raise ValueError("Archive doesn't contain a SQLite database.")
                tmpd = Path(tempfile.mkdtemp())
                z.extract(inner[0], tmpd)
                p = tmpd / inner[0]
                k = "sqlite"
        if k == "sqlite":
            details = await _do_import_sqlite(p, payload)
        elif k == "xlsx":
            details = await _do_import_xlsx(p, payload)
        else:
            raise ValueError(f"Cannot import file kind '{k}'. Please export an Excel file from Vyapar (Reports → Sale/Party/Item/Purchase Report → Excel icon).")
        await write_audit(user["name"], "vyapar_import", "import", payload.token,
                          {**details, "dry_run": payload.dry_run})
        bits = [f"{details.get(k2,0)} {k2.replace('_',' ')}" for k2 in
                ("parties","items","sales","purchases","payments_in","payments_out","quotations","sale_orders","purchase_orders","delivery_challans","job_work_out","sale_returns")
                if details.get(k2)]
        summary = " · ".join(bits) or "nothing matched"
        if payload.dry_run: summary += " (dry run)"
        if details.get("company_seeded"): summary += " · company details auto-filled"
        await db.vyapar_import_jobs.update_one({"id": job_id}, {"$set": {
            "status": "done", "finished_at": now_iso(),
            "result": {"ok": True, "summary": summary, "details": details, "dry_run": payload.dry_run},
        }})
    except Exception as e:
        await db.vyapar_import_jobs.update_one({"id": job_id}, {"$set": {
            "status": "error", "finished_at": now_iso(), "error": str(e),
        }})

@api.post("/integrations/vyapar/import")
async def vyapar_import(payload: VyaparImportIn, user=Depends(require_roles("admin"))):
    meta = await db.vyapar_uploads.find_one({"id": payload.token, "user_id": user["id"]}, {"_id": 0})
    if not meta:
        raise HTTPException(404, "Upload not found. Re-upload the file.")
    path = Path(meta["path"])
    if not path.exists():
        raise HTTPException(404, "Uploaded file is no longer on disk. Please re-upload.")
    job_id = new_id()
    await db.vyapar_import_jobs.insert_one({
        "id": job_id, "token": payload.token, "user_id": user["id"], "user_name": user["name"],
        "status": "running", "opts": payload.dict(), "started_at": now_iso(),
    })
    task = asyncio.create_task(_run_vyapar_import_job(job_id, path, meta.get("kind"), payload, user))
    _VYAPAR_BG_TASKS.add(task)
    task.add_done_callback(_VYAPAR_BG_TASKS.discard)
    return {"ok": True, "job_id": job_id, "status": "running"}

@api.get("/integrations/vyapar/import/jobs/{job_id}")
async def vyapar_import_job_status(job_id: str, user=Depends(require_roles("admin"))):
    job = await db.vyapar_import_jobs.find_one({"id": job_id, "user_id": user["id"]}, {"_id": 0})
    if not job:
        raise HTTPException(404, "Job not found")
    return job


class _VyaparBackfillIn(BaseModel):
    token: str

@api.post("/integrations/vyapar/backfill-outstanding")
async def vyapar_backfill_outstanding(payload: _VyaparBackfillIn, user=Depends(require_roles("admin"))):
    """One-off migration for invoices/bills imported before certain fields existed. Sets, via $set
    matched by vyapar_id (safe to re-run, no re-insert/duplication):
    - outstanding = Vyapar's own live txn_current_balance (ground truth)
    - po_date / po_number / ship_to_address / purchaser_name / payment_mode — these existed in the
      original .vyb all along (txn_po_date, txn_po_ref_number, txn_shipping_address, kb_udf_values
      "Purchaser Name", txn_payment_type_id) but weren't mapped by the importer at the time these
      docs were first created, so they were silently missing from the invoice/bill PDF."""
    import sqlite3, zipfile, tempfile
    meta = await db.vyapar_uploads.find_one({"id": payload.token, "user_id": user["id"]}, {"_id": 0})
    if not meta:
        raise HTTPException(404, "Upload not found. Re-upload the file.")
    path = Path(meta["path"])
    if not path.exists():
        raise HTTPException(404, "Uploaded file is no longer on disk. Please re-upload.")
    kind = meta.get("kind")
    if kind in ("zip", "zip_sqlite"):
        with zipfile.ZipFile(path) as z:
            inner = [n for n in z.namelist() if n.endswith(".db") or n.endswith(".sqlite") or n.endswith(".vyp")]
            if not inner:
                raise HTTPException(400, "Archive doesn't contain a SQLite database.")
            tmpd = Path(tempfile.mkdtemp())
            z.extract(inner[0], tmpd)
            path = tmpd / inner[0]

    def _run():
        con = sqlite3.connect(str(path)); con.row_factory = sqlite3.Row
        cur = con.cursor()

        purchaser_lookup: Dict[int, str] = {}
        try:
            cur.execute('SELECT udf_field_id FROM "kb_udf_fields" WHERE udf_field_name = ?', ("Purchaser Name",))
            _pf_ids = [int(r["udf_field_id"]) for r in cur.fetchall()]
            if _pf_ids:
                qmarks = ",".join("?" * len(_pf_ids))
                cur.execute(f'SELECT udf_ref_id, udf_value FROM "kb_udf_values" WHERE udf_value_field_id IN ({qmarks})', _pf_ids)
                for r in cur.fetchall():
                    if r["udf_value"]:
                        purchaser_lookup[int(r["udf_ref_id"])] = str(r["udf_value"]).strip()
        except Exception:
            pass

        ptype_lookup: Dict[int, str] = {}
        try:
            cur.execute('SELECT paymentType_id, paymentType_name FROM "kb_paymentTypes"')
            for r in cur.fetchall():
                ptype_lookup[int(r["paymentType_id"])] = (r["paymentType_name"] or "")
        except Exception:
            pass

        def map_ptype(pid):
            if pid is None:
                return "Credit"
            nm = (ptype_lookup.get(int(pid)) if pid is not None else "") or ""
            nmu = nm.upper()
            if "CASH" in nmu: return "Cash"
            if "CHEQUE" in nmu: return "Cheque"
            if "UPI" in nmu: return "UPI"
            if "CARD" in nmu: return "Card"
            return "Bank Transfer" if nm else "Credit"

        cur.execute("""SELECT txn_id, txn_type, txn_current_balance, txn_po_date, txn_po_ref_number,
                              txn_shipping_address, txn_payment_type_id
                       FROM kb_transactions WHERE txn_type IN (1,2)""")
        out_rows = []
        for r in cur.fetchall():
            txn_id = int(r["txn_id"])
            out_rows.append({
                "txn_id": txn_id, "txn_type": int(r["txn_type"]),
                "cur_bal": float(r["txn_current_balance"] or 0),
                "po_date": str(r["txn_po_date"] or "")[:10],
                "po_number": (r["txn_po_ref_number"] or "").strip(),
                "ship_to_address": (r["txn_shipping_address"] or "").strip(),
                "purchaser_name": purchaser_lookup.get(txn_id, ""),
                "payment_mode": map_ptype(r["txn_payment_type_id"]),
            })
        con.close()
        return out_rows

    rows = await asyncio.to_thread(_run)
    updated_inv = 0; updated_bill = 0
    for row in rows:
        txn_id = row["txn_id"]; t = row["txn_type"]; cur_bal = row["cur_bal"]
        out = round(max(cur_bal, 0.0), 2)
        status = ("paid" if cur_bal <= 0.01 else "sent") if t == 1 else ("received" if cur_bal <= 0.01 else "open")
        coll = db.invoices if t == 1 else db.vendor_bills
        updates = {"outstanding": out, "status": status}
        if row["po_date"]: updates["po_date"] = row["po_date"]
        if row["po_number"]: updates["po_number"] = row["po_number"]
        if row["ship_to_address"]: updates["ship_to_address"] = row["ship_to_address"]
        if row["purchaser_name"]: updates["purchaser_name"] = row["purchaser_name"]
        if row["payment_mode"]: updates["payment_mode"] = row["payment_mode"]
        res = await coll.update_one({"vyapar_id": str(txn_id)}, {"$set": updates})
        if res.matched_count:
            if t == 1: updated_inv += 1
            else: updated_bill += 1
    return {"ok": True, "invoices_updated": updated_inv, "bills_updated": updated_bill, "total_rows": len(rows)}


class _GstLineBackfillIn(BaseModel):
    dry_run: bool = True

_GST_SLABS = [0, 5, 12, 18, 28]

async def _backfill_line_gst_for_collection(coll, dry_run: bool):
    affected = 0
    examples = []
    cursor = coll.find({"source": "vyapar"}, {"_id": 0, "id": 1, "code": 1, "lines": 1, "cgst": 1, "sgst": 1, "igst": 1, "subtotal": 1})
    async for doc in cursor:
        lines = doc.get("lines") or []
        if not lines:
            continue
        doc_tax_total = float(doc.get("cgst", 0) or 0) + float(doc.get("sgst", 0) or 0) + float(doc.get("igst", 0) or 0)
        if doc_tax_total <= 0.01:
            continue  # genuinely a 0-tax / non-GST document — leave alone
        if any(float(l.get("gst_rate", 0) or 0) > 0 for l in lines):
            continue  # lines already carry a real rate — not the bug we're fixing
        subtotal = float(doc.get("subtotal", 0) or 0)
        if subtotal <= 0:
            continue
        raw_rate = doc_tax_total / subtotal * 100
        rate = min(_GST_SLABS, key=lambda s: abs(s - raw_rate))
        if rate <= 0:
            continue
        new_lines = []
        for l in lines:
            qty = float(l.get("qty", 0) or 0); r = float(l.get("rate", 0) or 0)
            disc = float(l.get("discount_amount", 0) or 0)
            taxable = max(qty * r - disc, 0)
            new_lines.append({**l, "gst_rate": rate, "gst_amount": round(taxable * rate / 100, 2)})
        affected += 1
        if len(examples) < 8:
            examples.append({"id": doc.get("id"), "code": doc.get("code"), "rate": rate})
        if not dry_run:
            await coll.update_one({"id": doc["id"]}, {"$set": {"lines": new_lines}})
    return affected, examples

@api.post("/integrations/vyapar/backfill-line-gst")
async def vyapar_backfill_line_gst(payload: _GstLineBackfillIn, user=Depends(require_roles("admin"))):
    """One-off fix for invoices/bills whose line items show 0% GST on the PDF Tax Summary even
    though the document's own cgst/sgst/igst totals (and hence the grand total) are correct.
    Root cause: Vyapar's kb_lineitems.lineitem_tax_amount was often unset, so the importer's
    per-line implied_gst_rate came out 0 even for fully-taxed invoices. This derives the real
    rate from the document-level tax total instead (snapped to the nearest GST slab) and writes
    it onto every line. Pure Mongo — does not need the original .vyb re-uploaded. Call with
    dry_run=true first to see the affected count with no writes; dry_run=false to apply."""
    inv_count, inv_examples = await _backfill_line_gst_for_collection(db.invoices, payload.dry_run)
    bill_count, bill_examples = await _backfill_line_gst_for_collection(db.vendor_bills, payload.dry_run)
    return {"dry_run": payload.dry_run,
            "invoices_affected": inv_count, "invoice_examples": inv_examples,
            "vendor_bills_affected": bill_count, "vendor_bill_examples": bill_examples}


class VyaparReconcileIn(BaseModel):
    token: str

@api.post("/integrations/vyapar/reconcile")
async def vyapar_reconcile(payload: VyaparReconcileIn, user=Depends(require_roles("admin"))):
    """Compare live ERP data against an uploaded Vyapar backup — the cutover safety check.
    Every row shows Vyapar value vs ERP value; all-green means nothing was lost in import."""
    import sqlite3
    meta = await db.vyapar_uploads.find_one({"id": payload.token, "user_id": user["id"]}, {"_id": 0})
    if not meta:
        raise HTTPException(404, "Upload not found. Re-upload the file.")
    path = Path(meta["path"])
    if not path.exists():
        raise HTTPException(404, "Uploaded file is no longer on disk. Please re-upload.")
    kind = meta.get("kind")
    if kind in ("zip", "zip_sqlite"):
        import zipfile, tempfile
        with zipfile.ZipFile(path) as z:
            inner = [n for n in z.namelist() if n.endswith(".db") or n.endswith(".sqlite") or n.endswith(".vyp")]
            if not inner:
                raise HTTPException(400, "Archive doesn't contain a SQLite database.")
            tmpd = Path(tempfile.mkdtemp())
            z.extract(inner[0], tmpd)
            path = tmpd / inner[0]
            kind = "sqlite"
    if kind != "sqlite":
        raise HTTPException(400, "Reconciliation needs a .vyb backup (SQLite), not an Excel export.")

    con = sqlite3.connect(str(path)); con.row_factory = sqlite3.Row
    cur = con.cursor()

    def vy_count_total(ttypes):
        q = ",".join(str(t) for t in ttypes)
        cur.execute(f"SELECT COUNT(*) n, ROUND(SUM(COALESCE(txn_cash_amount,0)+COALESCE(txn_balance_amount,0)),2) s FROM kb_transactions WHERE txn_type IN ({q})")
        r = cur.fetchone()
        return int(r["n"] or 0), float(r["s"] or 0)

    async def erp_count_total(coll):
        docs = await coll.find({"source": "vyapar"}, {"_id": 0, "total": 1, "amount": 1}).to_list(50000)
        return len(docs), round(sum(float(d.get("total", d.get("amount", 0)) or 0) for d in docs), 2)

    rows: List[Dict[str, Any]] = []
    async def add_row(metric, vy_n, vy_s, coll):
        n, s = await erp_count_total(coll)
        rows.append({"metric": metric, "vyapar_count": vy_n, "erp_count": n,
                     "vyapar_total": round(vy_s, 2), "erp_total": s,
                     "ok": (vy_n == n) and abs(vy_s - s) < 1.0})

    await add_row("Sale Invoices", *vy_count_total([1]), db.invoices)
    await add_row("Purchase Bills", *vy_count_total([2]), db.vendor_bills)
    await add_row("Credit Notes (Sale Returns)", *vy_count_total([21]), db.credit_notes)
    await add_row("Purchase Returns", *vy_count_total([23]), db.purchase_returns)
    await add_row("Quotations", *vy_count_total([27]), db.quotations)
    await add_row("Sale Orders", *vy_count_total([24]), db.sale_orders)
    await add_row("Purchase Orders", *vy_count_total([28]), db.purchase_orders)
    await add_row("Delivery Challans", *vy_count_total([30]), db.delivery_challans)
    await add_row("Proforma Invoices", *vy_count_total([65]), db.proforma_invoices)
    await add_row("Expenses", *vy_count_total([7]), db.expenses)

    # Payments: ERP side excludes synthetic cash-on-doc records (vyapar_id starts with "cash-")
    for label, tt, coll in (("Payments In", 3, db.payments_in), ("Payments Out", 4, db.payments_out)):
        vy_n, vy_s = vy_count_total([tt])
        docs = await coll.find({"source": "vyapar", "vyapar_id": {"$not": {"$regex": "^cash-"}}},
                               {"_id": 0, "amount": 1}).to_list(50000)
        s = round(sum(float(d.get("amount") or 0) for d in docs), 2)
        rows.append({"metric": label, "vyapar_count": vy_n, "erp_count": len(docs),
                     "vyapar_total": round(vy_s, 2), "erp_total": s,
                     "ok": (vy_n == len(docs)) and abs(vy_s - s) < 1.0})

    # Parties + items + stock
    cur.execute("SELECT COUNT(*) FROM kb_names WHERE name_type = 1")
    vy_parties = int(cur.fetchone()[0] or 0)
    erp_cust = await db.customers.count_documents({"source": "vyapar"})
    erp_supp = await db.suppliers.count_documents({"source": "vyapar"})
    rows.append({"metric": "Parties (ERP: customers + suppliers, dual-role counted twice)",
                 "vyapar_count": vy_parties, "erp_count": erp_cust + erp_supp,
                 "vyapar_total": None, "erp_total": None, "ok": (erp_cust + erp_supp) >= vy_parties})
    cur.execute("SELECT COUNT(*) n, ROUND(SUM(COALESCE(item_stock_quantity,0)),2) q FROM kb_items WHERE item_is_active IS NULL OR item_is_active != 0")
    r = cur.fetchone(); vy_items, vy_qty = int(r["n"] or 0), float(r["q"] or 0)
    idocs = await db.items.find({"source": "vyapar"}, {"_id": 0, "qty_on_hand": 1}).to_list(50000)
    erp_qty = round(sum(float(d.get("qty_on_hand") or 0) for d in idocs), 2)
    rows.append({"metric": "Items / total stock qty", "vyapar_count": vy_items, "erp_count": len(idocs),
                 "vyapar_total": round(vy_qty, 2), "erp_total": erp_qty,
                 "ok": abs(vy_items - len(idocs)) <= 1 and abs(vy_qty - erp_qty) < 1.0})

    # THE money check: receivable/payable (ERP calc incl. opening balances) vs Vyapar live party balances
    cur.execute("SELECT SUM(CASE WHEN amount>0 THEN amount ELSE 0 END) r, SUM(CASE WHEN amount<0 THEN -amount ELSE 0 END) p FROM kb_names WHERE name_type = 1")
    r = cur.fetchone(); vy_recv, vy_pay = float(r["r"] or 0), float(r["p"] or 0)
    rp = await dashboard_receivable_payable(user)  # reuse the live calc
    rows.append({"metric": "RECEIVABLE (outstanding)", "vyapar_count": None, "erp_count": None,
                 "vyapar_total": round(vy_recv, 2), "erp_total": rp["receivable_total"],
                 "ok": abs(vy_recv - rp["receivable_total"]) < 1.0})
    rows.append({"metric": "PAYABLE (outstanding)", "vyapar_count": None, "erp_count": None,
                 "vyapar_total": round(vy_pay, 2), "erp_total": rp["payable_total"],
                 "ok": abs(vy_pay - rp["payable_total"]) < 1.0})

    con.close()
    all_ok = all(x["ok"] for x in rows)
    await write_audit(user["name"], "vyapar_reconcile", "import", payload.token, {"all_ok": all_ok})
    return {"all_ok": all_ok, "rows": rows,
            "note": "Counts compare ERP records imported from Vyapar (source=vyapar). Money rows compare full ERP outstanding (incl. opening balances) to Vyapar's live party balances."}




# ---------------------------------------------------------------------------
# AI Fixture Concept Generator — part drawing / 3D (STL) → jig & fixture brief
# ---------------------------------------------------------------------------
DENPLEX_FIXTURE_KB = (
    "\n\nDENPLEX HOUSE STYLE (learned from real Denplex fixtures — follow this):\n"
    "CONSTRUCTION: a flat BASE PLATE (MS or aluminium, ~12-20 mm) cut by LASER or WATERJET with "
    "rectangular LIGHTENING CUTOUTS to keep weight low. Onto it dowel + bolt several VMC-machined "
    "VERTICAL CRADLE POSTS (zinc-plated steel) whose tops have a V-groove or forked/slotted profile "
    "that cradles the tube/part at each node or bend. Posts located by dowel pins, fixed with socket-head "
    "cap screws into tapped holes. Add a stamped PART-NUMBER TAG.\n"
    "IN-HOUSE PROCESSES (prefer these, keep it simple & cost-effective): laser cutting, waterjet, VMC "
    "milling, lathe/turning, drilling, tapping, cutting, shaping, milling, zinc electroplating.\n"
    "BRAZING-FIXTURE RULES (critical when the operation is brazing):\n"
    "1) Support the tube only at straights/nodes with V-cradles — keep every BRAZE JOINT in free air so "
    "the torch reaches all sides of the joint.\n"
    "2) Present each joint TILTED/ANGLED so molten flux and excess filler RUN OFF and never pool/leave "
    "residue at the joint.\n"
    "3) Minimise contact and heat-sinking near a joint; where contact is unavoidable use STAINLESS STEEL "
    "or coated contact so the fixture does not braze to the part or rob heat. Never clamp on a joint.\n"
    "4) Allow thermal expansion — locate from one datum end and let the rest float; do not over-constrain "
    "a part that will be heated.\n"
    "5) Keep it LIGHT (lightening cutouts) and LOW-COST (few machined parts, standard fasteners).\n"
    "BOM STYLE: base plate (laser/waterjet), set of V-cradle posts (VMC, zinc-plated, qty = number of "
    "support nodes), dowel pins, SHCS, optional rest pads, part-number tag.\n"
    "DRAWING STYLE: A3, dimensions in mm, ISO 2768 medium tolerance, title block (Checked: Neel), "
    "part-number naming like 'A163833T2A - <part> - FX<n> - PR<nnn>'."
)

FIXTURE_SYSTEM = (
    "You are a senior jig & fixture design engineer at Denplex Engineering Company, a precision "
    "machining / jigs & fixtures manufacturer in Ahmedabad, India. From the part drawing/photo and "
    "the inputs given, propose a practical FIRST-DRAFT FIXTURE CONCEPT (a brief, not CAD). "
    "Apply 3-2-1 locating principles. Prefer standard off-the-shelf components (toggle clamps, "
    "round + diamond locating pins, rest pads, dowel pins, hydraulic/pneumatic cylinders, "
    "support buttons, eye bolts). Be specific, realistic and cost-aware for an Indian SME shop. "
    "Money in INR. Return ONLY JSON, no prose." + DENPLEX_FIXTURE_KB + (
        "\n\nGROUND EVERYTHING IN THE ACTUAL PART. Study the attached rendered 3D views and the geometry "
        "data (bounding box in mm, cylindrical faces, hole/round diameters). Count the real bends / branches / "
        "tube ends of the part and provide ONE V-cradle support per node — so standard_components lists the "
        "correct QUANTITY of cradle posts (e.g. '6'), not a vague number. Size the base plate about 30-50 mm "
        "larger than the part's X-Y bounding box, and give approximate post heights from the Z extent. Quote the "
        "real dimensions (mm) from the geometry inside locating_scheme, supports and base_plate. Be specific to "
        "THIS part — never generic."
    )
)
FIXTURE_SCHEMA = (
    'Return ONLY this JSON: {"fixture_type":str, "summary":str, '
    '"geometry":{'
    '"base_plate":{"length_mm":number,"width_mm":number,"thickness_mm":number,'
    '"cutouts":[{"x_mm":number,"y_mm":number,"w_mm":number,"h_mm":number}]},'
    '"posts":[{"x_mm":number,"y_mm":number,"height_mm":number,"width_mm":number,'
    '"top":"v_groove"|"flat"|"pin","pin_diameter_mm":number}],'
    '"clamps":[{"x_mm":number,"y_mm":number,"height_mm":number}],'
    '"part_proxy":{"type":"tube"|"block"|"none","length_mm":number,"diameter_mm":number,'
    '"x_mm":number,"y_mm":number,"z_mm":number,"axis":"x"|"y"}'
    '}, '
    '"locating_scheme":str, "locators":[str], "clamping":[str], "supports":[str], '
    '"base_plate":str, "actuation":str, '
    '"standard_components":[{"item":str,"qty":str,"note":str}], '
    '"access_and_clearance":[str], "distortion_risks":[str], "inspection_points":[str], '
    '"estimated_build":{"cost_inr":number,"time_days":number}, "assumptions":str}. '
    "fixture_type = the recommended type (e.g. Milling fixture, Drilling jig, Welding jig, "
    "Hydraulic machining fixture, Leak-test fixture, Inspection fixture). "
    "locating_scheme = how the 3-2-1 datums are established. standard_components = a BOM of "
    "off-the-shelf parts with quantities. "
    "GEOMETRY is a SEPARATE, numeric, machine-usable version of the same concept, used to build a "
    "real (simplified) 3D render — never omit it, never leave it as placeholder zeros. KEEP IT SMALL: "
    "posts array MUST have AT MOST 6 entries and cutouts MUST have AT MOST 4 entries, total — if the "
    "part has more real support nodes than that, merge the least critical ones rather than listing "
    "more (this is a rough concept render, not a manufacturing model; going over this limit will "
    "truncate your response and break it, so stay under it even if standard_components mentions a "
    "higher post count in the text BOM). Coordinate convention: origin (0,0,0) at one corner of the "
    "base plate's bottom face; X runs along the plate's length, Y along its width, Z is up. "
    "base_plate.thickness_mm typically 12-20. Each post's x_mm/y_mm is the CENTRE of its footprint on "
    "the plate; height_mm is measured from the plate's TOP face. part_proxy approximates the real "
    "part resting in the fixture using its actual bounding box/geometry from the inputs — 'tube' for "
    "round/tubular parts, 'block' for prismatic parts; set type='none' only if neither approximation "
    "is reasonable. Write the geometry object BEFORE the longer text fields (access_and_clearance, "
    "distortion_risks, inspection_points, assumptions) in your JSON so it's never at risk of being cut "
    "off if you run long."
)

class FixtureConceptIn(BaseModel):
    image_base64: str = ""
    mime: str = "image/png"
    stl_base64: str = ""
    step_base64: str = ""
    part_name: str = ""
    material: str = ""
    fixture_type: str = ""
    operation: str = ""
    machine: str = ""
    qty: int = 1
    datums: str = ""
    notes: str = ""

def _stl_bbox(b64: str):
    import base64 as _b64, struct
    raw = _b64.b64decode((b64 or "").split(",")[-1])
    xs = []; ys = []; zs = []
    if len(raw) > 84:
        try:
            n = struct.unpack_from("<I", raw, 80)[0]
            if 84 + 50 * n == len(raw) and n > 0:
                off = 84
                for _ in range(n):
                    base = off + 12
                    for v in range(3):
                        x, y, z = struct.unpack_from("<fff", raw, base + v * 12)
                        xs.append(x); ys.append(y); zs.append(z)
                    off += 50
        except Exception:
            xs = []
    if not xs:
        txt = raw.decode("utf-8", "ignore")
        for line in txt.splitlines():
            line = line.strip()
            if line.startswith("vertex"):
                p = line.split()
                if len(p) >= 4:
                    try:
                        xs.append(float(p[1])); ys.append(float(p[2])); zs.append(float(p[3]))
                    except Exception:
                        pass
    if not xs:
        raise ValueError("Could not read STL vertices")
    dx = round(max(xs) - min(xs), 2); dy = round(max(ys) - min(ys), 2); dz = round(max(zs) - min(zs), 2)
    return {"x": dx, "y": dy, "z": dz, "bbox_volume_cm3": round(dx * dy * dz / 1000.0, 1)}

@api.post("/fixture/concept")
async def fixture_concept(inp: FixtureConceptIn, user=Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "AI fixture generator is not configured. Set ANTHROPIC_API_KEY in the backend environment.")
    dims = None
    cad_views = []
    if inp.stl_base64:
        try:
            dims = _stl_bbox(inp.stl_base64)
        except Exception:
            dims = None
    # STEP file → real geometry + rendered views via the CAD microservice (if configured)
    if inp.step_base64 and CAD_SERVICE_URL:
        try:
            async with httpx.AsyncClient(timeout=180) as cx:
                rr = await cx.post(f"{CAD_SERVICE_URL}/analyze", json={"step_base64": inp.step_base64, "views": 3})
            if rr.status_code < 400:
                cad = rr.json()
                g = cad.get("geometry") or {}
                if g.get("bbox_mm"):
                    bx = g["bbox_mm"]; dims = {"x": bx.get("x"), "y": bx.get("y"), "z": bx.get("z"),
                                               "bbox_volume_cm3": g.get("volume_cm3")}
                dims = dims or {}
                if isinstance(dims, dict):
                    dims["cad"] = g
                cad_views = (cad.get("views") or [])[:3]
        except Exception:
            cad_views = []
    ctx = (f"Part name: {inp.part_name or 'unknown'}\nMaterial: {inp.material or 'unknown'}\n"
           f"Desired fixture type: {inp.fixture_type or '(recommend the best)'}\n"
           f"Machining/usage operation: {inp.operation or 'unknown'}\nMachine: {inp.machine or 'unknown'}\n"
           f"Batch quantity: {inp.qty}\nDatums / critical features: {inp.datums or 'not specified'}\n"
           f"Extra notes: {inp.notes or 'none'}")
    if dims:
        ctx += f"\n3D model bounding box (mm): X={dims.get('x')}, Y={dims.get('y')}, Z={dims.get('z')} (approx envelope {dims.get('bbox_volume_cm3')} cc)"
        cg = dims.get("cad") or {}
        if cg.get("planar_faces") is not None:
            ctx += f"\nCAD geometry: {cg.get('planar_faces')} planar faces, {cg.get('cylindrical_faces')} cylindrical faces"
        if cg.get("hole_or_round_diameters_mm"):
            ctx += f"; round/hole diameters (mm): {cg['hole_or_round_diameters_mm']}"
        if cad_views:
            ctx += f"\n({len(cad_views)} rendered CAD views of the part are attached.)"
    content = []
    if inp.image_base64:
        b64 = inp.image_base64.split(",")[-1]
        mt = inp.mime or "image/png"
        if mt == "application/pdf":
            content.append({"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}})
        else:
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": b64}})
    for v in cad_views:
        content.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": v}})
    content.append({"type": "text", "text": FIXTURE_SCHEMA + "\n\nINPUTS:\n" + ctx})
    body = {"model": QC_VISION_MODEL, "max_tokens": 6000, "system": FIXTURE_SYSTEM, "messages": [{"role": "user", "content": content}]}
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=120) as cx:
            r = await cx.post(f"{ANTHROPIC_BASE_URL}/v1/messages", json=body, headers=headers)
    except Exception as e:
        raise HTTPException(502, f"Could not reach the AI API: {e}")
    if r.status_code >= 400:
        raise HTTPException(502, f"AI API error {r.status_code}: {r.text[:200]}")
    data = r.json()
    text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
    concept = _loads_tolerant(text) or {}

    # Turn the concept's numeric "geometry" into a real (simplified) 3D CAD render via the CAD
    # microservice — a genuine solid model of the PROPOSED FIXTURE, not just the part, so the PDF
    # shows real geometry instead of relying solely on the AI-drawn sketch.
    fixture_cad_views = []
    geo = concept.get("geometry") if isinstance(concept, dict) else None
    if geo and CAD_SERVICE_URL:
        try:
            # Defensive clamp: even with the prompt's cap, never send a runaway number of
            # features to the CAD builder (keeps boolean-op time/stability bounded).
            if isinstance(geo.get("posts"), list):
                geo["posts"] = geo["posts"][:6]
            if isinstance(geo.get("clamps"), list):
                geo["clamps"] = geo["clamps"][:2]
            bp = geo.get("base_plate")
            if isinstance(bp, dict) and isinstance(bp.get("cutouts"), list):
                bp["cutouts"] = bp["cutouts"][:4]
            async with httpx.AsyncClient(timeout=90) as cx:
                fr = await cx.post(f"{CAD_SERVICE_URL}/fixture-build", json={**geo, "views": 3})
            if fr.status_code < 400:
                fixture_cad_views = (fr.json().get("views") or [])[:3]
        except Exception:
            fixture_cad_views = []

    return {"concept": concept, "dims": dims, "views": cad_views,
            "fixture_cad_views": fixture_cad_views, "raw": text if not concept else ""}

@api.post("/fixture/concept/pdf")
async def fixture_concept_pdf(body: dict, user=Depends(get_current_user)):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    c = (body or {}).get("concept") or {}
    meta = (body or {}).get("meta") or {}
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=16 * mm, bottomMargin=16 * mm, leftMargin=18 * mm, rightMargin=18 * mm)
    ss = getSampleStyleSheet()
    h = ParagraphStyle("h", parent=ss["Heading2"], textColor=colors.HexColor("#DC2626"), spaceBefore=8, spaceAfter=3)
    el = [Paragraph("<b>Fixture Concept Brief</b>", ss["Title"]),
          Paragraph(f"Denplex Engineering Company &nbsp;·&nbsp; {meta.get('part_name','')} &nbsp;·&nbsp; {c.get('fixture_type','')}", ss["Normal"]),
          Spacer(1, 4 * mm)]
    # Real (simplified) 3D render of the PROPOSED FIXTURE, built from the concept's numeric geometry
    # via CadQuery — genuine solid geometry, not an AI hand-drawn sketch. Shown first as the primary
    # picture; the AI sketch below is only shown as a fallback when this isn't available.
    fixture_cad_views_b64 = (body or {}).get("fixture_cad_views_base64") or []
    if fixture_cad_views_b64:
        try:
            from reportlab.platypus import Image as RLImage
            import base64 as _b64
            labels = ["Isometric", "Top", "Right"]
            thumbs = []
            for i, v in enumerate(fixture_cad_views_b64[:3]):
                png = _b64.b64decode((v or "").split(",")[-1])
                img = RLImage(io.BytesIO(png))
                maxw = 82 * mm
                if img.drawWidth > maxw:
                    ratio = maxw / img.drawWidth; img.drawWidth = maxw; img.drawHeight *= ratio
                thumbs.append([img, Paragraph(labels[i] if i < len(labels) else f"View {i+1}", ss["Normal"])])
            if thumbs:
                trow = [[t[0] for t in thumbs], [t[1] for t in thumbs]]
                t = Table(trow, colWidths=[86 * mm] * len(thumbs))
                t.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
                el += [Paragraph("Proposed fixture — 3D concept render", h), t, Spacer(1, 4 * mm)]
        except Exception:
            fixture_cad_views_b64 = []
    # Real rendered views of the actual uploaded part (from the STEP/CAD service) — genuine geometry,
    # not an AI guess. Shown first since it's the most trustworthy picture in the document.
    cad_views_b64 = (body or {}).get("cad_views_base64") or []
    if cad_views_b64:
        try:
            from reportlab.platypus import Image as RLImage
            import base64 as _b64
            labels = ["Isometric", "Top", "Right"]
            thumbs = []
            for i, v in enumerate(cad_views_b64[:3]):
                png = _b64.b64decode((v or "").split(",")[-1])
                img = RLImage(io.BytesIO(png))
                maxw = 54 * mm
                if img.drawWidth > maxw:
                    ratio = maxw / img.drawWidth; img.drawWidth = maxw; img.drawHeight *= ratio
                thumbs.append([img, Paragraph(labels[i] if i < len(labels) else f"View {i+1}", ss["Normal"])])
            if thumbs:
                trow = [[t[0] for t in thumbs], [t[1] for t in thumbs]]
                t = Table(trow, colWidths=[58 * mm] * len(thumbs))
                t.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
                el += [Paragraph("Actual part — rendered from the uploaded 3D file", h), t, Spacer(1, 4 * mm)]
        except Exception:
            pass
    # Optional concept sketch (rasterized PNG passed from the browser) — AI-drawn schematic of the
    # fixture. Only included when the real 3D render above isn't available, to avoid showing a rough
    # hand-drawn sketch alongside genuine solid geometry.
    spng = (body or {}).get("sketch_png_base64") or "" if not fixture_cad_views_b64 else ""
    if spng:
        try:
            from reportlab.platypus import Image as RLImage
            import base64 as _b64
            png = _b64.b64decode(spng.split(",")[-1])
            img = RLImage(io.BytesIO(png));
            maxw = 170 * mm
            if img.drawWidth > maxw:
                ratio = maxw / img.drawWidth; img.drawWidth = maxw; img.drawHeight *= ratio
            el += [Paragraph("Concept sketch (AI schematic — not to scale)", h), img, Spacer(1, 4 * mm)]
        except Exception:
            pass
    if c.get("summary"):
        el += [Paragraph("Summary", h), Paragraph(str(c["summary"]), ss["Normal"])]
    def bullets(title, items):
        items = [str(x) for x in (items or []) if str(x).strip()]
        if not items:
            return
        el.append(Paragraph(title, h))
        el.append(ListFlowable([ListItem(Paragraph(i, ss["Normal"])) for i in items], bulletType="bullet", leftIndent=10))
    if c.get("locating_scheme"):
        el += [Paragraph("Locating scheme (3-2-1)", h), Paragraph(str(c["locating_scheme"]), ss["Normal"])]
    bullets("Locators", c.get("locators"))
    bullets("Clamping", c.get("clamping"))
    bullets("Supports", c.get("supports"))
    if c.get("base_plate"):
        el += [Paragraph("Base plate", h), Paragraph(str(c["base_plate"]), ss["Normal"])]
    if c.get("actuation"):
        el += [Paragraph("Actuation", h), Paragraph(str(c["actuation"]), ss["Normal"])]
    sc = c.get("standard_components") or []
    if sc:
        el.append(Paragraph("Standard components (BOM)", h))
        rows = [["Item", "Qty", "Note"]] + [[str(x.get("item", "")), str(x.get("qty", "")), str(x.get("note", ""))] for x in sc]
        t = Table(rows, colWidths=[70 * mm, 20 * mm, 80 * mm])
        t.setStyle(TableStyle([("FONTSIZE", (0, 0), (-1, -1), 8), ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
                               ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                               ("VALIGN", (0, 0), (-1, -1), "TOP")]))
        el.append(t)
    bullets("Access & clearance", c.get("access_and_clearance"))
    bullets("Distortion / risks", c.get("distortion_risks"))
    bullets("Inspection points", c.get("inspection_points"))
    eb = c.get("estimated_build") or {}
    if eb:
        el += [Paragraph("Estimated build", h), Paragraph(f"Approx cost ₹{eb.get('cost_inr','-')} &nbsp;·&nbsp; lead time {eb.get('time_days','-')} days", ss["Normal"])]
    if c.get("assumptions"):
        el += [Paragraph("Assumptions", h), Paragraph(str(c["assumptions"]), ss["Italic"])]
    el += [Spacer(1, 6 * mm), Paragraph("AI-generated concept — engineer review required before design release.", ss["Italic"])]
    doc.build(el); buf.seek(0)
    return Response(content=buf.getvalue(), media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="FixtureConcept_{meta.get("part_name","part")}.pdf"'})

class CadGlbIn(BaseModel):
    step_base64: str = ""

@api.post("/cad/glb")
async def cad_glb(inp: CadGlbIn, user=Depends(get_current_user)):
    """Proxy a STEP to the CAD microservice and return a GLB mesh for the in-ERP 3D viewer."""
    if not CAD_SERVICE_URL:
        raise HTTPException(503, "CAD viewer not configured. Set CAD_SERVICE_URL in the backend environment.")
    if not inp.step_base64:
        raise HTTPException(400, "No STEP file provided")
    try:
        async with httpx.AsyncClient(timeout=180) as cx:
            rr = await cx.post(f"{CAD_SERVICE_URL}/analyze", json={"step_base64": inp.step_base64, "views": 0})
    except Exception as e:
        raise HTTPException(502, f"Could not reach the CAD service: {e}")
    if rr.status_code >= 400:
        raise HTTPException(502, f"CAD service error {rr.status_code}: {rr.text[:200]}")
    data = rr.json()
    return {"mesh_base64": data.get("mesh_base64", ""), "mesh_format": data.get("mesh_format", "stl"), "geometry": data.get("geometry", {})}

SKETCH_SYSTEM = (
    "You are a jig & fixture designer. Produce a CLEAN, LABELLED pseudo-3D ISOMETRIC concept sketch of the "
    "whole proposed fixture ASSEMBLY as a SINGLE self-contained SVG (no external refs, no scripts) — a single "
    "picture that reads as a 3D object, NOT a flat orthographic plan/front pair. "
    "ISOMETRIC CONSTRUCTION (follow exactly): work on one shared isometric grid for the whole drawing — "
    "vertical axis straight up, and two horizontal axes each 30 degrees off horizontal (one rising to the "
    "left, one rising to the right). For EVERY box-like part of the fixture (base plate slab, posts, blocks, "
    "clamp bodies, tags) draw it as a small solid: its TOP face, its LEFT face and its RIGHT face, each as a "
    "separate <polygon> whose 4 corners are computed by moving along those isometric axes (never a plain "
    "rectangle) — TOP face lightest fill (#eee), LEFT face mid fill (#ccc), RIGHT face darkest fill (#999), "
    "all with a thin dark outline (#222). Round or cylindrical elements (pins, tube/pipe stock) may be drawn "
    "as simple ellipses/capsule shapes consistent with the same iso angle. Keep every element's size and "
    "position consistent with the shared grid so the assembly reads correctly: wide flat base plate at the "
    "bottom, posts/blocks/locators standing up from it, the part nested into the locators/clamps above the "
    "base plate, clamp bodies pressing down onto the part. "
    "Show the base plate, the part in place, locating elements (V-blocks / locating pins / rest pads), "
    "clamps, and any process-specific elements (for brazing: the nest and heat shield). Add short leader-line "
    "labels for each component (base plate, each locator/post, each clamp, part), a couple of short dimension "
    "hints, and a small title block (part name + fixture type) in a flat 2D corner box (title block itself "
    "does not need to be isometric). Use thin black strokes (#222), the grey iso-face fills above, an accent "
    "red (#DC2626) for clamps, white background, NO photorealism, no gradients/filters. Use viewBox='0 0 900 "
    "640' (allow extra height for iso depth) and orient the assembly so it reads left-to-right / low-to-high "
    "in a natural 3-quarter view. "
    "For a DENPLEX BRAZING FIXTURE specifically: iso-draw a long flat BASE PLATE slab with a few rectangular "
    "LIGHTENING CUTOUTS visible on its top face, and several VERTICAL CRADLE POSTS (small iso boxes with a "
    "V-notch cut into the top face) standing on it; draw the copper tube/header as an elongated capsule "
    "resting elevated across the V-cradles, tilted slightly, with its braze joints shown in free air between "
    "posts, plus a couple of small arrows marked 'torch access' / 'flux run-off'. Label the base, V-cradle "
    "posts, dowels and the part-number tag. "
    "Keep the SVG COMPACT — at most a few dozen shapes, plain <polygon>/<rect>/<line>/<path>/<ellipse>/<text>, "
    "no embedded fonts — so the whole drawing fits in one response and is valid. "
    "MATCH THE REAL PART: from the attached rendered views, keep the part's true proportions (use the "
    "bounding-box aspect ratio) when placing it in the isometric scene, make the base-plate footprint match "
    "the part's plan aspect ratio, and place a cradle post/locator under EACH visible node / bend / end of "
    "the part (same count as the brief). "
    "Return ONLY the SVG markup starting with <svg and ending with </svg> — no markdown, no commentary."
)

class FixtureSketchIn(BaseModel):
    concept: dict = {}
    part_name: str = ""
    material: str = ""
    dims: Optional[dict] = None
    image_base64: str = ""
    mime: str = "image/png"
    views: list = []   # rendered CAD view PNGs (base64) so the sketch matches the real part

@api.post("/fixture/sketch")
async def fixture_sketch(inp: FixtureSketchIn, user=Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "AI fixture generator is not configured. Set ANTHROPIC_API_KEY in the backend environment.")
    c = inp.concept or {}
    ctx = (f"Part: {inp.part_name or 'part'}; Material: {inp.material or 'n/a'}; "
           f"Fixture type: {c.get('fixture_type', '')}.\n"
           f"Locating scheme: {c.get('locating_scheme', '')}\n"
           f"Locators: {c.get('locators')}\nClamping: {c.get('clamping')}\n"
           f"Supports: {c.get('supports')}\nBase plate: {c.get('base_plate')}")
    if inp.dims:
        ctx += f"\nPart envelope (mm): {inp.dims}"
    content = []
    if inp.image_base64:
        b64 = inp.image_base64.split(",")[-1]; mt = inp.mime or "image/png"
        if mt == "application/pdf":
            content.append({"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}})
        else:
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": b64}})
    for v in (inp.views or [])[:3]:
        content.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": v}})
    if inp.views:
        ctx += "\n(Rendered views of the actual part are attached — trace the real part outline from them and build the fixture around that exact shape.)"
    content.append({"type": "text", "text": "Draw the fixture concept schematic for:\n" + ctx})
    body = {"model": QC_VISION_MODEL, "max_tokens": 8000, "system": SKETCH_SYSTEM, "messages": [{"role": "user", "content": content}]}
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=150) as cx:
            r = await cx.post(f"{ANTHROPIC_BASE_URL}/v1/messages", json=body, headers=headers)
    except Exception as e:
        raise HTTPException(502, f"Could not reach the AI API: {e}")
    if r.status_code >= 400:
        raise HTTPException(502, f"AI API error {r.status_code}: {r.text[:200]}")
    text = "".join(b.get("text", "") for b in r.json().get("content", []) if b.get("type") == "text")
    i = text.find("<svg")
    if i == -1:
        svg = ""
    else:
        j = text.rfind("</svg>")
        svg = text[i:j + 6] if j != -1 else (text[i:].rstrip() + "</svg>")   # best-effort close if truncated
    svg = re.sub(r"<script.*?</script>", "", svg, flags=re.S | re.I)
    return {"svg": svg, "ok": bool(svg), "hint": ("" if svg else text[:300])}

# ---------------- Trial Signup ----------------
@api.post("/trial/request")
async def submit_trial_request(payload: TrialRequestIn):
    if await db.trial_requests.find_one({"email": payload.email.lower(), "status": {"$in": ["pending", "approved"]}}):
        raise HTTPException(400, "A trial request for this email already exists. Please contact admin@denplex.co.")
    doc = {
        "id": new_id(),
        "name": payload.name,
        "company": payload.company,
        "phone": payload.phone,
        "email": payload.email.lower(),
        "gstin": payload.gstin or "",
        "business_type": payload.business_type or "",
        "purpose": payload.purpose or "",
        "status": "pending",
        "created_at": now_iso(),
        "reviewed_at": "",
        "reviewed_by": "",
        "review_note": "",
        "approved_user_id": "",
        "trial_expires_at": "",
        "temp_password": "",
    }
    await db.trial_requests.insert_one(doc)
    return {"ok": True, "id": doc["id"], "message": "Thank you! Your trial request has been received. We'll verify and email you within 24 hours."}

@api.get("/trial/requests")
async def list_trial_requests(status: Optional[str] = None, user=Depends(require_roles("admin"))):
    q = {"status": status} if status else {}
    rows = await db.trial_requests.find(q, {"_id": 0}).sort("created_at", -1).to_list(500)
    return rows

@api.post("/trial/requests/{rid}/approve")
async def approve_trial(rid: str, payload: TrialApproveIn, user=Depends(require_roles("admin"))):
    req = await db.trial_requests.find_one({"id": rid}, {"_id": 0})
    if not req:
        raise HTTPException(404, "Request not found")
    if req["status"] != "pending":
        raise HTTPException(400, f"Request already {req['status']}")
    if await db.users.find_one({"email": req["email"]}):
        raise HTTPException(400, "A user with this email already exists")
    temp_pw = "trial-" + secrets.token_urlsafe(6)
    expires = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
    new_user_id = new_id()
    await db.users.insert_one({
        "id": new_user_id,
        "name": req["name"],
        "email": req["email"],
        "role": "trial",
        "password": hash_password(temp_pw),
        "trial_expires_at": expires,
        "created_at": now_iso(),
        "company": req.get("company", ""),
        "phone": req.get("phone", ""),
    })
    await db.trial_requests.update_one(
        {"id": rid},
        {"$set": {
            "status": "approved",
            "reviewed_at": now_iso(),
            "reviewed_by": user["name"],
            "review_note": payload.note or "",
            "approved_user_id": new_user_id,
            "trial_expires_at": expires,
            "temp_password": temp_pw,
        }},
    )
    await write_audit(user["name"], "trial_approved", "trial_request", rid, {"email": req["email"]})
    return {"ok": True, "email": req["email"], "temp_password": temp_pw, "trial_expires_at": expires}

@api.post("/trial/requests/{rid}/reject")
async def reject_trial(rid: str, payload: TrialApproveIn, user=Depends(require_roles("admin"))):
    req = await db.trial_requests.find_one({"id": rid}, {"_id": 0})
    if not req:
        raise HTTPException(404, "Request not found")
    if req["status"] != "pending":
        raise HTTPException(400, f"Request already {req['status']}")
    await db.trial_requests.update_one(
        {"id": rid},
        {"$set": {"status": "rejected", "reviewed_at": now_iso(), "reviewed_by": user["name"], "review_note": payload.note or ""}},
    )
    await write_audit(user["name"], "trial_rejected", "trial_request", rid)
    return {"ok": True}

@api.delete("/trial/requests/{rid}")
async def del_trial_request(rid: str, user=Depends(require_roles("admin"))):
    await db.trial_requests.delete_one({"id": rid})
    return {"ok": True}


# ---------------- Payment In / Out / Expenses endpoints ----------------
def _recalc_payment_status(p: Dict[str, Any]) -> Dict[str, Any]:
    """Compute allocated_amount + status from allocations list."""
    allocs = p.get("allocations") or []
    total_alloc = sum(float(a.get("amount") or 0) for a in allocs)
    total = float(p.get("amount") or 0)
    p["allocated_amount"] = total_alloc
    if total_alloc <= 0.01:
        p["status"] = "Unused"
    elif total_alloc < total - 0.01:
        p["status"] = "Partially Used"
    else:
        p["status"] = "Used"
    return p

async def _settled_per_invoice(invoice_ids=None):
    """Map invoice_id -> total settled (cash allocations + TDS deducted) across all payments-in."""
    q = {"allocations.document_id": {"$in": list(invoice_ids)}} if invoice_ids else {}
    pays = await db.payments_in.find(q, {"_id": 0, "allocations": 1}).to_list(20000)
    settled: Dict[str, float] = {}
    for p in pays:
        for a in (p.get("allocations") or []):
            if a.get("document_type") == "invoice":
                settled[a["document_id"]] = settled.get(a["document_id"], 0) + float(a.get("amount") or 0) + float(a.get("tds_amount") or 0)
    return settled

async def _refresh_invoice_paid_status(invoice_ids):
    invoice_ids = [i for i in (invoice_ids or []) if i]
    if not invoice_ids:
        return
    settled = await _settled_per_invoice(invoice_ids)
    for iid in invoice_ids:
        inv = await db.invoices.find_one({"id": iid}, {"_id": 0, "total": 1})
        if not inv:
            continue
        if settled.get(iid, 0) >= float(inv.get("total", 0)) - 0.01:
            await db.invoices.update_one({"id": iid}, {"$set": {"status": "paid"}})
        else:
            await db.invoices.update_one({"id": iid, "status": "paid"}, {"$set": {"status": "sent"}})

@api.get("/payments-in/open-invoices/{party_id}")
async def open_invoices_for_party(party_id: str, user=Depends(get_current_user)):
    """A customer's invoices with a remaining balance — for allocating a payment (with optional TDS adjustment)."""
    invs = await db.invoices.find({"customer_id": party_id}, {"_id": 0}).to_list(10000)
    settled = await _settled_per_invoice()
    out = []
    for inv in invs:
        bal = float(inv.get("total", 0)) - settled.get(inv.get("id", ""), 0)
        if bal > 0.01:
            out.append({"id": inv.get("id"), "code": inv.get("code"), "date": str(inv.get("date", ""))[:10],
                        "total": round(float(inv.get("total", 0)), 2), "outstanding": round(bal, 2)})
    out.sort(key=lambda r: r["date"])
    return out

@api.post("/payments-in")
async def create_payment_in(p: PaymentIn, user=Depends(get_current_user)):
    doc = p.model_dump()
    doc["code"] = doc.get("code") or await gen_code("PMT-IN", "payment_in")
    _recalc_payment_status(doc)
    await db.payments_in.insert_one(doc)
    serialize(doc)
    await _refresh_invoice_paid_status([a.get("document_id") for a in (doc.get("allocations") or []) if a.get("document_type") == "invoice"])
    await write_audit(user.get("name", ""), "payment_in_created", "payment_in", doc["id"], {"amount": doc["amount"], "party": doc["party_name"]})
    return doc

@api.get("/payments-in")
async def list_payments_in(user=Depends(get_current_user)):
    return await list_collection(db.payments_in)

@api.put("/payments-in/{pid}")
async def update_payment_in(pid: str, p: PaymentIn, user=Depends(get_current_user)):
    data = p.model_dump(); data.pop("id", None); data.pop("created_at", None)
    _recalc_payment_status(data)
    await db.payments_in.update_one({"id": pid}, {"$set": data})
    await _refresh_invoice_paid_status([a.get("document_id") for a in (data.get("allocations") or []) if a.get("document_type") == "invoice"])
    return {"ok": True}

@api.delete("/payments-in/{pid}")
async def del_payment_in(pid: str, user=Depends(require_roles("admin", "accountant", "ca"))):
    await db.payments_in.delete_one({"id": pid})
    return {"ok": True}

async def _settled_per_bill(bill_ids=None):
    """Map vendor_bill id -> total settled (cash + TDS deducted) across all payments-out."""
    q = {"allocations.document_id": {"$in": list(bill_ids)}} if bill_ids else {}
    pays = await db.payments_out.find(q, {"_id": 0, "allocations": 1}).to_list(20000)
    settled: Dict[str, float] = {}
    for p in pays:
        for a in (p.get("allocations") or []):
            if a.get("document_type") == "vendor_bill":
                settled[a["document_id"]] = settled.get(a["document_id"], 0) + float(a.get("amount") or 0) + float(a.get("tds_amount") or 0)
    return settled

async def _refresh_bill_paid_status(bill_ids):
    bill_ids = [i for i in (bill_ids or []) if i]
    if not bill_ids:
        return
    settled = await _settled_per_bill(bill_ids)
    for bid in bill_ids:
        b = await db.vendor_bills.find_one({"id": bid}, {"_id": 0, "total": 1})
        if not b:
            continue
        if settled.get(bid, 0) >= float(b.get("total", 0)) - 0.01:
            await db.vendor_bills.update_one({"id": bid}, {"$set": {"status": "paid"}})
        else:
            await db.vendor_bills.update_one({"id": bid, "status": "paid"}, {"$set": {"status": "unpaid"}})

@api.get("/payments-out/open-bills/{party_id}")
async def open_bills_for_party(party_id: str, user=Depends(get_current_user)):
    """A supplier's purchase bills with a remaining balance — for allocating a payment (with optional TDS adjustment)."""
    bills = await db.vendor_bills.find({"supplier_id": party_id}, {"_id": 0}).to_list(10000)
    settled = await _settled_per_bill()
    out = []
    for b in bills:
        bal = float(b.get("total", 0)) - settled.get(b.get("id", ""), 0)
        if bal > 0.01:
            out.append({"id": b.get("id"), "code": b.get("code"), "date": str(b.get("date", ""))[:10],
                        "total": round(float(b.get("total", 0)), 2), "outstanding": round(bal, 2)})
    out.sort(key=lambda r: r["date"])
    return out

@api.post("/payments-out")
async def create_payment_out(p: PaymentOut, user=Depends(get_current_user)):
    doc = p.model_dump()
    doc["code"] = doc.get("code") or await gen_code("PMT-OUT", "payment_out")
    _recalc_payment_status(doc)
    await db.payments_out.insert_one(doc)
    await _refresh_bill_paid_status([a.get("document_id") for a in (doc.get("allocations") or []) if a.get("document_type") == "vendor_bill"])
    await write_audit(user.get("name", ""), "payment_out_created", "payment_out", doc["id"], {"amount": doc["amount"], "party": doc["party_name"]})
    return doc

@api.get("/payments-out")
async def list_payments_out(user=Depends(get_current_user)):
    return await list_collection(db.payments_out)

@api.put("/payments-out/{pid}")
async def update_payment_out(pid: str, p: PaymentOut, user=Depends(get_current_user)):
    data = p.model_dump(); data.pop("id", None); data.pop("created_at", None)
    _recalc_payment_status(data)
    await db.payments_out.update_one({"id": pid}, {"$set": data})
    await _refresh_bill_paid_status([a.get("document_id") for a in (data.get("allocations") or []) if a.get("document_type") == "vendor_bill"])
    return {"ok": True}

@api.delete("/payments-out/{pid}")
async def del_payment_out(pid: str, user=Depends(require_roles("admin", "accountant", "ca"))):
    await db.payments_out.delete_one({"id": pid})
    return {"ok": True}

# Expense categories
@api.post("/expense-categories")
async def create_expense_category(c: ExpenseCategory, user=Depends(require_roles("admin", "accountant", "ca", "manager"))):
    doc = c.model_dump()
    if await db.expense_categories.find_one({"name": doc["name"]}):
        raise HTTPException(400, "Category already exists")
    await db.expense_categories.insert_one(doc)
    return doc

@api.get("/expense-categories")
async def list_expense_categories(user=Depends(get_current_user)):
    rows = await db.expense_categories.find({}, {"_id": 0}).sort("name", 1).to_list(500)
    # Seed defaults if empty
    if not rows:
        defaults = ["Courier", "Salary", "Rent", "Petrol", "Tea", "Transport", "Labour Bill", "Other"]
        for name in defaults:
            d = ExpenseCategory(name=name).model_dump()
            await db.expense_categories.insert_one(d)
        rows = await db.expense_categories.find({}, {"_id": 0}).sort("name", 1).to_list(500)
    return rows

@api.delete("/expense-categories/{cid}")
async def del_expense_category(cid: str, user=Depends(require_roles("admin", "accountant", "ca"))):
    # Refuse if expenses exist under this category
    if await db.expenses.count_documents({"category_id": cid}) > 0:
        raise HTTPException(400, "Cannot delete: expenses exist under this category")
    await db.expense_categories.delete_one({"id": cid})
    return {"ok": True}

# Expenses
@api.post("/expenses")
async def create_expense(e: Expense, user=Depends(get_current_user)):
    doc = e.model_dump()
    doc["code"] = doc.get("code") or await gen_code("EXP", "expense")
    # Auto-set status from paid_amount
    if doc.get("paid_amount", 0) >= doc.get("amount", 0) - 0.01:
        doc["status"] = "Paid"
    elif doc.get("paid_amount", 0) > 0.01:
        doc["status"] = "Partial"
    else:
        doc["status"] = "Unpaid"
    # Fill category_name
    if doc.get("category_id") and not doc.get("category_name"):
        cat = await db.expense_categories.find_one({"id": doc["category_id"]}, {"_id": 0, "name": 1})
        doc["category_name"] = (cat or {}).get("name", "")
    await db.expenses.insert_one(doc)
    await write_audit(user.get("name", ""), "expense_created", "expense", doc["id"], {"amount": doc["amount"], "category": doc.get("category_name")})
    return doc

@api.get("/expenses")
async def list_expenses_v2(category_id: Optional[str] = None, user=Depends(get_current_user)):
    q = {"category_id": category_id} if category_id else {}
    return await list_collection(db.expenses, q)

@api.put("/expenses/{eid}")
async def update_expense(eid: str, e: Expense, user=Depends(get_current_user)):
    data = e.model_dump(); data.pop("id", None); data.pop("created_at", None)
    if data.get("paid_amount", 0) >= data.get("amount", 0) - 0.01:
        data["status"] = "Paid"
    elif data.get("paid_amount", 0) > 0.01:
        data["status"] = "Partial"
    else:
        data["status"] = "Unpaid"
    await db.expenses.update_one({"id": eid}, {"$set": data})
    return {"ok": True}

@api.delete("/expenses/{eid}")
async def del_expense_v2(eid: str, user=Depends(require_roles("admin", "accountant", "ca"))):
    await db.expenses.delete_one({"id": eid})
    return {"ok": True}

# Dashboard rollup endpoints
@api.get("/dashboard/receivable-payable")
async def dashboard_receivable_payable(user=Depends(get_current_user)):
    """Computes total receivable (from open invoices) and payable (from open vendor bills)."""
    # Receivable: invoices where status != 'paid'
    open_invoices = await db.invoices.find({"status": {"$ne": "paid"}}, {"_id": 0, "id": 1, "total": 1, "customer_id": 1, "outstanding": 1}).to_list(10000)
    # Subtract allocated payments
    payments_in = await db.payments_in.find({}, {"_id": 0, "allocations": 1}).to_list(10000)
    allocated_per_inv: Dict[str, float] = {}
    for p in payments_in:
        for a in (p.get("allocations") or []):
            if a.get("document_type") == "invoice":
                allocated_per_inv[a["document_id"]] = allocated_per_inv.get(a["document_id"], 0) + float(a.get("amount") or 0) + float(a.get("tds_amount") or 0)
    receivable_total = 0.0
    receivable_parties = set()
    for inv in open_invoices:
        # Prefer Vyapar's own live balance (ground truth, captures credit notes/returns/etc.) when present;
        # fall back to total-minus-allocations for invoices created outside the Vyapar import.
        if inv.get("outstanding") is not None:
            outstanding = float(inv.get("outstanding") or 0)
        else:
            outstanding = float(inv.get("total", 0)) - allocated_per_inv.get(inv.get("id", ""), 0)
        if outstanding > 0.01:
            receivable_total += outstanding
            if inv.get("customer_id"):
                receivable_parties.add(inv["customer_id"])

    # Payable: vendor bills where balance > 0
    bills = await db.vendor_bills.find({}, {"_id": 0, "total": 1, "supplier_id": 1, "status": 1, "id": 1, "outstanding": 1}).to_list(10000)
    payments_out = await db.payments_out.find({}, {"_id": 0, "allocations": 1}).to_list(10000)
    allocated_per_bill: Dict[str, float] = {}
    for p in payments_out:
        for a in (p.get("allocations") or []):
            if a.get("document_type") == "vendor_bill":
                allocated_per_bill[a["document_id"]] = allocated_per_bill.get(a["document_id"], 0) + float(a.get("amount") or 0) + float(a.get("tds_amount") or 0)
    payable_total = 0.0
    payable_parties = set()
    for b in bills:
        if b.get("status") == "paid":
            continue
        if b.get("outstanding") is not None:
            outstanding = float(b.get("outstanding") or 0)
        else:
            outstanding = float(b.get("total", 0)) - allocated_per_bill.get(b.get("id", ""), 0)
        if outstanding > 0.01:
            payable_total += outstanding
            if b.get("supplier_id"):
                payable_parties.add(b["supplier_id"])

    # Party opening balances carried from Vyapar import (not represented by open documents)
    opening_recv = 0.0
    for d in await db.customers.find({"opening_balance": {"$exists": True, "$ne": 0}},
                                     {"_id": 0, "opening_balance": 1}).to_list(5000):
        opening_recv += float(d.get("opening_balance") or 0)
    opening_pay = 0.0
    for d in await db.suppliers.find({"opening_balance": {"$exists": True, "$ne": 0}},
                                     {"_id": 0, "opening_balance": 1}).to_list(5000):
        opening_pay += float(d.get("opening_balance") or 0)
    receivable_total += opening_recv
    payable_total += opening_pay

    return {
        "receivable_total": round(receivable_total, 2),
        "receivable_parties_count": len(receivable_parties),
        "payable_total": round(payable_total, 2),
        "payable_parties_count": len(payable_parties),
        "opening_receivable": round(opening_recv, 2),
        "opening_payable": round(opening_pay, 2),
    }

@api.get("/reports/overdue")
async def overdue_invoices(user=Depends(get_current_user)):
    """Open invoices past their due date, with days overdue and outstanding amount — for payment follow-ups."""
    today = datetime.utcnow().date()
    invs = await db.invoices.find({"status": {"$ne": "paid"}}, {"_id": 0}).to_list(10000)
    payments_in = await db.payments_in.find({}, {"_id": 0, "allocations": 1}).to_list(10000)
    alloc: Dict[str, float] = {}
    for p in payments_in:
        for a in (p.get("allocations") or []):
            if a.get("document_type") == "invoice":
                alloc[a["document_id"]] = alloc.get(a["document_id"], 0) + float(a.get("amount") or 0) + float(a.get("tds_amount") or 0)
    rows = []
    total = 0.0
    for inv in invs:
        outstanding = float(inv.get("total", 0)) - alloc.get(inv.get("id", ""), 0)
        if outstanding <= 0.01:
            continue
        due = str(inv.get("due_date", "") or "")[:10]
        days = None
        try:
            if due:
                days = (today - datetime.strptime(due, "%Y-%m-%d").date()).days
        except Exception:
            days = None
        if days is None or days <= 0:
            continue   # not overdue yet
        rows.append({
            "id": inv.get("id"), "code": inv.get("code"), "customer_name": inv.get("customer_name"),
            "customer_id": inv.get("customer_id"), "total": round(float(inv.get("total", 0)), 2),
            "outstanding": round(outstanding, 2), "due_date": due, "days_overdue": days,
            "phone": inv.get("customer_phone", "") or inv.get("phone", ""),
        })
        total += outstanding
    rows.sort(key=lambda r: r["days_overdue"], reverse=True)
    return {"overdue": rows, "total_overdue": round(total, 2), "count": len(rows)}

# ---------------------------------------------------------------------------
# AI Assistant (read-only) — answers questions from a live data snapshot.
# ---------------------------------------------------------------------------
ASSISTANT_SYSTEM = (
    "You are ARIA, the AI assistant inside Denplex ERP (a precision-machining / jigs & fixtures company). "
    "If greeted or asked who you are, introduce yourself briefly as ARIA. "
    "Answer the user's question ONLY from the JSON data snapshot provided in the user message. "
    "Be concise and direct. Show money in Indian rupees (₹) with thousands separators. "
    "If the snapshot does not contain the answer, say you don't have that data yet rather than guessing. "
    "Never invent numbers, customers, or invoices. You are READ-ONLY: if asked to create, edit, send, or delete "
    "anything, explain that you can only answer questions for now and they should use the relevant ERP screen."
)

class AssistantIn(BaseModel):
    question: str
    history: list = []

async def _assistant_snapshot():
    invs = await db.invoices.find({}, {"_id": 0}).to_list(20000)
    bills = await db.vendor_bills.find({}, {"_id": 0}).to_list(20000)
    items = await db.items.find({}, {"_id": 0}).to_list(20000)
    settled = await _settled_per_invoice()
    today = datetime.utcnow().date()
    recv = 0.0; overdue = []; cust_tot: Dict[str, float] = {}
    for inv in invs:
        out = float(inv.get("total", 0)) - settled.get(inv.get("id", ""), 0)
        if out > 0.01:
            recv += out
        cust_tot[inv.get("customer_name", "?")] = cust_tot.get(inv.get("customer_name", "?"), 0) + float(inv.get("total", 0))
        due = str(inv.get("due_date", "") or "")[:10]
        if out > 0.01 and due:
            try:
                d = (today - datetime.strptime(due, "%Y-%m-%d").date()).days
                if d > 0:
                    overdue.append({"code": inv.get("code"), "customer": inv.get("customer_name"), "outstanding": round(out, 2), "days_overdue": d})
            except Exception:
                pass
    overdue.sort(key=lambda r: r["days_overdue"], reverse=True)
    recent = [{"code": i.get("code"), "customer": i.get("customer_name"), "date": str(i.get("date", ""))[:10],
               "total": round(float(i.get("total", 0)), 2), "status": i.get("status")}
              for i in sorted(invs, key=lambda x: str(x.get("date", "")), reverse=True)[:10]]
    pouts = await db.payments_out.find({}, {"_id": 0, "allocations": 1}).to_list(20000)
    paid_bill: Dict[str, float] = {}
    for p in pouts:
        for a in (p.get("allocations") or []):
            if a.get("document_type") == "vendor_bill":
                paid_bill[a["document_id"]] = paid_bill.get(a["document_id"], 0) + float(a.get("amount") or 0) + float(a.get("tds_amount") or 0)
    pay = 0.0
    for b in bills:
        if b.get("status") == "paid":
            continue
        out = float(b.get("total", 0)) - paid_bill.get(b.get("id", ""), 0)
        if out > 0.01:
            pay += out
    low = [{"sku": it.get("sku"), "name": it.get("name"), "qty_on_hand": it.get("qty_on_hand"), "reorder_level": it.get("reorder_level")}
           for it in items if float(it.get("reorder_level", 0) or 0) > 0 and float(it.get("qty_on_hand", 0) or 0) <= float(it.get("reorder_level", 0) or 0)][:15]
    await _ensure_accounts_seeded()
    accts = await db.fin_accounts.find({}, {"_id": 0}).to_list(100)
    acc = [{"name": a.get("name"), "type": a.get("type"), "balance": await _account_balance(a)} for a in accts]
    top_cust = sorted(cust_tot.items(), key=lambda kv: -kv[1])[:5]
    return {
        "as_of": now_iso()[:10],
        "counts": {"invoices": len(invs), "vendor_bills": len(bills), "items": len(items),
                   "customers": await db.customers.count_documents({}), "suppliers": await db.suppliers.count_documents({})},
        "sales_total_all_time": round(sum(float(i.get("total", 0)) for i in invs), 2),
        "purchase_total_all_time": round(sum(float(b.get("total", 0)) for b in bills), 2),
        "receivable_total": round(recv, 2), "payable_total": round(pay, 2),
        "overdue_count": len(overdue), "overdue_total": round(sum(o["outstanding"] for o in overdue), 2), "overdue_top": overdue[:10],
        "recent_invoices": recent, "top_customers_by_sales": [{"customer": k, "total": round(v, 2)} for k, v in top_cust],
        "low_stock_items": low, "cash_bank_accounts": acc,
    }

@api.post("/assistant")
async def ai_assistant(inp: AssistantIn, user=Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "AI assistant is not configured. Set ANTHROPIC_API_KEY in the backend environment.")
    import json as _json
    snap = await _assistant_snapshot()
    user_content = f"DATA SNAPSHOT (JSON, as of {snap['as_of']}):\n{_json.dumps(snap, default=str)}\n\nQUESTION: {inp.question}"
    msgs = []
    for h in (inp.history or [])[-6:]:
        if h.get("role") in ("user", "assistant") and h.get("content"):
            msgs.append({"role": h["role"], "content": str(h["content"])[:2000]})
    msgs.append({"role": "user", "content": user_content})
    body = {"model": QC_VISION_MODEL, "max_tokens": 1024, "system": ASSISTANT_SYSTEM, "messages": msgs}
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=60) as cx:
            r = await cx.post(f"{ANTHROPIC_BASE_URL}/v1/messages", json=body, headers=headers)
    except Exception as e:
        raise HTTPException(502, f"Could not reach the AI API: {e}")
    if r.status_code >= 400:
        raise HTTPException(502, f"AI API error {r.status_code}: {r.text[:200]}")
    data = r.json()
    text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
    return {"answer": text.strip() or "I couldn't generate an answer."}

# ---------------------------------------------------------------------------
# Accounting books — Profit & Loss + Balance Sheet (derived from transactions).
# Computed live (not a posted journal); the Balance Sheet balances by treating
# Owner's Capital & Reserves as the balancing figure.
# ---------------------------------------------------------------------------
def _in_period(d, frm, to):
    d = str(d or "")[:10]
    return bool(d) and frm <= d <= to

async def _compute_pnl(frm, to):
    invs = [i for i in await db.invoices.find({}, {"_id": 0}).to_list(50000) if _in_period(i.get("date"), frm, to)]
    bills = [b for b in await db.vendor_bills.find({}, {"_id": 0}).to_list(50000) if _in_period(b.get("date"), frm, to)]
    crns = [c for c in await db.credit_notes.find({}, {"_id": 0}).to_list(50000) if _in_period(c.get("date"), frm, to)]
    drns = [d for d in await db.debit_notes.find({}, {"_id": 0}).to_list(50000) if _in_period(d.get("date"), frm, to)]
    exps = [e for e in await db.expenses.find({}, {"_id": 0}).to_list(50000) if _in_period(e.get("date"), frm, to)]
    sales = sum(float(i.get("subtotal", 0) or 0) for i in invs)
    sales_ret = sum(float(c.get("subtotal", c.get("total", 0)) or 0) for c in crns)
    purchases = sum(float(b.get("subtotal", 0) or 0) for b in bills)
    purch_ret = sum(float(d.get("subtotal", d.get("total", 0)) or 0) for d in drns)
    cats = {c.get("name"): c.get("classification", "indirect") for c in await db.expense_categories.find({}, {"_id": 0}).to_list(2000)}
    direct = 0.0; indirect = 0.0; by_cat: Dict[str, float] = {}
    for e in exps:
        amt = float(e.get("amount", e.get("total", 0)) or 0)
        cn = e.get("category", "Other") or "Other"
        by_cat[cn] = by_cat.get(cn, 0) + amt
        if cats.get(cn) == "direct":
            direct += amt
        else:
            indirect += amt
    net_sales = sales - sales_ret
    net_purch = purchases - purch_ret
    # COGS = Opening Stock + Purchases - Closing Stock
    ob = await _get_setting("opening_balances", {}) or {}
    opening_stock = float(ob.get("opening_stock", 0) or 0)
    items = await db.items.find({}, {"_id": 0, "qty_on_hand": 1, "unit_cost": 1}).to_list(50000)
    closing_stock = sum(float(it.get("qty_on_hand", 0) or 0) * float(it.get("unit_cost", 0) or 0) for it in items)
    cogs = opening_stock + net_purch - closing_stock
    gross = net_sales - cogs - direct
    net = gross - indirect
    return {
        "range": {"from": frm, "to": to},
        "sales": round(net_sales, 2), "sales_gross": round(sales, 2), "sales_returns": round(sales_ret, 2),
        "purchases": round(net_purch, 2), "purchase_returns": round(purch_ret, 2),
        "opening_stock": round(opening_stock, 2), "closing_stock": round(closing_stock, 2), "cogs": round(cogs, 2),
        "direct_expenses": round(direct, 2), "indirect_expenses": round(indirect, 2),
        "expenses_by_category": [{"category": k, "amount": round(v, 2)} for k, v in sorted(by_cat.items(), key=lambda kv: -kv[1])],
        "gross_profit": round(gross, 2), "net_profit": round(net, 2),
    }

@api.get("/reports/pnl")
async def report_pnl(date_from: str = "", date_to: str = "", user=Depends(get_current_user)):
    if not date_from or not date_to:
        d1, d2 = _fy_default_range(); date_from = date_from or d1; date_to = date_to or d2
    return await _compute_pnl(date_from, date_to)

@api.get("/reports/balance-sheet")
async def report_balance_sheet(as_of: str = "", user=Depends(get_current_user)):
    as_of = as_of or now_iso()[:10]
    await _ensure_accounts_seeded()
    accts = await db.fin_accounts.find({}, {"_id": 0}).to_list(200)
    cash_bank = 0.0; acc_list = []
    for a in accts:
        bal = await _account_balance(a); cash_bank += bal
        acc_list.append({"name": a.get("name"), "balance": round(bal, 2)})
    invs = await db.invoices.find({}, {"_id": 0}).to_list(50000)
    sin = await _settled_per_invoice()
    receivable = sum(max(float(i.get("total", 0)) - sin.get(i.get("id", ""), 0), 0) for i in invs if str(i.get("date", ""))[:10] <= as_of)
    receivable += sum(float(d.get("opening_balance") or 0) for d in await db.customers.find(
        {"opening_balance": {"$exists": True, "$ne": 0}}, {"_id": 0, "opening_balance": 1}).to_list(5000))
    bills = await db.vendor_bills.find({}, {"_id": 0}).to_list(50000)
    sbl = await _settled_per_bill()
    payable = sum(max(float(b.get("total", 0)) - sbl.get(b.get("id", ""), 0), 0) for b in bills if str(b.get("date", ""))[:10] <= as_of)
    payable += sum(float(d.get("opening_balance") or 0) for d in await db.suppliers.find(
        {"opening_balance": {"$exists": True, "$ne": 0}}, {"_id": 0, "opening_balance": 1}).to_list(5000))
    items = await db.items.find({}, {"_id": 0}).to_list(50000)
    inventory = sum(float(it.get("qty_on_hand", 0) or 0) * float(it.get("unit_cost", 0) or 0) for it in items)
    out_gst = sum(float(i.get("cgst", 0)) + float(i.get("sgst", 0)) + float(i.get("igst", 0)) for i in invs if str(i.get("date", ""))[:10] <= as_of)
    in_gst = sum(float(b.get("cgst", 0)) + float(b.get("sgst", 0)) + float(b.get("igst", 0)) for b in bills if str(b.get("date", ""))[:10] <= as_of)
    gst_net = out_gst - in_gst
    gst_payable = max(gst_net, 0); gst_credit = max(-gst_net, 0)
    pin = await db.payments_in.find({}, {"_id": 0, "allocations": 1}).to_list(50000)
    pout = await db.payments_out.find({}, {"_id": 0, "allocations": 1}).to_list(50000)
    tds_recv = sum(float(a.get("tds_amount", 0) or 0) for p in pin for a in (p.get("allocations") or []))
    tds_pay = sum(float(a.get("tds_amount", 0) or 0) for p in pout for a in (p.get("allocations") or []))
    pnl_all = await _compute_pnl("0000-01-01", as_of)
    retained = pnl_all["net_profit"]
    ob = await _get_setting("opening_balances", {}) or {}
    opening_capital = float(ob.get("capital", 0) or 0)
    total_assets = round(cash_bank + receivable + inventory + tds_recv + gst_credit, 2)
    total_liab = round(payable + gst_payable + tds_pay, 2)
    # Adjustment is the residual once opening capital + retained earnings are accounted for.
    adjustment = round(total_assets - total_liab - opening_capital - retained, 2)
    return {
        "as_of": as_of,
        "assets": {"cash_bank": round(cash_bank, 2), "accounts": acc_list, "receivable": round(receivable, 2),
                   "inventory": round(inventory, 2), "tds_receivable": round(tds_recv, 2), "gst_credit": round(gst_credit, 2),
                   "total": total_assets},
        "liabilities": {"payable": round(payable, 2), "gst_payable": round(gst_payable, 2), "tds_payable": round(tds_pay, 2),
                        "total": total_liab},
        "equity": {"opening_capital": round(opening_capital, 2), "retained_earnings": round(retained, 2),
                   "adjustment": adjustment, "capital_balancing": adjustment,
                   "total": round(opening_capital + retained + adjustment, 2)},
        "net_profit_to_date": round(retained, 2),
    }

# ---- Statement exports (Excel + PDF) ----
def _stmt_xlsx(title, rows, sub=""):
    from openpyxl import Workbook
    from openpyxl.styles import Font
    wb = Workbook(); ws = wb.active; ws.title = title[:31]
    ws["A1"] = title; ws["A1"].font = Font(bold=True, size=14)
    if sub:
        ws["A2"] = sub; ws["A2"].font = Font(size=9, color="888888")
    r = 4
    for label, val, bold in rows:
        ws.cell(r, 1, label)
        if bold:
            ws.cell(r, 1).font = Font(bold=True)
        if val is not None and val != "":
            c = ws.cell(r, 2, float(val)); c.number_format = '#,##0.00'
            if bold:
                c.font = Font(bold=True)
        r += 1
    ws.column_dimensions["A"].width = 44; ws.column_dimensions["B"].width = 18
    buf = io.BytesIO(); wb.save(buf); buf.seek(0); return buf.getvalue()

def _stmt_pdf(title, rows, sub=""):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm, leftMargin=18 * mm, rightMargin=18 * mm)
    ss = getSampleStyleSheet()
    el = [Paragraph(f"<b>{title}</b>", ss["Title"])]
    if sub:
        el.append(Paragraph(sub, ss["Normal"]))
    el.append(Spacer(1, 6 * mm))
    data = [[label, ("" if val is None or val == "" else f"{float(val):,.2f}")] for label, val, bold in rows]
    t = Table(data, colWidths=[115*mm, 50*mm])
    style = [("FONTSIZE", (0, 0), (-1, -1), 10), ("ALIGN", (1, 0), (1, -1), "RIGHT"),
             ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
             ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]
    for i, (label, val, bold) in enumerate(rows):
        if bold:
            style.append(("FONTNAME", (0, i), (-1, i), "Helvetica-Bold"))
            style.append(("TEXTCOLOR", (0, i), (0, i), colors.HexColor("#DC2626")))
    t.setStyle(TableStyle(style))
    el.append(t)
    el.append(Spacer(1, 6 * mm))
    el.append(Paragraph("Computed live from posted documents — Denplex Engineering Company", ss["Italic"]))
    doc.build(el); buf.seek(0); return buf.getvalue()

def _pnl_rows(d):
    rows = [("Sales (net of returns)", d["sales"], False)]
    if d["sales_returns"]:
        rows.append(("Less: Sales returns", d["sales_returns"], False))
    rows.append(("Opening Stock", d.get("opening_stock", 0), False))
    rows.append(("Add: Purchases (net)", d["purchases"], False))
    rows.append(("Less: Closing Stock", d.get("closing_stock", 0), False))
    rows.append(("Cost of Goods Sold", d.get("cogs", d["purchases"]), False))
    if d["direct_expenses"]:
        rows.append(("Less: Direct expenses", d["direct_expenses"], False))
    rows.append(("Gross Profit", d["gross_profit"], True))
    for e in d["expenses_by_category"]:
        rows.append(("   " + e["category"], e["amount"], False))
    rows.append(("Total indirect expenses", d["indirect_expenses"], False))
    rows.append(("Net Profit", d["net_profit"], True))
    return rows

def _bs_rows(d):
    a, l, eq = d["assets"], d["liabilities"], d["equity"]
    rows = [("ASSETS", "", True), ("Cash & Bank", a["cash_bank"], False)]
    for x in a.get("accounts", []):
        rows.append(("   " + str(x.get("name", "")), x.get("balance", 0), False))
    rows.append(("Accounts Receivable", a["receivable"], False))
    rows.append(("Inventory (stock value)", a["inventory"], False))
    if a.get("tds_receivable"):
        rows.append(("TDS Receivable", a["tds_receivable"], False))
    if a.get("gst_credit"):
        rows.append(("GST Credit (ITC)", a["gst_credit"], False))
    rows.append(("Total Assets", a["total"], True))
    rows.append(("", "", False))
    rows.append(("LIABILITIES", "", True))
    rows.append(("Accounts Payable", l["payable"], False))
    if l.get("gst_payable"):
        rows.append(("GST Payable", l["gst_payable"], False))
    if l.get("tds_payable"):
        rows.append(("TDS Payable", l["tds_payable"], False))
    rows.append(("Total Liabilities", l["total"], True))
    rows.append(("EQUITY", "", True))
    rows.append(("Owner's Capital (opening)", eq.get("opening_capital", 0), False))
    rows.append(("Retained Earnings (net profit to date)", eq["retained_earnings"], False))
    if eq.get("adjustment"):
        rows.append(("Adjustment / Suspense", eq["adjustment"], False))
    rows.append(("Total Equity", eq["total"], True))
    rows.append(("Total Liabilities & Equity", round(l["total"] + eq["total"], 2), True))
    return rows

@api.get("/reports/pnl/export")
async def pnl_export(date_from: str = "", date_to: str = "", fmt: str = "xlsx", user=Depends(get_current_user)):
    if not date_from or not date_to:
        d1, d2 = _fy_default_range(); date_from = date_from or d1; date_to = date_to or d2
    d = await _compute_pnl(date_from, date_to)
    rows = _pnl_rows(d); sub = f"Period: {date_from} to {date_to}"
    fname = f"P&L_{date_from}_to_{date_to}"
    if fmt == "pdf":
        return Response(content=_stmt_pdf("Profit & Loss Statement", rows, sub), media_type="application/pdf",
                        headers={"Content-Disposition": f'attachment; filename="{fname}.pdf"'})
    return Response(content=_stmt_xlsx("Profit & Loss", rows, sub),
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": f'attachment; filename="{fname}.xlsx"'})

@api.get("/reports/balance-sheet/export")
async def bs_export(as_of: str = "", fmt: str = "xlsx", user=Depends(get_current_user)):
    d = await report_balance_sheet(as_of=as_of, user=user)
    rows = _bs_rows(d); sub = f"As of {d['as_of']}"
    fname = f"BalanceSheet_{d['as_of']}"
    if fmt == "pdf":
        return Response(content=_stmt_pdf("Balance Sheet", rows, sub), media_type="application/pdf",
                        headers={"Content-Disposition": f'attachment; filename="{fname}.pdf"'})
    return Response(content=_stmt_xlsx("Balance Sheet", rows, sub),
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": f'attachment; filename="{fname}.xlsx"'})

# ---------------------------------------------------------------------------
# GST Reports — GSTR-1 (outward), GSTR-3B (summary), GSTR-2/Purchase register.
# Computed directly from invoices + vendor_bills. Read-only.
# ---------------------------------------------------------------------------
def _fy_default_range():
    today = datetime.utcnow().date()
    y = today.year if today.month >= 4 else today.year - 1
    return f"{y}-04-01", f"{y+1}-03-31"

def _date10(v):
    return str(v or "")[:10]

def _in_range(d, frm, to):
    d = _date10(d)
    return bool(d) and frm <= d <= to

def _line_taxable(l):
    amt = float(l.get("qty", 0)) * float(l.get("rate", 0))
    amt -= amt * float(l.get("discount_pct", 0)) / 100.0
    amt -= float(l.get("discount_amount", 0))
    return amt if amt > 0 else 0.0

async def _gst_collect(frm, to):
    """Return (invoices, vendor_bills) within range."""
    invs = [i for i in await db.invoices.find({}, {"_id": 0}).to_list(50000) if _in_range(i.get("date"), frm, to)]
    bills = [b for b in await db.vendor_bills.find({}, {"_id": 0}).to_list(50000) if _in_range(b.get("date"), frm, to)]
    return invs, bills

def _build_gstr1(invs):
    b2b, b2c_map, hsn_map = [], {}, {}
    for inv in invs:
        itype = inv.get("invoice_type", "gst")
        inter = bool(inv.get("is_interstate"))
        taxable = float(inv.get("subtotal", 0))
        cgst, sgst, igst = float(inv.get("cgst", 0)), float(inv.get("sgst", 0)), float(inv.get("igst", 0))
        row = {
            "invoice_no": inv.get("code", ""), "date": _date10(inv.get("date")),
            "party": inv.get("customer_name", ""), "gstin": inv.get("customer_gstin", "") or "",
            "place_of_supply": inv.get("place_of_supply", ""), "type": itype,
            "taxable": round(taxable, 2), "cgst": round(cgst, 2), "sgst": round(sgst, 2),
            "igst": round(igst, 2), "total": round(float(inv.get("total", 0)), 2),
        }
        if (inv.get("customer_gstin") or "").strip():
            b2b.append(row)
        else:
            key = (round(_max_line_rate(inv), 2), inv.get("place_of_supply", ""))
            agg = b2c_map.setdefault(key, {"rate": key[0], "place_of_supply": key[1], "taxable": 0, "cgst": 0, "sgst": 0, "igst": 0, "total": 0, "count": 0})
            for k, v in (("taxable", taxable), ("cgst", cgst), ("sgst", sgst), ("igst", igst), ("total", float(inv.get("total", 0)))):
                agg[k] += v
            agg["count"] += 1
        # HSN summary (skip non-gst/export tax)
        for l in (inv.get("lines") or []):
            tx = _line_taxable(l)
            rate = float(l.get("gst_rate", 0)) if itype == "gst" else 0.0
            tax = tx * rate / 100.0
            hk = (l.get("hsn", "") or "", rate)
            h = hsn_map.setdefault(hk, {"hsn": hk[0], "rate": rate, "qty": 0, "taxable": 0, "cgst": 0, "sgst": 0, "igst": 0})
            h["qty"] += float(l.get("qty", 0)); h["taxable"] += tx
            if inter: h["igst"] += tax
            else: h["cgst"] += tax / 2; h["sgst"] += tax / 2
    for d in list(b2c_map.values()) + list(hsn_map.values()):
        for k in ("taxable", "cgst", "sgst", "igst", "total", "qty"):
            if k in d: d[k] = round(d[k], 2)
    return {"b2b": b2b, "b2c": list(b2c_map.values()), "hsn": list(hsn_map.values())}

def _max_line_rate(inv):
    rates = [float(l.get("gst_rate", 0)) for l in (inv.get("lines") or [])]
    return max(rates) if rates else 0.0

def _sum_docs(docs):
    t = {"taxable": 0, "cgst": 0, "sgst": 0, "igst": 0, "total": 0}
    for d in docs:
        t["taxable"] += float(d.get("subtotal", 0)); t["cgst"] += float(d.get("cgst", 0))
        t["sgst"] += float(d.get("sgst", 0)); t["igst"] += float(d.get("igst", 0))
        t["total"] += float(d.get("total", 0))
    return {k: round(v, 2) for k, v in t.items()}

@api.get("/reports/gst/summary")
async def gst_summary(date_from: str = "", date_to: str = "", user=Depends(get_current_user)):
    frm, to = (date_from or "", date_to or "")
    if not frm or not to:
        d1, d2 = _fy_default_range(); frm = frm or d1; to = to or d2
    invs, bills = await _gst_collect(frm, to)
    taxable_invs = [i for i in invs if i.get("invoice_type", "gst") == "gst"]
    export_invs = [i for i in invs if i.get("invoice_type") == "export"]
    nongst_invs = [i for i in invs if i.get("invoice_type") == "non_gst"]
    gstr1 = _build_gstr1(invs)
    outward = _sum_docs(taxable_invs)
    inward = _sum_docs(bills)
    out_tax = round(outward["cgst"] + outward["sgst"] + outward["igst"], 2)
    itc = round(inward["cgst"] + inward["sgst"] + inward["igst"], 2)
    gstr3b = {
        "outward": outward, "export_value": round(sum(float(i.get("total", 0)) for i in export_invs), 2),
        "nongst_value": round(sum(float(i.get("total", 0)) for i in nongst_invs), 2),
        "inward_itc": inward, "output_tax": out_tax, "input_tax_credit": itc,
        "net_payable": round(out_tax - itc, 2),
    }
    gstr2 = [{
        "bill_no": b.get("code", ""), "date": _date10(b.get("date")), "party": b.get("supplier_name", ""),
        "gstin": b.get("supplier_gstin", "") or "", "taxable": round(float(b.get("subtotal", 0)), 2),
        "cgst": round(float(b.get("cgst", 0)), 2), "sgst": round(float(b.get("sgst", 0)), 2),
        "igst": round(float(b.get("igst", 0)), 2), "total": round(float(b.get("total", 0)), 2),
    } for b in bills]
    return {
        "range": {"from": frm, "to": to},
        "counts": {"invoices": len(invs), "bills": len(bills)},
        "gstr1": gstr1, "gstr3b": gstr3b, "gstr2": gstr2,
    }

@api.get("/reports/gst/export.xlsx")
async def gst_export_xlsx(date_from: str = "", date_to: str = "", user=Depends(get_current_user)):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    data = await gst_summary(date_from, date_to, user)
    frm, to = data["range"]["from"], data["range"]["to"]
    wb = Workbook(); wb.remove(wb.active)
    hdr = Font(bold=True, color="FFFFFF"); fill = PatternFill("solid", fgColor="DC2626")
    def sheet(name, cols, rows):
        ws = wb.create_sheet(name[:31])
        for c, h in enumerate(cols, 1):
            cell = ws.cell(1, c, h); cell.font = hdr; cell.fill = fill
        for r, row in enumerate(rows, 2):
            for c, v in enumerate(row, 1):
                ws.cell(r, c, v)
        return ws
    g1 = data["gstr1"]
    sheet("GSTR-1 B2B", ["Invoice No", "Date", "Party", "GSTIN", "Place of Supply", "Type", "Taxable", "CGST", "SGST", "IGST", "Total"],
          [[r["invoice_no"], r["date"], r["party"], r["gstin"], r["place_of_supply"], r["type"], r["taxable"], r["cgst"], r["sgst"], r["igst"], r["total"]] for r in g1["b2b"]])
    sheet("GSTR-1 B2C", ["Rate%", "Place of Supply", "Invoices", "Taxable", "CGST", "SGST", "IGST", "Total"],
          [[r["rate"], r["place_of_supply"], r["count"], r["taxable"], r["cgst"], r["sgst"], r["igst"], r["total"]] for r in g1["b2c"]])
    sheet("HSN Summary", ["HSN/SAC", "Rate%", "Qty", "Taxable", "CGST", "SGST", "IGST"],
          [[r["hsn"], r["rate"], r["qty"], r["taxable"], r["cgst"], r["sgst"], r["igst"]] for r in g1["hsn"]])
    b = data["gstr3b"]
    sheet("GSTR-3B", ["Section", "Taxable", "CGST", "SGST", "IGST"],
          [["Outward taxable (3.1a)", b["outward"]["taxable"], b["outward"]["cgst"], b["outward"]["sgst"], b["outward"]["igst"]],
           ["Zero-rated/Export (3.1b)", b["export_value"], 0, 0, 0],
           ["Nil/Non-GST (3.1e)", b["nongst_value"], 0, 0, 0],
           ["Inward ITC (4a)", b["inward_itc"]["taxable"], b["inward_itc"]["cgst"], b["inward_itc"]["sgst"], b["inward_itc"]["igst"]],
           ["Output Tax", "", "", "", b["output_tax"]],
           ["Input Tax Credit", "", "", "", b["input_tax_credit"]],
           ["NET PAYABLE", "", "", "", b["net_payable"]]])
    sheet("GSTR-2 Purchases", ["Bill No", "Date", "Party", "GSTIN", "Taxable", "CGST", "SGST", "IGST", "Total"],
          [[r["bill_no"], r["date"], r["party"], r["gstin"], r["taxable"], r["cgst"], r["sgst"], r["igst"], r["total"]] for r in data["gstr2"]])
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": f'attachment; filename="GST_Report_{frm}_to_{to}.xlsx"'})

@api.get("/dashboard/sales-trend")
async def dashboard_sales_trend(days: int = 30, user=Depends(get_current_user)):
    """Daily sale totals for the last N days. For the Home chart."""
    from datetime import timedelta as _td
    cutoff = (datetime.now(timezone.utc) - _td(days=days)).isoformat()
    invoices = await db.invoices.find({"date": {"$gte": cutoff[:10]}}, {"_id": 0, "date": 1, "total": 1}).to_list(50000)
    by_day: Dict[str, float] = {}
    for inv in invoices:
        d = str(inv.get("date", ""))[:10]
        by_day[d] = by_day.get(d, 0) + float(inv.get("total", 0) or 0)
    return {"days": days, "series": [{"date": k, "total": round(v, 2)} for k, v in sorted(by_day.items())]}

@api.get("/dashboard/shopfloor")
async def dashboard_shopfloor(user=Depends(get_current_user)):
    """M.6 — factory-floor live picture. Workflow stage counts, delayed jobs,
    today dispatches, material shortages. Built from existing WO status data;
    will become more granular once M.5 operation routing lands."""
    from datetime import date as _date
    today_iso = _date.today().isoformat()

    active_wo = await db.work_orders.count_documents({"status": {"$in": ["planned", "in_progress", "qc"]}})
    qc_pending = await db.work_orders.count_documents({"status": "qc"})
    delayed_count = await db.work_orders.count_documents({
        "status": {"$nin": ["completed", "cancelled"]},
        "due_date": {"$lt": today_iso, "$ne": ""}
    })

    stages_pipeline = [
        {"$match": {"status": {"$in": ["planned", "in_progress", "qc", "completed", "on_hold"]}}},
        {"$group": {"_id": "$status", "count": {"$sum": 1}}},
    ]
    stage_rows = await db.work_orders.aggregate(stages_pipeline).to_list(20)
    stage_map = {r["_id"]: r["count"] for r in stage_rows}
    workflow_stages = [
        {"stage": "Planned",     "key": "planned",     "count": stage_map.get("planned", 0),     "color": "slate"},
        {"stage": "In Progress", "key": "in_progress", "count": stage_map.get("in_progress", 0), "color": "blue"},
        {"stage": "QC Hold",     "key": "qc",          "count": stage_map.get("qc", 0),          "color": "amber"},
        {"stage": "On Hold",     "key": "on_hold",     "count": stage_map.get("on_hold", 0),     "color": "red"},
        {"stage": "Completed",   "key": "completed",   "count": stage_map.get("completed", 0),   "color": "emerald"},
    ]

    try:
        challans_today = await db.delivery_challans.count_documents({"date": {"$regex": f"^{today_iso}"}})
    except Exception: challans_today = 0
    try:
        invoices_today = await db.invoices.count_documents({"date": {"$regex": f"^{today_iso}"}})
    except Exception: invoices_today = 0
    dispatches_today = challans_today + invoices_today

    items = await db.items.find(
        {}, {"_id": 0, "id": 1, "name": 1, "stock": 1, "reorder_level": 1, "uom": 1, "sku": 1, "code": 1}
    ).to_list(5000)
    low_stock = [i for i in items if float(i.get("stock", 0) or 0) <= float(i.get("reorder_level", 0) or 0)]

    delayed_list = await db.work_orders.find(
        {"status": {"$nin": ["completed", "cancelled"]}, "due_date": {"$lt": today_iso, "$ne": ""}},
        {"_id": 0, "code": 1, "customer_name": 1, "product": 1, "due_date": 1, "status": 1, "priority": 1, "id": 1}
    ).sort("due_date", 1).to_list(5)

    # M.5 — operation-level live counts (MES). Group active operations by name.
    operation_stages = []
    try:
        op_pipeline = [
            {"$match": {"status": {"$in": ["pending", "running", "hold", "done"]}}},
            {"$group": {
                "_id": "$operation",
                "running": {"$sum": {"$cond": [{"$eq": ["$status", "running"]}, 1, 0]}},
                "pending": {"$sum": {"$cond": [{"$eq": ["$status", "pending"]}, 1, 0]}},
                "hold":    {"$sum": {"$cond": [{"$eq": ["$status", "hold"]}, 1, 0]}},
                "done":    {"$sum": {"$cond": [{"$eq": ["$status", "done"]}, 1, 0]}},
            }},
        ]
        op_rows = await db.wo_operations.aggregate(op_pipeline).to_list(100)
        for r in op_rows:
            running = r.get("running", 0); pending = r.get("pending", 0); hold = r.get("hold", 0)
            operation_stages.append({
                "operation": r["_id"] or "—",
                "running": running, "pending": pending, "hold": hold,
                "done": r.get("done", 0),
                "active": running + pending + hold,
            })
        operation_stages.sort(key=lambda x: x["active"], reverse=True)
    except Exception:
        operation_stages = []

    # Machine utilization — machines actively running an operation vs total active machines
    try:
        running_machines = await db.wo_operations.distinct("machine", {"status": "running", "machine": {"$nin": ["", None]}})
        running_count = len([x for x in running_machines if x])
        machines_total = await db.machines.count_documents({"is_active": {"$ne": False}})
        if machines_total == 0:
            allm = await db.wo_operations.distinct("machine", {"machine": {"$nin": ["", None]}})
            machines_total = len([x for x in allm if x])
        machine_util = round(running_count / machines_total * 100) if machines_total else None
    except Exception:
        running_count = 0; machines_total = 0; machine_util = None

    return {
        "active_wo": active_wo,
        "delayed_jobs": delayed_count,
        "qc_pending": qc_pending,
        "today_dispatches": dispatches_today,
        "material_shortage": len(low_stock),
        "machine_utilization_pct": machine_util,
        "machines_running": running_count,
        "machines_total": machines_total,
        "operation_stages": operation_stages,
        "workflow_stages": workflow_stages,
        "delayed_list": delayed_list,
        "low_stock_top": low_stock[:5],
        "today": today_iso,
    }

@api.get("/parties/{pid}/transactions")
async def party_transactions(pid: str, kind: Optional[str] = None, user=Depends(get_current_user)):
    """Document-centric transaction list for one party (invoices/bills/returns/payments), the
    data behind the Party Details drill-down panel: Type/Number/Date/Total/Balance/Due Date/Status
    per row, newest first. Distinct from /parties/{pid}/statement (a chronological debit/credit
    running-balance ledger) — this mirrors each document's own status instead.

    `kind` ("customer" or "supplier") is an optional hint from the calling page. It matters because
    8 of 302 suppliers share the exact same `id` as a customer record of the same name (real
    dual-role parties Denplex both buys from and sells to, confirmed live) — without a hint we'd
    always resolve those ids to "customer" and silently show the wrong side's ledger when clicked
    from the Suppliers page. If kind is omitted (e.g. old cached frontend) we fall back to the old
    customer-first auto-detect so the endpoint never 404s."""
    customer = await db.customers.find_one({"id": pid}, {"_id": 0})
    supplier = await db.suppliers.find_one({"id": pid}, {"_id": 0})
    if kind == "supplier" and supplier:
        party, resolved_kind = supplier, "supplier"
    elif kind == "customer" and customer:
        party, resolved_kind = customer, "customer"
    elif customer:
        party, resolved_kind = customer, "customer"
    else:
        party, resolved_kind = supplier, "supplier"
    if not party:
        raise HTTPException(404, "Party not found")
    kind = resolved_kind
    # NOTE: Vyapar-imported documents (all 839 invoices / 1822 bills / 532+1032 payments / 8 credit
    # notes / 5 purchase returns, confirmed live) never got a real customer_id/supplier_id/party_id —
    # the importer only wrote a denormalized *_name string. So matching on id alone returns nothing
    # for every legacy record; match on name too (case-insensitive exact) as the primary path, with
    # id kept as an $or arm so any future non-imported document that DOES set the id still matches.
    name = (party.get("name") or "").strip()
    name_rx = {"$regex": f"^{re.escape(name)}$", "$options": "i"} if name else None
    rows = []
    if kind == "customer":
        cust_q = {"$or": [{"customer_id": pid}] + ([{"customer_name": name_rx}] if name_rx else [])}
        pay_q = {"$or": [{"party_id": pid}] + ([{"party_name": name_rx}] if name_rx else [])}
        for inv in await db.invoices.find(cust_q, {"_id": 0}).to_list(5000):
            rows.append({"type": "Sale", "number": inv.get("code"), "date": inv.get("date"),
                         "total": float(inv.get("total", 0) or 0),
                         "balance": float(inv.get("outstanding", inv.get("total", 0)) or 0),
                         "due_date": inv.get("due_date"), "status": inv.get("status"),
                         "doc_type": "invoice", "id": inv.get("id")})
        for cn in await db.credit_notes.find(cust_q, {"_id": 0}).to_list(5000):
            rows.append({"type": "Sale Return", "number": cn.get("code"), "date": cn.get("date"),
                         "total": float(cn.get("total", 0) or 0), "balance": 0,
                         "due_date": None, "status": cn.get("status"),
                         "doc_type": "credit_note", "id": cn.get("id")})
        for p in await db.payments_in.find(pay_q, {"_id": 0}).to_list(5000):
            amt = float(p.get("amount", 0) or 0); alloc = float(p.get("allocated_amount", 0) or 0)
            rows.append({"type": "Payment In", "number": p.get("code"), "date": p.get("date"),
                         "total": amt, "balance": round(amt - alloc, 2),
                         "due_date": None, "status": p.get("status"),
                         "doc_type": "payment_in", "id": p.get("id")})
    else:
        supp_q = {"$or": [{"supplier_id": pid}] + ([{"supplier_name": name_rx}] if name_rx else [])}
        pay_q = {"$or": [{"party_id": pid}] + ([{"party_name": name_rx}] if name_rx else [])}
        for b in await db.vendor_bills.find(supp_q, {"_id": 0}).to_list(5000):
            rows.append({"type": "Purchase", "number": b.get("code"), "date": b.get("date"),
                         "total": float(b.get("total", 0) or 0),
                         "balance": float(b.get("outstanding", b.get("total", 0)) or 0),
                         "due_date": b.get("due_date"), "status": b.get("status"),
                         "doc_type": "vendor_bill", "id": b.get("id")})
        for pr in await db.purchase_returns.find(supp_q, {"_id": 0}).to_list(5000):
            rows.append({"type": "Purchase Return", "number": pr.get("code"), "date": pr.get("date"),
                         "total": float(pr.get("total", 0) or 0), "balance": 0,
                         "due_date": None, "status": pr.get("status"),
                         "doc_type": "purchase_return", "id": pr.get("id")})
        for p in await db.payments_out.find(pay_q, {"_id": 0}).to_list(5000):
            amt = float(p.get("amount", 0) or 0); alloc = float(p.get("allocated_amount", 0) or 0)
            rows.append({"type": "Payment Out", "number": p.get("code"), "date": p.get("date"),
                         "total": amt, "balance": round(amt - alloc, 2),
                         "due_date": None, "status": p.get("status"),
                         "doc_type": "payment_out", "id": p.get("id")})
    rows.sort(key=lambda r: str(r.get("date") or ""), reverse=True)
    return {
        "party": {"id": party.get("id"), "name": party.get("name"), "phone": party.get("phone"),
                  "email": party.get("email"), "gstin": party.get("gstin"), "address": party.get("address"),
                  "contact_person": party.get("contact_person")},
        "kind": kind,
        "transactions": rows,
    }

@api.get("/parties/{pid}/statement")
async def party_statement(pid: str, period: str = "this_year", user=Depends(get_current_user)):
    """Per-party ledger: opening balance + transactions + closing balance."""
    party = await db.customers.find_one({"id": pid}, {"_id": 0}) or await db.suppliers.find_one({"id": pid}, {"_id": 0})
    if not party:
        raise HTTPException(404, "Party not found")
    invoices = await db.invoices.find({"customer_id": pid}, {"_id": 0}).sort("date", 1).to_list(5000)
    pmts = await db.payments_in.find({"party_id": pid}, {"_id": 0}).sort("date", 1).to_list(5000)
    credit_notes = await db.credit_notes.find({"customer_id": pid}, {"_id": 0}).sort("date", 1).to_list(5000) if hasattr(db, "credit_notes") else []
    # Build transaction list with running balance
    txns = []
    for inv in invoices:
        txns.append({"date": inv.get("date"), "type": "Sale", "ref": inv.get("code"), "debit": float(inv.get("total", 0)), "credit": 0})
    for p in pmts:
        txns.append({"date": p.get("date"), "type": "Payment In", "ref": p.get("code"), "debit": 0, "credit": float(p.get("amount", 0))})
    txns.sort(key=lambda t: str(t.get("date", "")))
    running = 0.0
    for t in txns:
        running += t["debit"] - t["credit"]
        t["running"] = round(running, 2)
    return {"party": party, "opening_balance": 0, "transactions": txns, "closing_balance": round(running, 2)}


# ---------------- Proforma Invoice (Phase B.1) ----------------
@api.post("/proforma-invoices")
async def create_proforma(p: ProformaInvoice, user=Depends(get_current_user)):
    doc = p.model_dump()
    doc["code"] = doc.get("code") or await gen_code("PFI", "proforma")
    doc.update(compute_invoice_totals(doc["lines"], doc.get("is_interstate", False)))
    await db.proforma_invoices.insert_one(doc)
    return serialize(doc)

@api.get("/proforma-invoices")
async def list_proforma(user=Depends(get_current_user)):
    return await list_collection(db.proforma_invoices)

@api.put("/proforma-invoices/{pid}")
async def update_proforma(pid: str, p: ProformaInvoice, user=Depends(get_current_user)):
    data = p.model_dump(); data.pop("id", None); data.pop("created_at", None)
    data.update(compute_invoice_totals(data["lines"], data.get("is_interstate", False)))
    await db.proforma_invoices.update_one({"id": pid}, {"$set": data})
    return {"ok": True}

@api.delete("/proforma-invoices/{pid}")
async def del_proforma(pid: str, user=Depends(require_roles("admin", "manager", "accountant", "ca", "sales"))):
    await db.proforma_invoices.delete_one({"id": pid})
    return {"ok": True}

@api.post("/proforma-invoices/{pid}/convert")
async def convert_proforma(pid: str, user=Depends(get_current_user)):
    pf = await db.proforma_invoices.find_one({"id": pid}, {"_id": 0})
    if not pf: raise HTTPException(404, "Proforma not found")
    if pf.get("status") == "converted": raise HTTPException(400, "Already converted")
    inv_data = {k: v for k, v in pf.items() if k not in ("id", "code", "created_at", "status", "converted_invoice_id", "valid_until")}
    inv_data["status"] = "draft"
    inv = Invoice(**inv_data).model_dump()
    inv["code"] = await gen_code("INV", "invoice")
    inv.update(compute_invoice_totals(inv["lines"], inv.get("is_interstate", False)))
    await db.invoices.insert_one(inv)
    await db.proforma_invoices.update_one({"id": pid}, {"$set": {"status": "converted", "converted_invoice_id": inv["id"]}})
    await write_audit(user.get("name", ""), "proforma_converted", "proforma_invoice", pid, {"invoice_id": inv["id"]})
    return {"ok": True, "invoice_id": inv["id"], "invoice_code": inv["code"]}

@api.get("/proforma-invoices/{pid}/pdf")
async def proforma_pdf(pid: str, user=Depends(get_current_user)):
    pf = await db.proforma_invoices.find_one({"id": pid}, {"_id": 0})
    if not pf: raise HTTPException(404, "Not found")
    company = await get_setting("integrations")
    tpl = await _tpl_for("proforma")
    party_extra = {}
    if pf.get("customer_id"):
        c = await db.customers.find_one({"id": pf["customer_id"]}, {"_id": 0}) or {}
        party_extra = {"address": c.get("address",""), "phone": c.get("phone",""), "gstin": c.get("gstin",""), "state": c.get("state","")}
    doc_meta = {
        "due_date": pf.get("valid_until", ""),
        "place_of_supply": pf.get("place_of_supply", ""),
        "po_number": pf.get("po_number", ""),
        "po_date": pf.get("po_date", ""),
        "purchaser_name": pf.get("purchaser_name", ""),
        "payment_mode": pf.get("payment_mode", ""),
        "is_interstate": bool(pf.get("is_interstate")),
    }
    pdf = _build_doc_pdf("Proforma Invoice", pf.get("code", ""), "To", pf.get("customer_name", ""), str(pf.get("date", ""))[:10],
                         pf.get("lines", []), {"total": pf.get("total", 0)},
                         gst_breakup={"cgst": pf.get("cgst", 0), "sgst": pf.get("sgst", 0), "igst": pf.get("igst", 0)},
                         company=company, notes=pf.get("notes", ""), tpl=tpl,
                         party_extra=party_extra, doc_meta=doc_meta,
                         copy_label="PROFORMA INVOICE")
    return Response(content=pdf, media_type="application/pdf",
                    headers={"Content-Disposition": f'inline; filename="{pf.get("code","proforma")}.pdf"'})

# ---------------- Sale Return / Purchase Return (Phase B.2) ----------------
@api.post("/sale-returns")
async def create_sale_return(r: SaleReturn, user=Depends(get_current_user)):
    doc = r.model_dump()
    doc["code"] = doc.get("code") or await gen_code("SR", "sale_return")
    subtotal = sum(float(l.get("qty", 0) or 0) * float(l.get("rate", 0) or 0) for l in doc["lines"])
    gst_total = sum(float(l.get("qty", 0) or 0) * float(l.get("rate", 0) or 0) * float(l.get("gst_rate", 0) or 0) / 100 for l in doc["lines"])
    doc["subtotal"] = subtotal; doc["total"] = subtotal + gst_total
    is_interstate = False
    if doc.get("original_invoice_id"):
        orig = await db.invoices.find_one({"id": doc["original_invoice_id"]}, {"_id": 0}) or {}
        is_interstate = bool(orig.get("is_interstate"))
    if is_interstate:
        doc["igst"] = gst_total; doc["cgst"] = 0; doc["sgst"] = 0
    else:
        doc["cgst"] = gst_total / 2; doc["sgst"] = gst_total / 2; doc["igst"] = 0
    if doc.get("restore_inventory"):
        for line in doc["lines"]:
            code = (line.get("item_code") or "").strip()
            if not code: continue
            await db.items.update_one({"code": code}, {"$inc": {"stock": float(line.get("qty", 0) or 0)}})
    cn_doc = {
        "id": new_id(),
        "code": await gen_code("CN", "credit_note"),
        "customer_id": doc.get("customer_id"),
        "customer_name": doc.get("customer_name"),
        "date": doc.get("date"),
        "lines": [{"description": l.get("description"), "qty": l.get("qty"), "rate": l.get("rate"), "gst_rate": l.get("gst_rate")} for l in doc["lines"]],
        "total": doc["total"],
        "cgst": doc.get("cgst"), "sgst": doc.get("sgst"), "igst": doc.get("igst"),
        "notes": f"Auto-generated from Sale Return {doc['code']}",
        "source_doc_id": doc["id"], "source_doc_type": "sale_return",
        "created_at": now_iso(),
    }
    await db.credit_notes.insert_one(cn_doc)
    doc["credit_note_id"] = cn_doc["id"]
    doc["status"] = "issued"
    await db.sale_returns.insert_one(doc)
    await write_audit(user.get("name", ""), "sale_return_created", "sale_return", doc["id"], {"customer": doc.get("customer_name"), "total": doc["total"]})
    return serialize(doc)

@api.get("/sale-returns")
async def list_sale_returns(user=Depends(get_current_user)):
    return await list_collection(db.sale_returns)

@api.delete("/sale-returns/{rid}")
async def del_sale_return(rid: str, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    r = await db.sale_returns.find_one({"id": rid}, {"_id": 0})
    if r and r.get("credit_note_id"):
        await db.credit_notes.delete_one({"id": r["credit_note_id"]})
    await db.sale_returns.delete_one({"id": rid})
    return {"ok": True}

@api.post("/purchase-returns")
async def create_purchase_return(r: PurchaseReturn, user=Depends(get_current_user)):
    doc = r.model_dump()
    doc["code"] = doc.get("code") or await gen_code("PR", "purchase_return")
    subtotal = sum(float(l.get("qty", 0) or 0) * float(l.get("rate", 0) or 0) for l in doc["lines"])
    gst_total = sum(float(l.get("qty", 0) or 0) * float(l.get("rate", 0) or 0) * float(l.get("gst_rate", 0) or 0) / 100 for l in doc["lines"])
    doc["subtotal"] = subtotal; doc["total"] = subtotal + gst_total
    doc["cgst"] = gst_total / 2; doc["sgst"] = gst_total / 2; doc["igst"] = 0
    if doc.get("reduce_inventory"):
        for line in doc["lines"]:
            code = (line.get("item_code") or "").strip()
            if not code: continue
            await db.items.update_one({"code": code}, {"$inc": {"stock": -float(line.get("qty", 0) or 0)}})
    dn_doc = {
        "id": new_id(),
        "code": await gen_code("DN", "debit_note"),
        "supplier_id": doc.get("supplier_id"),
        "supplier_name": doc.get("supplier_name"),
        "date": doc.get("date"),
        "lines": [{"description": l.get("description"), "qty": l.get("qty"), "rate": l.get("rate"), "gst_rate": l.get("gst_rate")} for l in doc["lines"]],
        "total": doc["total"],
        "notes": f"Auto-generated from Purchase Return {doc['code']}",
        "source_doc_id": doc["id"], "source_doc_type": "purchase_return",
        "created_at": now_iso(),
    }
    await db.debit_notes.insert_one(dn_doc)
    doc["debit_note_id"] = dn_doc["id"]
    doc["status"] = "issued"
    await db.purchase_returns.insert_one(doc)
    await write_audit(user.get("name", ""), "purchase_return_created", "purchase_return", doc["id"], {"supplier": doc.get("supplier_name"), "total": doc["total"]})
    return serialize(doc)

@api.get("/purchase-returns")
async def list_purchase_returns(user=Depends(get_current_user)):
    return await list_collection(db.purchase_returns)

@api.delete("/purchase-returns/{rid}")
async def del_purchase_return(rid: str, user=Depends(require_roles("admin", "manager", "accountant", "ca"))):
    r = await db.purchase_returns.find_one({"id": rid}, {"_id": 0})
    if r and r.get("debit_note_id"):
        await db.debit_notes.delete_one({"id": r["debit_note_id"]})
    await db.purchase_returns.delete_one({"id": rid})
    return {"ok": True}

# ---------------- Universal export (CSV / XLSX) — Phase F ----------------
EXPORT_COLLECTIONS = {
    "invoices": ("invoices", ["code", "date", "customer_name", "total", "status", "po_number"]),
    "proforma-invoices": ("proforma_invoices", ["code", "date", "customer_name", "total", "status"]),
    "quotations": ("quotations", ["code", "date", "customer_name", "total", "status"]),
    "purchase-orders": ("purchase_orders", ["code", "date", "supplier_name", "total", "status"]),
    "vendor-bills": ("vendor_bills", ["code", "date", "supplier_name", "total", "status"]),
    "credit-notes": ("credit_notes", ["code", "date", "customer_name", "total"]),
    "debit-notes": ("debit_notes", ["code", "date", "supplier_name", "total"]),
    "sale-orders": ("sale_orders", ["code", "date", "customer_name", "total", "status"]),
    "delivery-challans": ("delivery_challans", ["code", "date", "customer_name", "total", "status"]),
    "customers": ("customers", ["code", "name", "gstin", "phone", "email", "address"]),
    "suppliers": ("suppliers", ["code", "name", "gstin", "phone", "email", "address"]),
    "items": ("items", ["code", "name", "hsn", "rate", "stock", "reorder_level"]),
    "parts": ("parts", ["part_number", "customer_part_number", "name", "customer_name", "material", "sourcing", "current_revision", "cycle_time_minutes", "weight_kg", "is_active"]),
    "bom": ("bom", ["code", "parent_part_number", "product_name", "bom_type", "revision", "is_active"]),
    "payments-in": ("payments_in", ["code", "date", "party_name", "amount", "payment_type", "status"]),
    "payments-out": ("payments_out", ["code", "date", "party_name", "amount", "payment_type", "status"]),
    "expenses": ("expenses", ["code", "date", "category_name", "party_name", "amount", "paid_amount", "status"]),
    "sale-returns": ("sale_returns", ["code", "date", "customer_name", "total", "original_invoice_code", "status"]),
    "purchase-returns": ("purchase_returns", ["code", "date", "supplier_name", "total", "original_bill_code", "status"]),
}

@api.get("/export/{collection}.{format}")
async def export_collection(collection: str, format: str, q: str = "", date_from: str = "", date_to: str = "", status: str = "", user=Depends(get_current_user)):
    """Export any registered collection as CSV or XLSX. Streams binary back.
    Honors the same filters the list UI shows (search q, date range, status) so a
    download reflects what the user is looking at, not the whole collection."""
    if collection not in EXPORT_COLLECTIONS:
        raise HTTPException(404, f"Unknown collection '{collection}'")
    if format not in ("csv", "xlsx"):
        raise HTTPException(400, "format must be csv or xlsx")
    coll_name, cols = EXPORT_COLLECTIONS[collection]
    rows = await db[coll_name].find({}, {"_id": 0}).sort("created_at", -1).to_list(50000)
    # ---- apply UI filters in-memory (collections vary in shape) ----
    df = (date_from or "").strip()[:10]; dt = (date_to or "").strip()[:10]
    st = (status or "").strip().lower(); ql = (q or "").strip().lower()
    def _keep(r):
        d = str(r.get("date", ""))[:10]
        if df and d and d < df: return False
        if dt and d and d > dt: return False
        if st and str(r.get("status", "")).strip().lower() != st: return False
        if ql:
            hay = " ".join(str(r.get(c, "")) for c in cols).lower()
            if ql not in hay: return False
        return True
    rows = [r for r in rows if _keep(r)]
    if format == "csv":
        import csv as _csv
        buf = io.StringIO()
        w = _csv.DictWriter(buf, fieldnames=cols, extrasaction="ignore")
        w.writeheader()
        for r in rows:
            w.writerow({c: r.get(c, "") for c in cols})
        return Response(content=buf.getvalue(), media_type="text/csv",
                        headers={"Content-Disposition": f'attachment; filename="{collection}.csv"'})
    # xlsx
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise HTTPException(500, "openpyxl not available")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = collection[:31]
    # Header row
    ws.append([c.replace("_", " ").title() for c in cols])
    header_fill = PatternFill(start_color="DC2626", end_color="DC2626", fill_type="solid")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="left", vertical="center")
    for r in rows:
        ws.append([r.get(c, "") for c in cols])
    # Auto-fit-ish: widen columns based on header
    for i, c in enumerate(cols, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = max(15, len(c) + 4)
    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    return Response(content=buf.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": f'attachment; filename="{collection}.xlsx"'})

# ---------------- Global search ----------------
SEARCH_TARGETS = [
    # (collection, type_label, frontend_route, [search_fields], label_field, sub_field)
    ("parts", "Part", "/app/parts", ["part_number", "customer_part_number", "name", "customer_name"], "part_number", "name"),
    ("items", "Item", "/app/inventory", ["sku", "name", "hsn"], "name", "sku"),
    ("customers", "Customer", "/app/customers", ["name", "gstin", "phone", "email"], "name", "gstin"),
    ("suppliers", "Supplier", "/app/suppliers", ["name", "gstin", "phone", "email"], "name", "gstin"),
    ("invoices", "Invoice", "/app/invoices", ["code", "customer_name", "po_number"], "code", "customer_name"),
    ("quotations", "Quotation", "/app/quotations", ["code", "customer_name"], "code", "customer_name"),
    ("work_orders", "Work Order", "/app/work-orders", ["code", "part_name", "part_number", "customer_name"], "code", "part_name"),
    ("purchase_orders", "Purchase Order", "/app/purchase-orders", ["code", "supplier_name"], "code", "supplier_name"),
    ("vendor_bills", "Purchase Bill", "/app/docs/vendor-bills", ["code", "supplier_name"], "code", "supplier_name"),
]

@api.get("/search")
async def global_search(q: str = "", limit: int = 6, user=Depends(get_current_user)):
    """Lightweight cross-collection search for the global header search box."""
    ql = (q or "").strip()
    if len(ql) < 2:
        return {"results": [], "query": ql}
    rx = {"$regex": re.escape(ql), "$options": "i"}
    lim = max(1, min(int(limit or 6), 15))
    results = []
    for coll, type_label, route, fields, label_f, sub_f in SEARCH_TARGETS:
        try:
            docs = await db[coll].find({"$or": [{f: rx} for f in fields]}, {"_id": 0}).limit(lim).to_list(lim)
            for d in docs:
                results.append({
                    "type": type_label, "route": route, "id": d.get("id", ""),
                    "label": str(d.get(label_f) or d.get("name") or d.get("code") or "—"),
                    "sub": str(d.get(sub_f) or ""),
                })
        except Exception:
            continue
    return {"results": results, "query": ql}

# ---------------- BOM Revisions + Drawings (Phase M.2) ----------------
@api.post("/bom/{bid}/revisions")
async def add_bom_revision(bid: str, rev: BOMRevision, user=Depends(get_current_user)):
    """Promote a new BOM revision. Snapshots the current lines into the history,
    bumps the active revision label, optionally updates the assembly drawing."""
    bom = await db.boms.find_one({"id": bid}, {"_id": 0})
    if not bom: raise HTTPException(404, "BOM not found")
    rev_doc = rev.model_dump()
    if not rev_doc.get("created_by"):
        rev_doc["created_by"] = user.get("name", "")
    # If caller didn't provide a lines snapshot, capture current BOM lines
    if not rev_doc.get("lines_snapshot"):
        rev_doc["lines_snapshot"] = list(bom.get("lines", []))
    history = bom.get("revision_history", []) + [rev_doc]
    update = {
        "revision_history": history,
        "revision": rev_doc["revision"],
    }
    if rev_doc.get("drawing_pdf_b64"):
        update["drawing_pdf_b64"] = rev_doc["drawing_pdf_b64"]
        update["drawing_filename"] = rev_doc.get("drawing_filename", "")
    await db.boms.update_one({"id": bid}, {"$set": update})
    await write_audit(user.get("name", ""), "bom_revision_added", "bom", bid,
                      {"code": bom.get("code"), "revision": rev_doc["revision"], "reason": rev_doc.get("change_reason")})
    return {"ok": True, "revision": rev_doc["revision"]}

@api.get("/bom/{bid}/drawing")
async def download_bom_drawing(bid: str, revision: Optional[str] = None, user=Depends(get_current_user)):
    """Stream the BOM's assembly drawing PDF (current or a specific historical revision)."""
    bom = await db.boms.find_one({"id": bid}, {"_id": 0})
    if not bom: raise HTTPException(404, "BOM not found")
    if revision:
        rev = next((r for r in bom.get("revision_history", []) if r.get("revision") == revision), None)
        if not rev: raise HTTPException(404, f"Revision {revision} not found")
        b64 = rev.get("drawing_pdf_b64", "")
        filename = rev.get("drawing_filename") or f"{bom.get('code')}_{revision}.pdf"
    else:
        b64 = bom.get("drawing_pdf_b64", "")
        filename = bom.get("drawing_filename") or f"{bom.get('code')}.pdf"
    if not b64: raise HTTPException(404, "No drawing on file for this BOM/revision")
    try:
        data = base64.b64decode(b64.split(",", 1)[1] if b64.startswith("data:") else b64)
    except Exception:
        raise HTTPException(500, "Failed to decode drawing")
    return Response(content=data, media_type="application/pdf",
                    headers={"Content-Disposition": f'inline; filename="{filename}"'})

# ---------------- BOM auto-extraction from uploaded files (Phase M.3b) ----------------
@api.post("/bom/extract")
async def bom_extract(file: UploadFile = File(...), user=Depends(get_current_user)):
    """Extract candidate BOM lines from an uploaded file.
    Supports: PDF (assembly drawings with parts list), STEP/STP (CAD component tree),
    XLSX/XLS/CSV (BOM exports from SolidWorks/Solid Edge/any tool).
    Returns a list of candidates the user can tick to add to a BOM."""
    name = (file.filename or "").lower()
    raw = await file.read()
    candidates = []
    notes = []

    if name.endswith(".pdf"):
        try:
            import pdfplumber
            buf = io.BytesIO(raw)
            with pdfplumber.open(buf) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables() or []
                    for t_idx, table in enumerate(tables):
                        if not table or len(table) < 2: continue
                        header = [str(c or "").strip().lower() for c in table[0]]
                        def col(*aliases):
                            for a in aliases:
                                for i, h in enumerate(header):
                                    if a in h: return i
                            return -1
                        col_no = col("item no", "sno", "sl no", "#", "no.")
                        col_part = col("part no", "part #", "drawing", "part number")
                        col_desc = col("description", "item name", "part name", "name")
                        col_qty = col("qty", "quantity", "qty.")
                        col_mat = col("material", "matl")
                        if col_desc < 0 and col_part < 0:
                            continue
                        for row in table[1:]:
                            if not row or all((c or "").strip() == "" for c in row): continue
                            cells = [str(c or "").strip() for c in row]
                            def get(i): return cells[i] if 0 <= i < len(cells) else ""
                            qty_raw = get(col_qty).replace(",", "")
                            try: qty_val = float(qty_raw) if qty_raw else 1.0
                            except: qty_val = 1.0
                            cand = {
                                "source": f"PDF page {page_num} - table {t_idx + 1}",
                                "part_number": get(col_part),
                                "name": get(col_desc),
                                "qty": qty_val,
                                "uom": "Nos",
                                "material": get(col_mat),
                                "sourcing_guess": "manufactured",
                            }
                            if cand["name"] or cand["part_number"]:
                                candidates.append(cand)
            if not candidates:
                notes.append("No BOM table detected. The PDF may be image-based (scanned) - try uploading the original CAD/Excel BOM, or use AI vision (Phase I.2) once available.")
        except Exception as e:
            raise HTTPException(500, f"PDF parse failed: {e}")

    elif name.endswith(".step") or name.endswith(".stp"):
        try:
            text = raw.decode("utf-8", errors="ignore")
            import re as _re
            seen = set()
            for m in _re.finditer(r"PRODUCT\s*\(\s*'([^']+)'\s*,\s*'([^']*)'", text):
                pname = m.group(1).strip()
                pdesc = m.group(2).strip()
                key = (pname, pdesc)
                if key in seen: continue
                seen.add(key)
                candidates.append({
                    "source": "STEP file - PRODUCT entry",
                    "part_number": pname,
                    "name": pdesc or pname,
                    "qty": 1,
                    "uom": "Nos",
                    "material": "",
                    "sourcing_guess": "manufactured",
                })
            notes.append(f"STEP contained {len(candidates)} unique PRODUCT entries. Note: STEP files don't carry quantities - all defaulted to 1.")
        except Exception as e:
            raise HTTPException(500, f"STEP parse failed: {e}")

    elif name.endswith(".xlsx") or name.endswith(".xls") or name.endswith(".csv"):
        try:
            import pandas as pd
            buf = io.BytesIO(raw)
            if name.endswith(".csv"):
                df = pd.read_csv(buf)
            else:
                df = pd.read_excel(buf)
            cols = [str(c).strip().lower() for c in df.columns]
            def col(*aliases):
                for a in aliases:
                    for i, c in enumerate(cols):
                        if a in c: return df.columns[i]
                return None
            col_part = col("part no", "part #", "part number", "drawing")
            col_desc = col("description", "part name", "item name", "name")
            col_qty = col("qty", "quantity")
            col_mat = col("material", "matl")
            col_unit = col("uom", "unit")
            for _, row in df.iterrows():
                def get(c): return str(row.get(c, "")).strip() if c else ""
                qty_raw = get(col_qty).replace(",", "")
                try: qty_val = float(qty_raw) if qty_raw else 1.0
                except: qty_val = 1.0
                cand = {
                    "source": "Excel/CSV row",
                    "part_number": get(col_part),
                    "name": get(col_desc),
                    "qty": qty_val,
                    "uom": get(col_unit) or "Nos",
                    "material": get(col_mat),
                    "sourcing_guess": "manufactured",
                }
                if cand["name"] or cand["part_number"]:
                    candidates.append(cand)
            notes.append(f"Detected columns mapped: part={col_part}, desc={col_desc}, qty={col_qty}, material={col_mat}")
        except Exception as e:
            raise HTTPException(500, f"Excel/CSV parse failed: {e}")

    elif any(name.endswith(ext) for ext in (".sldprt", ".sldasm", ".par", ".asm", ".ipt", ".prt")):
        return {
            "candidates": [],
            "notes": [
                f"Proprietary CAD format ({name.split('.')[-1].upper()}) - requires the source software to read.",
                "Workaround: open the file in SolidWorks/Solid Edge - File menu - 'Save BOM as Excel' - upload that Excel here.",
                "Direct CAD-API integration is on the Phase E roadmap (SolidWorks PDM / Solid Edge connectors).",
            ],
        }

    else:
        raise HTTPException(400, f"Unsupported file type: {name}. Accepted: PDF, STEP/STP, XLSX/XLS/CSV.")

    return {"candidates": candidates, "notes": notes, "count": len(candidates)}

# ---------------- BOM Hierarchy explosion (Phase M.3) ----------------
@api.get("/bom/by-part/{part_id}")
async def bom_by_part(part_id: str, user=Depends(get_current_user)):
    """Find the active BOM for a given part."""
    bom = await db.boms.find_one({"parent_part_id": part_id, "is_active": True, "is_default": True}, {"_id": 0})
    if not bom:
        bom = await db.boms.find_one({"parent_part_id": part_id, "is_active": True}, {"_id": 0})
    if not bom:
        raise HTTPException(404, "No BOM found for this part")
    return bom

@api.get("/bom/{bid}/explode")
async def bom_explode(bid: str, levels: int = 3, user=Depends(get_current_user)):
    """Recursively flatten a BOM down to `levels` deep.
    Returns a tree of {part, qty, level, sourcing, children: [...]}.
    Useful for full material requirements planning."""
    root = await db.boms.find_one({"id": bid}, {"_id": 0})
    if not root:
        raise HTTPException(404, "BOM not found")

    async def expand(bom: Dict[str, Any], level: int, qty_multiplier: float) -> List[Dict[str, Any]]:
        out = []
        for line in bom.get("lines", []):
            effective_qty = float(line.get("qty", 0) or 0) * qty_multiplier * (1 + float(line.get("scrap_factor_pct", 0) or 0) / 100)
            entry = {
                "level": level,
                "line_seq": bom.get("lines", []).index(line),
                "part_id": line.get("component_part_id") or line.get("item_id"),
                "part_number": line.get("component_part_number") or line.get("item_name"),
                "part_name": line.get("component_part_name") or line.get("item_name"),
                "qty": effective_qty,
                "uom": line.get("uom", "pcs"),
                "scrap_factor_pct": line.get("scrap_factor_pct", 0),
                "sourcing": line.get("sourcing", ""),
                "children": [],
            }
            # If this component itself has a BOM and we haven't exhausted levels, recurse
            if level < levels and line.get("component_part_id"):
                sub_bom = await db.boms.find_one({"parent_part_id": line["component_part_id"], "is_active": True, "is_default": True}, {"_id": 0})
                if not sub_bom:
                    sub_bom = await db.boms.find_one({"parent_part_id": line["component_part_id"], "is_active": True}, {"_id": 0})
                if sub_bom:
                    entry["children"] = await expand(sub_bom, level + 1, effective_qty)
            out.append(entry)
        return out

    tree = await expand(root, 1, 1.0)
    return {
        "bom": {"id": root["id"], "code": root.get("code"), "parent_part_number": root.get("parent_part_number"), "product_name": root.get("product_name"), "revision": root.get("revision")},
        "levels": levels,
        "lines": tree,
    }

# ---------------- Part Master (Phase M.1) ----------------
@api.post("/parts")
async def create_part(p: PartMaster, user=Depends(get_current_user)):
    doc = p.model_dump()
    doc["drawing_pdf_b64"] = await _drive_offload(doc.get("drawing_pdf_b64"), doc.get("drawing_filename") or f"{doc.get('part_number','part')}.pdf", "application/pdf", "Drawings")
    doc["step_file_b64"] = await _drive_offload(doc.get("step_file_b64"), doc.get("step_filename") or f"{doc.get('part_number','part')}.step", "application/octet-stream", "STEP Files")
    # Auto-snapshot current files into a revision entry if revision is set
    if doc.get("current_revision") and not doc.get("revisions"):
        rev = {
            "revision": doc["current_revision"],
            "effective_date": doc.get("created_at") or now_iso(),
            "change_reason": "Initial release",
            "drawing_pdf_b64": doc.get("drawing_pdf_b64", ""),
            "step_file_b64": doc.get("step_file_b64", ""),
            "drawing_filename": doc.get("drawing_filename", ""),
            "step_filename": doc.get("step_filename", ""),
            "created_by": user.get("name", ""),
            "notes": "",
        }
        doc["revisions"] = [rev]
    if await db.parts.find_one({"part_number": doc["part_number"]}):
        raise HTTPException(400, f"Part number {doc['part_number']} already exists")
    await db.parts.insert_one(doc)
    await write_audit(user.get("name", ""), "part_created", "part", doc["id"], {"part_number": doc["part_number"]})
    return serialize(doc)

class PartCandidate(BaseModel):
    """Lightweight candidate row from BOM extraction — used by bulk-from-candidates."""
    part_number: Optional[str] = ""
    name: Optional[str] = ""
    description: Optional[str] = ""
    material: Optional[str] = ""
    sourcing_guess: Optional[str] = "manufactured"

@api.post("/parts/bulk-from-candidates")
async def bulk_parts_from_candidates(items: List[PartCandidate], user=Depends(get_current_user)):
    """Given a list of extracted candidates (from BOM file extract), match each to an existing Part
    or create a new one. Returns a list mirroring input order with the resolved part_id and status."""
    result = []
    created = 0
    matched = 0
    skipped = 0
    for item in items:
        pn = (item.part_number or "").strip()
        nm = (item.name or item.description or "").strip()
        # Derive a part_number if missing — use sanitized name; fall back to AUTO-<id>
        if not pn:
            if nm:
                pn = "EXT-" + "".join(c if c.isalnum() else "-" for c in nm.upper())[:40].strip("-")
            else:
                pn = "AUTO-" + new_id()[:8].upper()
        if not pn:
            skipped += 1
            result.append({"part_id": "", "part_number": "", "status": "skipped", "reason": "empty"})
            continue
        # Check for existing (case-insensitive match on part_number)
        existing = await db.parts.find_one(
            {"part_number": {"$regex": f"^{re.escape(pn)}$", "$options": "i"}},
            {"_id": 0, "id": 1, "part_number": 1}
        )
        if existing:
            result.append({
                "part_id": existing["id"],
                "part_number": existing["part_number"],
                "status": "matched",
            })
            matched += 1
            continue
        # Create
        doc = {
            "id": new_id(),
            "part_number": pn,
            "customer_part_number": "",
            "name": nm or pn,
            "description": "",
            "customer_id": "",
            "customer_name": "",
            "material": (item.material or ""),
            "material_grade": "",
            "process": [],
            "cycle_time_minutes": 0,
            "weight_kg": 0,
            "raw_material_size": "",
            "raw_material_qty_per_part": 0,
            "inspection_plan": "",
            "critical_dimensions": [],
            "tools_required": [],
            "current_revision": "Rev A",
            "revisions": [],
            "drawing_pdf_b64": "",
            "step_file_b64": "",
            "drawing_filename": "",
            "step_filename": "",
            "sourcing": item.sourcing_guess or "manufactured",
            "is_active": True,
            "notes": f"Auto-created from BOM extraction at {now_iso()}",
            "created_at": now_iso(),
        }
        await db.parts.insert_one(doc)
        result.append({
            "part_id": doc["id"],
            "part_number": pn,
            "status": "created",
        })
        created += 1
    await write_audit(user.get("name", ""), "parts_bulk_from_candidates", "parts", "",
                      {"created": created, "matched": matched, "skipped": skipped, "total": len(items)})
    return {"items": result, "created": created, "matched": matched, "skipped": skipped}

@api.get("/parts")
async def list_parts(customer_id: Optional[str] = None, search: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if customer_id: q["customer_id"] = customer_id
    if search:
        q["$or"] = [
            {"part_number": {"$regex": search, "$options": "i"}},
            {"customer_part_number": {"$regex": search, "$options": "i"}},
            {"name": {"$regex": search, "$options": "i"}},
        ]
    return await list_collection(db.parts, q)

@api.get("/parts/{pid}")
async def get_part(pid: str, user=Depends(get_current_user)):
    p = await db.parts.find_one({"id": pid}, {"_id": 0})
    if not p: raise HTTPException(404, "Part not found")
    return p

@api.put("/parts/{pid}")
async def update_part(pid: str, p: PartMaster, user=Depends(get_current_user)):
    data = p.model_dump(); data.pop("id", None); data.pop("created_at", None)
    data["drawing_pdf_b64"] = await _drive_offload(data.get("drawing_pdf_b64"), data.get("drawing_filename") or f"{data.get('part_number','part')}.pdf", "application/pdf", "Drawings")
    data["step_file_b64"] = await _drive_offload(data.get("step_file_b64"), data.get("step_filename") or f"{data.get('part_number','part')}.step", "application/octet-stream", "STEP Files")
    await db.parts.update_one({"id": pid}, {"$set": data})
    await write_audit(user.get("name", ""), "part_updated", "part", pid, {"part_number": data.get("part_number")})
    return {"ok": True}

@api.delete("/parts/{pid}")
async def delete_part(pid: str, force: bool = False, user=Depends(require_roles("admin", "manager", "design", "production"))):
    """Soft delete by default (is_active=False). Pass ?force=true for permanent removal."""
    if force:
        p = await db.parts.find_one({"id": pid}, {"_id": 0})
        await _recycle("parts", "Part", p, user)
        await db.parts.delete_one({"id": pid})
        await write_audit(user.get("name", ""), "part_deleted_hard", "part", pid,
                          {"part_number": p.get("part_number") if p else "", "name": p.get("name") if p else ""})
        return {"ok": True, "deleted": True}
    await db.parts.update_one({"id": pid}, {"$set": {"is_active": False}})
    await write_audit(user.get("name", ""), "part_deactivated", "part", pid, {})
    return {"ok": True, "soft_deleted": True}

@api.post("/parts/{pid}/revisions")
async def add_part_revision(pid: str, rev: PartRevision, user=Depends(get_current_user)):
    """Promote a new revision. Mirrors the new revision's files into the part's current_revision fields."""
    p = await db.parts.find_one({"id": pid}, {"_id": 0})
    if not p: raise HTTPException(404, "Part not found")
    rev_doc = rev.model_dump()
    if not rev_doc.get("created_by"):
        rev_doc["created_by"] = user.get("name", "")
    revisions = p.get("revisions", []) + [rev_doc]
    update = {
        "revisions": revisions,
        "current_revision": rev_doc["revision"],
        "drawing_pdf_b64": rev_doc.get("drawing_pdf_b64") or p.get("drawing_pdf_b64", ""),
        "step_file_b64": rev_doc.get("step_file_b64") or p.get("step_file_b64", ""),
        "drawing_filename": rev_doc.get("drawing_filename") or p.get("drawing_filename", ""),
        "step_filename": rev_doc.get("step_filename") or p.get("step_filename", ""),
    }
    await db.parts.update_one({"id": pid}, {"$set": update})
    await write_audit(user.get("name", ""), "part_revision_added", "part", pid,
                      {"part_number": p.get("part_number"), "revision": rev_doc["revision"], "reason": rev_doc.get("change_reason")})
    return {"ok": True, "revision": rev_doc["revision"]}

@api.get("/parts/{pid}/drawing")
async def download_part_drawing(pid: str, revision: Optional[str] = None, user=Depends(get_current_user)):
    """Stream the drawing PDF (current revision by default, or a specific historical revision)."""
    p = await db.parts.find_one({"id": pid}, {"_id": 0})
    if not p: raise HTTPException(404, "Part not found")
    if revision:
        rev = next((r for r in p.get("revisions", []) if r.get("revision") == revision), None)
        if not rev: raise HTTPException(404, f"Revision {revision} not found")
        b64 = rev.get("drawing_pdf_b64", "")
        filename = rev.get("drawing_filename") or f"{p.get('part_number')}_{revision}.pdf"
    else:
        b64 = p.get("drawing_pdf_b64", "")
        filename = p.get("drawing_filename") or f"{p.get('part_number')}.pdf"
    if not b64: raise HTTPException(404, "No drawing on file for this revision")
    try:
        data = await _resolve_b64_or_drive(b64)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(500, "Failed to read drawing")
    return Response(content=data, media_type="application/pdf",
                    headers={"Content-Disposition": f'inline; filename="{filename}"'})

@api.get("/parts/{pid}/step")
async def download_part_step(pid: str, revision: Optional[str] = None, user=Depends(get_current_user)):
    """Stream the STEP / CAD file (current revision by default, or a specific historical revision)."""
    p = await db.parts.find_one({"id": pid}, {"_id": 0})
    if not p: raise HTTPException(404, "Part not found")
    if revision:
        rev = next((r for r in p.get("revisions", []) if r.get("revision") == revision), None)
        if not rev: raise HTTPException(404, f"Revision {revision} not found")
        b64 = rev.get("step_file_b64", "")
        filename = rev.get("step_filename") or f"{p.get('part_number')}_{revision}.step"
    else:
        b64 = p.get("step_file_b64", "")
        filename = p.get("step_filename") or f"{p.get('part_number')}.step"
    if not b64: raise HTTPException(404, "No STEP/CAD file on file for this revision")
    try:
        data = await _resolve_b64_or_drive(b64)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(500, "Failed to read STEP file")
    return Response(content=data, media_type="application/step",
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'})
# ==================== Material States (M.4a) ====================
DEFAULT_MATERIAL_STATES = [
    {"key": "raw",             "name": "Raw Material",     "color": "slate",   "sort_order": 10, "is_system": True, "description": "Material received from supplier, not yet issued."},
    {"key": "wip",             "name": "WIP",              "color": "blue",    "sort_order": 20, "is_system": True, "description": "Issued to shop floor, being machined."},
    {"key": "inspection_hold", "name": "Inspection Hold",  "color": "amber",   "sort_order": 30, "is_system": True, "description": "Off the machine, awaiting QC."},
    {"key": "heat_treatment",  "name": "Heat Treatment",   "color": "amber",   "sort_order": 40, "is_system": True, "description": "Sent out for heat treatment."},
    {"key": "plating",         "name": "Plating",          "color": "amber",   "sort_order": 50, "is_system": True, "description": "Sent out for plating/coating."},
    {"key": "vendor_out",      "name": "Vendor Out",       "color": "amber",   "sort_order": 60, "is_system": True, "description": "Sent for any other vendor job-work."},
    {"key": "fg",              "name": "Finished Goods",   "color": "emerald", "sort_order": 70, "is_system": True, "description": "Passed QC, ready to dispatch."},
    {"key": "rejected",        "name": "Rejected",         "color": "red",     "sort_order": 80, "is_system": True, "description": "Failed QC, on hold for review."},
    {"key": "scrap",           "name": "Scrap",            "color": "red",     "sort_order": 90, "is_system": True, "description": "Written off (cannot recover)."},
]

class MaterialState(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    key: str = ""
    name: str
    description: Optional[str] = ""
    color: str = "slate"
    sort_order: int = 100
    is_active: bool = True
    is_system: bool = False
    created_at: str = Field(default_factory=now_iso)

class MaterialStateMovement(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    item_id: Optional[str] = ""
    item_sku: Optional[str] = ""
    item_name: Optional[str] = ""
    part_id: Optional[str] = ""
    part_number: Optional[str] = ""
    part_name: Optional[str] = ""
    qty: float
    from_state: Optional[str] = ""
    to_state: Optional[str] = ""
    ref_type: Optional[str] = ""
    ref_id: Optional[str] = ""
    ref_code: Optional[str] = ""
    location: Optional[str] = ""
    lot_no: Optional[str] = ""
    note: Optional[str] = ""
    created_by: Optional[str] = ""
    created_at: str = Field(default_factory=now_iso)

async def _seed_material_states_if_needed():
    existing = await db.material_states.count_documents({})
    if existing > 0:
        return
    for s in DEFAULT_MATERIAL_STATES:
        doc = MaterialState(**s).model_dump()
        await db.material_states.insert_one(doc)

def _slugify(text: str) -> str:
    out = (text or "").strip().lower()
    out = "".join(c if c.isalnum() else "_" for c in out)
    while "__" in out:
        out = out.replace("__", "_")
    return out.strip("_")

@api.get("/material-states")
async def list_material_states(user=Depends(get_current_user)):
    await _seed_material_states_if_needed()
    items = await db.material_states.find({}).sort("sort_order", 1).to_list(500)
    return [serialize(s) for s in items]

@api.post("/material-states")
async def create_material_state(s: MaterialState, user=Depends(get_current_user)):
    doc = s.model_dump()
    if not doc.get("key"):
        doc["key"] = _slugify(doc.get("name", ""))
    if not doc.get("key"):
        raise HTTPException(400, "State name required")
    existing = await db.material_states.find_one({"key": doc["key"]})
    if existing:
        raise HTTPException(400, f"State key '{doc['key']}' already exists")
    doc["is_system"] = False
    await db.material_states.insert_one(doc)
    return serialize(doc)

@api.put("/material-states/{sid}")
async def update_material_state(sid: str, s: MaterialState, user=Depends(get_current_user)):
    existing = await db.material_states.find_one({"id": sid})
    if not existing:
        raise HTTPException(404, "State not found")
    data = s.model_dump()
    data.pop("id", None); data.pop("created_at", None); data.pop("is_system", None); data.pop("key", None)
    await db.material_states.update_one({"id": sid}, {"$set": data})
    updated = await db.material_states.find_one({"id": sid}, {"_id": 0})
    return serialize(updated)

@api.delete("/material-states/{sid}")
async def delete_material_state(sid: str, user=Depends(get_current_user)):
    existing = await db.material_states.find_one({"id": sid})
    if not existing:
        raise HTTPException(404, "State not found")
    if existing.get("is_system"):
        raise HTTPException(400, "System state cannot be deleted — deactivate instead")
    in_use = await db.material_state_movements.count_documents({
        "$or": [{"from_state": existing["key"]}, {"to_state": existing["key"]}]
    })
    if in_use > 0:
        raise HTTPException(400, f"State '{existing['name']}' used in {in_use} movements — deactivate instead")
    await db.material_states.delete_one({"id": sid})
    return {"ok": True}

async def record_state_movement(*, item_id="", item_sku="", item_name="",
                                 part_id="", part_number="", part_name="",
                                 qty=0, from_state="", to_state="",
                                 ref_type="", ref_id="", ref_code="",
                                 location="", lot_no="", note="", user_email=""):
    if qty is None or float(qty) <= 0:
        raise ValueError("Quantity must be positive")
    if not from_state and not to_state:
        raise ValueError("Movement must have either from_state or to_state")
    if from_state and from_state == to_state:
        raise ValueError("From and To states must differ")
    mv = MaterialStateMovement(
        item_id=item_id, item_sku=item_sku, item_name=item_name,
        part_id=part_id, part_number=part_number, part_name=part_name,
        qty=float(qty), from_state=from_state, to_state=to_state,
        ref_type=ref_type, ref_id=ref_id, ref_code=ref_code,
        location=location, lot_no=lot_no, note=note,
        created_by=user_email,
    )
    doc = mv.model_dump()
    await db.material_state_movements.insert_one(doc)
    return serialize(doc)

@api.post("/material-states/move")
async def post_state_movement(body: MaterialStateMovement, user=Depends(get_current_user)):
    try:
        email = ""
        if isinstance(user, dict):
            email = user.get("email", "")
        else:
            email = getattr(user, "email", "") or ""
        result = await record_state_movement(
            item_id=body.item_id or "", item_sku=body.item_sku or "", item_name=body.item_name or "",
            part_id=body.part_id or "", part_number=body.part_number or "", part_name=body.part_name or "",
            qty=body.qty, from_state=body.from_state or "", to_state=body.to_state or "",
            ref_type=body.ref_type or "", ref_id=body.ref_id or "", ref_code=body.ref_code or "",
            location=body.location or "", lot_no=body.lot_no or "", note=body.note or "",
            user_email=email,
        )
        return result
    except ValueError as e:
        raise HTTPException(400, str(e))

@api.get("/material-states/movements")
async def list_state_movements(item_id: Optional[str] = None, state: Optional[str] = None,
                                ref_type: Optional[str] = None, limit: int = 300,
                                user=Depends(get_current_user)):
    q = {}
    if item_id: q["item_id"] = item_id
    if state:   q["$or"] = [{"from_state": state}, {"to_state": state}]
    if ref_type: q["ref_type"] = ref_type
    moves = await db.material_state_movements.find(q).sort("created_at", -1).to_list(limit)
    return [serialize(m) for m in moves]

@api.get("/material-states/balance")
async def get_state_balances(user=Depends(get_current_user)):
    pipeline = [
        {"$facet": {
            "inward": [
                {"$match": {"to_state": {"$ne": ""}}},
                {"$group": {
                    "_id": {"item_id": "$item_id", "item_sku": "$item_sku", "item_name": "$item_name", "state": "$to_state"},
                    "qty": {"$sum": "$qty"}
                }}
            ],
            "outward": [
                {"$match": {"from_state": {"$ne": ""}}},
                {"$group": {
                    "_id": {"item_id": "$item_id", "item_sku": "$item_sku", "item_name": "$item_name", "state": "$from_state"},
                    "qty": {"$sum": "$qty"}
                }}
            ]
        }}
    ]
    cursor = db.material_state_movements.aggregate(pipeline)
    result = await cursor.to_list(1)
    if not result:
        return []
    def keyof(r):
        m = r["_id"]
        return (m.get("item_id", ""), m.get("item_sku", ""), m.get("item_name", ""), m.get("state", ""))
    inward = {keyof(r): r for r in result[0].get("inward", [])}
    outward = {keyof(r): r for r in result[0].get("outward", [])}
    all_keys = set(inward.keys()) | set(outward.keys())
    balances = []
    for k in all_keys:
        in_qty = inward[k]["qty"] if k in inward else 0
        out_qty = outward[k]["qty"] if k in outward else 0
        net = in_qty - out_qty
        if abs(net) < 0.0001:
            continue
        balances.append({
            "item_id": k[0],
            "item_sku": k[1],
            "item_name": k[2],
            "state": k[3],
            "qty": net
        })
    return balances

@api.get("/material-states/summary")
async def get_state_summary(user=Depends(get_current_user)):
    balances = await get_state_balances(user)
    summary = {}
    for b in balances:
        summary[b["state"]] = summary.get(b["state"], 0) + b["qty"]
    return summary
# ==================== End Material States (M.4a) ====================


# ---------------- App config ----------------
app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
